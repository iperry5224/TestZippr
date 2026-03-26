"""SOPRA ISSO — Risk Acceptance & Deviation Log (AI-Enhanced)"""
import streamlit as st
from datetime import datetime
import io

from sopra.persistence import _load_json, _save_json, _next_id, load_results_from_file
from sopra.isso.ai_engine import ai_draft_risk_acceptance

def render_risk_acceptance():
    """Risk Acceptance and Deviation Log page with AI-drafted justifications."""
    st.markdown("## ⚖️ Risk Acceptance & Deviation Log")
    st.markdown("Document accepted risks, operational justifications, compensating controls, and approval authority.")

    items = _load_json("risk_acceptances.json", [])
    results = st.session_state.get("opra_assessment_results")
    if not results or not results.get("findings"):
        results = load_results_from_file()
        if results:
            st.session_state.opra_assessment_results = results
    failed_ctrls = []
    if results and results.get("findings"):
        failed_ctrls = [f for f in results["findings"] if f["status"] == "Failed"]

    # New entry form
    with st.expander("➕ Log New Risk Acceptance", expanded=not bool(items)):
        ctrl_options = [f"{f['control_id']} — {f['control_name']}" for f in failed_ctrls] if failed_ctrls else ["No failed controls available"]
        selected = st.selectbox("Control / Finding", ctrl_options, key="ra_ctrl")

        # ── AI Justification Drafting ────────────────────────────────
        if selected and " — " in selected:
            if st.button("🤖 AI Draft Justification", key="ra_ai_draft"):
                cid = selected.split(" — ")[0]
                cname = selected.split(" — ")[1] if " — " in selected else selected
                # Find the matching finding for severity/category
                matched = next((f for f in failed_ctrls if f.get("control_id") == cid), {})
                severity = matched.get("severity", "Medium")
                category = matched.get("category", "")

                with st.spinner("AI is drafting risk acceptance justification..."):
                    draft = ai_draft_risk_acceptance(cid, cname, severity, category)

                if draft:
                    st.session_state["_ra_ai_draft"] = draft
                    st.rerun()

        # Pre-fill from AI if available
        ai_draft = st.session_state.pop("_ra_ai_draft", "")
        default_just = ""
        default_comp = ""
        if ai_draft:
            # Parse sections from AI draft
            sections = ai_draft.split("**")
            just_text = ""
            comp_text = ""
            for i, sec in enumerate(sections):
                if "JUSTIFICATION" in sec.upper():
                    just_text = sections[i + 1].strip().strip(":").strip() if i + 1 < len(sections) else ""
                elif "COMPENSATING" in sec.upper():
                    comp_text = sections[i + 1].strip().strip(":").strip() if i + 1 < len(sections) else ""
            default_just = just_text or ai_draft
            default_comp = comp_text
            st.success("🤖 AI draft loaded. Review and edit before submitting.")

        ra_justification = st.text_area("Operational Justification", value=default_just, placeholder="Why can't this be remediated? What is the operational requirement?", key="ra_just")
        ra_compensating = st.text_area("Compensating Controls", value=default_comp, placeholder="What alternative measures mitigate this risk?", key="ra_comp")
        c1, c2 = st.columns(2)
        with c1:
            ra_authority = st.text_input("Approving Authority", placeholder="Name, Title", key="ra_auth")
        with c2:
            ra_expiry = st.text_input("Acceptance Expiry", value=(datetime.now().replace(year=datetime.now().year + 1)).strftime("%Y-%m-%d"), key="ra_exp")
        ra_level = st.selectbox("Risk Level", ["Low", "Moderate", "High", "Very High"], key="ra_level")

        if st.button("✅ Log Risk Acceptance", type="primary", key="ra_submit"):
            if ra_justification and ra_authority:
                cid = selected.split(" — ")[0] if " — " in selected else selected
                items.append({
                    "id": _next_id(items, "RA"),
                    "control_id": cid,
                    "finding": selected,
                    "justification": ra_justification,
                    "compensating_controls": ra_compensating,
                    "approved_by": ra_authority,
                    "approved_date": datetime.now().strftime("%Y-%m-%d"),
                    "expiry_date": ra_expiry,
                    "risk_level": ra_level,
                    "status": "Active",
                    "created": datetime.now().strftime("%Y-%m-%d"),
                    "ai_assisted": bool(default_just),
                })
                _save_json("risk_acceptances.json", items)
                st.success("Risk acceptance logged.")
                st.rerun()
            else:
                st.warning("Justification and Approving Authority are required.")

    if not items:
        st.info("No risk acceptances logged yet.")
        return

    # Active vs Expired
    today = datetime.now().strftime("%Y-%m-%d")
    active = [i for i in items if i["status"] == "Active" and i.get("expiry_date", "9999") >= today]
    expired = [i for i in items if i.get("expiry_date", "9999") < today or i["status"] == "Expired"]
    for e in expired:
        e["status"] = "Expired"
    _save_json("risk_acceptances.json", items)

    ai_count = len([i for i in items if i.get("ai_assisted")])
    k1, k2, k3 = st.columns(3)
    with k1:
        st.metric("Total Acceptances", len(items))
    with k2:
        st.metric("Active", len(active))
    with k3:
        st.metric("Expired", len(expired), delta=f"-{len(expired)}" if expired else "0", delta_color="inverse")

    if ai_count:
        st.caption(f"🤖 {ai_count} acceptance(s) drafted with AI assistance")

    # Create Acceptance of Risk document (DOCX)
    st.markdown("---")
    if st.button("📄 Create Acceptance of Risk Document (.docx)", use_container_width=True, type="primary", key="ra_export_docx"):
        _export_acceptance_of_risk_docx(items)
    st.markdown("---")
    for item in items:
        level_color = {"Low": "#00d9ff", "Moderate": "#ffc107", "High": "#ff6b6b", "Very High": "#e94560"}.get(item["risk_level"], "#888")
        status_badge = "🟢 Active" if item["status"] == "Active" else "🔴 Expired"
        ai_tag = " 🤖" if item.get("ai_assisted") else ""
        with st.expander(f"{item['id']} — {item['control_id']} | Risk: {item['risk_level']} | {status_badge}{ai_tag}"):
            st.markdown(f"**Finding:** {item['finding']}")
            st.markdown(f"**Justification:** {item['justification']}")
            st.markdown(f"**Compensating Controls:** {item.get('compensating_controls', 'None')}")
            st.markdown(f"**Approved By:** {item['approved_by']} on {item['approved_date']}")
            st.markdown(f"**Expiry:** {item.get('expiry_date', 'N/A')}")
            st.markdown(f"<span style='color:{level_color};font-weight:700;'>Risk Level: {item['risk_level']}</span>", unsafe_allow_html=True)
            if item["status"] == "Active":
                if st.button("🚫 Revoke Acceptance", key=f"ra_revoke_{item['id']}"):
                    item["status"] = "Revoked"
                    _save_json("risk_acceptances.json", items)
                    st.rerun()


def _export_acceptance_of_risk_docx(items):
    """Generate and offer download of Acceptance of Risk document (DOCX)."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Acceptance of Risk")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)
    doc.add_paragraph()
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Total risk acceptances: {len(items)}")
    doc.add_paragraph()

    for item in items:
        doc.add_heading(f"{item['id']} — {item['control_id']}", level=2)
        doc.add_paragraph(f"Finding: {item.get('finding', '')}")
        doc.add_paragraph(f"Risk Level: {item.get('risk_level', 'N/A')}")
        doc.add_paragraph(f"Status: {item.get('status', 'N/A')}")
        doc.add_paragraph(f"Approved by: {item.get('approved_by', '')} on {item.get('approved_date', '')}")
        doc.add_paragraph(f"Expiry: {item.get('expiry_date', 'N/A')}")
        doc.add_paragraph("Operational Justification:")
        doc.add_paragraph(item.get("justification", ""))
        doc.add_paragraph("Compensating Controls:")
        doc.add_paragraph(item.get("compensating_controls", "None"))
        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        "📥 Download Acceptance of Risk (.docx)",
        data=buf.getvalue(),
        file_name=f"SOPRA_Acceptance_of_Risk_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        key="ra_dl_docx",
    )
