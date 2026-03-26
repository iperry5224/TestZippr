"""
Create SAELAR-53 Overview Document and upload to S3
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
import boto3
import os

def create_saelar_overview():
    """Create SAELAR-53 Overview Word Document."""
    
    doc = Document()
    
    # Title
    title = doc.add_heading('SAELAR-53', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Security Assessment Engine for Live AWS Resources')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Document info
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f'NIST 800-53 Rev 5 Compliance Assessment Platform\nVersion 1.0 | {datetime.now().strftime("%B %d, %Y")}')
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(128, 128, 128)
    
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'SAELAR-53 (Security Assessment Engine for Live AWS Resources) is a comprehensive '
        'NIST 800-53 Rev 5 security controls assessment application designed for real-time '
        'automated security evaluation of AWS cloud environments. Built specifically for '
        'Information System Security Officers (ISSOs) and security teams, SAELAR provides '
        'assessment, documentation, and remediation capabilities in a single integrated platform.'
    )
    
    doc.add_paragraph(
        'Unlike traditional compliance tools that rely on static snapshots or manual checklists, '
        'SAELAR performs live queries against AWS resources to provide accurate, real-time '
        'security posture assessments. The platform automatically generates audit-ready '
        'documentation including System Security Plans (SSPs) and Plans of Action & Milestones (POA&Ms).'
    )
    
    # Key Capabilities
    doc.add_heading('Core Capabilities', level=1)
    
    capabilities = [
        ('NIST 800-53 Assessment', 'Automated evaluation of 300+ NIST 800-53 Rev 5 security controls against live AWS resources with real-time status reporting.'),
        ('Risk Score Calculator', 'Quantitative risk scoring based on control failures, vulnerability severity, and exposure metrics to prioritize remediation efforts.'),
        ('SSP Generator', 'Automatic generation of System Security Plans with pre-populated control implementation statements based on assessment results.'),
        ('POA&M Auto-Generation', 'Creates Plans of Action & Milestones for failed or partially implemented controls with suggested remediation timelines.'),
        ('AWS Console Integration', 'Direct connection to AWS environments via secure credential management with real-time resource querying.'),
        ('Gemini AI Remediation', 'AI-powered remediation guidance via Google Gemini CLI integration for intelligent fix recommendations.'),
        ('Document Management', 'S3-integrated documentation storage for SSPs, POA&Ms, assessment evidence, and compliance artifacts.'),
    ]
    
    for cap_name, cap_desc in capabilities:
        p = doc.add_paragraph()
        p.add_run(f'• {cap_name}: ').bold = True
        p.add_run(cap_desc)
    
    # Key Differentiators
    doc.add_heading('Key Differentiators', level=1)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    differentiators = [
        ('Real-Time Assessment', 'Evaluates live AWS resources, not static snapshots or outdated configurations'),
        ('ISSO-Friendly Design', 'Purpose-built for security practitioners with intuitive workflows'),
        ('Quality Control-Ready', 'Produces audit-ready documentation meeting FedRAMP and agency requirements'),
        ('Automated Documentation', 'Eliminates manual SSP and POA&M creation with intelligent auto-generation'),
        ('Integrated Risk Scoring', 'Quantifies security posture with numerical scores for executive reporting'),
        ('AI-Powered Guidance', 'Leverages Gemini AI for context-aware remediation recommendations'),
    ]
    
    for i, (feature, description) in enumerate(differentiators):
        row = table.rows[i]
        row.cells[0].text = feature
        row.cells[1].text = description
        # Bold the feature name
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Architecture
    doc.add_heading('System Architecture', level=1)
    
    doc.add_paragraph(
        'SAELAR-53 is built on a modular architecture using Python and Streamlit, '
        'with direct AWS integration via Boto3. The application consists of the following components:'
    )
    
    components = [
        ('nist_setup.py', 'Main application controller and entry point'),
        ('nist_dashboard.py', 'Dashboard UI components and styling'),
        ('nist_800_53_rev5_full.py', 'NIST 800-53 Rev 5 control assessor engine'),
        ('ssp_generator.py', 'System Security Plan document generator'),
        ('wordy.py', 'Word document conversion utility'),
    ]
    
    for comp_name, comp_desc in components:
        p = doc.add_paragraph()
        p.add_run(f'• {comp_name}: ').bold = True
        p.add_run(comp_desc)
    
    # Functional Tabs
    doc.add_heading('Application Modules', level=1)
    
    modules_table = doc.add_table(rows=7, cols=2)
    modules_table.style = 'Table Grid'
    
    modules = [
        ('Tab', 'Function'),
        ('NIST Assessment', 'Run automated NIST 800-53 control assessments against AWS'),
        ('AWS Console', 'Manage AWS credentials and connection settings'),
        ('Gemini CLI', 'Access AI-powered remediation guidance'),
        ('Risk Calculator', 'Calculate and visualize security risk scores'),
        ('SSP Generator', 'Generate System Security Plans and POA&Ms'),
        ('Documentation', 'Manage and store compliance documents in S3'),
    ]
    
    for i, (tab, function) in enumerate(modules):
        row = modules_table.rows[i]
        row.cells[0].text = tab
        row.cells[1].text = function
        if i == 0:  # Header row
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    # Deployment
    doc.add_heading('Deployment & Launch', level=1)
    
    doc.add_paragraph('SAELAR-53 can be launched using the following PowerShell commands:')
    
    commands = doc.add_paragraph()
    commands.add_run('saelar').bold = True
    commands.add_run(' - Start SAELAR with ngrok tunnel (public URL)\n')
    commands.add_run('saelar -nongrok').bold = True
    commands.add_run(' - Start SAELAR locally only\n')
    commands.add_run('killsaelar').bold = True
    commands.add_run(' - Stop all SAELAR processes')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'The application runs on Streamlit and is accessible via web browser at '
        'http://localhost:8501 by default, or via the ngrok public URL when tunneling is enabled.'
    )
    
    # Security Considerations
    doc.add_heading('Security Considerations', level=1)
    
    security_items = [
        'AWS credentials are stored securely and validated before use',
        'Supports HTTPS via SSL certificate configuration',
        'Session-based authentication for credential management',
        'S3 bucket integration uses IAM-based access controls',
        'All assessment data can be encrypted at rest in S3',
    ]
    
    for item in security_items:
        doc.add_paragraph(f'• {item}')
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f'SAELAR-53 | NIST 800-53 Rev 5 Compliance Platform\n'
        f'Document Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}'
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    output_path = r'C:\Users\iperr\OneDrive\Desktop\AI-Guides\SAELAR-53_Overview.docx'
    doc.save(output_path)
    print(f'[OK] Document saved to: {output_path}')
    
    return output_path


def upload_to_s3(file_path, bucket_name='saegrctest1'):
    """Upload document to S3 bucket."""
    
    try:
        s3_client = boto3.client('s3')
        
        file_name = os.path.basename(file_path)
        s3_key = f'Documentation/{file_name}'
        
        s3_client.upload_file(
            file_path,
            bucket_name,
            s3_key,
            ExtraArgs={
                'Metadata': {
                    'document_type': 'SAELAR Overview',
                    'generated_by': 'SAELAR Document Generator',
                    'generated_date': datetime.now().isoformat()
                }
            }
        )
        
        print(f'[OK] Uploaded to S3: s3://{bucket_name}/{s3_key}')
        return f's3://{bucket_name}/{s3_key}'
        
    except Exception as e:
        print(f'[ERROR] S3 upload failed: {e}')
        return None


if __name__ == '__main__':
    # Create the document
    doc_path = create_saelar_overview()
    
    # Upload to S3
    s3_location = upload_to_s3(doc_path)
    
    if s3_location:
        print(f'\n✅ SAELAR-53 Overview document created and uploaded!')
        print(f'   Local: {doc_path}')
        print(f'   S3: {s3_location}')

