#!/usr/bin/env python3
"""
Create Acceptance of Risk (AOR) document for SAELAR/SOPRA sandbox testing.
Usage: python create_aor_document.py
Output: AOR_SAELAR_SOPRA_Sandbox_YYYYMMDD.docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import datetime
from pathlib import Path

ROOT = Path(__file__).resolve().parent


def create_aor():
    doc = Document()

    # Title
    title = doc.add_heading("Acceptance of Risk (AOR)", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph("Potential Data Leakage/Spillage During Sandbox Testing")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    meta = doc.add_paragraph(
        f"SAELAR & SOPRA Sandbox Environment | {datetime.now().strftime('%B %d, %Y')}"
    )
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    doc.add_paragraph()

    doc.add_paragraph("─" * 70)
    doc.add_paragraph()

    # Purpose
    doc.add_heading("Purpose", level=1)
    doc.add_paragraph(
        "This document describes potential data leakage and spillage risks associated with "
        "testing SAELAR (Security Assessment Engine for Live AWS Resources) and SOPRA "
        "(Security Operations Platform for Risk Assessment) in a controlled sandbox environment. "
        "The organization acknowledges and accepts these risks for the duration of the sandbox testing period."
    )
    doc.add_paragraph()

    # Scope
    doc.add_heading("Scope", level=1)
    doc.add_paragraph(
        "This AOR applies to sandbox testing activities involving SAELAR and SOPRA applications "
        "deployed in non-production environments. It does not cover production deployments."
    )
    doc.add_paragraph()

    # Data Leakage/Spillage Potentials
    doc.add_heading("Potential Data Leakage/Spillage Risks", level=1)

    risks = [
        (
            "AI Model Interaction Exposure (Bedrock/Ollama)",
            "User prompts and AI responses sent to AWS Bedrock or local Ollama may include "
            "sensitive assessment context, system names, account identifiers, or remediation details. "
            "Cloud providers may retain or log this data for model improvement or support. "
            "Even with local Ollama, prompts and responses may be written to disk or logs during testing."
        ),
        (
            "AWS Credentials in Assessment Flows",
            "SAELAR prompts for IAM access keys, secret keys, and account IDs when configuring "
            "AWS-based assessments. These values may be stored in session state, logs, or temporary "
            "files. In a sandbox environment, test credentials should be used rather than production credentials."
        ),
        (
            "Third-Party API and External Service Transmission",
            "Integration with CISA KEV, Nessus, Splunk, and similar services transmits product names, "
            "CVE identifiers, and other context over the network. Third-party providers may log "
            "requests and responses. Data may pass through ngrok or other tunneling services with "
            "their own logging and retention policies."
        ),
        (
            "Exported Documents and Reports",
            "SSPs, POA&Ms, risk assessments, and other exports may contain system names, owners, "
            "findings, remediation details, and organizational information. These files may be "
            "stored on local or shared drives in the sandbox with less restrictive access controls "
            "than production environments."
        ),
        (
            "Authentication and Configuration Data at Rest",
            "User credentials (hashed) and API tokens are stored in configuration files such as "
            ".nist_users.json and application settings. These files may be backed up, copied, or "
            "accessible to other processes or users on the same sandbox host."
        ),
    ]

    for i, (risk_title, risk_desc) in enumerate(risks, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {risk_title}").bold = True
        p.add_run("\n" + risk_desc)
        doc.add_paragraph()

    # Mitigating Factors
    doc.add_heading("Mitigating Factors", level=1)
    doc.add_paragraph(
        "The following compensating controls are in place to offset the risks identified above:"
    )
    doc.add_paragraph()

    mitigations = [
        (
            "AI Model Interaction Exposure",
            "Use air-gapped Ollama (SAELAR_AIRGAPPED=true) for testing when feasible to keep prompts and responses on-premises. When using Bedrock, restrict prompts to synthetic or de-identified data only. Avoid pasting production system names or real account identifiers into AI conversations."
        ),
        (
            "AWS Credentials in Assessment Flows",
            "Dedicated test IAM users and roles with minimal permissions (read-only where possible) will be used for sandbox assessments. No production credentials will be entered. Test credentials will be rotated or deleted after the testing period."
        ),
        (
            "Third-Party API and External Service Transmission",
            "Sandbox or trial API keys will be used for Nessus, Splunk, and similar integrations where available. Network egress from the sandbox can be restricted to approved endpoints. CISA KEV and similar public APIs transmit only product/CVE identifiers, not organization-specific data."
        ),
        (
            "Exported Documents and Reports",
            "Sandbox storage will be isolated from production systems. Exports will use synthetic system names, placeholder owners, and de-identified findings. Access to sandbox file shares will be limited to authorized testers. Exported files will be purged at the end of the testing period."
        ),
        (
            "Authentication and Configuration Data at Rest",
            "Restrictive file permissions (e.g., 600) on credential and configuration files. Sandbox hosts will be dedicated to testing and not used for production workloads. No automated backups of sandbox credential files; if backups exist, they will be excluded from off-site replication."
        ),
    ]

    for i, (risk_ref, mitig_desc) in enumerate(mitigations, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {risk_ref}").bold = True
        p.add_run("\n" + mitig_desc)
        doc.add_paragraph()

    doc.add_paragraph()

    # Acceptance
    doc.add_heading("Acceptance Statement", level=1)
    doc.add_paragraph(
        "The organization accepts the risk of potential data leakage and spillage associated with "
        "the items listed above during sandbox testing of SAELAR and SOPRA."
    )
    doc.add_paragraph(
        "Testing will be conducted with the understanding that no production data or production "
        "credentials will be used. Synthetic, de-identified, or test data only will be employed "
        "where practicable."
    )
    doc.add_paragraph()

    # Sign-off placeholder
    doc.add_paragraph("─" * 70)
    doc.add_paragraph()
    doc.add_paragraph("Authorized by: _________________________").italic = True
    doc.add_paragraph("Title: _________________________").italic = True
    doc.add_paragraph("Date: _________________________").italic = True
    doc.add_paragraph()

    return doc


def main():
    doc = create_aor()
    date_str = datetime.now().strftime("%Y%m%d")
    out_path = ROOT / f"AOR_SAELAR_SOPRA_Sandbox_{date_str}.docx"
    doc.save(out_path)
    print(f"Created: {out_path}")
    return out_path


if __name__ == "__main__":
    main()
