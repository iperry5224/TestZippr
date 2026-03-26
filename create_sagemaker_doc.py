"""Generate AWS SageMaker Integration Guide Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
from pathlib import Path


def _bold_cell(cell):
    """Set cell text to bold."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _set_bold_header(cells, texts):
    """Set header cells with bold text."""
    for i, t in enumerate(texts):
        if i < len(cells):
            cells[i].text = t
            _bold_cell(cells[i])


def create_sagemaker_doc():
    doc = Document()

    # Title
    title = doc.add_heading("Leveraging AWS SageMaker in SAELAR and SOPRA", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    # SageMaker vs Bedrock
    doc.add_heading("SageMaker vs. Bedrock (Current Setup)", level=1)
    doc.add_paragraph(
        "SAELAR and SOPRA currently use AWS Bedrock for AI. Here's how SageMaker differs and when it adds value:"
    )
    table = doc.add_table(rows=5, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    _set_bold_header(hdr, ["Aspect", "Bedrock (Current)", "SageMaker"])
    data = [
        ("Purpose", "Invoke pre-built models via API", "Train, fine-tune, and host custom models"),
        ("Ops Burden", "Zero — fully managed", "You manage endpoints, scaling, updates"),
        ("Flexibility", "Fixed model catalog", "Any custom or fine-tuned model"),
        ("Cost Model", "Pay per token", "Instance hours + inference"),
    ]
    for i, (a, b, c) in enumerate(data, 1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
    doc.add_paragraph()

    # Use Cases
    doc.add_heading("Ways to Leverage SageMaker in SAELAR/SOPRA", level=1)

    doc.add_heading("1. Fine-Tuned Models for Domain-Specific Tasks", level=2)
    doc.add_paragraph(
        "Train or fine-tune a model on your organization's compliance language. Use your own POA&M templates, "
        "SSP narrative formats, and risk acceptance examples so the model speaks your vocabulary."
    )
    doc.add_paragraph("Examples:")
    for ex in [
        "POA&M language from past assessments",
        "SSP control implementation narratives",
        "Risk acceptance justifications and compensating controls",
    ]:
        doc.add_paragraph(ex, style="List Bullet")
    doc.add_paragraph(
        "Integration: Add a SageMaker runtime client; when a specific feature (e.g., POA&M generation) is triggered, "
        "call your SageMaker endpoint instead of Bedrock. Same prompts, different inference backend."
    )
    doc.add_paragraph()

    doc.add_heading("2. SageMaker JumpStart for Model Hosting", level=2)
    doc.add_paragraph(
        "Use JumpStart to deploy foundation models or Hugging Face models that aren't yet in Bedrock. "
        "Useful if you need a specific model, different region, or more control over deployment."
    )
    doc.add_paragraph("Integration: Invoke the JumpStart endpoint instead of Bedrock when that model is preferred.")
    doc.add_paragraph()

    doc.add_heading("3. Custom Compliance / Mapping Models", level=2)
    doc.add_paragraph("Train dedicated models for structured, repetitive compliance tasks:")
    table2 = doc.add_table(rows=4, cols=3)
    table2.style = "Table Grid"
    h2 = table2.rows[0].cells
    _set_bold_header(h2, ["Task", "Training Data", "Benefit"])
    data2 = [
        ("STIG → NIST mapping", "Historical STIG findings + NIST control labels", "More accurate, consistent mapping"),
        ("Evidence sufficiency", "Past evidence + control pass/fail outcomes", "Automated triage of evidence quality"),
        ("Control inheritance", "System boundaries + control types", "Consistent Inherited/Common/System-Specific classification"),
    ]
    for i, row in enumerate(data2, 1):
        for j, val in enumerate(row[:3]):
            table2.rows[i].cells[j].text = val
    doc.add_paragraph()

    doc.add_heading("4. Hybrid Bedrock + SageMaker", level=2)
    doc.add_paragraph("Use the right engine for each use case:")
    table3 = doc.add_table(rows=5, cols=4)
    table3.style = "Table Grid"
    h3 = table3.rows[0].cells
    _set_bold_header(h3, ["Use Case", "Engine", "Why"])
    data3 = [
        ("General chat, remediation suggestions", "Bedrock", "Fast, managed, good default"),
        ("Org-specific POA&M/SSP language", "SageMaker", "Fine-tuned on your templates"),
        ("STIG/NIST mapping", "SageMaker", "Domain-specific classifier"),
        ("Batch SSP narratives (200 controls)", "SageMaker Batch Transform", "Offline bulk processing"),
    ]
    for i, row in enumerate(data3, 1):
        for j, val in enumerate(row):
            table3.rows[i].cells[j].text = val
    doc.add_paragraph()

    doc.add_heading("5. Batch Processing", level=2)
    doc.add_paragraph(
        "Use SageMaker Batch Transform for offline, bulk inference. Generate SSP narratives for 200 controls overnight, "
        "analyze hundreds of evidence artifacts in batch, or map historical findings to controls."
    )
    doc.add_paragraph(
        "Integration: Trigger a batch job; results land in S3. SAELAR/SOPRA can read from S3 or a queue instead of real-time inference."
    )
    doc.add_paragraph()

    doc.add_heading("6. Reinforcement / Continuous Learning", level=2)
    doc.add_paragraph(
        "Learn from user behavior: when users edit AI-generated POA&M entries, accept or reject risk justifications, "
        "or mark evidence as sufficient. Use that feedback to fine-tune models periodically."
    )
    doc.add_paragraph()

    # Integration Options
    doc.add_heading("Practical Integration Options", level=1)

    doc.add_heading("Option A — Add SageMaker as Alternative Backend", level=2)
    doc.add_paragraph(
        "Add a config flag: USE_SAGEMAKER_ENDPOINT = True/False. When True, route AI requests to your SageMaker "
        "endpoint instead of Bedrock. Fall back to Bedrock if the endpoint is unavailable."
    )
    doc.add_paragraph()

    doc.add_heading("Option B — Use SageMaker for Specific Features", level=2)
    table4 = doc.add_table(rows=5, cols=3)
    table4.style = "Table Grid"
    h4 = table4.rows[0].cells
    _set_bold_header(h4, ["Feature", "Current", "SageMaker Option"])
    data4 = [
        ("POA&M Generation", "Bedrock", "Fine-tuned model on your POA&M format"),
        ("STIG Mapping", "Bedrock", "Trained classifier on your mapping history"),
        ("Evidence Analysis", "Bedrock", "Model trained on past audit evidence outcomes"),
        ("Chad (general chat)", "Bedrock", "Keep Bedrock — general purpose"),
    ]
    for i, row in enumerate(data4, 1):
        for j, val in enumerate(row):
            table4.rows[i].cells[j].text = val
    doc.add_paragraph()

    # Considerations
    doc.add_heading("Considerations", level=1)
    doc.add_paragraph(
        "Cost: SageMaker adds endpoint + instance costs; Bedrock is pay-per-token. Ops: SageMaker requires endpoint deployment, "
        "monitoring, scaling. Data: Fine-tuning needs labeled data (POA&M examples, SSP narratives). Latency: SageMaker endpoints "
        "can have cold starts; use provisioned concurrency for real-time UX. Air-gapped / GovCloud: SageMaker can run in VPC with "
        "no public endpoint — useful for air-gapped or FedRAMP High."
    )
    doc.add_paragraph()

    # Recommendation
    doc.add_heading("Recommended Approach", level=1)
    doc.add_paragraph("Start with Bedrock (keep current). Identify 1–2 high-value, repetitive tasks (e.g., STIG mapping, "
                     "POA&M format). Gather training data. Pilot a SageMaker endpoint for that task. If quality improves "
                     "meaningfully, roll out. Use SageMaker only where it clearly beats Bedrock — don't over-engineer.")
    doc.add_paragraph()

    # Footer
    doc.add_paragraph("— End of Document —")

    out = Path(__file__).parent / "AWS_SageMaker_Integration_SAELAR_SOPRA.docx"
    doc.save(out)
    print(f"Created: {out}")
    return out


if __name__ == "__main__":
    create_sagemaker_doc()
