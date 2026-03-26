"""Generate SAELAR Demo Notes Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime


def create_demo_notes():
    doc = Document()

    # Title
    title = doc.add_heading("SAELAR Demo Notes", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()

    # ========== 12 PLATFORM HIGHLIGHTS ==========
    doc.add_heading("12 Platform Highlights", level=1)

    highlights = [
        (
            "Real-Time NIST 800-53 Rev 5 Assessment",
            "Live evaluation of 300+ security controls against your AWS infrastructure—not static snapshots. Covers all 13 cloud-relevant control families with pass/fail status and actionable findings.",
        ),
        (
            "Zero Licensing Cost",
            "Self-hosted platform with no per-resource or per-seat fees. A cost-effective alternative to AWS Audit Manager (~$1.25/resource/month) when NIST 800-53 is your primary framework.",
        ),
        (
            "AWS Security Hub Integration",
            "Imports and maps findings from GuardDuty, Inspector, Macie, IAM Access Analyzer, and Firewall Manager. Automatic mapping of Security Hub findings to NIST 800-53 control families.",
        ),
        (
            "Quantitative Risk Scoring",
            "Risk calculation engine using NIST SP 800-30 Rev 1 methodology (Likelihood × Impact). 5×5 risk matrix with LOW/MEDIUM/HIGH/CRITICAL severity levels for executive reporting.",
        ),
        (
            "Automated SSP & POA&M Generation",
            "One-click generation of System Security Plans and Plans of Action & Milestones. Pre-populated control implementation statements based on assessment results with suggested remediation timelines.",
        ),
        (
            "Automated Vulnerability Management",
            "AI-powered remediation center with attack chain detection, remediation plans, validation scripts, impact analysis, and ticket generation for ServiceNow or Jira.",
        ),
        (
            "ISSO Toolkit",
            "12 specialized tools for ATO support: POA&M Tracker, Risk Acceptance, Evidence Collection, Control Inheritance, Assessment Scheduling, STIG Import, Incident Correlation, and more.",
        ),
        (
            "BOD 22-01 Compliance",
            "CISA Known Exploited Vulnerabilities catalog integration. Flags KEV-listed vulnerabilities for mandatory remediation timelines.",
        ),
        (
            "MITRE ATT&CK & Threat Modeling",
            "Control-to-threat mapping, asset-threat relationship matrix, and MITRE ATT&CK framework integration for threat intelligence correlation.",
        ),
        (
            "FIPS 199 Support",
            "Low (75 controls), Moderate (155), and High (200) baseline filtering. System categorization aligned to FedRAMP and agency requirements.",
        ),
        (
            "Bidirectional Framework Crosswalks",
            "NIST 800-53 Rev 5 and CIS Controls v8 mapping with coverage stats and CSV export. Natural language queries across control frameworks.",
        ),
        (
            "Air-Gapped & Flexible Deployment",
            "Docker containerization and EC2 deployment. Air-gapped mode with local Ollama for fully isolated operation—no internet required. AWS Bedrock AI when connected.",
        ),
    ]

    for i, (name, desc) in enumerate(highlights, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {name} — ").bold = True
        p.add_run(desc)

    doc.add_paragraph()

    # ========== 10 AI AUTOMATIONS FOR ISSO FUNCTIONS ==========
    doc.add_heading("10 Ways We Use AI to Automate ISSO Functions", level=1)

    ai_items = [
        (
            "AI-Powered POA&M Generation",
            "One click generates complete POA&M entries for all failed findings—smart milestones (assess → implement → verify), severity-calibrated due dates (Critical: 14 days, High: 30, Medium: 60, Low: 90), and category-aware responsible party assignments. Saves 4–6 hours per assessment cycle.",
        ),
        (
            "AI-Drafted Risk Acceptance Justifications",
            "Produces AO-ready risk acceptance packages with operational justifications, 3–5 compensating controls, residual risk assessments, and recommended re-evaluation dates. Saves 15–30 minutes per acceptance.",
        ),
        (
            "AI Evidence Sufficiency Analysis",
            "Reviews each evidence artifact against its mapped control—assesses Yes/Partial/No sufficiency, identifies gaps, and recommends improvements. Saves 6–10 hours per audit prep.",
        ),
        (
            "AI Control Inheritance Auto-Classification",
            "Classifies all 200 controls as Inherited, Common, or System-Specific with provider attribution and rationale in seconds. Saves 4–6 hours per system boundary review.",
        ),
        (
            "AI STIG-to-SOPRA Auto-Mapping",
            "Maps DISA STIG or CIS Benchmark scan results to internal controls with confidence scores (High/Medium/Low). Saves 2–4 hours per STIG import.",
        ),
        (
            "AI Incident-to-Finding Correlation",
            "NLP analysis of incident descriptions suggests which failed controls are related, with relevance scores and plain-language explanations. Saves 20–40 minutes per incident.",
        ),
        (
            "AI Natural Language Crosswalk Queries",
            "Ask plain-English questions about control mappings (e.g., 'Which NIST families have the most failures?') and get immediate, data-driven answers. Saves 15–25 minutes per query.",
        ),
        (
            "AI-Written SSP Control Narratives",
            "Generates 150–250 word implementation narratives for all controls using formal RMF language. Output inserted directly into the SSP .docx. Saves 150–300+ hours per ATO cycle.",
        ),
        (
            "AI Remediation Plans & Attack Chain Detection",
            "Generates detailed remediation playbooks per finding, detects multi-control attack chains, provides validation scripts, and performs change impact analysis. Saves 8–15 hours per remediation cycle.",
        ),
        (
            "AI Automated Ticket Generation",
            "Pre-filled ServiceNow or Jira tickets from failed findings with severity-mapped priority, descriptions, and remediation steps. Bulk generation for entire assessment cycles. Saves 2–4 hours per assessment.",
        ),
    ]

    for i, (name, desc) in enumerate(ai_items, 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {name} — ").bold = True
        p.add_run(desc)

    doc.add_paragraph()
    doc.add_paragraph("— End of SAELAR Demo Notes —")
    doc.add_paragraph()

    out_path = r"C:\Users\iperr\TestZippr\SAELAR_Demo_notes.docx"
    doc.save(out_path)
    print(f"Saved: {out_path}")
    return out_path


if __name__ == "__main__":
    create_demo_notes()
