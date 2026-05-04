"""
NIST 800-53 Rev 5 Application Controller
=========================================
Main setup and controller file that orchestrates all application components.
This is the entry point for running the NIST assessment application.

Usage:
    streamlit run nist_setup.py
    
    Or with HTTPS:
    streamlit run nist_setup.py --server.sslCertFile=ssl_certs/streamlit.crt --server.sslKeyFile=ssl_certs/streamlit.key

Air-Gapped Mode:
    Set environment variable: SAELAR_AIRGAPPED=true
    Requires Ollama running locally: ollama serve
"""

import streamlit as st
import base64
import boto3
import json
import os
import requests
from pathlib import Path
from typing import Optional, Tuple
from botocore.exceptions import ClientError, NoCredentialsError

# =============================================================================
# AI CONFIGURATION - Supports both Ollama (air-gapped) and Bedrock (cloud)
# =============================================================================

# Check if running in air-gapped mode
AIRGAPPED_MODE = os.environ.get('SAELAR_AIRGAPPED', 'false').lower() == 'true'

# Ollama configuration (for air-gapped environments)
OLLAMA_BASE_URL = os.environ.get('OLLAMA_BASE_URL', 'http://localhost:11434')
OLLAMA_MODELS = [
    "llama3:8b",           # Llama 3 8B - good balance of speed/quality
    "llama3.1:8b",         # Llama 3.1 8B
    "mistral:7b",          # Mistral 7B - fast and capable
    "mixtral:8x7b",        # Mixtral 8x7B - higher quality, more resources
    "phi3:mini",           # Phi-3 Mini - lightweight
    "gemma:7b",            # Google Gemma 7B
]

# AWS Bedrock models (for cloud environments) - Claude removed per policy
BEDROCK_MODELS = [
    "nvidia.nemotron-nano-12b-v2",                # NVIDIA Nemotron Nano 12B
    "amazon.titan-text-express-v1",               # Titan Express
    "amazon.titan-text-lite-v1",                  # Titan Lite
    "meta.llama3-8b-instruct-v1:0",               # Llama 3
    "mistral.mistral-7b-instruct-v0:2",           # Mistral
]


def check_ollama_available() -> Tuple[bool, list]:
    """
    Check if Ollama is running and what models are available.
    
    Returns:
        Tuple of (is_available: bool, available_models: list)
    """
    try:
        response = requests.get(f"{OLLAMA_BASE_URL}/api/tags", timeout=2)
        if response.status_code == 200:
            models = response.json().get('models', [])
            model_names = [m['name'] for m in models]
            return True, model_names
        return False, []
    except:
        return False, []


def call_ollama(messages: list, system_prompt: str, max_tokens: int = 4096, model: str = None) -> str:
    """
    Call AI via local Ollama instance (fully air-gapped, no internet required).
    
    Args:
        messages: List of conversation messages [{"role": "user/assistant", "content": "..."}]
        system_prompt: System prompt
        max_tokens: Maximum tokens in response
        model: Specific model to use (optional)
    
    Returns:
        AI response text
    """
    is_available, available_models = check_ollama_available()
    
    if not is_available:
        raise Exception("""
⚠️ **Ollama Not Running**

SAELAR is configured for air-gapped mode but Ollama is not available.

**To start Ollama:**
1. Install Ollama from https://ollama.ai
2. Run: `ollama serve`
3. Pull a model: `ollama pull llama3:8b`

**Or disable air-gapped mode:**
- Remove/unset the SAELAR_AIRGAPPED environment variable
- Restart SAELAR
""")
    
    # Find a suitable model
    selected_model = model
    if not selected_model:
        # Try preferred models in order
        for preferred in OLLAMA_MODELS:
            if any(preferred in m for m in available_models):
                selected_model = preferred
                break
        # Fall back to first available
        if not selected_model and available_models:
            selected_model = available_models[0]
    
    if not selected_model:
        raise Exception(f"No Ollama models available. Pull a model with: ollama pull llama3:8b")
    
    # Build the prompt with system context
    full_messages = []
    if system_prompt:
        full_messages.append({"role": "system", "content": system_prompt})
    full_messages.extend(messages)
    
    # Call Ollama API
    try:
        response = requests.post(
            f"{OLLAMA_BASE_URL}/api/chat",
            json={
                "model": selected_model,
                "messages": full_messages,
                "stream": False,
                "options": {
                    "num_predict": max_tokens,
                    "temperature": 0.7
                }
            },
            timeout=120  # Allow up to 2 minutes for response
        )
        
        if response.status_code == 200:
            return response.json().get('message', {}).get('content', 'No response generated')
        else:
            raise Exception(f"Ollama error: {response.status_code} - {response.text}")
            
    except requests.exceptions.Timeout:
        raise Exception("Ollama request timed out. The model may be loading or the query is too complex.")
    except requests.exceptions.ConnectionError:
        raise Exception("Cannot connect to Ollama. Make sure it's running: ollama serve")


def call_bedrock_ai(messages: list, system_prompt: str, max_tokens: int = 4096, region: str = "us-east-1") -> str:
    """
    Call AI via Amazon Bedrock (data stays within AWS).
    Tries multiple models until one works.
    
    Args:
        messages: List of conversation messages [{"role": "user/assistant", "content": "..."}]
        system_prompt: System prompt
        max_tokens: Maximum tokens in response
        region: AWS region for Bedrock
    
    Returns:
        AI response text
    """
    try:
        bedrock = boto3.client(
            service_name='bedrock-runtime',
            region_name=region
        )
        
        # Format messages for Bedrock Converse API
        formatted_messages = []
        for msg in messages:
            formatted_messages.append({
                "role": msg["role"],
                "content": [{"text": msg["content"]}]
            })
        
        # Try each model until one works
        last_error = None
        for model_id in BEDROCK_MODELS:
            try:
                response = bedrock.converse(
                    modelId=model_id,
                    messages=formatted_messages,
                    system=[{"text": system_prompt}],
                    inferenceConfig={
                        "maxTokens": min(max_tokens, 4096),  # Some models have lower limits
                        "temperature": 0.7
                    }
                )
                # If we get here, the call succeeded
                return response['output']['message']['content'][0]['text']
            except ClientError as e:
                last_error = e
                error_code = e.response.get('Error', {}).get('Code', 'Unknown')
                # Continue to next model on these errors
                if error_code in ['ResourceNotFoundException', 'ValidationException', 'AccessDeniedException']:
                    continue
                else:
                    raise
        
        # If we get here, no models worked
        raise Exception(f"No Bedrock models available. Please enable a model in the Bedrock console. Last error: {last_error}")
        
    except NoCredentialsError:
        raise Exception("AWS credentials not configured. Please configure AWS credentials first.")
    except Exception as e:
        error_msg = str(e)
        if "No Bedrock models available" in error_msg:
            raise Exception("""
⚠️ **Bedrock Not Configured**

To keep data within AWS, you need to enable a model in Amazon Bedrock:

1. Go to **AWS Console** → **Amazon Bedrock**
2. Look for **Model catalog** or **Foundation models**
3. Select a model (Titan, Llama, or Mistral)
4. Click **Request access** or **Enable**

**Supported Regions:** us-east-1, us-west-2, eu-west-1, ap-northeast-1

Run this to check available models:
```
aws bedrock list-foundation-models --region us-east-1 --query 'modelSummaries[*].modelId'
```
""")
        raise Exception(f"Bedrock error: {error_msg}")


def get_ai_mode_status() -> dict:
    """
    Get the current AI mode status for display in UI.
    
    Returns:
        dict with 'mode', 'status', 'details' keys
    """
    if AIRGAPPED_MODE:
        ollama_ok, models = check_ollama_available()
        if ollama_ok:
            return {
                "mode": "Air-Gapped (Ollama)",
                "status": "✅ Connected",
                "details": f"Models: {', '.join(models[:3])}{'...' if len(models) > 3 else ''}",
                "color": "#10b981"
            }
        else:
            return {
                "mode": "Air-Gapped (Ollama)",
                "status": "❌ Not Connected",
                "details": "Run: ollama serve",
                "color": "#ef4444"
            }
    else:
        # Check if Ollama is available even in standard mode
        ollama_ok, _ = check_ollama_available()
        if ollama_ok:
            return {
                "mode": "Hybrid (Ollama Primary)",
                "status": "✅ Ollama Active",
                "details": "Bedrock as fallback",
                "color": "#10b981"
            }
        else:
            return {
                "mode": "Cloud (AWS Bedrock)",
                "status": "☁️ Cloud Mode",
                "details": "Internet required",
                "color": "#3b82f6"
            }


def render_ai_mode_indicator():
    """Render the AI mode indicator in the sidebar."""
    status = get_ai_mode_status()
    st.sidebar.markdown(f"""
    <div style="background: linear-gradient(135deg, {status['color']}22, {status['color']}11); 
                border: 1px solid {status['color']}; 
                border-radius: 8px; 
                padding: 0.75rem; 
                margin-bottom: 1rem;">
        <div style="font-weight: 600; color: {status['color']}; font-size: 0.85rem;">
            🤖 AI Mode: {status['mode']}
        </div>
        <div style="color: {status['color']}; font-size: 0.8rem; margin-top: 0.25rem;">
            {status['status']}
        </div>
        <div style="color: #6b7280; font-size: 0.75rem; margin-top: 0.25rem;">
            {status['details']}
        </div>
    </div>
    """, unsafe_allow_html=True)


def call_ai(messages: list, system_prompt: str, max_tokens: int = 4096, region: str = "us-east-1") -> str:
    """
    Unified AI call function - automatically chooses Ollama (air-gapped) or Bedrock (cloud).
    
    In air-gapped mode (SAELAR_AIRGAPPED=true):
        - Uses Ollama exclusively
        - No internet connectivity required
    
    In standard mode:
        - Tries Ollama first (if available)
        - Falls back to Bedrock
    
    Args:
        messages: List of conversation messages
        system_prompt: System prompt
        max_tokens: Maximum tokens
        region: AWS region for Bedrock (only used if Bedrock is selected)
    
    Returns:
        AI response text
    """
    if AIRGAPPED_MODE:
        # Air-gapped mode: Ollama only, no fallback to cloud
        return call_ollama(messages, system_prompt, max_tokens)
    
    # Standard mode: Try Ollama first, fall back to Bedrock
    ollama_available, _ = check_ollama_available()
    
    if ollama_available:
        try:
            return call_ollama(messages, system_prompt, max_tokens)
        except Exception as e:
            # Log the error but try Bedrock
            print(f"Ollama failed, falling back to Bedrock: {e}")
    
    # Fall back to Bedrock
    return call_bedrock_ai(messages, system_prompt, max_tokens, region)


# =============================================================================
# SPLASH PAGE ASSETS
# =============================================================================

# Use relative path or environment variable for containerization
import platform
LOGO_PATH = Path(os.environ.get('SAELAR_LOGO_PATH', Path(__file__).parent / 'assets' / 'saelar_logo.png'))
DISCLAIMER_TEXT = "Use of this tool implies affirmation and compliance with the terms and conditions outlined in governing policy."


# =============================================================================
# AWS CREDENTIALS MANAGEMENT
# =============================================================================

# Path to store credentials securely
# For containers: mount ~/.aws or set environment variables
AWS_CREDENTIALS_DIR = Path(os.environ.get('AWS_SHARED_CREDENTIALS_FILE', Path.home() / ".aws")).parent if os.environ.get('AWS_SHARED_CREDENTIALS_FILE') else Path.home() / ".aws"
AWS_CREDENTIALS_FILE = AWS_CREDENTIALS_DIR / "credentials"
AWS_CONFIG_FILE = AWS_CREDENTIALS_DIR / "config"
# SAELAR config - use app directory or environment variable
SAELAR_CREDENTIALS_FILE = Path(os.environ.get('SAELAR_CONFIG_FILE', Path(__file__).parent / '.saelar_aws_config.json'))


def validate_aws_credentials(access_key: str, secret_key: str, account_id: str, region: str = "us-east-1") -> Tuple[bool, str]:
    """
    Validate AWS credentials by attempting to get caller identity.
    
    Returns:
        Tuple of (success: bool, message: str)
    """
    try:
        # Create a session with the provided credentials
        session = boto3.Session(
            aws_access_key_id=access_key,
            aws_secret_access_key=secret_key,
            region_name=region
        )
        
        # Try to get caller identity
        sts = session.client('sts')
        identity = sts.get_caller_identity()
        
        # Verify account ID matches
        actual_account_id = identity.get('Account', '')
        
        if account_id and actual_account_id != account_id:
            return False, f"Account ID mismatch. Expected: {account_id}, Got: {actual_account_id}"
        
        return True, f"✅ Successfully authenticated to AWS Account: {actual_account_id}"
        
    except NoCredentialsError:
        return False, "No credentials provided"
    except ClientError as e:
        error_code = e.response.get('Error', {}).get('Code', 'Unknown')
        error_msg = e.response.get('Error', {}).get('Message', str(e))
        return False, f"AWS Error ({error_code}): {error_msg}"
    except Exception as e:
        return False, f"Error: {str(e)}"


def set_aws_credentials(access_key: str, secret_key: str, region: str = "us-east-1"):
    """Set AWS credentials as environment variables for boto3."""
    os.environ['AWS_ACCESS_KEY_ID'] = access_key
    os.environ['AWS_SECRET_ACCESS_KEY'] = secret_key
    os.environ['AWS_DEFAULT_REGION'] = region


def save_aws_credentials_to_file(
    access_key: str, 
    secret_key: str, 
    region: str,
    account_id: str,
    iam_username: str = "",
    iam_password: str = "",
    profile_name: str = "saelar"
) -> Tuple[bool, str]:
    """
    Save AWS credentials to AWS credentials file and SAELAR config.
    This allows them to be used with aws configure and other CLI tools.
    
    Returns:
        Tuple of (success: bool, message: str)
    """
    import configparser
    import json
    
    try:
        # Create .aws directory if it doesn't exist
        AWS_CREDENTIALS_DIR.mkdir(parents=True, exist_ok=True)
        
        # Read existing credentials file or create new
        credentials = configparser.ConfigParser()
        if AWS_CREDENTIALS_FILE.exists():
            credentials.read(AWS_CREDENTIALS_FILE)
        
        # Add/update the profile
        if profile_name not in credentials:
            credentials[profile_name] = {}
        
        credentials[profile_name]['aws_access_key_id'] = access_key
        credentials[profile_name]['aws_secret_access_key'] = secret_key
        
        # Write credentials file
        with open(AWS_CREDENTIALS_FILE, 'w') as f:
            credentials.write(f)
        
        # Read existing config file or create new
        config = configparser.ConfigParser()
        if AWS_CONFIG_FILE.exists():
            config.read(AWS_CONFIG_FILE)
        
        # Add/update the profile config
        profile_section = f"profile {profile_name}" if profile_name != "default" else "default"
        if profile_section not in config:
            config[profile_section] = {}
        
        config[profile_section]['region'] = region
        config[profile_section]['output'] = 'json'
        
        # Write config file
        with open(AWS_CONFIG_FILE, 'w') as f:
            config.write(f)
        
        # Save SAELAR-specific config with additional info (IAM username/password for reference)
        saelar_config = {
            'account_id': account_id,
            'iam_username': iam_username,
            'iam_password': iam_password,  # Note: stored for user reference only
            'access_key_id': access_key,
            'region': region,
            'profile_name': profile_name,
            'configured_at': str(Path.home()),
        }
        
        with open(SAELAR_CREDENTIALS_FILE, 'w') as f:
            json.dump(saelar_config, f, indent=2)
        
        return True, f"Credentials saved to AWS profile '{profile_name}'"
        
    except Exception as e:
        return False, f"Error saving credentials: {str(e)}"


def load_saved_credentials() -> Optional[dict]:
    """Load previously saved SAELAR credentials including secret key from AWS credentials file."""
    import json
    import configparser
    
    try:
        if SAELAR_CREDENTIALS_FILE.exists():
            with open(SAELAR_CREDENTIALS_FILE, 'r') as f:
                creds = json.load(f)
            
            # Load the secret key from AWS credentials file
            profile_name = creds.get('profile_name', 'saelar')
            if AWS_CREDENTIALS_FILE.exists():
                aws_creds = configparser.ConfigParser()
                aws_creds.read(str(AWS_CREDENTIALS_FILE))
                
                # Try the profile name, then 'default' as fallback
                for section in [profile_name, 'saelar', 'default']:
                    if section in aws_creds.sections() or section in aws_creds:
                        try:
                            secret = aws_creds.get(section, 'aws_secret_access_key', fallback='')
                            access = aws_creds.get(section, 'aws_access_key_id', fallback='')
                            if secret:
                                creds['secret_key'] = secret
                            if access and not creds.get('access_key_id'):
                                creds['access_key_id'] = access
                            if creds.get('secret_key'):
                                break
                        except Exception:
                            continue
            
            # Validate we have the required credentials
            if creds.get('access_key_id') and creds.get('secret_key'):
                return creds
            else:
                print(f"Missing credentials: access_key={bool(creds.get('access_key_id'))}, secret={bool(creds.get('secret_key'))}")
                return None
                
    except Exception as e:
        print(f"Error loading credentials: {e}")
    return None


def clear_aws_credentials():
    """Clear AWS credentials from environment."""
    for key in ['AWS_ACCESS_KEY_ID', 'AWS_SECRET_ACCESS_KEY', 'AWS_SESSION_TOKEN']:
        if key in os.environ:
            del os.environ[key]


def render_aws_config_page() -> bool:
    """
    Render AWS credentials configuration page.
    
    Returns:
        True if credentials are configured and valid, False otherwise
    """
    # Get logo
    logo_base64 = get_base64_image(LOGO_PATH)
    
    st.markdown("""
    <style>
        .aws-config-container {
            max-width: 600px;
            margin: 0 auto;
            padding: 2rem;
        }
        .aws-header {
            text-align: center;
            margin-bottom: 2rem;
        }
        .aws-logo {
            max-width: 200px;
            margin-bottom: 1rem;
        }
        .credential-info {
            background: #f0f9ff;
            border: 1px solid #0ea5e9;
            border-radius: 8px;
            padding: 1rem;
            margin-bottom: 1.5rem;
        }
        .credential-warning {
            background: #fef3c7;
            border: 1px solid #f59e0b;
            border-radius: 8px;
            padding: 1rem;
            margin-top: 1rem;
        }
        .credential-success {
            background: #ecfdf5;
            border: 1px solid #10b981;
            border-radius: 8px;
            padding: 1rem;
            margin-top: 1rem;
        }
        .section-divider {
            border-top: 1px solid #e5e7eb;
            margin: 1.5rem 0;
            position: relative;
        }
        .section-divider span {
            background: white;
            padding: 0 0.5rem;
            position: absolute;
            top: -0.7rem;
            left: 50%;
            transform: translateX(-50%);
            color: #9ca3af;
            font-size: 0.85rem;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Check for saved credentials
    saved_creds = load_saved_credentials()
    
    # Center content
    col1, col2, col3 = st.columns([1, 2, 1])
    
    with col2:
        # Logo
        if logo_base64:
            st.markdown(f"""
            <div class="aws-header">
                <img src="data:image/png;base64,{logo_base64}" class="aws-logo" alt="SAELAR Logo"/>
                <h2>☁️ AWS Configuration</h2>
                <p style="color: #64748b;">Configure your AWS credentials to run security assessments</p>
            </div>
            """, unsafe_allow_html=True)
        else:
            st.markdown("""
            <div class="aws-header">
                <h2>☁️ AWS Configuration</h2>
                <p style="color: #64748b;">Configure your AWS credentials to run security assessments</p>
            </div>
            """, unsafe_allow_html=True)
        
        # Show saved credentials option if available
        if saved_creds:
            st.markdown("""
            <div class="credential-success">
                <strong>✅ Previously Saved Credentials Found</strong>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown(f"""
            - **Account ID:** `{saved_creds.get('account_id', 'N/A')}`
            - **IAM User:** `{saved_creds.get('iam_username', 'N/A')}`
            - **Profile:** `{saved_creds.get('profile_name', 'saelar')}`
            - **Region:** `{saved_creds.get('region', 'us-east-1')}`
            """)
            
            col_a, col_b = st.columns(2)
            with col_a:
                if st.button("🔗 Saved Credentials", use_container_width=True):
                    # Set environment variables
                    set_aws_credentials(
                        saved_creds.get('access_key_id', ''),
                        saved_creds.get('secret_key', ''),
                        saved_creds.get('region', 'us-east-1')
                    )
                    st.session_state.aws_configured = True
                    st.session_state.aws_account_id = saved_creds.get('account_id', '')
                    st.session_state.aws_region = saved_creds.get('region', 'us-east-1')
                    st.session_state.aws_iam_username = saved_creds.get('iam_username', '')
                    st.rerun()
                    return True
            with col_b:
                if st.button("🔄 Enter New Credentials", use_container_width=True):
                    saved_creds = None  # Force new entry
                    st.rerun()
            
            st.markdown("---")
        
        # Info box
        st.markdown("""
        <div class="credential-info">
            <strong>ℹ️ AWS Credentials Required</strong><br/>
            To assess your AWS environment, please provide:
            <ul>
                <li><strong>Account ID</strong> - Your 12-digit AWS Account ID</li>
                <li><strong>IAM Username</strong> - Your IAM user name</li>
                <li><strong>IAM Password</strong> - Your IAM console password (stored securely for reference)</li>
                <li><strong>Access Key ID</strong> - From IAM user security credentials</li>
                <li><strong>Secret Access Key</strong> - From IAM user security credentials</li>
            </ul>
        </div>
        """, unsafe_allow_html=True)
        
        with st.form("aws_credentials_form"):
            st.markdown("### 🔐 AWS Account & IAM Credentials")
            
            # AWS Account Info
            account_id = st.text_input(
                "AWS Account ID *",
                placeholder="123456789012",
                max_chars=12,
                help="Your 12-digit AWS Account ID (found in top-right of AWS Console)"
            )
            
            st.markdown('<div class="section-divider"><span>IAM User Info</span></div>', unsafe_allow_html=True)
            
            # IAM User Credentials
            col_user, col_pass = st.columns(2)
            with col_user:
                iam_username = st.text_input(
                    "IAM Username *",
                    placeholder="security-auditor",
                    help="Your IAM username for AWS Console login"
                )
            with col_pass:
                iam_password = st.text_input(
                    "IAM Password *",
                    type="password",
                    placeholder="••••••••••",
                    help="Your IAM password (stored locally for your reference)"
                )
            
            st.markdown('<div class="section-divider"><span>API Access Keys</span></div>', unsafe_allow_html=True)
            
            # API Access Keys
            access_key = st.text_input(
                "Access Key ID *",
                placeholder="AKIAIOSFODNN7EXAMPLE",
                help="Your IAM Access Key ID (starts with AKIA...)"
            )
            
            secret_key = st.text_input(
                "Secret Access Key *",
                type="password",
                placeholder="wJalrXUtnFEMI/K7MDENG/bPxRfiCYEXAMPLEKEY",
                help="Your IAM Secret Access Key"
            )
            
            st.markdown('<div class="section-divider"><span>Configuration</span></div>', unsafe_allow_html=True)
            
            # Region and options
            col_region, col_profile = st.columns(2)
            with col_region:
                region = st.selectbox(
                    "Default Region",
                    options=["us-east-1", "us-east-2", "us-west-1", "us-west-2",
                            "eu-west-1", "eu-west-2", "eu-central-1",
                            "ap-southeast-1", "ap-southeast-2", "ap-northeast-1"],
                    index=0
                )
            with col_profile:
                profile_name = st.text_input(
                    "AWS Profile Name",
                    value="saelar",
                    help="Name for the AWS CLI profile"
                )
            
            save_to_aws = st.checkbox(
                "💾 Save credentials to AWS CLI profile",
                value=True,
                help="Saves credentials to ~/.aws/credentials for use with AWS CLI and other tools"
            )
            
            st.markdown("""
            <div class="credential-warning">
                ⚠️ <strong>Security Note:</strong> Credentials will be saved locally for use with AWS CLI tools.
                Your IAM password is stored only for your reference and is not transmitted.
            </div>
            """, unsafe_allow_html=True)
            
            submitted = st.form_submit_button("🔗 Connect to AWS", use_container_width=True)
            
            if submitted:
                # Validate required fields
                if not account_id or not access_key or not secret_key:
                    st.error("❌ Please fill in Account ID, Access Key ID, and Secret Access Key")
                    return False
                
                if not iam_username:
                    st.error("❌ Please enter your IAM Username")
                    return False
                
                if len(account_id) != 12 or not account_id.isdigit():
                    st.error("❌ Account ID must be exactly 12 digits")
                    return False
                
                with st.spinner("🔍 Validating AWS credentials..."):
                    success, message = validate_aws_credentials(access_key, secret_key, account_id, region)
                
                if success:
                    st.success(message)
                    
                    # Save credentials to AWS CLI profile if requested
                    if save_to_aws:
                        save_success, save_msg = save_aws_credentials_to_file(
                            access_key=access_key,
                            secret_key=secret_key,
                            region=region,
                            account_id=account_id,
                            iam_username=iam_username,
                            iam_password=iam_password,
                            profile_name=profile_name
                        )
                        if save_success:
                            st.info(f"💾 {save_msg}")
                            st.code(f"# Use with AWS CLI:\naws --profile {profile_name} sts get-caller-identity", language="bash")
                        else:
                            st.warning(f"⚠️ {save_msg}")
                    
                    # Set environment variables for this session
                    set_aws_credentials(access_key, secret_key, region)
                    
                    # Store in session state
                    st.session_state.aws_configured = True
                    st.session_state.aws_account_id = account_id
                    st.session_state.aws_region = region
                    st.session_state.aws_iam_username = iam_username
                    st.session_state.aws_access_key = access_key
                    st.session_state.aws_profile = profile_name
                    
                    st.success("✅ AWS credentials configured successfully!")
                    st.rerun()
                    return True
                else:
                    st.error(f"❌ {message}")
                    return False
        
        st.markdown("---")
        
        # Existing credentials option
        with st.expander("🔄 Use Existing AWS Profile"):
            st.info("🔍 Checking for existing AWS credentials in environment or AWS profile...")
            
            try:
                sts = boto3.client('sts')
                identity = sts.get_caller_identity()
                found_account_id = identity.get('Account', 'Unknown')
                arn = identity.get('Arn', 'Unknown')
                
                st.success(f"""
                ✅ **Found valid AWS credentials!**
                - **Account ID:** {found_account_id}
                - **Identity:** {arn}
                """)
                
                if st.button("🚀 Continue with These Credentials", use_container_width=True, key="use_existing"):
                    st.session_state.aws_configured = True
                    st.session_state.aws_account_id = found_account_id
                    st.session_state.aws_region = boto3.Session().region_name or "us-east-1"
                    st.rerun()
                    return True
                    
            except NoCredentialsError:
                st.warning("No existing AWS credentials found in environment.")
            except Exception as e:
                st.error(f"Error: {str(e)}")
        
        st.markdown("---")
        
        # Back button
        if st.button("← Back to Login", use_container_width=True):
            st.session_state.authenticated = False
            st.rerun()
    
    return False

# =============================================================================
# IMPORT ALL MODULES
# =============================================================================

# Dashboard module - dependencies, config, CSS, and main dashboard components
from nist_dashboard import (
    # Configuration
    S3_BUCKET_NAME,
    S3_PREFIX,
    CERT_DIR,
    CERT_FILE,
    KEY_FILE,
    CONTROL_FAMILIES,
    AWS_REGIONS,
    
    # Functions
    apply_custom_css,
    get_status_emoji,
    get_status_class,
    run_openssl_command,
    render_header,
    render_sidebar,
    render_metrics,
    render_family_summary,
    render_welcome_screen,
    render_controls_selection,
    render_threat_modeling_section,
    render_security_hub_findings_only,
)

# Pages module - secondary pages and features
from nist_pages import (
    render_certificate_viewer,
    render_control_result,
    render_results,
    create_export_data,
    save_results_to_s3,
    render_s3_files,
    render_export_section,
    run_assessment_with_progress,
)

# Authentication module - user auth and session management
from nist_auth import (
    init_session_state,
    is_session_valid,
    get_current_user,
    get_current_role,
    has_permission,
    login_user,
    logout_user,
    render_login_form,
    render_user_info,
    render_user_management,
    require_auth,
    require_permission,
)

# Core assessor from the assessment library
from nist_800_53_rev5_full import NIST80053Rev5Assessor, ControlResult, ControlStatus

# Import Risk Calculator classes
from risk_score_app import Finding, Likelihood, Impact, RiskScoreCalculator


# =============================================================================
# NIST TO RISK CALCULATOR CONVERSION (Module-level for reuse)
# =============================================================================

def convert_nist_to_risk_findings(nist_results, assessor_obj=None):
    """
    Convert NIST assessment results to Risk Calculator findings.
    This function is at module level for reuse across the application.
    
    Args:
        nist_results: List of ControlResult objects from NIST assessment
        assessor_obj: Optional assessor object (not currently used)
    
    Returns:
        List of Finding objects for the Risk Score Calculator
    """
    findings = []
    
    # Likelihood mapping based on control family criticality
    likelihood_map = {
        'AC': Likelihood.HIGH, 'IA': Likelihood.HIGH, 'SC': Likelihood.HIGH,
        'AU': Likelihood.MEDIUM, 'SI': Likelihood.MEDIUM, 'CM': Likelihood.MEDIUM,
        'IR': Likelihood.MEDIUM, 'RA': Likelihood.MEDIUM, 'MP': Likelihood.MEDIUM,
        'SR': Likelihood.MEDIUM, 'CA': Likelihood.LOW, 'CP': Likelihood.LOW,
        'SA': Likelihood.LOW,
    }
    
    # Impact mapping based on status
    impact_map = {
        ControlStatus.FAIL: Impact.SIGNIFICANT,
        ControlStatus.WARNING: Impact.MODERATE,
        ControlStatus.ERROR: Impact.SEVERE,
    }
    
    for i, result in enumerate(nist_results):
        # ControlResult is a dataclass, access attributes directly
        status = result.status if hasattr(result, 'status') else ControlStatus.PASS
        
        # Skip PASS results
        if status == ControlStatus.PASS:
            continue
        
        family = result.family if hasattr(result, 'family') else 'Unknown'
        control_id = result.control_id if hasattr(result, 'control_id') else f'CTRL-{i+1}'
        control_name = result.control_name if hasattr(result, 'control_name') else 'Unknown Control'
        
        # Get findings and recommendations (these are lists in the dataclass)
        result_findings = result.findings if hasattr(result, 'findings') else []
        recommendations = result.recommendations if hasattr(result, 'recommendations') else []
        
        description = "; ".join(result_findings) if result_findings else f"Issue detected with {control_name}"
        remediation = "; ".join(recommendations) if recommendations else "Review and remediate according to NIST guidelines"
        
        finding = Finding(
            finding_id=f"NIST-{control_id}-{i+1:03d}",
            title=f"{control_id}: {control_name}",
            description=description,
            control_family=family,
            control_id=control_id,
            likelihood=likelihood_map.get(family, Likelihood.MEDIUM),
            impact=impact_map.get(status, Impact.MODERATE),
            remediation=remediation,
            status="OPEN"
        )
        findings.append(finding)
    
    return findings


# =============================================================================
# PAGE CONFIGURATION
# =============================================================================

def configure_page():
    """Configure the Streamlit page settings."""
    st.set_page_config(
        page_title="SAELAR-53 - NIST 800-53 Rev 5 Assessment",
        page_icon="🛡️",
        layout="wide",
        initial_sidebar_state="collapsed"
    )
    # ── SAELAR Light Theme Override ──────────────────────────────────
    st.markdown("""
    <style>
        [data-testid="stAppViewContainer"], .stApp {
            background-color: #ffffff !important;
        }
        [data-testid="stSidebar"] {
            background-color: #f0f2f6 !important;
        }
        [data-testid="stSidebar"] * {
            color: #1a1a2e !important;
        }
        header[data-testid="stHeader"] {
            background-color: #ffffff !important;
        }
        [data-testid="stAppViewContainer"] p,
        [data-testid="stAppViewContainer"] span,
        [data-testid="stAppViewContainer"] li,
        [data-testid="stAppViewContainer"] td,
        [data-testid="stAppViewContainer"] th,
        [data-testid="stAppViewContainer"] label {
            color: #1a1a2e !important;
        }
        [data-testid="stAppViewContainer"] code {
            color: #0b3d6e !important;
            background-color: #e8f0fe !important;
        }
        [data-testid="stAppViewContainer"] .stMarkdown div {
            color: #1a1a2e;
        }
        [data-testid="stAppViewContainer"] h1,
        [data-testid="stAppViewContainer"] h2,
        [data-testid="stAppViewContainer"] h3,
        [data-testid="stAppViewContainer"] h4 {
            color: #0b3d6e !important;
        }
        [data-testid="stMetricValue"] {
            color: #0b3d6e !important;
        }
        [data-testid="stMetricDelta"] {
            color: #1a1a2e !important;
        }
        .stTabs [data-baseweb="tab-list"] {
            background-color: #f0f2f6 !important;
        }
        .stTabs [data-baseweb="tab"] {
            color: #1a1a2e !important;
        }
        .stSelectbox > div > div,
        .stTextInput > div > div > input,
        .stTextArea > div > div > textarea {
            background-color: #f8f9fa !important;
            color: #1a1a2e !important;
            border-color: #d0d5dd !important;
        }
        .stButton > button {
            color: #ffffff !important;
            background-color: #0b3d6e !important;
            border-color: #0b3d6e !important;
        }
        .stButton > button:hover {
            background-color: #0a5299 !important;
            border-color: #0a5299 !important;
        }
        /* Sidebar buttons: light background + dark text for readability */
        [data-testid="stSidebar"] .stButton > button,
        [data-testid="stSidebar"] .stButton > button p {
            background-color: #ffffff !important;
            color: #1e293b !important;
            border: 1px solid #94a3b8 !important;
        }
        [data-testid="stSidebar"] .stButton > button:hover {
            background-color: #f1f5f9 !important;
            color: #0f172a !important;
        }
        [data-testid="stExpander"] {
            background-color: #f8f9fa !important;
            border-color: #d0d5dd !important;
        }
        [data-testid="stExpander"] summary span {
            color: #0b3d6e !important;
        }
        /* ── Clean checkboxes: grayscale + darken (main content only - exclude sidebar Control Families) ── */
        [data-testid="stAppViewContainer"] [data-testid="stCheckbox"] div[role="checkbox"] {
            filter: grayscale(1) brightness(0.4) !important;
        }
        [data-testid="stAppViewContainer"] [data-testid="stCheckbox"] div[role="checkbox"][aria-checked="false"] {
            filter: grayscale(1) brightness(1) opacity(0.5) !important;
        }

        /* ── Visible scrollbars on all panels ── */
        ::-webkit-scrollbar {
            width: 12px !important;
            height: 12px !important;
        }
        ::-webkit-scrollbar-thumb {
            background: #0b3d6e !important;
            border-radius: 6px !important;
            border: 2px solid #f0f2f6 !important;
        }
        ::-webkit-scrollbar-thumb:hover {
            background: #0a5299 !important;
        }
        ::-webkit-scrollbar-track {
            background: #e2e6ec !important;
            border-radius: 6px !important;
        }
        /* Sidebar scrollbar */
        [data-testid="stSidebar"] ::-webkit-scrollbar {
            width: 10px !important;
        }
        [data-testid="stSidebar"] ::-webkit-scrollbar-thumb {
            background: #7a8ba8 !important;
            border-radius: 5px !important;
        }
        [data-testid="stSidebar"] ::-webkit-scrollbar-track {
            background: #dce1e8 !important;
        }
        /* Firefox scrollbar */
        * {
            scrollbar-width: auto !important;
            scrollbar-color: #0b3d6e #e2e6ec !important;
        }
    </style>
    """, unsafe_allow_html=True)


# =============================================================================
# SPLASH PAGE
# =============================================================================

def get_base64_image(image_path: Path) -> str:
    """Convert image to base64 for embedding in HTML."""
    try:
        with open(image_path, "rb") as img_file:
            return base64.b64encode(img_file.read()).decode()
    except Exception:
        return ""


def render_splash_page() -> bool:
    """
    Render the splash page with logo, disclaimer, and ENTER button.
    
    Returns:
        True if user clicked ENTER, False otherwise
    """
    # Scroll to top on page load
    st.markdown("""
    <style>
        html, body, [data-testid="stAppViewContainer"], .main {
            scroll-behavior: auto !important;
        }
    </style>
    <script>
        // Immediate scroll
        window.scrollTo({top: 0, left: 0, behavior: 'instant'});
        document.documentElement.scrollTop = 0;
        document.body.scrollTop = 0;
        
        // Delayed scroll to catch after render
        setTimeout(function() {
            window.scrollTo({top: 0, left: 0, behavior: 'instant'});
            var mainContainer = document.querySelector('[data-testid="stAppViewContainer"]');
            if (mainContainer) mainContainer.scrollTop = 0;
        }, 100);
    </script>
    """, unsafe_allow_html=True)
    
    # Hide sidebar on splash page
    st.markdown("""
    <style>
        [data-testid="stSidebar"] {
            display: none;
        }
        [data-testid="stSidebarNav"] {
            display: none;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Get logo as base64
    logo_base64 = get_base64_image(LOGO_PATH)
    
    # Clean splash page CSS - white background with thick blue border
    st.markdown("""
    <style>
        /* Air-gapped compatible: Using system fonts instead of Google Fonts */
        /* Fallback font stack provides similar aesthetic to Orbitron/Rajdhani */
        
        .stApp {
            background: #ffffff;
        }
        
        .stApp > header {
            background: transparent;
            display: none !important;
        }
        
        [data-testid="stHeader"] {
            display: none !important;
        }
        
        [data-testid="stAppViewContainer"] {
            padding-top: 0 !important;
        }
        
        [data-testid="stAppViewBlockContainer"] {
            padding-top: 0 !important;
        }
        
        .stMainBlockContainer, [data-testid="stMainBlockContainer"] {
            padding-top: 0 !important;
            min-height: auto !important;
        }
        
        section.main {
            padding-top: 0 !important;
        }
        
        section.main > div:first-child {
            padding-top: 0 !important;
        }
        
        .block-container, [data-testid="block-container"] {
            padding-top: 0 !important;
            max-width: 100% !important;
        }
        
        .main .block-container {
            border: 8px solid #1e3a5f;
            border-radius: 16px;
            padding: 0.5rem 2rem 2rem 2rem !important;
            margin: 1rem;
            margin-top: 0 !important;
            background: #ffffff;
            box-shadow: 0 4px 20px rgba(30, 58, 95, 0.2);
        }
        
        .splash-logo-container {
            margin-top: -1rem !important;
        }
        
        .splash-disclaimer {
            background: linear-gradient(135deg, #dc2626 0%, #b91c1c 50%, #991b1b 100%);
            border: 2px solid #7f1d1d;
            border-radius: 8px;
            padding: 0.6rem 1rem;
            margin: 0.5rem auto;
            text-align: center;
            box-shadow: 0 2px 8px rgba(220, 38, 38, 0.3);
        }
        
        .splash-footer {
            text-align: center;
            margin-top: 2rem;
            color: #64748b;
            font-size: 0.85rem;
        }
        
        .stButton > button {
            font-family: 'Segoe UI', 'SF Pro Display', -apple-system, BlinkMacSystemFont, 'Helvetica Neue', Arial, sans-serif !important;
            font-size: 1.2rem !important;
            font-weight: 600 !important;
            letter-spacing: 4px !important;
            padding: 1rem 4rem !important;
            background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%) !important;
            color: #ffffff !important;
            border: none !important;
            border-radius: 8px !important;
            cursor: pointer !important;
            transition: all 0.3s ease !important;
            box-shadow: 0 4px 20px rgba(30, 58, 95, 0.4) !important;
        }
        
        .stButton > button:hover {
            transform: translateY(-2px) !important;
            box-shadow: 0 8px 30px rgba(30, 58, 95, 0.6) !important;
            background: linear-gradient(135deg, #2d5a87 0%, #3d7ab5 100%) !important;
        }
    </style>
    """, unsafe_allow_html=True)
    
    # Display logo centered using HTML for precise control
    logo_base64 = get_base64_image(LOGO_PATH)
    
    if logo_base64:
        st.markdown(f"""
        <div class="splash-logo-container" style="display: flex; justify-content: center; align-items: center; width: 100%; margin-top: 0; margin-bottom: 1rem;">
            <img src="data:image/png;base64,{logo_base64}" style="width: 100%; max-width: 420px; height: auto;" />
        </div>
        """, unsafe_allow_html=True)
    else:
        st.markdown("<h1 style='text-align: center; font-size: 5rem;'>🛡️</h1>", unsafe_allow_html=True)
    
    # SAELAR-53 acronym explanation with box - centered (same width as logo)
    st.markdown("""
    <div style="display: flex; justify-content: center; width: 100%; margin: 0.5rem 0 1.5rem 0;">
        <div style="
            border: 4px solid #1e3a5f;
            border-radius: 17px;
            padding: 1.5rem 2rem;
            background: #f8fafc;
            box-shadow: 0 3px 11px rgba(30, 58, 95, 0.15);
            width: 100%;
            max-width: 420px;
        ">
            <div style="text-align: left; color: #1e3a5f; font-size: 1.2rem; line-height: 1.8;">
                <span style="color: #059669; font-weight: bold; font-size: 1.5rem;">S</span> - Security<br>
                <span style="color: #059669; font-weight: bold; font-size: 1.5rem;">A</span> - Architecture and<br>
                <span style="color: #059669; font-weight: bold; font-size: 1.5rem;">E</span> - Evaluation<br>
                <span style="color: #059669; font-weight: bold; font-size: 1.5rem;">L</span> - Linear<br>
                <span style="color: #059669; font-weight: bold; font-size: 1.5rem;">A</span> - Assessment<br>
                <span style="color: #059669; font-weight: bold; font-size: 1.5rem;">R</span> - Reporting Tool<br>
                <span style="color: #1e3a5f; font-weight: bold; font-size: 1.5rem;">53</span> - NIST 800-53
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Center the ENTER button using columns
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("ENTER", use_container_width=True, key="splash_enter"):
            return True
    
    # Disclaimer box with red gradient - centered (below ENTER button, same width as other elements)
    st.markdown(f"""
    <div style="display: flex; justify-content: center; width: 100%; margin: 0.5rem 0;">
        <div class="splash-disclaimer" style="width: 100%; max-width: 420px;">
            <div style="font-size: 0.8rem; margin-bottom: 0.15rem;">⚠️</div>
            <div style="font-size: 0.65rem; color: #ffffff; line-height: 1.3; font-weight: 500;">
                {DISCLAIMER_TEXT}
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Footer - centered
    st.markdown("""
    <div style="text-align: center; margin-top: 2rem; color: #64748b; font-size: 0.85rem;">
        © 2026 Security Architecture & Evaluation | v1.0.0
    </div>
    """, unsafe_allow_html=True)
    
    return False


# =============================================================================
# APPLICATION CONTROLLER
# =============================================================================

class NISTApplicationController:
    """
    Main application controller that orchestrates all components.
    Handles routing, state management, and component coordination.
    """
    
    def __init__(self):
        """Initialize the application controller."""
        self.assessor: Optional[NIST80053Rev5Assessor] = None
        self._init_session_state()
    
    def _init_session_state(self):
        """Initialize all session state variables."""
        # Splash page state
        if 'splash_accepted' not in st.session_state:
            st.session_state.splash_accepted = False
        
        # Authentication state
        init_session_state()
        
        # AWS configuration state
        if 'aws_configured' not in st.session_state:
            st.session_state.aws_configured = False
        if 'aws_validated' not in st.session_state:
            st.session_state.aws_validated = False
        if 'aws_account_id' not in st.session_state:
            st.session_state.aws_account_id = None
        if 'aws_region' not in st.session_state:
            st.session_state.aws_region = "us-east-1"
        if 'aws_iam_username' not in st.session_state:
            st.session_state.aws_iam_username = None
        if 'aws_profile' not in st.session_state:
            st.session_state.aws_profile = "saelar"
        
        # Assessment state
        if 'results' not in st.session_state:
            st.session_state.results = None
        if 'assessor' not in st.session_state:
            st.session_state.assessor = None
        if 'families_assessed' not in st.session_state:
            st.session_state.families_assessed = []
        if 's3_saved' not in st.session_state:
            st.session_state.s3_saved = False
        
        # Navigation state
        if 'current_page' not in st.session_state:
            st.session_state.current_page = 'assessment'
    
    def check_authentication(self) -> bool:
        """
        Check if user is authenticated.
        Returns True if authenticated, False otherwise.
        """
        # For demo purposes, authentication can be disabled by setting this to True
        AUTH_ENABLED = True
        
        if not AUTH_ENABLED:
            return True
        
        return is_session_valid()
    
    def run_assessment(self, region: str, selected_families: list, include_security_hub: bool = False) -> bool:
        """
        Run the NIST assessment for selected families.
        
        Args:
            region: AWS region to assess
            selected_families: List of family codes to assess
            include_security_hub: Whether to import Security Hub findings
            
        Returns:
            True if assessment completed successfully, False otherwise
        """
        # Check permission
        username = get_current_user()
        if username and not has_permission(username, 'can_run_assessment'):
            st.error("⚠️ You don't have permission to run assessments.")
            return False
        
        # Verify AWS credentials are set
        if not os.environ.get('AWS_ACCESS_KEY_ID') or not os.environ.get('AWS_SECRET_ACCESS_KEY'):
            # Try to re-load credentials from saved config
            saved_creds = load_saved_credentials()
            if saved_creds and saved_creds.get('access_key_id') and saved_creds.get('secret_key'):
                set_aws_credentials(
                    saved_creds.get('access_key_id', ''),
                    saved_creds.get('secret_key', ''),
                    saved_creds.get('region', 'us-east-1')
                )
            else:
                st.error("❌ AWS credentials not found. Please re-authenticate.")
                st.session_state.aws_configured = False
                st.rerun()
                return False
        
        try:
            with st.spinner("🔐 Connecting to AWS..."):
                selected_region = None if region == "Default" else region
                self.assessor = NIST80053Rev5Assessor(region=selected_region)
                st.session_state.assessor = self.assessor
            
            st.success(f"✅ Connected to AWS Account: **{self.assessor.account_id}**")
            
            # Show which families will be assessed
            families_str = ", ".join(selected_families)
            st.info(f"📋 Assessing families: **{families_str}**")
            
            # Run assessment
            st.session_state.results = run_assessment_with_progress(self.assessor, selected_families)
            
            # Import Security Hub findings if enabled
            if include_security_hub:
                with st.spinner("🔗 Importing Security Hub findings..."):
                    try:
                        sh_results = self.assessor.import_security_hub_findings(max_findings=100)
                        if sh_results:
                            st.session_state.results.extend(sh_results)
                            st.success(f"📥 Imported {len(sh_results)} Security Hub finding groups!")
                    except Exception as e:
                        st.warning(f"⚠️ Could not import Security Hub findings: {str(e)}")
            st.session_state.families_assessed = selected_families
            
            # Auto-save results to S3 (store status for display at bottom of page)
            family_label = "-".join(selected_families) if len(selected_families) <= 4 else f"{len(selected_families)}_families"
            with st.spinner("💾 Saving results..."):
                s3_key = save_results_to_s3(st.session_state.results, self.assessor, family_label)
                if s3_key:
                    st.session_state.s3_saved = True
                    st.session_state.s3_save_path = f"s3://{S3_BUCKET_NAME}/{S3_PREFIX}"
                else:
                    st.session_state.s3_saved = False
                    st.session_state.s3_save_path = None
            
            # Auto-populate Risk Score Calculator with NIST findings
            with st.spinner("📊 Updating Risk Calculator..."):
                try:
                    imported_findings = convert_nist_to_risk_findings(st.session_state.results, self.assessor)
                    st.session_state.risk_nist_findings = imported_findings
                    st.session_state.risk_nist_imported = True
                    st.session_state.risk_use_demo = False
                    st.success(f"📊 Risk Calculator auto-updated with {len(imported_findings)} findings!")
                except Exception as e:
                    st.warning(f"⚠️ Could not auto-update Risk Calculator: {str(e)}")
            
            return True
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.info("💡 Ensure your AWS credentials are configured correctly.")
            return False
    
    def render_assessment_page(self, region, selected_families, run_assessment, include_security_hub=False):
        """Render the main assessment page."""
        family_names = NIST80053Rev5Assessor.CONTROL_FAMILIES
        
        if run_assessment and selected_families:
            self.run_assessment(region, selected_families, include_security_hub)
        
        if st.session_state.results and st.session_state.assessor:
            results = st.session_state.results
            assessor = st.session_state.assessor
            summary = assessor.generate_summary(results)
            
            st.markdown("---")
            
            # RESULTS header
            st.markdown("""
            <div style="text-align: center; margin: 1.5rem 0;">
                <h1 style="font-size: 22pt; font-weight: bold; color: #1e3a5f; margin: 0;">RESULTS</h1>
            </div>
            """, unsafe_allow_html=True)
            
            render_metrics(summary)
            
            # Risk level indicator with clickable links to findings
            score = (summary['passed'] / summary['total_controls'] * 100) if summary['total_controls'] > 0 else 0
            
            if summary['failed'] > 0:
                st.markdown(f"""
                <a href="#failed-findings" style="text-decoration: none;">
                    <div style="background: linear-gradient(135deg, #fecaca 0%, #fca5a5 100%); 
                                border: 1px solid #dc2626; border-radius: 8px; padding: 1rem; 
                                margin: 0.5rem 0; cursor: pointer; transition: all 0.2s ease;"
                         onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 12px rgba(220,38,38,0.3)';"
                         onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='none';">
                        <span style="color: #991b1b; font-weight: 600; font-size: 1rem;">
                            🚨 {summary['failed']} control(s) FAILED - Immediate attention required!
                        </span>
                        <span style="color: #b91c1c; float: right;">↓ Click to view</span>
                    </div>
                </a>
                """, unsafe_allow_html=True)
            elif summary['warnings'] > 5:
                st.markdown(f"""
                <a href="#warning-findings" style="text-decoration: none;">
                    <div style="background: linear-gradient(135deg, #fef3c7 0%, #fde68a 100%); 
                                border: 1px solid #d97706; border-radius: 8px; padding: 1rem; 
                                margin: 0.5rem 0; cursor: pointer; transition: all 0.2s ease;"
                         onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 12px rgba(217,119,6,0.3)';"
                         onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='none';">
                        <span style="color: #92400e; font-weight: 600; font-size: 1rem;">
                            ⚠️ {summary['warnings']} warning(s) - Review recommended
                        </span>
                        <span style="color: #b45309; float: right;">↓ Click to view</span>
                    </div>
                </a>
                """, unsafe_allow_html=True)
            elif score >= 80:
                st.markdown(f"""
                <a href="#passed-findings" style="text-decoration: none;">
                    <div style="background: linear-gradient(135deg, #dcfce7 0%, #bbf7d0 100%); 
                                border: 1px solid #22c55e; border-radius: 8px; padding: 1rem; 
                                margin: 0.5rem 0; cursor: pointer; transition: all 0.2s ease;"
                         onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 12px rgba(34,197,94,0.3)';"
                         onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='none';">
                        <span style="color: #166534; font-weight: 600; font-size: 1rem;">
                            ✅ Strong compliance posture - {score:.0f}% score
                        </span>
                        <span style="color: #15803d; float: right;">↓ View details</span>
                    </div>
                </a>
                """, unsafe_allow_html=True)
            
            # Link to Risk Calculator
            findings_count = summary.get('failed', 0) + summary.get('warnings', 0)
            if findings_count > 0:
                st.markdown(f"""
                <div style="background: linear-gradient(135deg, #f59e0b15 0%, #d9770615 100%); 
                            border: 1px solid #f59e0b; border-radius: 8px; padding: 1rem; margin: 1rem 0;
                            display: flex; align-items: center; justify-content: space-between;">
                    <div>
                        <span style="font-weight: 600; color: #92400e;">📊 Risk Analysis Available</span>
                        <span style="color: #78716c; margin-left: 0.5rem;">
                            {findings_count} findings ready for risk scoring
                        </span>
                    </div>
                    <div style="color: #92400e; font-size: 0.9rem;">
                        → Go to <strong>Risk Calculator</strong> tab to analyze
                    </div>
                </div>
                """, unsafe_allow_html=True)
            
            # Security Hub Findings Box - matching Risk Analysis format exactly
            # Hidden button triggered by JavaScript click on the div
            view_sh = st.button("view_sh_hidden", key="view_security_hub_results", type="secondary")
            
            st.markdown("""
            <style>
                /* Hide the actual button */
                div[data-testid="stButton"]:has(button[kind="secondary"]) { display: none; }
            </style>
            <div onclick="document.querySelector('button[kind=secondary]').click();" 
                 style="background: linear-gradient(135deg, #3b82f615 0%, #1d4ed815 100%); 
                        border: 1px solid #3b82f6; border-radius: 8px; padding: 1rem; margin: 1rem 0;
                        display: flex; align-items: center; justify-content: space-between;
                        cursor: pointer; transition: all 0.2s ease;"
                 onmouseover="this.style.transform='translateY(-2px)'; this.style.boxShadow='0 4px 12px rgba(59,130,246,0.3)';"
                 onmouseout="this.style.transform='translateY(0)'; this.style.boxShadow='none';">
                <div>
                    <span style="font-weight: 600; color: #1e40af;">🔗 AWS Security Hub Findings</span>
                    <span style="color: #64748b; margin-left: 0.5rem;">
                        Aggregated findings from GuardDuty, Inspector, Macie & more
                    </span>
                </div>
                <div style="color: #1e40af; font-size: 0.9rem;">
                    → Click to <strong>View Findings</strong>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            if view_sh:
                st.session_state.show_security_hub_only = True
                st.rerun()
            
            # Threat Modeling Section
            st.markdown("---")
            render_threat_modeling_section(assessment_data=results, findings=None)
            
            st.markdown("---")
            render_family_summary(summary, family_names)
            
            st.markdown("---")
            render_results(results, family_names)
            
            # Chad (AI) interactive box - right above SAVED RESULTS & EXPORT
            st.markdown("---")
            from pathlib import Path
            chad_avatar_path = Path(__file__).parent / 'assets' / 'chad2.jpg'
            failed_count = summary.get('failed', 0)
            warning_count = summary.get('warnings', 0)
            passed_count = summary.get('passed', 0)
            if failed_count > 5:
                chad_message = f"🚨 <strong>Critical attention needed!</strong> I found {failed_count} failed controls. Click on me to get a remediation plan."
                chad_bg, chad_border = "#fef2f2", "#fecaca"
            elif failed_count > 0:
                chad_message = f"⚠️ <strong>{failed_count} controls need attention.</strong> I can help you prioritize and fix these issues."
                chad_bg, chad_border = "#fffbeb", "#fde68a"
            elif warning_count > 5:
                chad_message = f"📋 <strong>{warning_count} warnings to review.</strong> Let me help you address these before they become issues."
                chad_bg, chad_border = "#f0f9ff", "#bae6fd"
            else:
                chad_message = f"✅ <strong>Great job!</strong> {passed_count} controls passed. Ask me about maintaining this security posture."
                chad_bg, chad_border = "#f0fdf4", "#86efac"
            if chad_avatar_path.exists():
                import base64
                with open(chad_avatar_path, "rb") as img_file:
                    chad_img_data = base64.b64encode(img_file.read()).decode()
                chad_avatar_html = f'<img src="data:image/jpeg;base64,{chad_img_data}" style="width: 160px; height: 160px; border-radius: 12px; object-fit: cover; border: 2px solid #3b82f6; box-shadow: 0 4px 12px rgba(59,130,246,0.25);" alt="Chad (AI)">'
            else:
                chad_avatar_html = '<div style="width: 160px; height: 160px; background: linear-gradient(135deg, #3b82f6 0%, #1e3a5f 100%); border-radius: 12px; display: flex; align-items: center; justify-content: center; font-size: 3.2rem; border: 2px solid #3b82f6;">🛡️</div>'
            chad_spacer, chad_col = st.columns([1, 3])
            with chad_col:
                st.markdown(f"""
                <div style="display: flex; align-items: center; gap: 0.75rem; background: {chad_bg}; border: 1px solid {chad_border}; border-radius: 12px; padding: 0.75rem 1rem;">
                    {chad_avatar_html}
                    <div style="max-width: 100%; flex: 1;">
                        <div style="font-weight: 700; color: #1e3a5f; font-size: 1rem; margin-bottom: 0.25rem;">Chad (AI) - Security Analyst</div>
                        <div style="color: #374151; font-size: 0.9rem; line-height: 1.4;">{chad_message}</div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
                st.markdown("<div style='margin-top: 0.5rem;'></div>", unsafe_allow_html=True)
                if 'chad_inline_history' not in st.session_state:
                    st.session_state.chad_inline_history = []
                with st.expander("💬 **Chat with Chad (AI)**", expanded=True):
                    st.markdown("**Quick Actions:**")
                    qcol1, qcol2, qcol3 = st.columns(3)
                    with qcol1:
                        if st.button("📊 Summarize", key="chad_inline_summary", use_container_width=True):
                            st.session_state.chad_inline_query = "Give me a quick summary of my assessment results and the top priorities I should focus on."
                    with qcol2:
                        if st.button("🔧 Remediate", key="chad_inline_remediate", use_container_width=True):
                            st.session_state.chad_inline_query = "What are the specific remediation steps for my failed controls? Provide AWS CLI commands."
                    with qcol3:
                        if st.button("📋 POA&M", key="chad_inline_poam", use_container_width=True):
                            st.session_state.chad_inline_query = "Generate a POA&M document for my failed controls with realistic timelines."
                    st.markdown("---")
                    chat_container = st.container(height=800)
                    with chat_container:
                        if not st.session_state.chad_inline_history:
                            st.markdown("""
                            <div style="background: #f0f9ff; border-radius: 12px; padding: 1rem; margin: 0.5rem 0;">
                                <strong style="color: #0369a1;">🛡️ Chad (AI):</strong>
                                <p style="color: #0c4a6e; margin: 0.25rem 0 0 0;">Hi! I've analyzed your assessment results. Ask me anything about your findings, remediation steps, or compliance status!</p>
                            </div>
                            """, unsafe_allow_html=True)
                        else:
                            for msg in st.session_state.chad_inline_history:
                                if msg["role"] == "user":
                                    st.markdown(f"""
                                    <div style="background: #e0e7ff; border-radius: 12px; padding: 0.75rem 1rem; margin: 0.5rem 0; margin-left: 20%;">
                                        <strong style="color: #3730a3;">You:</strong>
                                        <p style="color: #1e1b4b; margin: 0.25rem 0 0 0;">{msg["content"]}</p>
                                    </div>
                                    """, unsafe_allow_html=True)
                                else:
                                    st.markdown(f"""
                                    <div style="background: #f0fdf4; border-radius: 12px; padding: 0.75rem 1rem; margin: 0.5rem 0; margin-right: 10%;">
                                        <strong style="color: #166534;">🛡️ Chad (AI):</strong>
                                        <div style="color: #14532d; margin: 0.25rem 0 0 0;">{msg["content"]}</div>
                                    </div>
                                    """, unsafe_allow_html=True)
                    st.markdown("---")
                    input_col, send_col, clear_col = st.columns([4, 1, 1])
                    with input_col:
                        user_question = st.text_input("Ask Chad (AI):", placeholder="e.g., 'How do I fix the MFA finding?'", key="chad_inline_input", label_visibility="collapsed")
                    with send_col:
                        send_clicked = st.button("Send", key="chad_inline_send", type="primary", use_container_width=True)
                    with clear_col:
                        if st.button("Clear", key="chad_inline_clear", use_container_width=True):
                            st.session_state.chad_inline_history = []
                            st.rerun()
                    if send_clicked and user_question:
                        st.session_state.chad_inline_query = user_question
                    if st.session_state.get('chad_inline_query'):
                        query = st.session_state.chad_inline_query
                        st.session_state.chad_inline_query = None
                        st.session_state.chad_inline_history.append({"role": "user", "content": query})
                        with st.spinner("🤔 Chad (AI) is thinking..."):
                            try:
                                context_lines = [f"Assessment Summary: {summary['passed']} passed, {summary['failed']} failed, {summary['warnings']} warnings out of {summary['total_controls']} controls."]
                                failed_controls = [r for r in results if hasattr(r, 'status') and '❌' in str(r.status.value)][:5]
                                if failed_controls:
                                    context_lines.append("Failed controls:")
                                    for ctrl in failed_controls:
                                        context_lines.append(f"- {ctrl.control_id}: {ctrl.control_name}")
                                messages = [{"role": m["role"], "content": m["content"]} for m in st.session_state.chad_inline_history]
                                system_prompt = f"You are Chad (AI), a friendly security analyst. Be concise, helpful, and actionable. Use markdown for formatting. Context: {' '.join(context_lines)}"
                                response_text = call_ai(messages, system_prompt, max_tokens=4096)
                                st.session_state.chad_inline_history.append({"role": "assistant", "content": response_text})
                                st.rerun()
                            except Exception as e:
                                st.session_state.chad_inline_history.append({"role": "assistant", "content": f"❌ Error: {str(e)}"})
                                st.rerun()
            
            # Export and saved files section at the bottom
            st.markdown("---")
            st.markdown("""
            <div style="text-align: center; margin: 1.5rem 0;">
                <h2 style="font-size: 18pt; font-weight: bold; color: #1e3a5f; margin: 0;">📁 SAVED RESULTS & EXPORT</h2>
            </div>
            """, unsafe_allow_html=True)
            
            # Show save status from auto-save
            if st.session_state.get('s3_saved'):
                st.success(f"📦 Results automatically saved to: `{st.session_state.get('s3_save_path', 'S3')}`")
            elif st.session_state.get('s3_save_path') is None and st.session_state.get('s3_saved') == False:
                st.warning("⚠️ Could not auto-save to S3 - use manual download options below")
            
            render_export_section(results, assessor)
        
        else:
            render_welcome_screen()
    
    def render_aws_console_page(self):
        """Render the AWS Console link page."""
        st.markdown("""
        <div style="text-align: center; padding: 3rem;">
            <h2 style="color: #1e3a5f;">☁️ AWS Management Console</h2>
            <p style="color: #64748b; margin-bottom: 2rem;">Access your AWS account directly in the AWS Management Console</p>
            <a href="https://console.aws.amazon.com" target="_blank" style="
                display: inline-block;
                background: linear-gradient(135deg, #ff9900 0%, #ff8000 100%);
                color: #ffffff;
                padding: 1rem 3rem;
                border-radius: 8px;
                text-decoration: none;
                font-weight: 600;
                font-size: 1.1rem;
                box-shadow: 0 4px 15px rgba(255, 153, 0, 0.4);
                transition: all 0.3s ease;
            " onmouseover="this.style.transform='translateY(-2px)'" onmouseout="this.style.transform='translateY(0)'">
                🔗 Open AWS Console
            </a>
        </div>
        """, unsafe_allow_html=True)
    
    def render_chad_security_analyst_page(self):
        """Render Chad (AI) - the AI Security Analyst chat interface powered by AWS Bedrock."""
        
        # Initialize chat history in session state
        if 'chad_messages' not in st.session_state:
            st.session_state.chad_messages = []
        # Note: Chad (AI) now uses AWS Bedrock - no API key needed, uses AWS credentials
        
        # Header with Chad (AI)'s avatar
        chad_avatar_path = Path(__file__).parent / 'assets' / 'chad_avatar.png'
        
        # Create header with avatar - compact so Quick Questions and Chat show above the fold
        st.markdown("""
        <style>
        .chad-header {
            display: flex;
            align-items: center;
            gap: 1rem;
            padding: 0.75rem 1rem;
            background: linear-gradient(135deg, #f0f9ff 0%, #e0f2fe 100%);
            border-radius: 12px;
            border: 1px solid #bae6fd;
            margin-bottom: 0.75rem;
        }
        .chad-avatar {
            width: 56px;
            height: 56px;
            border-radius: 8px;
            object-fit: cover;
            border: 2px solid #3b82f6;
            flex-shrink: 0;
        }
        .chad-info h1 {
            color: #1e3a5f;
            margin: 0;
            font-size: 1.5rem;
        }
        .chad-info .title {
            color: #3b82f6;
            margin: 0;
            font-size: 0.9rem;
            font-weight: 600;
        }
        .chad-info .tagline {
            color: #64748b;
            margin: 0.2rem 0 0 0;
            font-size: 0.8rem;
        }
        </style>
        """, unsafe_allow_html=True)
        
        # Check if avatar exists and display header
        if chad_avatar_path.exists():
            # Read and encode the image
            import base64
            with open(chad_avatar_path, "rb") as img_file:
                img_data = base64.b64encode(img_file.read()).decode()
            
            st.markdown(f"""
            <div class="chad-header">
                <img src="data:image/png;base64,{img_data}" class="chad-avatar" alt="Chad (AI)">
                <div class="chad-info">
                    <h1>Chad (AI)</h1>
                    <p class="title">Cyber Security Analyst</p>
                    <p class="tagline">Your AI-powered security analyst • Powered by AWS Bedrock 🔒</p>
                </div>
            </div>
            """, unsafe_allow_html=True)
        else:
            # Fallback with emoji avatar
            st.markdown("""
            <div class="chad-header">
                <div style="width: 56px; height: 56px; background: linear-gradient(135deg, #3b82f6 0%, #1e3a5f 100%); 
                            border-radius: 8px; display: flex; align-items: center; justify-content: center;
                            font-size: 1.8rem; border: 2px solid #3b82f6; flex-shrink: 0;">
                    🛡️
                </div>
                <div class="chad-info">
                    <h1>Chad (AI)</h1>
                    <p class="title">Cyber Security Analyst</p>
                    <p class="tagline">Your AI-powered security analyst • Powered by AWS Bedrock 🔒</p>
                </div>
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Chad (AI) uses AWS Bedrock - no API key configuration needed
        # Quick action buttons - contextual based on whether assessment exists
        has_assessment_data = st.session_state.get('results') is not None and len(st.session_state.get('results', [])) > 0
        
        st.markdown("### 💡 Quick Questions")
        st.caption("Click a button to ask Chad (AI) a question, or type in the chat below.")
        
        if has_assessment_data:
            # Assessment-focused buttons
            cols = st.columns(3)
            with cols[0]:
                if st.button("🔴 Analyze Findings", use_container_width=True, key="chad_btn_risks_v2"):
                    self._chad_send_message("Analyze my assessment findings. What are the most critical security gaps I need to address immediately?")
            with cols[1]:
                if st.button("📊 Assessment Summary", use_container_width=True, key="chad_btn_compliance_v2"):
                    self._chad_send_message("Give me a detailed summary of my assessment results. Which control families are weakest and why?")
            with cols[2]:
                if st.button("🔧 Remediation Plan", use_container_width=True, key="chad_btn_remediation_v2"):
                    self._chad_send_message("Create a prioritized remediation plan for my failed controls. Include specific AWS CLI commands or scripts I can use.")
        else:
            # General security buttons
            cols = st.columns(3)
            with cols[0]:
                if st.button("📚 NIST Controls", use_container_width=True, key="chad_btn_nist_v2"):
                    self._chad_send_message("Explain the NIST 800-53 Rev 5 control families and which ones are most critical for AWS environments.")
            with cols[1]:
                if st.button("🔍 Security Best Practices", use_container_width=True, key="chad_btn_bestprac_v2"):
                    self._chad_send_message("What are the top security best practices I should implement in my AWS environment?")
            with cols[2]:
                if st.button("☁️ AWS Security", use_container_width=True, key="chad_btn_aws_v2"):
                    self._chad_send_message("What are the most common AWS security misconfigurations and how do I prevent them?")
        
        st.markdown("---")
        
        # Report Generation Section (only show if assessment data exists)
        if has_assessment_data:
            st.markdown("### 📄 Generate Reports")
            st.caption("Click a report type to have Chad (AI) generate it based on your assessment data")
            
            # Row 1 of report buttons
            row1 = st.columns(4)
            with row1[0]:
                if st.button("📊 Executive Summary", use_container_width=True, key="rpt_exec_v2"):
                    self._chad_send_message("""Generate an Executive Summary Report based on my assessment. Include:
1. Overall security posture rating (Critical/High/Medium/Low risk)
2. Key metrics: total controls assessed, pass rate, critical findings count
3. Top 5 most critical findings with business impact
4. Recommended immediate actions (prioritized)
5. 30/60/90 day improvement roadmap
Format this as a professional executive briefing.""")
            with row1[1]:
                if st.button("📋 POA&M", use_container_width=True, key="rpt_poam_v2"):
                    self._chad_send_message("""Generate a Plan of Action and Milestones (POA&M) document based on my failed controls. For each finding include:
1. Weakness ID and Description
2. Point of Contact (placeholder)
3. Resource Requirements
4. Scheduled Completion Date (suggest realistic timelines)
5. Milestones with Milestone Changes
6. Status
Format as a structured POA&M table that can be submitted to auditors.""")
            with row1[2]:
                if st.button("⚠️ Risk Assessment", use_container_width=True, key="rpt_risk_v2"):
                    self._chad_send_message("""Generate a Risk Assessment Report using NIST SP 800-30 methodology. Include:
1. Executive Summary of risk posture
2. Risk assessment methodology explanation
3. Threat sources and events identified
4. Vulnerability analysis with likelihood ratings
5. Impact analysis (Confidentiality, Integrity, Availability)
6. Risk determination matrix (Likelihood × Impact)
7. Risk response recommendations (Accept, Mitigate, Transfer, Avoid)
8. Residual risk after proposed mitigations""")
            with row1[3]:
                if st.button("🔧 Technical Report", use_container_width=True, key="rpt_tech_v2"):
                    self._chad_send_message("""Generate a detailed Technical Findings Report. For each failed control include:
1. Control ID and Name
2. Control Family
3. Technical finding description with evidence
4. Affected AWS resources/services
5. Root cause analysis
6. Step-by-step remediation instructions with AWS CLI commands
7. Verification steps to confirm remediation
8. References to AWS documentation""")
            
            # Row 2 of report buttons
            row2 = st.columns(4)
            with row2[0]:
                if st.button("👔 Board Briefing", use_container_width=True, key="rpt_board_v2"):
                    self._chad_send_message("""Generate a Board of Directors Security Briefing. Use non-technical language and include:
1. Security Program Health Score (visual rating)
2. Year-over-year improvement trends (if available, otherwise baseline)
3. Key risk areas in business terms
4. Regulatory compliance status (NIST, FedRAMP readiness)
5. Resource/investment needs
6. Peer comparison context (industry benchmarks)
7. Strategic recommendations
Keep it concise, high-level, and focused on business risk.""")
            with row2[1]:
                if st.button("📑 Audit Package", use_container_width=True, key="rpt_audit_v2"):
                    self._chad_send_message("""Generate an Audit-Ready Documentation Package. Include:
1. System Security Plan (SSP) summary
2. Control implementation status matrix
3. Evidence of compliance for passing controls
4. POA&M for non-compliant controls
5. Risk acceptance documentation template
6. Continuous monitoring summary
7. Authorization recommendation statement
Format for FedRAMP/FISMA audit submission.""")
            with row2[2]:
                if st.button("🛠️ Remediation Plan", use_container_width=True, key="rpt_remed_v2"):
                    self._chad_send_message("""Generate a Comprehensive Remediation Plan. Include:
1. Prioritized remediation roadmap (Critical → High → Medium → Low)
2. Resource requirements (time, personnel, tools)
3. Dependencies between remediations
4. Quick wins (items fixable in < 1 hour)
5. Detailed remediation scripts for each finding:
   - AWS CLI commands
   - Terraform code snippets
   - CloudFormation templates where applicable
6. Testing/validation procedures
7. Rollback procedures""")
            with row2[3]:
                if st.button("✅ Compliance Status", use_container_width=True, key="rpt_comply_v2"):
                    self._chad_send_message("""Generate a NIST 800-53 Rev 5 Compliance Status Report. Include:
1. Overall compliance percentage
2. Compliance by control family (AC, AU, CM, IA, etc.)
3. Heat map description (which families need attention)
4. Gap analysis: what's needed for full compliance
5. Comparison to baseline requirements (Low/Moderate/High)
6. Timeline to achieve target compliance level
7. Control inheritance from AWS (shared responsibility)""")
            
            st.markdown("---")
        
        # Chat interface
        st.markdown("### 💬 Chat with Chad (AI)")
        
        # Display assessment context if available
        has_assessment = st.session_state.get('results') is not None and len(st.session_state.get('results', [])) > 0
        
        if has_assessment:
            # Build assessment summary for display
            results = st.session_state.results
            failed_count = len([r for r in results if hasattr(r, 'status') and '❌' in str(r.status.value)])
            warning_count = len([r for r in results if hasattr(r, 'status') and '⚠️' in str(r.status.value)])
            passed_count = len([r for r in results if hasattr(r, 'status') and '✅' in str(r.status.value)])
            total_count = len(results)
            
            # Determine severity color
            if failed_count > 5:
                severity_color = "#dc2626"  # Red
                severity_bg = "#fef2f2"
                severity_border = "#fecaca"
                severity_text = "High Risk"
            elif failed_count > 0:
                severity_color = "#d97706"  # Orange
                severity_bg = "#fffbeb"
                severity_border = "#fde68a"
                severity_text = "Medium Risk"
            else:
                severity_color = "#16a34a"  # Green
                severity_bg = "#f0fdf4"
                severity_border = "#86efac"
                severity_text = "Low Risk"
            
            st.markdown(f"""
            <div style="background: {severity_bg}; border: 2px solid {severity_border}; border-radius: 12px; padding: 1.5rem; margin: 1rem 0;">
                <div style="display: flex; align-items: center; gap: 0.75rem; margin-bottom: 1rem;">
                    <span style="font-size: 1.5rem;">📊</span>
                    <span style="font-weight: 700; font-size: 1.1rem; color: {severity_color};">Assessment Data Loaded - {severity_text}</span>
                </div>
                <div style="display: grid; grid-template-columns: repeat(4, 1fr); gap: 1rem; margin-bottom: 1rem;">
                    <div style="text-align: center; padding: 0.75rem; background: white; border-radius: 8px; border: 1px solid #e5e7eb;">
                        <div style="font-size: 1.5rem; font-weight: 700; color: #1e3a5f;">{total_count}</div>
                        <div style="font-size: 0.85rem; color: #64748b;">Controls</div>
                    </div>
                    <div style="text-align: center; padding: 0.75rem; background: white; border-radius: 8px; border: 1px solid #e5e7eb;">
                        <div style="font-size: 1.5rem; font-weight: 700; color: #16a34a;">{passed_count}</div>
                        <div style="font-size: 0.85rem; color: #64748b;">Passed</div>
                    </div>
                    <div style="text-align: center; padding: 0.75rem; background: white; border-radius: 8px; border: 1px solid #e5e7eb;">
                        <div style="font-size: 1.5rem; font-weight: 700; color: #d97706;">{warning_count}</div>
                        <div style="font-size: 0.85rem; color: #64748b;">Warnings</div>
                    </div>
                    <div style="text-align: center; padding: 0.75rem; background: white; border-radius: 8px; border: 1px solid #e5e7eb;">
                        <div style="font-size: 1.5rem; font-weight: 700; color: #dc2626;">{failed_count}</div>
                        <div style="font-size: 0.85rem; color: #64748b;">Failed</div>
                    </div>
                </div>
                <p style="color: #374151; margin: 0; font-size: 0.95rem;">
                    💡 <strong>I've reviewed your assessment.</strong> Ask me about specific findings, remediation steps, or compliance questions!
                </p>
            </div>
            """, unsafe_allow_html=True)
        
        # Display chat messages
        chat_container = st.container()
        with chat_container:
            if not st.session_state.chad_messages:
                if not has_assessment:
                    # No assessment yet - show generic welcome
                    st.markdown("""
                    <div style="background: #f0f9ff; border: 1px solid #bae6fd; border-radius: 8px; padding: 1.5rem; margin: 1rem 0;">
                        <p style="color: #0369a1; margin: 0; font-size: 1rem;">
                            👋 <strong>Hi, I'm Chad (AI)!</strong> I'm your AI security analyst. I can help you with:
                        </p>
                        <ul style="color: #0369a1; margin: 0.5rem 0 0 1.5rem;">
                            <li>Analyzing your security findings and risks</li>
                            <li>Explaining NIST 800-53 controls and requirements</li>
                            <li>Generating remediation scripts and guidance</li>
                            <li>Creating compliance documentation</li>
                            <li>Answering security questions</li>
                        </ul>
                        <p style="color: #0369a1; margin: 0.5rem 0 0 0;">
                            <strong>Tip:</strong> Run an assessment first, and I'll automatically analyze your findings!
                        </p>
                    </div>
                    """, unsafe_allow_html=True)
            
            for message in st.session_state.chad_messages:
                if message["role"] == "user":
                    st.markdown(f"""
                    <div style="background: #e0e7ff; border-radius: 12px; padding: 1rem; margin: 0.5rem 0; margin-left: 2rem;">
                        <strong style="color: #3730a3;">You:</strong>
                        <p style="color: #1e1b4b; margin: 0.25rem 0 0 0;">{message["content"]}</p>
                    </div>
                    """, unsafe_allow_html=True)
                else:
                    st.markdown(f"""
                    <div style="background: #f0fdf4; border: 1px solid #86efac; border-radius: 12px; padding: 1rem; margin: 0.5rem 0; margin-right: 2rem;">
                        <strong style="color: #166534;">🛡️ Chad (AI):</strong>
                        <div style="color: #14532d; margin: 0.25rem 0 0 0;">{message["content"]}</div>
                    </div>
                    """, unsafe_allow_html=True)
        
        # Chat input
        st.markdown("---")
        chat_col1, chat_col2 = st.columns([5, 1])
        with chat_col1:
            user_input = st.text_input(
                "Ask Chad (AI) a question...",
                key="chad_input",
                placeholder="e.g., 'How do I fix the root MFA finding?' or 'Explain AC-2 control requirements'",
                label_visibility="collapsed"
            )
        with chat_col2:
            send_clicked = st.button("Send 📤", use_container_width=True, key="chad_send")
        
        if send_clicked and user_input:
            self._chad_send_message(user_input)
        
        # Clear chat button
        if st.session_state.chad_messages:
            if st.button("🗑️ Clear Chat", key="chad_clear"):
                st.session_state.chad_messages = []
                st.rerun()
    
    def _chad_send_message(self, user_message: str):
        """Send a message to Chad (AI) and get a response via AWS Bedrock."""
        # Add user message to history
        st.session_state.chad_messages.append({"role": "user", "content": user_message})
        
        # Build context from assessment data
        context = self._chad_build_context()
        
        # System prompt for Chad (AI)
        system_prompt = f"""You are Chad (AI), a friendly and knowledgeable AI security analyst working within SAELAR - a NIST 800-53 Rev 5 security assessment platform. 

Your personality:
- Professional but approachable
- Clear and concise in explanations
- Proactive about suggesting next steps
- Uses emojis sparingly to be friendly

Your capabilities:
- Analyze security findings and explain risks
- Explain NIST 800-53 controls and requirements
- Generate remediation scripts (AWS CLI, Terraform, CloudFormation, Python)
- Create POA&M entries and compliance documentation
- Provide security best practices guidance
- **Generate comprehensive cybersecurity reports on demand**

Current Assessment Context:
{context}

=== REPORT GENERATION CAPABILITY ===

When the user asks you to create, generate, or write any type of report, you MUST produce a complete, professional document. 
Recognize these report request patterns: "create a report", "generate a report", "write a report", "give me a report", "I need a report", "make a report", etc.

AVAILABLE REPORT TYPES (generate based on user request):

1. **Executive Summary Report**
   - Overall security posture rating (Critical/High/Medium/Low)
   - Key metrics: controls assessed, pass rate, critical findings
   - Top 5 findings with business impact
   - Recommended immediate actions
   - 30/60/90 day roadmap

2. **Risk Assessment Report** (NIST SP 800-30)
   - Executive summary
   - Methodology (NIST SP 800-30 Rev 1)
   - Threat sources and events
   - Vulnerability analysis with likelihood ratings
   - Impact analysis (C/I/A triad)
   - Risk determination matrix
   - Risk response recommendations
   - Residual risk assessment

3. **POA&M (Plan of Action & Milestones)**
   - Weakness ID and description
   - Affected control family
   - Point of contact placeholder
   - Resource requirements
   - Scheduled completion dates
   - Milestones
   - Current status

4. **Technical Findings Report**
   - Control ID and name
   - Finding description with evidence
   - Affected AWS resources
   - Root cause analysis
   - Remediation steps with CLI commands
   - Verification procedures

5. **Board/Executive Briefing**
   - Non-technical language
   - Security health score
   - Business risk translation
   - Investment needs
   - Strategic recommendations

6. **Compliance Status Report**
   - Overall compliance percentage
   - Breakdown by control family
   - Gap analysis
   - Baseline comparison (Low/Moderate/High)
   - AWS shared responsibility mapping

7. **Remediation Plan**
   - Prioritized roadmap
   - Resource requirements
   - Quick wins list
   - Detailed scripts (AWS CLI, Terraform)
   - Validation procedures

8. **Audit Package**
   - SSP summary
   - Control implementation matrix
   - Evidence documentation
   - POA&M
   - Authorization recommendation

FORMAT REQUIREMENTS:
- Use clear headers and sections
- Use markdown formatting for readability
- Include tables where appropriate
- Base ALL content on the actual assessment data provided
- If no assessment data, inform user to run an assessment first
- Make reports detailed and professional-grade

=== END REPORT CAPABILITY ===

Guidelines:
- Always base your answers on the provided assessment data when available
- If you don't have specific data, say so and provide general guidance
- For remediation, provide specific actionable steps
- Format code blocks properly for easy copying
- Be security-focused and compliance-aware
- When asked to create a report, ALWAYS generate a complete document, not just an outline"""
        
        try:
            with st.spinner("🤔 Chad (AI) is thinking..."):
                # Build messages for API
                messages = []
                for msg in st.session_state.chad_messages:
                    messages.append({"role": msg["role"], "content": msg["content"]})
                
                # Call AI (uses Ollama in air-gapped mode, Bedrock otherwise)
                assistant_message = call_ai(
                    messages=messages,
                    system_prompt=system_prompt,
                    max_tokens=8192  # Increased for comprehensive report generation
                )
                
                # Add assistant response to history
                st.session_state.chad_messages.append({"role": "assistant", "content": assistant_message})
                
                st.rerun()
                
        except Exception as e:
            st.error(f"❌ Error communicating with Bedrock: {str(e)}")
            st.session_state.chad_messages.pop()  # Remove the user message since we couldn't respond
    
    def _chad_draft_risk_acceptance(self, finding_label: str):
        """Use Chad (AI) to draft Operational Justification and Compensating Controls for the selected finding. Sets session state and reruns."""
        if not finding_label or finding_label.startswith("No POA&M"):
            st.warning("Select a POA&M finding first, then generate an SSP so Control / Finding has options.")
            return
        context = self._chad_build_context()
        system_prompt = """You are Chad (AI), a security analyst. Generate exactly two short sections for a Risk Acceptance form.
Output format (use these exact headers so the parser can split):
**Operational Justification:**
(2-4 sentences: why this risk is being accepted, operational or business need, and why remediation is deferred.)

**Compensating Controls:**
(2-4 sentences: what alternative or compensating measures are in place to mitigate this risk.)

Be concise and professional. Base content on the finding and any assessment context."""
        user_msg = f"""Draft an Operational Justification and Compensating Controls for this finding:

Finding: {finding_label}

Assessment context (if any):
{context[:1500] if context else 'None'}"""
        try:
            with st.spinner("🤔 Chad (AI) is drafting justification and compensating controls..."):
                response = call_ai(
                    messages=[{"role": "user", "content": user_msg}],
                    system_prompt=system_prompt,
                    max_tokens=1024,
                )
            # Parse **Operational Justification:** ... **Compensating Controls:** ...
            just_text = ""
            comp_text = ""
            if "**Operational Justification:**" in response and "**Compensating Controls:**" in response:
                parts = response.split("**Compensating Controls:**", 1)
                just_part = parts[0].split("**Operational Justification:**", 1)[-1].strip()
                just_text = just_part.strip().strip(":").strip()
                comp_text = parts[1].strip().strip(":").strip() if len(parts) > 1 else ""
            elif "Operational Justification:" in response and "Compensating Controls:" in response:
                parts = response.split("Compensating Controls:", 1)
                just_text = parts[0].split("Operational Justification:", 1)[-1].strip().strip(":").strip()
                comp_text = parts[1].strip().strip(":").strip() if len(parts) > 1 else ""
            else:
                # Fallback: first paragraph = justification, second = compensating
                blocks = [b.strip() for b in response.replace("**", "").split("\n\n") if b.strip()]
                if len(blocks) >= 2:
                    just_text, comp_text = blocks[0], blocks[1]
                else:
                    just_text = response[:1500]
            st.session_state.ra_chad_justification = just_text
            st.session_state.ra_chad_compensating = comp_text
            st.success("Chad (AI) draft loaded. Review and edit the fields below before submitting.")
            st.rerun()
        except Exception as e:
            st.error(f"Chad (AI) could not generate draft: {str(e)}")
    
    def _chad_build_context(self) -> str:
        """Build context string from current assessment data for Chad (AI)."""
        context_parts = []
        
        # Add assessment results if available
        if st.session_state.get('results'):
            results = st.session_state.results
            assessor = st.session_state.get('assessor')
            
            if assessor:
                summary = assessor.generate_summary(results)
                context_parts.append(f"Latest Assessment Summary:")
                context_parts.append(f"- Account: {summary.get('account_id', 'Unknown')}")
                context_parts.append(f"- Total Controls Assessed: {summary.get('total_controls', 0)}")
                context_parts.append(f"- Passed: {summary.get('passed', 0)}")
                context_parts.append(f"- Failed: {summary.get('failed', 0)}")
                context_parts.append(f"- Warnings: {summary.get('warnings', 0)}")
                
                # Add failed controls
                failed_controls = [r for r in results if r.status.value.startswith("❌")]
                if failed_controls:
                    context_parts.append(f"\nFailed Controls ({len(failed_controls)}):")
                    for ctrl in failed_controls[:10]:  # Limit to 10
                        context_parts.append(f"- {ctrl.control_id}: {ctrl.control_name}")
                        if ctrl.findings:
                            context_parts.append(f"  Findings: {'; '.join(ctrl.findings[:3])}")
                
                # Add warnings
                warning_controls = [r for r in results if r.status.value.startswith("⚠️")]
                if warning_controls:
                    context_parts.append(f"\nWarning Controls ({len(warning_controls)}):")
                    for ctrl in warning_controls[:10]:  # Limit to 10
                        context_parts.append(f"- {ctrl.control_id}: {ctrl.control_name}")
        else:
            context_parts.append("No assessment has been run yet. The user may want to run an assessment first.")
        
        return "\n".join(context_parts) if context_parts else "No assessment data available."
    
    def render_risk_calculator_page(self):
        """Render the Risk Calculator page with full visualizations."""
        from risk_score_calculator import (
            RiskScoreCalculator, RiskAssessment, Finding, 
            Likelihood, Impact, RiskLevel
        )
        from datetime import datetime
        import json
        
        try:
            import plotly.express as px
            import plotly.graph_objects as go
            import pandas as pd
            PLOTLY_AVAILABLE = True
        except ImportError:
            PLOTLY_AVAILABLE = False
        
        def get_risk_color(level):
            return {
                RiskLevel.CRITICAL: "#dc2626",
                RiskLevel.HIGH: "#f59e0b",
                RiskLevel.MEDIUM: "#eab308",
                RiskLevel.LOW: "#22c55e"
            }.get(level, "#6b7280")
        
        def render_risk_gauge(score, max_score=25):
            percentage = (score / max_score) * 100
            if percentage >= 68:
                color, level = "#dc2626", "CRITICAL"
            elif percentage >= 40:
                color, level = "#f59e0b", "HIGH"
            elif percentage >= 20:
                color, level = "#eab308", "MEDIUM"
            else:
                color, level = "#22c55e", "LOW"
            
            fig = go.Figure(go.Indicator(
                mode="gauge+number",
                value=score,
                domain={'x': [0, 1], 'y': [0, 1]},
                title={'text': f"Risk Score<br><span style='font-size:0.8em;color:{color}'>{level}</span>"},
                gauge={
                    'axis': {'range': [0, 25], 'tickwidth': 1},
                    'bar': {'color': color},
                    'steps': [
                        {'range': [0, 4], 'color': "#dcfce7"},
                        {'range': [4, 9], 'color': "#fef9c3"},
                        {'range': [9, 16], 'color': "#fed7aa"},
                        {'range': [16, 25], 'color': "#fecaca"}
                    ],
                    'threshold': {'line': {'color': "black", 'width': 4}, 'thickness': 0.75, 'value': score}
                }
            ))
            fig.update_layout(height=250, margin=dict(l=20, r=20, t=50, b=20), paper_bgcolor='rgba(0,0,0,0)')
            return fig
        
        def render_risk_distribution(findings):
            risk_counts = {
                'CRITICAL': len([f for f in findings if f.risk_level == RiskLevel.CRITICAL]),
                'HIGH': len([f for f in findings if f.risk_level == RiskLevel.HIGH]),
                'MEDIUM': len([f for f in findings if f.risk_level == RiskLevel.MEDIUM]),
                'LOW': len([f for f in findings if f.risk_level == RiskLevel.LOW])
            }
            colors = ['#dc2626', '#f59e0b', '#eab308', '#22c55e']
            fig = px.pie(values=list(risk_counts.values()), names=list(risk_counts.keys()),
                        color_discrete_sequence=colors, hole=0.4)
            fig.update_layout(height=300, margin=dict(l=20, r=20, t=30, b=20), paper_bgcolor='rgba(0,0,0,0)',
                            legend=dict(orientation="h", yanchor="bottom", y=-0.2, xanchor="center", x=0.5))
            return fig
        
        def render_risk_by_family(findings):
            family_data = {}
            for f in findings:
                if f.control_family not in family_data:
                    family_data[f.control_family] = {'count': 0, 'total_risk': 0}
                family_data[f.control_family]['count'] += 1
                family_data[f.control_family]['total_risk'] += f.risk_score
            
            families = list(family_data.keys())
            avg_risks = [family_data[f]['total_risk'] / family_data[f]['count'] for f in families]
            sorted_data = sorted(zip(families, avg_risks), key=lambda x: x[1], reverse=True)
            if sorted_data:
                families, avg_risks = zip(*sorted_data)
            else:
                families, avg_risks = [], []
            
            colors = ['#dc2626' if r > 16 else '#f59e0b' if r > 9 else '#eab308' if r > 4 else '#22c55e' for r in avg_risks]
            fig = go.Figure(data=[go.Bar(x=list(families), y=list(avg_risks), marker_color=colors,
                                        text=[f"{r:.1f}" for r in avg_risks], textposition='outside')])
            fig.update_layout(title="Average Risk by Control Family", xaxis_title="Control Family",
                            yaxis_title="Average Risk Score", height=350, margin=dict(l=20, r=20, t=50, b=20),
                            paper_bgcolor='rgba(0,0,0,0)', plot_bgcolor='rgba(0,0,0,0)', yaxis=dict(range=[0, 25]))
            return fig
        
        def render_risk_matrix_heatmap():
            impacts = ['Negligible', 'Minor', 'Moderate', 'Significant', 'Severe']
            likelihoods = ['Very Low', 'Low', 'Medium', 'High', 'Very High']
            matrix_data = []
            for i, likelihood in enumerate(likelihoods, 1):
                row = {'Likelihood': likelihood}
                for j, impact in enumerate(impacts, 1):
                    row[impact] = i * j
                matrix_data.append(row)
            df = pd.DataFrame(matrix_data).set_index('Likelihood')
            
            fig = go.Figure(data=go.Heatmap(
                z=df.values, x=impacts, y=likelihoods,
                colorscale=[[0, '#22c55e'], [0.16, '#22c55e'], [0.16, '#eab308'], [0.36, '#eab308'],
                           [0.36, '#f59e0b'], [0.64, '#f59e0b'], [0.64, '#dc2626'], [1, '#dc2626']],
                showscale=True, text=df.values, texttemplate="%{text}",
                textfont={"size": 14, "color": "white"},
                hovertemplate="Likelihood: %{y}<br>Impact: %{x}<br>Score: %{z}<extra></extra>"
            ))
            fig.update_layout(xaxis_title="Impact", yaxis_title="Likelihood", height=400,
                            margin=dict(l=20, r=20, t=30, b=20), paper_bgcolor='rgba(0,0,0,0)')
            return fig
        
        def create_sample_assessment():
            assessment = RiskAssessment(
                assessment_id=f"RA-{datetime.now().strftime('%Y%m%d%H%M%S')}",
                assessment_name="AWS Security Assessment",
                organization=st.session_state.get('aws_account_id', 'Security Architecture and Evaluation'),
                assessor="SAE Security Team", date=datetime.now()
            )
            sample_findings = [
                Finding(finding_id="FIND-001", title="Missing MFA on Root Account",
                       description="Root account does not have MFA enabled.", control_family="AC", control_id="AC-2",
                       likelihood=Likelihood.HIGH, impact=Impact.SEVERE,
                       remediation="Enable MFA on root account immediately."),
                Finding(finding_id="FIND-002", title="CloudTrail Not Enabled in All Regions",
                       description="CloudTrail logging is not enabled in all AWS regions.", control_family="AU", control_id="AU-2",
                       likelihood=Likelihood.MEDIUM, impact=Impact.SIGNIFICANT,
                       remediation="Enable CloudTrail in all regions with multi-region trail."),
                Finding(finding_id="FIND-003", title="Unencrypted S3 Buckets",
                       description="Several S3 buckets do not have default encryption enabled.", control_family="SC", control_id="SC-8",
                       likelihood=Likelihood.MEDIUM, impact=Impact.MODERATE,
                       remediation="Enable default encryption on all S3 buckets."),
                Finding(finding_id="FIND-004", title="Weak Password Policy",
                       description="IAM password policy does not meet complexity requirements.", control_family="IA", control_id="IA-5",
                       likelihood=Likelihood.HIGH, impact=Impact.SIGNIFICANT,
                       remediation="Update IAM password policy to require minimum 14 characters with complexity."),
                Finding(finding_id="FIND-005", title="No Baseline Configuration",
                       description="EC2 instances lack documented baseline configurations.", control_family="CM", control_id="CM-2",
                       likelihood=Likelihood.LOW, impact=Impact.MODERATE,
                       remediation="Implement AWS Systems Manager State Manager for configuration baselines."),
            ]
            for finding in sample_findings:
                assessment.add_finding(finding)
            return assessment
        
        # Note: convert_nist_to_risk_findings is now at module level for reuse
        
        # Header
        st.markdown("""
        <div style="background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 50%, #3d7ab5 100%); padding: 1.5rem;
                    border-radius: 12px; margin-bottom: 1.5rem; border-left: 5px solid #f59e0b;">
            <h2 style="color: #ffffff; margin: 0;">📊 SAELAR-53 Risk Score Calculator</h2>
            <p style="color: #b8d4e8; margin: 0.5rem 0 0 0;">Calculate, visualize, and prioritize security risks</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Initialize session state
        if 'risk_custom_findings' not in st.session_state:
            st.session_state.risk_custom_findings = []
        if 'risk_use_demo' not in st.session_state:
            st.session_state.risk_use_demo = True
        if 'risk_nist_imported' not in st.session_state:
            st.session_state.risk_nist_imported = False
        if 'risk_nist_findings' not in st.session_state:
            st.session_state.risk_nist_findings = []
        
        # Check if NIST assessment results are available
        nist_results_available = (st.session_state.get('results') is not None and 
                                  st.session_state.get('assessor') is not None)
        
        # Show NIST Assessment connection banner if results available
        if nist_results_available:
            nist_results = st.session_state.results
            nist_assessor = st.session_state.assessor
            summary = nist_assessor.generate_summary(nist_results)
            failed_warnings = summary.get('failed', 0) + summary.get('warnings', 0)
            
            # Check if already auto-imported
            already_imported = st.session_state.get('risk_nist_imported', False)
            import_status = "✅ Auto-imported" if already_imported else "📥 Ready to import"
            
            st.markdown(f"""
            <div style="background: linear-gradient(135deg, #059669 0%, #047857 100%); padding: 1rem 1.5rem;
                        border-radius: 8px; margin-bottom: 1rem; display: flex; align-items: center; justify-content: space-between;">
                <div>
                    <span style="color: #ffffff; font-weight: 600;">🔗 NIST Assessment Connected</span>
                    <span style="color: #d1fae5; margin-left: 1rem;">
                        {failed_warnings} findings | {import_status}
                    </span>
                </div>
            </div>
            """, unsafe_allow_html=True)
            
            col_import, col_poam, col_status = st.columns([1, 1, 2])
            with col_import:
                button_label = "🔄 Refresh NIST Data" if already_imported else "📥 Import NIST Findings"
                if st.button(button_label, use_container_width=True, type="primary", key="import_nist_btn"):
                    # Convert NIST results to risk findings
                    imported_findings = convert_nist_to_risk_findings(nist_results, nist_assessor)
                    st.session_state.risk_nist_findings = imported_findings
                    st.session_state.risk_nist_imported = True
                    st.session_state.risk_use_demo = False
                    st.success(f"✅ {'Refreshed' if already_imported else 'Imported'} {len(imported_findings)} findings from NIST Assessment!")
                    st.rerun()
            
            with col_poam:
                # POA&M toggle button
                if st.button("📋 View POA&Ms", use_container_width=True, key="toggle_poam_view"):
                    st.session_state.show_risk_poams = not st.session_state.get('show_risk_poams', False)
                    st.rerun()
            
            with col_status:
                if st.session_state.risk_nist_imported:
                    st.markdown(f"""
                    <div style="background: #dcfce7; padding: 0.5rem 1rem; border-radius: 6px; border-left: 4px solid #22c55e;">
                        <span style="color: #166534;">✅ {len(st.session_state.risk_nist_findings)} NIST findings loaded</span>
                    </div>
                    """, unsafe_allow_html=True)
            
            # Show POA&Ms if toggled
            if st.session_state.get('show_risk_poams', False):
                st.markdown("---")
                st.markdown("### 📋 POA&M Items from Risk Assessment")
                
                ssp_data = st.session_state.get('ssp_data')
                if ssp_data and ssp_data.get('poam'):
                    poam_items = ssp_data.get('poam', [])
                    
                    # Summary
                    high_count = sum(1 for p in poam_items if p.get('risk_level') == 'High')
                    medium_count = sum(1 for p in poam_items if p.get('risk_level') == 'Medium')
                    low_count = sum(1 for p in poam_items if p.get('risk_level') == 'Low')
                    
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total", len(poam_items))
                    with col2:
                        st.metric("🔴 High", high_count)
                    with col3:
                        st.metric("🟡 Medium", medium_count)
                    with col4:
                        st.metric("🟢 Low", low_count)
                    
                    # Display items in a table-like format
                    for item in poam_items[:15]:  # Limit to 15 items
                        risk_level = item.get('risk_level', 'Unknown')
                        risk_colors = {'High': '#ef4444', 'Medium': '#f59e0b', 'Low': '#10b981'}
                        color = risk_colors.get(risk_level, '#6b7280')
                        
                        with st.expander(f"**{item.get('control_id', 'N/A')}** - {risk_level} Risk"):
                            st.markdown(f"""
                            - **POA&M ID:** {item.get('poam_id', 'N/A')}
                            - **Weakness:** {item.get('weakness', 'N/A')}
                            - **Remediation:** {item.get('remediation_plan', 'N/A')}
                            - **Due Date:** {item.get('scheduled_completion', 'TBD')}
                            - **Status:** {item.get('status', 'Open')}
                            """)
                    
                    if len(poam_items) > 15:
                        st.info(f"Showing 15 of {len(poam_items)} items. View all in the SSP Generator → POA&Ms tab.")
                else:
                    st.info("📭 No POA&M items available. Generate an SSP first to create POA&M items from assessment findings.")
                
                # Close button
                if st.button("✖️ Close POA&Ms", use_container_width=True, key="close_poam_view"):
                    st.session_state.show_risk_poams = False
                    st.rerun()
        else:
            st.info("💡 **Tip:** Run a NIST Assessment first to automatically import findings into the Risk Calculator!")
        
        st.markdown("---")
        
        # Configuration row
        col_cfg1, col_cfg2, col_cfg3 = st.columns([2, 2, 1])
        with col_cfg1:
            uploaded_file = st.file_uploader("📂 Upload NIST Assessment JSON", type=['json'], key="risk_file_upload")
        with col_cfg2:
            # Only show demo checkbox if no NIST findings imported
            if not st.session_state.risk_nist_imported:
                use_demo = st.checkbox("Use Demo Data", value=st.session_state.risk_use_demo, key="risk_demo_check")
                st.session_state.risk_use_demo = use_demo
            else:
                use_demo = False
                st.markdown("*Using NIST Assessment data*")
        with col_cfg3:
            if st.button("🗑️ Clear All", use_container_width=True):
                st.session_state.risk_custom_findings = []
                st.session_state.risk_nist_findings = []
                st.session_state.risk_nist_imported = False
                st.session_state.risk_use_demo = True
                st.rerun()
        
        # Add custom finding expander
        with st.expander("➕ Add Custom Finding"):
            col1, col2 = st.columns(2)
            with col1:
                finding_title = st.text_input("Title", key="risk_title")
                control_family = st.selectbox("Control Family", 
                    ['AC', 'AU', 'CA', 'CM', 'CP', 'IA', 'IR', 'MA', 'MP', 'PE', 'PL', 'PS', 'RA', 'SA', 'SC', 'SI', 'SR'], key="risk_family")
                likelihood = st.selectbox("Likelihood", [l.name for l in Likelihood], index=2, key="risk_likelihood")
            with col2:
                finding_desc = st.text_area("Description", height=68, key="risk_desc")
                control_id = st.text_input("Control ID", placeholder="e.g., AC-2", key="risk_control_id")
                impact = st.selectbox("Impact", [i.name for i in Impact], index=2, key="risk_impact")
            remediation = st.text_input("Remediation", key="risk_remediation")
            
            if st.button("➕ Add Finding", use_container_width=True, key="risk_add_btn"):
                if finding_title:
                    new_finding = Finding(
                        finding_id=f"FIND-CUSTOM-{len(st.session_state.risk_custom_findings)+1}",
                        title=finding_title, description=finding_desc, control_family=control_family,
                        control_id=control_id or control_family, likelihood=Likelihood[likelihood],
                        impact=Impact[impact], remediation=remediation
                    )
                    st.session_state.risk_custom_findings.append(new_finding)
                    st.success(f"✅ Added: {finding_title}")
                    st.rerun()
        
        # Load assessment - Priority: 1) NIST imported, 2) Uploaded file, 3) Demo data
        if st.session_state.risk_nist_imported and st.session_state.risk_nist_findings:
            # Use imported NIST Assessment findings
            assessment = RiskAssessment(
                assessment_id=f"RA-NIST-{datetime.now().strftime('%Y%m%d%H%M%S')}",
                assessment_name="NIST 800-53 Assessment Risk Analysis",
                organization=st.session_state.get('aws_account_id', 'AWS Account'),
                assessor="SAELAR-53 Risk Calculator", date=datetime.now()
            )
            for finding in st.session_state.risk_nist_findings:
                assessment.add_finding(finding)
        elif uploaded_file:
            try:
                data = json.load(uploaded_file)
                calculator = RiskScoreCalculator()
                assessment = RiskAssessment(
                    assessment_id=f"RA-{datetime.now().strftime('%Y%m%d%H%M%S')}",
                    assessment_name=data.get('assessment_info', {}).get('scope', 'Uploaded Assessment'),
                    organization=data.get('assessment_info', {}).get('account_id', 'Unknown'),
                    assessor="SAELAR-53", date=datetime.now()
                )
                for result in data.get('results', []):
                    if result.get('status') != 'PASS':
                        finding = calculator.create_finding_from_nist_result(
                            control_id=result.get('control_id', 'Unknown'),
                            control_name=result.get('control_name', 'Unknown'),
                            family=result.get('family', 'Unknown'),
                            status=result.get('status', 'FAIL'),
                            findings_list=result.get('findings', []),
                            recommendations=result.get('recommendations', [])
                        )
                        if finding:
                            assessment.add_finding(finding)
                st.success("✅ Assessment loaded!")
            except Exception as e:
                st.error(f"Error: {e}")
                assessment = create_sample_assessment()
        elif use_demo:
            assessment = create_sample_assessment()
        else:
            assessment = RiskAssessment(
                assessment_id=f"RA-{datetime.now().strftime('%Y%m%d%H%M%S')}",
                assessment_name="Custom Assessment",
                organization=st.session_state.get('aws_account_id', 'Organization'),
                assessor="SAELAR-53", date=datetime.now()
            )
        
        # Add custom findings
        for finding in st.session_state.risk_custom_findings:
            assessment.add_finding(finding)
        
        if not assessment.findings:
            st.info("No findings to display. Enable demo data or add custom findings.")
            return
        
        # Create tabs - Added NIST 800-30 Enhanced tab
        tab1, tab2, tab3, tab4, tab5_risk = st.tabs(["📊 Dashboard", "📋 Findings", "📐 Risk Matrix", "📄 Report", "🔬 NIST 800-30"])
        
        with tab1:
            # Summary metrics
            col1, col2, col3, col4, col5 = st.columns(5)
            risk_color = get_risk_color(assessment.overall_risk_level)
            
            with col1:
                st.markdown(f"""<div style="background:#fff;padding:1rem;border-radius:10px;text-align:center;border:1px solid #e5e7eb;">
                    <div style="font-size:2rem;font-weight:700;color:{risk_color};">{assessment.overall_risk_level.value}</div>
                    <div style="color:#4a5568;font-size:0.8rem;">OVERALL RISK</div></div>""", unsafe_allow_html=True)
            with col2:
                st.markdown(f"""<div style="background:#fff;padding:1rem;border-radius:10px;text-align:center;border:1px solid #e5e7eb;">
                    <div style="font-size:2rem;font-weight:700;">{assessment.average_risk_score:.1f}</div>
                    <div style="color:#4a5568;font-size:0.8rem;">AVG SCORE</div></div>""", unsafe_allow_html=True)
            with col3:
                st.markdown(f"""<div style="background:#fff;padding:1rem;border-radius:10px;text-align:center;border:1px solid #e5e7eb;">
                    <div style="font-size:2rem;font-weight:700;">{assessment.total_findings}</div>
                    <div style="color:#4a5568;font-size:0.8rem;">FINDINGS</div></div>""", unsafe_allow_html=True)
            with col4:
                st.markdown(f"""<div style="background:#fff;padding:1rem;border-radius:10px;text-align:center;border:1px solid #e5e7eb;">
                    <div style="font-size:2rem;font-weight:700;">{assessment.compliance_score:.0f}%</div>
                    <div style="color:#4a5568;font-size:0.8rem;">COMPLIANCE</div></div>""", unsafe_allow_html=True)
            with col5:
                st.markdown(f"""<div style="background:#fff;padding:1rem;border-radius:10px;text-align:center;border:1px solid #e5e7eb;">
                    <div style="font-size:2rem;font-weight:700;color:#dc2626;">{assessment.open_findings}</div>
                    <div style="color:#4a5568;font-size:0.8rem;">OPEN ISSUES</div></div>""", unsafe_allow_html=True)
            
            st.markdown("---")
            
            if PLOTLY_AVAILABLE:
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("### 📈 Risk Gauge")
                    st.plotly_chart(render_risk_gauge(assessment.average_risk_score), use_container_width=True)
                with col2:
                    st.markdown("### 📊 Risk Distribution")
                    st.plotly_chart(render_risk_distribution(assessment.findings), use_container_width=True)
                
                # Risk Distribution Details - Clickable expandable sections
                st.markdown("<p style='color: #64748b; font-size: 0.9rem;'>Click any risk level below for detailed findings</p>", unsafe_allow_html=True)
                
                # Group findings by risk level
                risk_level_findings = {
                    'CRITICAL': [f for f in assessment.findings if f.risk_level == RiskLevel.CRITICAL],
                    'HIGH': [f for f in assessment.findings if f.risk_level == RiskLevel.HIGH],
                    'MEDIUM': [f for f in assessment.findings if f.risk_level == RiskLevel.MEDIUM],
                    'LOW': [f for f in assessment.findings if f.risk_level == RiskLevel.LOW]
                }
                
                # Risk level metadata
                risk_level_details = {
                    'CRITICAL': {
                        'emoji': '🔴',
                        'color': '#dc2626',
                        'score_range': '17-25',
                        'description': 'Immediate action required. These findings pose severe risk to the organization and should be remediated within 24-72 hours.',
                        'priority': 'P1 - Emergency',
                        'sla': '24-72 hours',
                        'escalation': 'CISO, System Owner, Authorizing Official',
                        'actions': [
                            'Immediately notify security leadership',
                            'Initiate incident response if exploitation suspected',
                            'Implement emergency compensating controls',
                            'Document risk acceptance if remediation delayed'
                        ]
                    },
                    'HIGH': {
                        'emoji': '🟠',
                        'color': '#f59e0b',
                        'score_range': '10-16',
                        'description': 'Urgent attention needed. These findings represent significant risk and should be prioritized in the next sprint or maintenance window.',
                        'priority': 'P2 - High',
                        'sla': '7-14 days',
                        'escalation': 'ISSO, System Owner',
                        'actions': [
                            'Add to remediation backlog with high priority',
                            'Assess compensating controls',
                            'Include in next POA&M update',
                            'Monitor for exploitation attempts'
                        ]
                    },
                    'MEDIUM': {
                        'emoji': '🟡',
                        'color': '#eab308',
                        'score_range': '5-9',
                        'description': 'Should be addressed in planned maintenance cycles. These findings represent moderate risk that should be tracked and remediated.',
                        'priority': 'P3 - Medium',
                        'sla': '30-60 days',
                        'escalation': 'ISSO',
                        'actions': [
                            'Add to standard remediation queue',
                            'Include in quarterly security review',
                            'Document in POA&M',
                            'Consider during next architecture review'
                        ]
                    },
                    'LOW': {
                        'emoji': '🟢',
                        'color': '#22c55e',
                        'score_range': '1-4',
                        'description': 'Address as resources permit. These findings represent minimal risk but should still be tracked for continuous improvement.',
                        'priority': 'P4 - Low',
                        'sla': '90+ days',
                        'escalation': 'None required',
                        'actions': [
                            'Track for future remediation',
                            'Consider during major upgrades',
                            'Include in annual security review',
                            'May accept risk with documentation'
                        ]
                    }
                }
                
                # Display expandable details for each risk level
                for level in ['CRITICAL', 'HIGH', 'MEDIUM', 'LOW']:
                    findings_list = risk_level_findings[level]
                    details = risk_level_details[level]
                    count = len(findings_list)
                    
                    if count > 0:
                        with st.expander(f"{details['emoji']} **{level}** — {count} finding{'s' if count != 1 else ''} | Score Range: {details['score_range']} | {details['priority']}"):
                            st.markdown(f"*{details['description']}*")
                            
                            col1, col2 = st.columns(2)
                            
                            with col1:
                                st.markdown("**Response Requirements:**")
                                st.markdown(f"- ⏱️ **SLA:** {details['sla']}")
                                st.markdown(f"- 📢 **Escalation:** {details['escalation']}")
                                
                                st.markdown("**Recommended Actions:**")
                                for action in details['actions']:
                                    st.markdown(f"- {action}")
                            
                            with col2:
                                st.markdown("**Findings in this Category:**")
                                for finding in findings_list[:10]:  # Show up to 10
                                    st.markdown(f"- `{finding.control_id}`: {finding.title[:50]}{'...' if len(finding.title) > 50 else ''}")
                                    st.caption(f"  Score: {finding.risk_score:.0f} | {finding.control_family}")
                                
                                if count > 10:
                                    st.info(f"... and {count - 10} more findings")
                
                st.markdown("---")
                st.plotly_chart(render_risk_by_family(assessment.findings), use_container_width=True)
                
                # Control Family Details - Clickable expandable sections
                st.markdown("<p style='color: #64748b; font-size: 0.9rem;'>Click any control family below for detailed risk analysis</p>", unsafe_allow_html=True)
                
                # Build family risk data from current findings
                family_risk_data = {}
                for finding in assessment.findings:
                    if finding.control_family not in family_risk_data:
                        family_risk_data[finding.control_family] = {
                            'findings': [],
                            'total_risk': 0,
                            'count': 0
                        }
                    family_risk_data[finding.control_family]['findings'].append(finding)
                    family_risk_data[finding.control_family]['total_risk'] += finding.risk_score
                    family_risk_data[finding.control_family]['count'] += 1
                
                # Control family metadata
                control_family_details = {
                    "AC": {
                        "name": "Access Control",
                        "description": "Policies and mechanisms to limit system access to authorized users.",
                        "key_controls": ["AC-2 Account Management", "AC-3 Access Enforcement", "AC-6 Least Privilege"],
                        "threats_addressed": ["Unauthorized Access", "Privilege Escalation", "Insider Threats"],
                        "aws_services": ["IAM", "Organizations", "SSO", "Resource Access Manager"]
                    },
                    "AU": {
                        "name": "Audit and Accountability",
                        "description": "Creation, protection, and analysis of audit records.",
                        "key_controls": ["AU-2 Audit Events", "AU-6 Audit Review", "AU-9 Protection of Audit Info"],
                        "threats_addressed": ["Log Tampering", "Forensic Evasion", "Compliance Gaps"],
                        "aws_services": ["CloudTrail", "CloudWatch Logs", "Config", "Security Hub"]
                    },
                    "CA": {
                        "name": "Assessment and Authorization",
                        "description": "Security assessments, authorizations, and continuous monitoring.",
                        "key_controls": ["CA-2 Security Assessments", "CA-7 Continuous Monitoring"],
                        "threats_addressed": ["Unassessed Risks", "Compliance Drift", "Shadow IT"],
                        "aws_services": ["Inspector", "Audit Manager", "Security Hub"]
                    },
                    "CM": {
                        "name": "Configuration Management",
                        "description": "Baseline configurations and change management processes.",
                        "key_controls": ["CM-2 Baseline Configuration", "CM-6 Configuration Settings", "CM-7 Least Functionality"],
                        "threats_addressed": ["Misconfigurations", "Configuration Drift", "Vulnerable Defaults"],
                        "aws_services": ["Config", "Systems Manager", "Service Catalog"]
                    },
                    "CP": {
                        "name": "Contingency Planning",
                        "description": "Business continuity and disaster recovery capabilities.",
                        "key_controls": ["CP-9 System Backup", "CP-10 System Recovery"],
                        "threats_addressed": ["Service Disruption", "Data Loss", "Ransomware"],
                        "aws_services": ["Backup", "S3 Replication", "RDS Snapshots", "DRS"]
                    },
                    "IA": {
                        "name": "Identification and Authentication",
                        "description": "User and device identity verification mechanisms.",
                        "key_controls": ["IA-2 Identification and Authentication", "IA-5 Authenticator Management"],
                        "threats_addressed": ["Credential Theft", "Identity Spoofing", "Brute Force"],
                        "aws_services": ["IAM", "Cognito", "Directory Service", "SSO"]
                    },
                    "IR": {
                        "name": "Incident Response",
                        "description": "Incident detection, response, and recovery procedures.",
                        "key_controls": ["IR-4 Incident Handling", "IR-5 Incident Monitoring", "IR-6 Incident Reporting"],
                        "threats_addressed": ["Delayed Detection", "Inadequate Response", "Evidence Loss"],
                        "aws_services": ["GuardDuty", "Detective", "Security Hub", "EventBridge"]
                    },
                    "MA": {
                        "name": "Maintenance",
                        "description": "System maintenance and update procedures.",
                        "key_controls": ["MA-2 Controlled Maintenance", "MA-4 Nonlocal Maintenance"],
                        "threats_addressed": ["Unpatched Vulnerabilities", "Maintenance Access Abuse"],
                        "aws_services": ["Systems Manager Patch Manager", "Maintenance Windows"]
                    },
                    "MP": {
                        "name": "Media Protection",
                        "description": "Protection of digital and physical media.",
                        "key_controls": ["MP-2 Media Access", "MP-4 Media Storage"],
                        "threats_addressed": ["Data Leakage", "Unauthorized Media Access"],
                        "aws_services": ["S3", "EBS Encryption", "Macie"]
                    },
                    "PE": {
                        "name": "Physical and Environmental Protection",
                        "description": "Physical access controls and environmental safeguards.",
                        "key_controls": ["PE-2 Physical Access Authorizations", "PE-3 Physical Access Control"],
                        "threats_addressed": ["Physical Intrusion", "Environmental Hazards"],
                        "aws_services": ["Managed by AWS (Shared Responsibility)"]
                    },
                    "PL": {
                        "name": "Planning",
                        "description": "Security planning and system security plans.",
                        "key_controls": ["PL-2 System Security Plan", "PL-4 Rules of Behavior"],
                        "threats_addressed": ["Inadequate Planning", "Policy Gaps"],
                        "aws_services": ["Artifact", "Audit Manager"]
                    },
                    "PS": {
                        "name": "Personnel Security",
                        "description": "Personnel screening and security awareness.",
                        "key_controls": ["PS-2 Position Risk Designation", "PS-3 Personnel Screening"],
                        "threats_addressed": ["Insider Threats", "Social Engineering"],
                        "aws_services": ["IAM", "Organizations SCPs"]
                    },
                    "RA": {
                        "name": "Risk Assessment",
                        "description": "Risk identification, analysis, and management.",
                        "key_controls": ["RA-3 Risk Assessment", "RA-5 Vulnerability Scanning"],
                        "threats_addressed": ["Unknown Vulnerabilities", "Unassessed Risks"],
                        "aws_services": ["Inspector", "Security Hub", "GuardDuty"]
                    },
                    "SA": {
                        "name": "System and Services Acquisition",
                        "description": "Secure system development and acquisition processes.",
                        "key_controls": ["SA-3 System Development Life Cycle", "SA-4 Acquisition Process"],
                        "threats_addressed": ["Supply Chain Risks", "Insecure Development"],
                        "aws_services": ["CodePipeline", "CodeBuild", "ECR Scanning"]
                    },
                    "SC": {
                        "name": "System and Communications Protection",
                        "description": "Network and data protection mechanisms.",
                        "key_controls": ["SC-7 Boundary Protection", "SC-8 Transmission Confidentiality", "SC-28 Protection at Rest"],
                        "threats_addressed": ["Data Interception", "Network Attacks", "Data Exfiltration"],
                        "aws_services": ["VPC", "Security Groups", "KMS", "ACM", "WAF"]
                    },
                    "SI": {
                        "name": "System and Information Integrity",
                        "description": "Flaw remediation, malware protection, and monitoring.",
                        "key_controls": ["SI-2 Flaw Remediation", "SI-3 Malware Protection", "SI-4 System Monitoring"],
                        "threats_addressed": ["Malware", "Unpatched Systems", "Integrity Violations"],
                        "aws_services": ["GuardDuty", "Inspector", "Macie", "Config Rules"]
                    }
                }
                
                # Display expandable details for each family with findings
                for family_code, data in sorted(family_risk_data.items(), key=lambda x: x[1]['total_risk'] / x[1]['count'], reverse=True):
                    avg_risk = data['total_risk'] / data['count']
                    finding_count = data['count']
                    
                    # Get family details
                    family_info = control_family_details.get(family_code, {
                        "name": family_code,
                        "description": "Control family details",
                        "key_controls": [],
                        "threats_addressed": [],
                        "aws_services": []
                    })
                    
                    # Risk level indicator
                    if avg_risk > 16:
                        risk_emoji = "🔴"
                        risk_level = "Critical"
                    elif avg_risk > 9:
                        risk_emoji = "🟠"
                        risk_level = "High"
                    elif avg_risk > 4:
                        risk_emoji = "🟡"
                        risk_level = "Medium"
                    else:
                        risk_emoji = "🟢"
                        risk_level = "Low"
                    
                    with st.expander(f"{risk_emoji} **{family_code} - {family_info['name']}** | Avg Risk: {avg_risk:.1f} | {finding_count} findings | {risk_level}"):
                        st.markdown(f"*{family_info['description']}*")
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            st.markdown("**Key Controls:**")
                            for control in family_info.get('key_controls', []):
                                st.markdown(f"- `{control}`")
                            
                            st.markdown("**Threats Addressed:**")
                            for threat in family_info.get('threats_addressed', []):
                                st.markdown(f"- 🎯 {threat}")
                        
                        with col2:
                            st.markdown("**AWS Services:**")
                            for service in family_info.get('aws_services', []):
                                st.markdown(f"- ☁️ {service}")
                            
                            st.markdown("**Findings in This Family:**")
                            for finding in data['findings'][:5]:
                                f_emoji = "🔴" if finding.risk_score > 16 else "🟠" if finding.risk_score > 9 else "🟡"
                                st.markdown(f"- {f_emoji} {finding.control_id}: {finding.title[:40]}...")
                
                # Add Threat Modeling Section - integrate NIST assessment data
                st.markdown("---")
                # Pass NIST assessment results and Risk Calculator findings
                nist_results = st.session_state.get('results', None)
                render_threat_modeling_section(assessment_data=nist_results, findings=assessment.findings)
            else:
                st.warning("Install plotly for interactive charts: `pip install plotly`")
        
        with tab2:
            st.markdown(f"### 📋 Detailed Findings ({len(assessment.findings)} total)")
            sorted_findings = sorted(assessment.findings, key=lambda f: f.risk_score, reverse=True)
            
            for finding in sorted_findings:
                risk_emoji = "🔴" if finding.risk_level == RiskLevel.CRITICAL else \
                            "🟠" if finding.risk_level == RiskLevel.HIGH else \
                            "🟡" if finding.risk_level == RiskLevel.MEDIUM else "🟢"
                
                with st.expander(f"{risk_emoji} **{finding.control_id}**: {finding.title} (Score: {finding.risk_score:.0f})"):
                    c1, c2, c3 = st.columns(3)
                    with c1:
                        st.metric("Risk Score", f"{finding.risk_score:.0f}/25")
                    with c2:
                        st.metric("Likelihood", finding.likelihood.name)
                    with c3:
                        st.metric("Impact", finding.impact.name)
                    st.markdown("**Description:**")
                    st.info(finding.description)
                    if finding.remediation:
                        st.markdown("**Remediation:**")
                        st.success(finding.remediation)
        
        with tab3:
            st.markdown("### 📐 Risk Matrix")
            if PLOTLY_AVAILABLE:
                st.plotly_chart(render_risk_matrix_heatmap(), use_container_width=True)
            else:
                st.markdown("""
                <div style="font-family:monospace;background:#1e293b;color:#e2e8f0;padding:1rem;border-radius:8px;">
<pre>                    I M P A C T
          NEG    MIN    MOD    SIG    SEV
       +------+------+------+------+------+
V.HIGH |  5   |  10  |  15  |  20  |  25  | ← CRITICAL
       +------+------+------+------+------+
 HIGH  |  4   |  8   |  12  |  16  |  20  | ← HIGH
       +------+------+------+------+------+
 MED   |  3   |  6   |  9   |  12  |  15  | ← MEDIUM
       +------+------+------+------+------+
 LOW   |  2   |  4   |  6   |  8   |  10  | ← LOW
       +------+------+------+------+------+
V.LOW  |  1   |  2   |  3   |  4   |  5   | ← LOW
       +------+------+------+------+------+</pre></div>""", unsafe_allow_html=True)
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.markdown("🟢 **LOW** (1-4)")
            with col2:
                st.markdown("🟡 **MEDIUM** (5-9)")
            with col3:
                st.markdown("🟠 **HIGH** (10-16)")
            with col4:
                st.markdown("🔴 **CRITICAL** (17-25)")
        
        with tab4:
            st.markdown("### 📄 Risk Assessment Report")
            calculator = RiskScoreCalculator()
            report = calculator.generate_risk_report(assessment)
            st.code(report, language="text")
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.download_button("📥 Download Report (TXT)", data=report,
                    file_name=f"risk_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                    mime="text/plain", use_container_width=True)
            with col2:
                export_data = {
                    'assessment_id': assessment.assessment_id, 'date': assessment.date.isoformat(),
                    'summary': {'overall_risk': assessment.overall_risk_level.value,
                               'avg_score': assessment.average_risk_score, 'compliance': assessment.compliance_score},
                    'findings': [{'id': f.finding_id, 'title': f.title, 'risk_score': f.risk_score,
                                 'risk_level': f.risk_level.value} for f in assessment.findings]
                }
                st.download_button("📥 Download Report (JSON)", data=json.dumps(export_data, indent=2),
                    file_name=f"risk_assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json", use_container_width=True)
            with col3:
                if st.button("📝 Generate RAR Document", use_container_width=True, key="generate_rar_btn"):
                    try:
                        from wordy import create_rar_document
                        import os
                        
                        # Build RAR data from assessment
                        system_info = {
                            'system_name': st.session_state.get('system_info', {}).get('system_name', 'Information System'),
                            'system_acronym': st.session_state.get('system_info', {}).get('system_acronym', 'NOAA50xx'),
                            'organization': st.session_state.get('system_info', {}).get('organization', 'Organization'),
                            'categorization': st.session_state.get('system_info', {}).get('categorization', 'Moderate'),
                            'system_owner': st.session_state.get('system_info', {}).get('system_owner', 'TBD'),
                            'isso_name': st.session_state.get('system_info', {}).get('isso', 'TBD'),
                            'system_description': st.session_state.get('system_info', {}).get('system_description', 'System description to be provided.'),
                            'preparer': 'SAELAR-53 Risk Assessment Generator',
                        }
                        
                        # Convert findings to RAR format
                        controls_data = []
                        poam_items = []
                        recommendations = []
                        
                        for finding in assessment.findings:
                            risk_level_map = {
                                RiskLevel.CRITICAL: 'High',
                                RiskLevel.HIGH: 'High',
                                RiskLevel.MEDIUM: 'Medium',
                                RiskLevel.LOW: 'Low'
                            }
                            control_data = {
                                'control_id': finding.control_id,
                                'control_name': finding.title,
                                'family': finding.control_family,
                                'finding': finding.description,
                                'risk_level': risk_level_map.get(finding.risk_level, 'Medium'),
                                'risk_score': finding.risk_score,
                            }
                            controls_data.append(control_data)
                            
                            # Create POA&M item for non-low findings
                            if finding.risk_level != RiskLevel.LOW:
                                poam_items.append({
                                    'poam_id': f"POAM-{finding.finding_id}",
                                    'control_id': finding.control_id,
                                    'weakness': finding.description,
                                    'remediation_plan': finding.remediation or 'Remediation to be determined',
                                    'risk_level': risk_level_map.get(finding.risk_level, 'Medium'),
                                    'status': 'Open',
                                })
                            
                            if finding.remediation:
                                recommendations.append({
                                    'control_id': finding.control_id,
                                    'recommendation': finding.remediation,
                                })
                        
                        # Determine overall risk level
                        overall_risk = 'Moderate'
                        if assessment.average_risk_score >= 16:
                            overall_risk = 'High'
                        elif assessment.average_risk_score >= 9:
                            overall_risk = 'Moderate'
                        else:
                            overall_risk = 'Low'
                        
                        rar_data = {
                            'system_info': system_info,
                            'metadata': {'version': '1.0'},
                            'assessment_results': {
                                'controls': controls_data,
                                'family_summary': {},
                            },
                            'vulnerabilities': [],
                            'risk_summary': {
                                'overall_risk_level': overall_risk,
                                'total_risk_score': assessment.average_risk_score,
                            },
                            'recommendations': recommendations,
                            'poam': poam_items,
                            'statistics': {
                                'total_controls': len(assessment.findings),
                                'implemented': 0,
                                'partial': 0,
                                'not_implemented': len(assessment.findings),
                            },
                        }
                        
                        # Generate document
                        doc_path = create_rar_document(rar_data)
                        
                        # Store in session state for download/open buttons
                        st.session_state.rar_generated_path = doc_path
                        st.session_state.rar_generated = True
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ Error generating RAR: {str(e)}")
            
            # Show download and open buttons if RAR was generated
            if st.session_state.get('rar_generated', False) and st.session_state.get('rar_generated_path'):
                doc_path = st.session_state.rar_generated_path
                st.success(f"✅ RAR Document generated!")
                
                col_dl1, col_dl2 = st.columns(2)
                with col_dl1:
                    # Read the file for download
                    try:
                        with open(doc_path, 'rb') as f:
                            doc_bytes = f.read()
                        st.download_button(
                            "📥 Download RAR Document",
                            data=doc_bytes,
                            file_name=os.path.basename(doc_path),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True,
                            key="download_rar_btn"
                        )
                    except Exception as e:
                        st.error(f"Could not read file: {e}")
                
                with col_dl2:
                    if st.button("📂 Open RAR Document", use_container_width=True, key="open_rar_btn"):
                        try:
                            import subprocess
                            import sys
                            if sys.platform == 'win32':
                                os.startfile(doc_path)
                            elif sys.platform == 'darwin':
                                subprocess.Popen(['open', doc_path])
                            else:
                                subprocess.Popen(['xdg-open', doc_path])
                            st.info(f"Opening: {os.path.basename(doc_path)}")
                        except Exception as e:
                            st.error(f"Could not open file: {e}")
                
                st.caption(f"📁 Saved to: {doc_path}")
        
        # NIST 800-30 Enhanced Analysis Tab
        with tab5_risk:
            st.markdown("""
            <div style="background: linear-gradient(135deg, #7c3aed 0%, #4c1d95 100%); padding: 1.5rem;
                        border-radius: 12px; margin-bottom: 1.5rem; border-left: 5px solid #fbbf24;">
                <h2 style="color: #ffffff; margin: 0;">🔬 NIST SP 800-30 Rev 1 Enhanced Analysis</h2>
                <p style="color: #ddd6fe; margin: 0.5rem 0 0 0;">
                    Threat Modeling • ALE Calculator • Multi-Dimensional Impact • MITRE ATT&CK Coverage
                </p>
            </div>
            """, unsafe_allow_html=True)
            
            # Import enhanced classes
            try:
                from risk_score_calculator import (
                    THREAT_SOURCE_CATALOG, MITRE_ATTACK_MAPPING,
                    ThreatSource, ThreatSourceType, ThreatCapability, ThreatIntent,
                    MultiDimensionalImpact, QuantitativeRiskMetrics, EnhancedLikelihood
                )
                
                # Create sub-tabs for NIST 800-30 sections
                nist_tab1, nist_tab2, nist_tab3, nist_tab4, nist_tab5 = st.tabs([
                    "🎯 Threat Sources", "💰 ALE Analysis", "📊 Impact (C/I/A)", "🗡️ MITRE ATT&CK", "📈 Enhanced Metrics"
                ])
                
                # Tab 1: Threat Source Catalog
                with nist_tab1:
                    st.markdown("### 🎯 Threat Source Identification (NIST 800-30 Section 2.1.1)")
                    st.markdown("*Characterizing threat sources by capability, intent, and targeting likelihood*")
                    
                    # Threat source metrics
                    col1, col2, col3, col4 = st.columns(4)
                    adversarial = len([s for s in THREAT_SOURCE_CATALOG.values() if s.source_type == ThreatSourceType.ADVERSARIAL])
                    accidental = len([s for s in THREAT_SOURCE_CATALOG.values() if s.source_type == ThreatSourceType.ACCIDENTAL])
                    structural = len([s for s in THREAT_SOURCE_CATALOG.values() if s.source_type == ThreatSourceType.STRUCTURAL])
                    environmental = len([s for s in THREAT_SOURCE_CATALOG.values() if s.source_type == ThreatSourceType.ENVIRONMENTAL])
                    
                    with col1:
                        st.metric("⚔️ Adversarial", adversarial)
                    with col2:
                        st.metric("🙈 Accidental", accidental)
                    with col3:
                        st.metric("⚙️ Structural", structural)
                    with col4:
                        st.metric("🌪️ Environmental", environmental)
                    
                    st.markdown("---")
                    
                    # Display threat sources
                    for source_id, source in THREAT_SOURCE_CATALOG.items():
                        type_colors = {
                            ThreatSourceType.ADVERSARIAL: "#dc2626",
                            ThreatSourceType.ACCIDENTAL: "#f59e0b",
                            ThreatSourceType.STRUCTURAL: "#3b82f6",
                            ThreatSourceType.ENVIRONMENTAL: "#10b981"
                        }
                        color = type_colors.get(source.source_type, "#6b7280")
                        
                        threat_score = source.capability.value * source.intent.value
                        threat_bar = "█" * (threat_score // 3)
                        
                        with st.expander(f"**{source.name}** | Threat Score: {threat_score}/25 | Type: {source.source_type.value}"):
                            col1, col2 = st.columns(2)
                            with col1:
                                st.markdown(f"**Description:** {source.description}")
                                st.markdown(f"**Capability:** {source.capability.name} ({source.capability.value}/5)")
                                st.markdown(f"**Intent:** {source.intent.name} ({source.intent.value}/5)")
                            with col2:
                                st.markdown(f"**Targeting Likelihood:** {source.targeting_likelihood:.0%}")
                                st.markdown(f"**Effective Threat:** {source.effective_threat:.1f}")
                                st.progress(threat_score / 25, text=f"Threat Level: {threat_bar}")
                
                # Tab 2: ALE Analysis
                with nist_tab2:
                    st.markdown("### 💰 Annualized Loss Expectancy (ALE) Analysis")
                    st.markdown("*Quantitative risk analysis using SLE × ARO = ALE methodology*")
                    
                    # Calculate total ALE from findings if using enhanced calculator
                    calculator = RiskScoreCalculator()
                    
                    # Estimate ALE for current findings
                    total_ale = 0
                    ale_by_family = {}
                    
                    for finding in assessment.findings:
                        # Estimate quantitative metrics
                        metrics = calculator.estimate_quantitative_metrics(finding)
                        finding_ale = metrics.annualized_loss_expectancy
                        total_ale += finding_ale
                        
                        family = finding.control_family
                        if family not in ale_by_family:
                            ale_by_family[family] = 0
                        ale_by_family[family] += finding_ale
                    
                    # Summary metrics
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.markdown(f"""
                        <div style="background:#fff;padding:1.5rem;border-radius:10px;text-align:center;border:2px solid #dc2626;">
                            <div style="font-size:0.9rem;color:#6b7280;">Total ALE</div>
                            <div style="font-size:2rem;font-weight:700;color:#dc2626;">${total_ale:,.0f}</div>
                            <div style="color:#9ca3af;font-size:0.8rem;">per year</div>
                        </div>
                        """, unsafe_allow_html=True)
                    with col2:
                        risk_reduction = total_ale * 0.70
                        st.markdown(f"""
                        <div style="background:#fff;padding:1.5rem;border-radius:10px;text-align:center;border:2px solid #22c55e;">
                            <div style="font-size:0.9rem;color:#6b7280;">Potential Savings</div>
                            <div style="font-size:2rem;font-weight:700;color:#22c55e;">${risk_reduction:,.0f}</div>
                            <div style="color:#9ca3af;font-size:0.8rem;">with remediation (70% reduction)</div>
                        </div>
                        """, unsafe_allow_html=True)
                    with col3:
                        roi_breakeven = total_ale * 0.70
                        st.markdown(f"""
                        <div style="background:#fff;padding:1.5rem;border-radius:10px;text-align:center;border:2px solid #3b82f6;">
                            <div style="font-size:0.9rem;color:#6b7280;">ROI Breakeven</div>
                            <div style="font-size:2rem;font-weight:700;color:#3b82f6;">${roi_breakeven:,.0f}</div>
                            <div style="color:#9ca3af;font-size:0.8rem;">max security investment</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    st.markdown("#### ALE by Control Family")
                    
                    # Sort by ALE
                    sorted_families = sorted(ale_by_family.items(), key=lambda x: x[1], reverse=True)
                    
                    for family, ale in sorted_families:
                        pct_of_total = (ale / total_ale * 100) if total_ale > 0 else 0
                        st.markdown(f"""
                        <div style="display:flex;align-items:center;margin-bottom:0.5rem;">
                            <div style="width:60px;font-weight:600;">{family}</div>
                            <div style="flex:1;background:#e5e7eb;border-radius:4px;height:24px;margin:0 1rem;">
                                <div style="width:{min(pct_of_total, 100)}%;background:linear-gradient(90deg, #dc2626, #f59e0b);
                                            height:100%;border-radius:4px;"></div>
                            </div>
                            <div style="width:150px;text-align:right;font-weight:600;">${ale:,.0f}</div>
                            <div style="width:60px;text-align:right;color:#6b7280;">{pct_of_total:.1f}%</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    st.info("""
                    **ALE Formula:** `ALE = SLE × ARO`
                    - **SLE (Single Loss Expectancy)** = Asset Value × Exposure Factor
                    - **ARO (Annual Rate of Occurrence)** = Expected incidents per year
                    """)
                
                # Tab 3: Multi-Dimensional Impact
                with nist_tab3:
                    st.markdown("### 📊 Multi-Dimensional Impact Analysis (NIST 800-30 Section 2.3)")
                    st.markdown("*Assessing impact across Confidentiality, Integrity, and Availability*")
                    
                    # Calculate aggregate C/I/A impacts
                    c_total, i_total, a_total, count = 0, 0, 0, 0
                    
                    # Map findings to C/I/A based on control family
                    cia_mapping = {
                        'AC': ('C', 5), 'IA': ('C', 5), 'MP': ('C', 4), 'SC': ('C', 4),
                        'AU': ('I', 4), 'SI': ('I', 5), 'CM': ('I', 4),
                        'CP': ('A', 5), 'IR': ('A', 4),
                    }
                    
                    for finding in assessment.findings:
                        if finding.control_family in ['AC', 'IA', 'MP', 'SC']:
                            c_total += finding.impact.value
                        elif finding.control_family in ['AU', 'SI', 'CM']:
                            i_total += finding.impact.value
                        elif finding.control_family in ['CP', 'IR']:
                            a_total += finding.impact.value
                        else:
                            c_total += finding.impact.value * 0.33
                            i_total += finding.impact.value * 0.33
                            a_total += finding.impact.value * 0.33
                        count += 1
                    
                    # Normalize to 0-5 scale
                    c_avg = (c_total / count) if count > 0 else 0
                    i_avg = (i_total / count) if count > 0 else 0
                    a_avg = (a_total / count) if count > 0 else 0
                    
                    # Display CIA gauges
                    col1, col2, col3 = st.columns(3)
                    
                    def get_impact_color(val):
                        if val >= 4:
                            return "#dc2626"
                        elif val >= 3:
                            return "#f59e0b"
                        elif val >= 2:
                            return "#eab308"
                        return "#22c55e"
                    
                    with col1:
                        c_color = get_impact_color(c_avg)
                        st.markdown(f"""
                        <div style="background:#fff;padding:1.5rem;border-radius:10px;text-align:center;border:2px solid {c_color};">
                            <div style="font-size:3rem;">🔐</div>
                            <div style="font-size:0.9rem;color:#6b7280;">CONFIDENTIALITY</div>
                            <div style="font-size:2.5rem;font-weight:700;color:{c_color};">{c_avg:.1f}/5</div>
                            <div style="color:#9ca3af;">Data disclosure risk</div>
                        </div>
                        """, unsafe_allow_html=True)
                    with col2:
                        i_color = get_impact_color(i_avg)
                        st.markdown(f"""
                        <div style="background:#fff;padding:1.5rem;border-radius:10px;text-align:center;border:2px solid {i_color};">
                            <div style="font-size:3rem;">✅</div>
                            <div style="font-size:0.9rem;color:#6b7280;">INTEGRITY</div>
                            <div style="font-size:2.5rem;font-weight:700;color:{i_color};">{i_avg:.1f}/5</div>
                            <div style="color:#9ca3af;">Data modification risk</div>
                        </div>
                        """, unsafe_allow_html=True)
                    with col3:
                        a_color = get_impact_color(a_avg)
                        st.markdown(f"""
                        <div style="background:#fff;padding:1.5rem;border-radius:10px;text-align:center;border:2px solid {a_color};">
                            <div style="font-size:3rem;">⚡</div>
                            <div style="font-size:0.9rem;color:#6b7280;">AVAILABILITY</div>
                            <div style="font-size:2.5rem;font-weight:700;color:{a_color};">{a_avg:.1f}/5</div>
                            <div style="color:#9ca3af;">Service disruption risk</div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    st.markdown("#### Impact Levels Reference")
                    st.markdown("""
                    | Level | Value | Description |
                    |-------|-------|-------------|
                    | SEVERE | 5 | Catastrophic impact to mission/organization |
                    | SIGNIFICANT | 4 | Major impact requiring significant recovery |
                    | MODERATE | 3 | Noticeable impact with manageable recovery |
                    | MINOR | 2 | Limited impact with minimal disruption |
                    | NEGLIGIBLE | 1 | Minimal or no noticeable impact |
                    """)
                
                # Tab 4: MITRE ATT&CK Coverage
                with nist_tab4:
                    st.markdown("### 🗡️ MITRE ATT&CK Threat Coverage")
                    st.markdown("*Mapping control gaps to adversary tactics and techniques*")
                    
                    # Collect techniques from findings
                    tactics_found = {}
                    for finding in assessment.findings:
                        family = finding.control_family
                        techniques = MITRE_ATTACK_MAPPING.get(family, [])
                        for tech in techniques:
                            if tech.tactic not in tactics_found:
                                tactics_found[tech.tactic] = []
                            if tech.technique_id not in [t['id'] for t in tactics_found[tech.tactic]]:
                                tactics_found[tech.tactic].append({
                                    'id': tech.technique_id,
                                    'name': tech.technique_name,
                                    'family': family
                                })
                    
                    # Display by tactic
                    if tactics_found:
                        # Summary
                        total_techniques = sum(len(t) for t in tactics_found.values())
                        st.metric("Total Techniques Identified", total_techniques)
                        
                        st.markdown("---")
                        
                        for tactic, techniques in sorted(tactics_found.items()):
                            with st.expander(f"**{tactic}** — {len(techniques)} technique(s)"):
                                for tech in techniques:
                                    st.markdown(f"- **{tech['id']}**: {tech['name']} *(from {tech['family']})*")
                    else:
                        st.info("Run an assessment to see MITRE ATT&CK mapping for your findings.")
                
                # Tab 5: Enhanced Metrics
                with nist_tab5:
                    st.markdown("### 📈 Enhanced Risk Metrics")
                    st.markdown("*Advanced risk calculations with confidence levels*")
                    
                    # Show enhanced metrics for top findings
                    sorted_findings = sorted(assessment.findings, key=lambda f: f.risk_score, reverse=True)
                    
                    st.markdown("#### Top Risk Findings with Enhanced Analysis")
                    
                    for i, finding in enumerate(sorted_findings[:5], 1):
                        # Calculate enhanced likelihood
                        enhanced_likelihood = EnhancedLikelihood(
                            threat_capability=ThreatCapability(min(finding.likelihood.value + 1, 5)),
                            threat_intent=ThreatIntent(finding.likelihood.value),
                            vulnerability_severity=finding.impact.value * 2,
                            predisposing_conditions=0.5
                        )
                        
                        calc_likelihood = enhanced_likelihood.calculated_likelihood
                        confidence = enhanced_likelihood.confidence_level
                        
                        # Estimate ALE
                        metrics = calculator.estimate_quantitative_metrics(finding)
                        
                        risk_color = "#dc2626" if finding.risk_level.value == "CRITICAL" else \
                                    "#f59e0b" if finding.risk_level.value == "HIGH" else \
                                    "#eab308" if finding.risk_level.value == "MEDIUM" else "#22c55e"
                        
                        st.markdown(f"""
                        <div style="background:#fff;padding:1rem;border-radius:8px;margin-bottom:1rem;
                                    border-left:4px solid {risk_color};">
                            <div style="font-weight:600;font-size:1.1rem;">{i}. {finding.title}</div>
                            <div style="display:flex;gap:2rem;margin-top:0.5rem;">
                                <div>
                                    <span style="color:#6b7280;">Base Score:</span> 
                                    <span style="font-weight:600;">{finding.risk_score:.0f}/25</span>
                                </div>
                                <div>
                                    <span style="color:#6b7280;">Enhanced Likelihood:</span> 
                                    <span style="font-weight:600;">{calc_likelihood.name}</span>
                                </div>
                                <div>
                                    <span style="color:#6b7280;">Confidence:</span> 
                                    <span style="font-weight:600;">{confidence}</span>
                                </div>
                                <div>
                                    <span style="color:#6b7280;">ALE:</span> 
                                    <span style="font-weight:600;color:#dc2626;">${metrics.annualized_loss_expectancy:,.0f}</span>
                                </div>
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    st.markdown("#### Methodology Notes")
                    st.info("""
                    **NIST SP 800-30 Rev 1 Enhanced Features:**
                    - **Threat Source Identification** - Catalog of adversarial and non-adversarial threats
                    - **Multi-Factor Likelihood** - Considers capability, intent, vulnerability severity, and conditions
                    - **Multi-Dimensional Impact** - C/I/A plus financial, mission, and reputation impact
                    - **Quantitative Analysis** - ALE calculations for ROI justification
                    - **Confidence Levels** - Uncertainty tracking for risk estimates
                    - **MITRE ATT&CK Mapping** - Threat event identification via technique mapping
                    """)
                    
            except ImportError as e:
                st.error(f"Could not load enhanced NIST 800-30 modules: {e}")
                st.info("Make sure risk_score_calculator.py has been updated with NIST 800-30 enhancements.")
    
    def render_admin_page(self):
        """Render the admin/user management page."""
        render_user_management()
    
    def render_main_content(self):
        """Render the main content area with tabs."""
        # Render sidebar FIRST before any tabs
        region, selected_families, run_assessment, include_security_hub = render_sidebar()
        
        # Show AI mode indicator in sidebar
        render_ai_mode_indicator()
        
        # Store sidebar values in session state for use in assessment
        st.session_state.current_region = region
        st.session_state.run_assessment_clicked = run_assessment
        st.session_state.include_security_hub = include_security_hub
        
        # Check if user clicked to view Security Hub findings only
        if st.session_state.get('show_security_hub_only', False):
            render_security_hub_findings_only()
            return
        
        # Determine which tabs to show based on permissions
        username = get_current_user()
        
        # Custom CSS for larger, more pronounced tab styling
        st.markdown("""
        <style>
            /* Style the tab buttons to be larger and more pronounced */
            .stTabs [data-baseweb="tab-list"] {
                gap: 8px;
            }
            .stTabs [data-baseweb="tab"] {
                font-size: 1.1rem !important;
                font-weight: 600 !important;
                padding: 0.75rem 1.25rem !important;
                white-space: nowrap !important;
            }
            .stTabs [data-baseweb="tab-list"] button {
                font-size: 1.1rem !important;
                font-weight: 600 !important;
            }
            .stTabs [data-baseweb="tab-list"] button p {
                font-size: 1.1rem !important;
                font-weight: 600 !important;
            }
        </style>
        """, unsafe_allow_html=True)
        
        # Define tabs - NIST Assessment, AWS Console, Chad (AI) (AI Analyst), Risk Calculator, SSP Generator, BOD 22-01, Threat Modeling, Artifacts
        tabs = ["🛡️ NIST Assessment", "☁️ AWS Console", "🤖 Chad (AI)", "📊 Risk Calculator", "📋 SSP Generator", "🚨 BOD 22-01", "🎯 Threat Modeling", "📁 Artifacts"]
        
        # Create tabs
        tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8 = st.tabs(tabs)
        with tab1:
            self.render_assessment_page(region, selected_families, run_assessment, include_security_hub)
        with tab2:
            self.render_aws_console_page()
        with tab3:
            self.render_chad_security_analyst_page()
        with tab4:
            self.render_risk_calculator_page()
        with tab5:
            self.render_ssp_generator_page()
        with tab6:
            self.render_bod_2201_page()
        with tab7:
            self.render_threat_modeling_page()
        with tab8:
            self.render_documentation_page()

    def render_ssp_generator_page(self):
        """Render the System Security Plan Generator page."""
        from datetime import datetime
        from ssp_generator import SSPGenerator, SystemCategorization
        from wordy import create_ssp_document, create_poam_document
        
        st.markdown("""
        <div style="background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 50%, #3d7ab5 100%); padding: 1.5rem;
                    border-radius: 12px; margin-bottom: 1.5rem; border-left: 5px solid #10b981;">
            <h2 style="color: #ffffff; margin: 0;">📋 System Security Plan (SSP) Generator</h2>
            <p style="color: #b8d4e8; margin: 0.5rem 0 0 0;">
                Generate NIST 800-53 compliant SSP documents from assessment results
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Check if assessment results exist
        has_assessment = st.session_state.get('results') is not None
        has_risk_data = st.session_state.get('risk_assessment') is not None
        has_system_info = bool(st.session_state.get('system_info_name'))
        
        # Status indicators
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            if has_system_info:
                st.success("✅ System Info: Complete")
            else:
                st.warning("⚠️ System Info: Not entered")
        with col2:
            if has_assessment:
                st.success("✅ NIST Assessment: Available")
            else:
                st.warning("⚠️ NIST Assessment: Not yet run")
        with col3:
            if has_risk_data:
                st.success("✅ Risk Calculator: Available")
            else:
                st.info("ℹ️ Risk Calculator: Optional")
        with col4:
            st.info("📄 Output: Word Document (.docx)")
        
        st.markdown("---")
        
        # Auto-populate from sidebar System section
        system_name = st.session_state.get('system_info_name', '')
        system_acronym = st.session_state.get('system_info_acronym', '')
        system_owner = st.session_state.get('system_info_owner', '')
        system_owner_email = st.session_state.get('system_info_owner_email', '')
        authorizing_official = st.session_state.get('system_info_ao', '')
        authorizing_official_email = st.session_state.get('system_info_ao_email', '')
        isso_name = st.session_state.get('system_info_isso', '')
        isso_email = st.session_state.get('system_info_isso_email', '')
        categorization = st.session_state.get('system_info_categorization', 'Moderate')
        operational_status = st.session_state.get('system_info_status', 'Operational')
        system_description = st.session_state.get('system_info_description', '')
        authorization_boundary = st.session_state.get('system_info_boundary', 'AWS Commercial Cloud')
        
        # Show system info summary (read-only, pulled from sidebar)
        if has_system_info:
            st.markdown("### 📋 System Information (from sidebar)")
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown(f"""
                <div style="background: #f0fdf4; border: 1px solid #10b981; border-radius: 8px; padding: 1rem;">
                    <p><strong>System Name:</strong> {system_name}</p>
                    <p><strong>Acronym:</strong> {system_acronym or 'N/A'}</p>
                    <p><strong>Owner:</strong> {system_owner}</p>
                    <p><strong>Categorization:</strong> {categorization}</p>
                </div>
                """, unsafe_allow_html=True)
            with col2:
                st.markdown(f"""
                <div style="background: #f0fdf4; border: 1px solid #10b981; border-radius: 8px; padding: 1rem;">
                    <p><strong>ISSO:</strong> {isso_name or 'N/A'}</p>
                    <p><strong>Authorizing Official:</strong> {authorizing_official or 'N/A'}</p>
                    <p><strong>Status:</strong> {operational_status}</p>
                    <p><strong>Boundary:</strong> {authorization_boundary[:50]}...</p>
                </div>
                """, unsafe_allow_html=True)
            
            st.info("💡 To modify system information, use the **System** section in the sidebar.")
        else:
            st.warning("⚠️ **No system information entered.** Please use the **System** section in the sidebar to enter system details before generating an SSP.")
        
        st.markdown("---")
        
        # Create sub-tabs for Generate SSP, POA&Ms, and Risk Acceptance
        ssp_tab1, ssp_tab2, ssp_tab3 = st.tabs(["📄 Generate SSP", "📋 POA&Ms", "⚖️ Risk Acceptance"])
        
        with ssp_tab1:
            # SSP Generation Options
            st.markdown("### ⚙️ Generation Options")
        
        col1, col2 = st.columns(2)
        with col1:
            include_poam = st.checkbox("Include POA&M Section", value=True)
            include_risk = st.checkbox("Include Risk Assessment Data", value=True)
        with col2:
            include_evidence = st.checkbox("Include Evidence Details", value=True)
            include_recommendations = st.checkbox("Include Remediation Recommendations", value=True)
        
        st.markdown("---")
        
        # Generate Button
        if st.button("📋 Generate System Security Plan", type="primary", use_container_width=True):
            if not system_name:
                st.error("❌ System Name is required")
            elif not system_owner:
                st.error("❌ System Owner is required")
            elif not system_description:
                st.error("❌ System Description is required")
            elif not has_assessment:
                st.error("❌ Please run a NIST Assessment first (NIST Assessment tab)")
            else:
                with st.spinner("🔄 Generating System Security Plan..."):
                    try:
                        # Create SSP Generator
                        ssp_gen = SSPGenerator(
                            system_name=system_name,
                            system_owner=system_owner,
                            categorization=categorization,
                            system_acronym=system_acronym,
                            system_owner_email=system_owner_email,
                            authorizing_official=authorizing_official,
                            authorizing_official_email=authorizing_official_email,
                            isso_name=isso_name,
                            isso_email=isso_email,
                            system_description=system_description,
                            authorization_boundary=authorization_boundary,
                            operational_status=operational_status
                        )
                        
                        # Load assessment results
                        ssp_gen.load_assessment_results(st.session_state.results)
                        
                        # Load risk data if available
                        if has_risk_data and include_risk:
                            ssp_gen.load_risk_findings(st.session_state.risk_assessment)
                        
                        # Generate SSP data
                        ssp_data = ssp_gen.to_dict()
                        
                        # Store in session state
                        st.session_state.ssp_data = ssp_data
                        st.session_state.ssp_system_name = system_name
                        st.session_state.ssp_acronym = system_acronym
                        st.session_state.ssp_owner = system_owner
                        st.session_state.ssp_isso = isso_name
                        st.session_state.ssp_description = system_description
                        st.session_state.ssp_boundary = authorization_boundary
                        
                        # Generate Word document
                        safe_name = system_name.replace(' ', '_').replace('/', '-')[:30]
                        doc_path = create_ssp_document(ssp_data)
                        
                        # Store path for download button
                        st.session_state.ssp_doc_path = doc_path
                        
                        # Read the file for download
                        with open(doc_path, 'rb') as f:
                            doc_bytes = f.read()
                        
                        # Upload to S3 Documentation/SSPs/ folder
                        s3_uploaded = False
                        s3_location = ""
                        poam_s3_uploaded = False
                        poam_s3_location = ""
                        try:
                            s3_bucket = st.session_state.get('s3_bucket') or os.environ.get('S3_BUCKET_NAME') or S3_BUCKET_NAME
                            if s3_bucket:
                                aws_access_key = st.session_state.get('aws_access_key') or os.environ.get('AWS_ACCESS_KEY_ID')
                                aws_secret_key = st.session_state.get('aws_secret_key') or os.environ.get('AWS_SECRET_ACCESS_KEY')
                                aws_region = st.session_state.get('aws_region') or os.environ.get('AWS_DEFAULT_REGION', 'us-east-1')
                                
                                if aws_access_key and aws_secret_key:
                                    s3_client = boto3.client('s3',
                                        aws_access_key_id=aws_access_key,
                                        aws_secret_access_key=aws_secret_key,
                                        region_name=aws_region
                                    )
                                else:
                                    s3_client = boto3.client('s3', region_name=aws_region)
                                
                                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                                
                                # Upload SSP to Documentation/SSPs/ folder
                                s3_key = f"Documentation/SSPs/SSP_{safe_name}_{timestamp}.docx"
                                
                                with open(doc_path, 'rb') as ssp_file:
                                    s3_client.upload_fileobj(
                                        ssp_file,
                                        s3_bucket,
                                        s3_key,
                                        ExtraArgs={
                                            'Metadata': {
                                                'document_type': 'System Security Plan (SSP)',
                                                'system_name': system_name,
                                                'generated_by': 'SAELAR-53 SSP Generator'
                                            }
                                        }
                                    )
                                s3_uploaded = True
                                s3_location = f"s3://{s3_bucket}/{s3_key}"
                                
                                # Create and upload POA&M to Documentation/POA&Ms/ folder
                                if include_poam and ssp_data.get('poam'):
                                    poam_doc_path = create_poam_document(ssp_data)
                                    
                                    poam_s3_key = f"Documentation/POA&Ms/POAM_{safe_name}_{timestamp}.docx"
                                    
                                    with open(poam_doc_path, 'rb') as poam_file:
                                        s3_client.upload_fileobj(
                                            poam_file,
                                            s3_bucket,
                                            poam_s3_key,
                                            ExtraArgs={
                                                'Metadata': {
                                                    'document_type': 'Plan of Action & Milestones (POA&M)',
                                                    'system_name': system_name,
                                                    'generated_by': 'SAELAR-53 POA&M Generator',
                                                    'poam_count': str(len(ssp_data.get('poam', [])))
                                                }
                                            }
                                        )
                                    poam_s3_uploaded = True
                                    poam_s3_location = f"s3://{s3_bucket}/{poam_s3_key}"
                        except Exception as s3_error:
                            st.warning(f"⚠️ Could not upload to S3: {str(s3_error)}")
                        
                        # Success message with download button
                        st.markdown("""
                        <div style="background: #10b981; color: white; padding: 1rem; border-radius: 8px; margin-bottom: 1rem;">
                            <h3 style="margin: 0; color: white;">✅ SSP Generated Successfully!</h3>
                            <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">Click the button below to download your System Security Plan</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                        st.download_button(
                            label="📥 Download SSP Document (.docx)",
                            data=doc_bytes,
                            file_name=f"SSP_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary",
                            use_container_width=True
                        )
                        
                        st.info(f"📄 Local copy saved to: `{doc_path}`")
                        
                        if s3_uploaded:
                            st.success(f"☁️ SSP uploaded to S3: `{s3_location}`")
                        
                        if poam_s3_uploaded:
                            st.success(f"📋 POA&M uploaded to S3: `{poam_s3_location}`")
                        
                        # Show summary
                        summary = ssp_data.get('summary', {})
                        
                        st.markdown("### 📊 SSP Summary")
                        
                        col1, col2, col3, col4 = st.columns(4)
                        with col1:
                            st.metric("Controls Assessed", summary.get('total_controls', 0))
                        with col2:
                            st.metric("Implemented", summary.get('implemented', 0))
                        with col3:
                            st.metric("Compliance %", f"{summary.get('compliance_percentage', 0)}%")
                        with col4:
                            st.metric("POA&M Items", summary.get('poam_items', 0))
                        
                        # Security Posture
                        posture = summary.get('security_posture', 'Unknown')
                        posture_colors = {
                            'Strong': '#10b981',
                            'Satisfactory': '#3b82f6',
                            'Needs Improvement': '#f59e0b',
                            'Unsatisfactory': '#ef4444'
                        }
                        color = posture_colors.get(posture, '#6b7280')
                        
                        st.markdown(f"""
                        <div style="background: {color}; color: white; padding: 1rem; border-radius: 8px; text-align: center; margin-top: 1rem;">
                            <h3 style="margin: 0; color: white;">Security Posture: {posture}</h3>
                            <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">{summary.get('posture_description', '')}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        
                    except Exception as e:
                        st.error(f"❌ Error generating SSP: {str(e)}")
                        import traceback
                        st.code(traceback.format_exc())
        
            # Show previous SSP data if available
            if st.session_state.get('ssp_data'):
                st.markdown("---")
                with st.expander("📄 View Last Generated SSP Data (JSON)"):
                    st.json(st.session_state.ssp_data)
        
        with ssp_tab2:
            # POA&Ms Tab - Display POA&M items from SSP data
            st.markdown("### 📋 Plan of Action & Milestones (POA&M)")
            st.markdown("*View and manage POA&M items that will be included in the SSP document.*")
            
            # Check if we have SSP data with POA&M items
            ssp_data = st.session_state.get('ssp_data')
            
            if ssp_data and ssp_data.get('poam'):
                poam_items = ssp_data.get('poam', [])
                
                # Summary metrics
                high_count = sum(1 for p in poam_items if p.get('risk_level') == 'High')
                medium_count = sum(1 for p in poam_items if p.get('risk_level') == 'Medium')
                low_count = sum(1 for p in poam_items if p.get('risk_level') == 'Low')
                
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Total POA&M Items", len(poam_items))
                with col2:
                    st.metric("🔴 High Risk", high_count)
                with col3:
                    st.metric("🟡 Medium Risk", medium_count)
                with col4:
                    st.metric("🟢 Low Risk", low_count)
                
                # Risk Acceptance tile — perform Acceptance of Risk process and create document (see Risk Acceptance tab)
                st.markdown("---")
                st.markdown(
                    '<div style="background: linear-gradient(145deg, rgba(255,193,7,0.12), rgba(255,193,7,0.04)); '
                    'border: 1px solid rgba(255,193,7,0.35); border-radius: 12px; padding: 1rem; margin-bottom: 0.5rem; display: inline-block;">'
                    '<div style="font-size: 1.8rem; margin-bottom: 0.25rem;">⚖️</div>'
                    '<div style="color: #f59e0b; font-weight: 700; font-size: 1rem;">Risk Acceptance</div>'
                    '<div style="color: #6b7280; font-size: 0.75rem;">Perform Acceptance of Risk process and create document → use the <strong>Risk Acceptance</strong> tab</div>'
                    '</div>',
                    unsafe_allow_html=True
                )
                st.markdown("---")
                
                # Filter by risk level
                risk_filter = st.selectbox(
                    "Filter by Risk Level",
                    ["All", "High", "Medium", "Low"],
                    key="poam_risk_filter"
                )
                
                # Apply filter
                if risk_filter != "All":
                    filtered_items = [p for p in poam_items if p.get('risk_level') == risk_filter]
                else:
                    filtered_items = poam_items
                
                st.markdown(f"**Showing {len(filtered_items)} of {len(poam_items)} items**")
                
                # Display POA&M items
                for idx, item in enumerate(filtered_items):
                    risk_level = item.get('risk_level', 'Unknown')
                    risk_colors = {'High': '#ef4444', 'Medium': '#f59e0b', 'Low': '#10b981'}
                    color = risk_colors.get(risk_level, '#6b7280')
                    
                    with st.expander(f"**{item.get('poam_id', f'POAM-{idx+1}')}** - {item.get('control_id', 'N/A')} ({risk_level} Risk)"):
                        st.markdown(f"""
                        <div style="border-left: 4px solid {color}; padding-left: 1rem;">
                            <p><strong>Control ID:</strong> {item.get('control_id', 'N/A')}</p>
                            <p><strong>Weakness:</strong> {item.get('weakness', 'N/A')}</p>
                            <p><strong>Remediation Plan:</strong> {item.get('remediation_plan', 'N/A')}</p>
                            <p><strong>Scheduled Completion:</strong> {item.get('scheduled_completion', 'TBD')}</p>
                            <p><strong>Status:</strong> {item.get('status', 'Open')}</p>
                        </div>
                        """, unsafe_allow_html=True)
                        if st.button("⚖️ Accept Risk", key=f"poam_accept_risk_{idx}_{item.get('control_id', '')}", help="Open Risk Acceptance tab and pre-select this finding"):
                            pending = f"{item.get('control_id', 'N/A')} — {item.get('weakness', 'N/A')[:50]}..."
                            st.session_state.ra_pending_finding = pending
                            st.success("Go to the **Risk Acceptance** tab to log this acceptance. This finding has been pre-selected.")
                            st.rerun()
                
                # Export options
                st.markdown("---")
                st.markdown("### 📤 Export POA&M")
                
                col1, col2 = st.columns(2)
                with col1:
                    # Download as JSON
                    import json
                    poam_json = json.dumps(poam_items, indent=2, default=str)
                    st.download_button(
                        "📥 Download POA&M (JSON)",
                        data=poam_json,
                        file_name=f"POAM_{st.session_state.get('system_info_name', 'System').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.json",
                        mime="application/json",
                        use_container_width=True
                    )
                with col2:
                    # Link to S3 POA&Ms folder
                    s3_bucket = st.session_state.get('s3_bucket') or os.environ.get('S3_BUCKET_NAME') or S3_BUCKET_NAME
                    aws_region = st.session_state.get('aws_region') or os.environ.get('AWS_DEFAULT_REGION', 'us-east-1')
                    poam_url = f"https://s3.console.aws.amazon.com/s3/buckets/{s3_bucket}?region={aws_region}&prefix=Documentation/POA%26Ms/"
                    st.markdown(f"""
                    <a href="{poam_url}" target="_blank" style="text-decoration: none;">
                        <div style="background: #f59e0b; color: white; padding: 0.6rem; border-radius: 6px; text-align: center; font-weight: 600;">
                            ☁️ View POA&Ms in S3
                        </div>
                    </a>
                    """, unsafe_allow_html=True)
            else:
                st.info("📭 **No POA&M items available yet.**")
                st.markdown("""
                To generate POA&M items:
                1. Run a **NIST Assessment** in the NIST Assessment tab
                2. Go to the **Generate SSP** tab
                3. Click **Generate System Security Plan**
                
                POA&M items are automatically created from failed or warning control findings.
                """)

        with ssp_tab3:
            # Risk Acceptance tab — Acceptance of Risk process and report all Risk Acceptance POA&Ms
            st.markdown("### ⚖️ Risk Acceptance")
            st.markdown("*Perform the Acceptance of Risk process and create an Acceptance of Risk document. All Risk Acceptance POA&Ms are reported here.*")
            
            if 'risk_acceptances' not in st.session_state:
                st.session_state.risk_acceptances = []
            ra_items = st.session_state.risk_acceptances
            
            # New Risk Acceptance form
            with st.expander("➕ Log New Risk Acceptance", expanded=len(ra_items) == 0):
                ssp_data = st.session_state.get('ssp_data')
                poam_options = []
                if ssp_data and ssp_data.get('poam'):
                    poam_options = [f"{p.get('control_id', 'N/A')} — {p.get('weakness', 'N/A')[:50]}..." for p in ssp_data.get('poam', [])]
                if not poam_options:
                    poam_options = ["No POA&M items — generate an SSP first"]
                ra_default_idx = 0
                if st.session_state.get('ra_pending_finding') and poam_options and not poam_options[0].startswith("No POA&M"):
                    try:
                        ra_default_idx = poam_options.index(st.session_state.ra_pending_finding)
                    except ValueError:
                        pass
                    st.session_state.pop('ra_pending_finding', None)
                ra_finding = st.selectbox("Control / Finding", poam_options, index=ra_default_idx, key="ra_finding")
                # Chad (AI) draft — autopopulate Operational Justification and Compensating Controls
                if poam_options and not poam_options[0].startswith("No POA&M"):
                    if st.button("🤖 Chad AI Draft", type="secondary", key="ra_chad_draft", help="Have Chad (AI) draft the justification and compensating controls for the selected finding"):
                        self._chad_draft_risk_acceptance(st.session_state.get("ra_finding", ""))
                ra_just_kw = {"placeholder": "Why is this risk being accepted? What is the operational requirement?", "key": "ra_just"}
                ra_comp_kw = {"placeholder": "What alternative measures mitigate this risk?", "key": "ra_comp"}
                if st.session_state.get("ra_chad_justification"):
                    ra_just_kw["value"] = st.session_state["ra_chad_justification"]
                if st.session_state.get("ra_chad_compensating"):
                    ra_comp_kw["value"] = st.session_state["ra_chad_compensating"]
                ra_justification = st.text_area("Operational Justification", **ra_just_kw)
                ra_compensating = st.text_area("Compensating Controls", **ra_comp_kw)
                ra_c1, ra_c2 = st.columns(2)
                with ra_c1:
                    ra_authority = st.text_input("Approving Authority", placeholder="Name, Title", key="ra_auth")
                with ra_c2:
                    ra_expiry = st.text_input("Acceptance Expiry", value=(datetime.now().replace(year=datetime.now().year + 1)).strftime("%Y-%m-%d"), key="ra_exp")
                ra_level = st.selectbox("Risk Level", ["Low", "Moderate", "High", "Very High"], key="ra_level")
                if st.button("✅ Log Risk Acceptance", type="primary", key="ra_submit"):
                    if ra_justification and ra_authority:
                        ra_items.append({
                            "id": f"RA-{len(ra_items)+1:03d}",
                            "finding": ra_finding,
                            "justification": ra_justification,
                            "compensating_controls": ra_compensating,
                            "approved_by": ra_authority,
                            "approved_date": datetime.now().strftime("%Y-%m-%d"),
                            "expiry_date": ra_expiry,
                            "risk_level": ra_level,
                            "status": "Active",
                        })
                        st.session_state.risk_acceptances = ra_items
                        st.session_state.pop("ra_chad_justification", None)
                        st.session_state.pop("ra_chad_compensating", None)
                        st.success("Risk acceptance logged.")
                        st.rerun()
                    else:
                        st.warning("Justification and Approving Authority are required.")
            
            if not ra_items:
                st.info("No risk acceptances logged yet. Use the form above to add one.")
            else:
                st.metric("Total Risk Acceptances", len(ra_items))
                st.markdown("---")
                st.markdown("### 📋 Comprehensive listing of all controls with acceptance of risk")
                st.caption("Each entry is hyperlink-able; use **View** to open the details.")
                # Query param to deep-link to a specific risk acceptance (e.g. ?risk_acceptance_id=RA-001)
                ra_link_id = st.query_params.get("risk_acceptance_id", "")
                # Table header
                h1, h2, h3, h4, h5 = st.columns([1, 3, 1, 1, 1])
                with h1:
                    st.markdown("**ID**")
                with h2:
                    st.markdown("**Control / Finding**")
                with h3:
                    st.markdown("**Risk Level**")
                with h4:
                    st.markdown("**Status**")
                with h5:
                    st.markdown("**Link**")
                st.markdown("---")
                for item in ra_items:
                    level_color = {"Low": "#10b981", "Moderate": "#f59e0b", "High": "#ef4444", "Very High": "#7f1d1d"}.get(item.get("risk_level", ""), "#6b7280")
                    c1, c2, c3, c4, c5 = st.columns([1, 3, 1, 1, 1])
                    with c1:
                        st.markdown(f"`{item.get('id', '')}`")
                    with c2:
                        finding_short = (item.get('finding', '')[:60] + '...') if len(item.get('finding', '')) > 60 else item.get('finding', '')
                        st.markdown(finding_short)
                    with c3:
                        st.markdown(f"<span style='color:{level_color};font-weight:600;'>{item.get('risk_level', '')}</span>", unsafe_allow_html=True)
                    with c4:
                        st.markdown(item.get('status', ''))
                    with c5:
                        q = f"?risk_acceptance_id={item.get('id', '')}"
                        st.link_button("View", q, key=f"ra_view_{item.get('id', '')}")
                st.markdown("---")
                st.markdown("**Details** (expand to view full justification and compensating controls)")
                for item in ra_items:
                    level_color = {"Low": "#10b981", "Moderate": "#f59e0b", "High": "#ef4444", "Very High": "#7f1d1d"}.get(item.get("risk_level", ""), "#6b7280")
                    is_linked = (ra_link_id == item.get("id", ""))
                    with st.expander(f"{item.get('id', '')} — {item.get('finding', '')[:50]}... | {item.get('risk_level', '')} | {item.get('status', '')}", expanded=is_linked):
                        st.markdown(f"**Finding:** {item.get('finding', '')}")
                        st.markdown(f"**Justification:** {item.get('justification', '')}")
                        st.markdown(f"**Compensating Controls:** {item.get('compensating_controls', 'None')}")
                        st.markdown(f"**Approved By:** {item.get('approved_by', '')} on {item.get('approved_date', '')}")
                        st.markdown(f"**Expiry:** {item.get('expiry_date', 'N/A')}")
                        st.markdown(f"<span style='color:{level_color};font-weight:700;'>Risk Level: {item.get('risk_level', '')}</span>", unsafe_allow_html=True)
                
                # Create Acceptance of Risk document
                st.markdown("---")
                if st.button("📄 Create Acceptance of Risk Document (.docx)", use_container_width=True, type="primary", key="ra_export_docx"):
                    try:
                        from docx import Document as DocxDocument
                        from docx.shared import Pt
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        from docx.shared import RGBColor
                        import io
                        doc = DocxDocument()
                        style = doc.styles["Normal"]
                        style.font.name = "Calibri"
                        style.font.size = Pt(10)
                        title = doc.add_paragraph()
                        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = title.add_run("Acceptance of Risk")
                        r.font.size = Pt(22)
                        r.bold = True
                        r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)
                        doc.add_paragraph()
                        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}  |  Total risk acceptances: {len(ra_items)}")
                        doc.add_paragraph()
                        for item in ra_items:
                            doc.add_heading(f"{item.get('id', '')} — {item.get('finding', '')[:60]}", level=2)
                            doc.add_paragraph(f"Risk Level: {item.get('risk_level', 'N/A')}")
                            doc.add_paragraph(f"Status: {item.get('status', 'N/A')}")
                            doc.add_paragraph(f"Approved by: {item.get('approved_by', '')} on {item.get('approved_date', '')}")
                            doc.add_paragraph(f"Expiry: {item.get('expiry_date', 'N/A')}")
                            doc.add_paragraph("Operational Justification:")
                            doc.add_paragraph(item.get("justification", ""))
                            doc.add_paragraph("Compensating Controls:")
                            doc.add_paragraph(item.get("compensating_controls", "None"))
                            doc.add_paragraph()
                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        st.download_button(
                            "📥 Download Acceptance of Risk (.docx)",
                            data=buf.getvalue(),
                            file_name=f"Acceptance_of_Risk_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="ra_dl_docx",
                        )
                    except Exception as e:
                        st.error(f"Could not create document: {e}")

    def render_bod_2201_page(self):
        """Render the CISA BOD 22-01 Known Exploited Vulnerabilities page."""
        from datetime import datetime
        import json
        
        st.markdown("""
        <div style="background: linear-gradient(135deg, #7f1d1d 0%, #991b1b 50%, #b91c1c 100%); padding: 1.5rem;
                    border-radius: 12px; margin-bottom: 1.5rem; border-left: 5px solid #fbbf24;">
            <h2 style="color: #ffffff; margin: 0;">🔴 CISA BOD 22-01 - Known Exploited Vulnerabilities</h2>
            <p style="color: #fecaca; margin: 0.5rem 0 0 0;">
                Track and remediate Known Exploited Vulnerabilities per Binding Operational Directive 22-01
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Initialize KEV checker
        try:
            from cisa_kev_checker import CISAKEVChecker, BOD2201Assessment, Severity
            kev_checker = CISAKEVChecker()
        except ImportError as e:
            st.error(f"KEV Checker module not found: {e}")
            return
        
        # Info boxes
        col1, col2, col3 = st.columns(3)
        with col1:
            st.info("📋 **BOD 22-01** requires federal agencies to remediate KEVs within specified timeframes")
        with col2:
            st.warning("⏰ **2021+ CVEs:** 2 weeks to remediate")
        with col3:
            st.error("⏰ **Pre-2021 CVEs:** 6 months to remediate")
        
        st.markdown("---")
        
        # Tabs for different views
        kev_tab1, kev_tab2, kev_tab3, kev_tab4 = st.tabs(["📊 Dashboard", "🔍 Check CVEs", "📋 Full Catalog", "📄 Report"])
        
        with kev_tab1:
            st.markdown("### KEV Catalog Overview")
            
            with st.spinner("Loading CISA KEV Catalog..."):
                if kev_checker.download_kev_catalog():
                    total_kevs = len(kev_checker.kev_catalog)
                    overdue = kev_checker.get_overdue_kevs()
                    due_soon = kev_checker.get_kevs_due_soon(7)
                    ransomware = kev_checker.get_ransomware_kevs()
                    
                    # Metrics
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        st.metric("Total KEVs in Catalog", f"{total_kevs:,}")
                    with col2:
                        st.metric("🔴 Overdue", len(overdue))
                    with col3:
                        st.metric("🟡 Due in 7 Days", len(due_soon))
                    with col4:
                        st.metric("💀 Ransomware-Related", len(ransomware))
                    
                    st.markdown("---")
                    
                    # Recent additions
                    st.markdown("### 🆕 Recently Added KEVs")
                    recent_kevs = sorted(
                        kev_checker.kev_catalog.values(),
                        key=lambda x: x.date_added,
                        reverse=True
                    )[:10]
                    
                    for kev in recent_kevs:
                        severity_color = "#dc2626" if kev.is_overdue else "#f59e0b" if kev.days_until_due <= 7 else "#22c55e"
                        status = "OVERDUE" if kev.is_overdue else f"{kev.days_until_due}d remaining"
                        
                        with st.expander(f"**{kev.cve_id}** - {kev.vendor_project} {kev.product}"):
                            st.markdown(f"""
                            - **Vulnerability:** {kev.vulnerability_name}
                            - **Date Added:** {kev.date_added.strftime('%Y-%m-%d')}
                            - **Due Date:** {kev.due_date.strftime('%Y-%m-%d')} ({status})
                            - **Ransomware Use:** {kev.known_ransomware_campaign_use}
                            - **Required Action:** {kev.required_action}
                            """)
                            st.caption(kev.short_description)
                else:
                    st.error("Failed to load KEV catalog. Check internet connection.")
        
        with kev_tab2:
            st.markdown("### 🔍 Check CVEs Against KEV Catalog")
            st.markdown("Enter CVE IDs to check if they are Known Exploited Vulnerabilities")
            
            # Demo data section
            with st.expander("🧪 **Load Demo Data** (for testing without real scan data)", expanded=False):
                st.markdown("Select a demo scenario to test the KEV checker:")
                
                try:
                    from sample_vuln_data import DEMO_SCENARIOS, get_kev_only_cves
                    
                    demo_cols = st.columns(3)
                    
                    with demo_cols[0]:
                        if st.button("🔴 Critical Outbreak", use_container_width=True, help="Log4Shell, ProxyLogon, Zerologon"):
                            st.session_state.demo_cves = DEMO_SCENARIOS["critical_outbreak"]["cves"]
                            st.rerun()
                    
                    with demo_cols[1]:
                        if st.button("📧 Exchange Compromise", use_container_width=True, help="Exchange-related KEVs"):
                            st.session_state.demo_cves = DEMO_SCENARIOS["exchange_compromise"]["cves"]
                            st.rerun()
                    
                    with demo_cols[2]:
                        if st.button("💀 Ransomware Risk", use_container_width=True, help="KEVs used in ransomware"):
                            st.session_state.demo_cves = DEMO_SCENARIOS["ransomware_indicators"]["cves"]
                            st.rerun()
                    
                    demo_cols2 = st.columns(3)
                    
                    with demo_cols2[0]:
                        if st.button("✅ Clean Scan", use_container_width=True, help="No KEVs - only regular vulns"):
                            st.session_state.demo_cves = DEMO_SCENARIOS["clean_scan"]["cves"]
                            st.rerun()
                    
                    with demo_cols2[1]:
                        if st.button("📊 Mixed Findings", use_container_width=True, help="Typical scan with some KEVs"):
                            st.session_state.demo_cves = DEMO_SCENARIOS["mixed_findings"]["cves"]
                            st.rerun()
                    
                    with demo_cols2[2]:
                        if st.button("🗑️ Clear Demo Data", use_container_width=True):
                            if 'demo_cves' in st.session_state:
                                del st.session_state.demo_cves
                            if 'kev_findings' in st.session_state:
                                del st.session_state.kev_findings
                            st.rerun()
                    
                    if 'demo_cves' in st.session_state:
                        st.success(f"✅ Demo data loaded: {len(st.session_state.demo_cves)} CVEs")
                        st.code("\n".join(st.session_state.demo_cves))
                        
                except ImportError:
                    st.warning("Demo data module not available")
            
            st.markdown("---")
            
            # Input methods
            input_method = st.radio("Input Method", ["Manual Entry", "Upload File", "Paste from Tenable/Scanner"], horizontal=True)
            
            cve_list = []
            
            # Check for demo data first
            if 'demo_cves' in st.session_state:
                cve_list = st.session_state.demo_cves
                st.info(f"📋 Using demo data: {len(cve_list)} CVEs loaded")
            elif input_method == "Manual Entry":
                cve_input = st.text_area(
                    "Enter CVE IDs (one per line or comma-separated)",
                    placeholder="CVE-2021-44228\nCVE-2021-34527\nCVE-2023-23397",
                    height=150
                )
                if cve_input:
                    # Parse CVEs
                    cve_list = [
                        cve.strip().upper() 
                        for cve in cve_input.replace(',', '\n').split('\n') 
                        if cve.strip() and cve.strip().upper().startswith('CVE-')
                    ]
                    
            elif input_method == "Upload File":
                uploaded_file = st.file_uploader("Upload CSV or JSON with CVE data", type=['csv', 'json', 'txt'])
                if uploaded_file:
                    content = uploaded_file.read().decode('utf-8')
                    # Extract CVE IDs using regex
                    import re
                    cve_pattern = r'CVE-\d{4}-\d{4,}'
                    cve_list = list(set(re.findall(cve_pattern, content.upper())))
                    st.success(f"Found {len(cve_list)} CVE IDs in file")
                    
            elif input_method == "Paste from Tenable/Scanner":
                scanner_input = st.text_area(
                    "Paste vulnerability scan output",
                    placeholder="Paste Tenable, Nessus, or other scanner output here...",
                    height=200
                )
                if scanner_input:
                    import re
                    cve_pattern = r'CVE-\d{4}-\d{4,}'
                    cve_list = list(set(re.findall(cve_pattern, scanner_input.upper())))
                    st.success(f"Extracted {len(cve_list)} CVE IDs")
            
            if cve_list:
                st.markdown(f"**CVEs to check:** {len(cve_list)}")
                
                if st.button("🔍 Check Against KEV Catalog", type="primary", use_container_width=True):
                    with st.spinner("Checking CVEs..."):
                        if kev_checker.download_kev_catalog():
                            found_kevs = kev_checker.check_cve_list(cve_list)
                            
                            if found_kevs:
                                st.error(f"⚠️ **{len(found_kevs)} KEVs FOUND** in your vulnerability list!")
                                
                                # Summary
                                overdue = [k for k in found_kevs if k.is_overdue]
                                due_soon = [k for k in found_kevs if 0 <= k.days_until_due <= 7]
                                
                                col1, col2, col3 = st.columns(3)
                                with col1:
                                    st.metric("Total KEVs Found", len(found_kevs))
                                with col2:
                                    st.metric("🔴 Overdue", len(overdue))
                                with col3:
                                    st.metric("🟡 Due Soon", len(due_soon))
                                
                                st.markdown("---")
                                
                                # Details
                                for kev in sorted(found_kevs, key=lambda x: x.days_until_due):
                                    if kev.is_overdue:
                                        st.error(f"**🔴 {kev.cve_id}** - OVERDUE by {abs(kev.days_until_due)} days")
                                    elif kev.days_until_due <= 7:
                                        st.warning(f"**🟡 {kev.cve_id}** - Due in {kev.days_until_due} days")
                                    else:
                                        st.info(f"**🟢 {kev.cve_id}** - Due in {kev.days_until_due} days")
                                    
                                    st.markdown(f"""
                                    > **{kev.vendor_project} {kev.product}** - {kev.vulnerability_name}  
                                    > Due: {kev.due_date.strftime('%Y-%m-%d')} | Ransomware: {kev.known_ransomware_campaign_use}  
                                    > **Action:** {kev.required_action}
                                    """)
                                
                                # Store for report generation
                                st.session_state.kev_findings = found_kevs
                            else:
                                st.success("✅ No KEVs found in your CVE list!")
                        else:
                            st.error("Failed to load KEV catalog")
        
        with kev_tab3:
            st.markdown("### 📋 Full KEV Catalog")
            
            if kev_checker.download_kev_catalog():
                # Search/filter options
                col1, col2 = st.columns(2)
                with col1:
                    vendor_search = st.text_input("🔍 Search by Vendor", placeholder="e.g., Microsoft, Apache")
                with col2:
                    product_search = st.text_input("🔍 Search by Product", placeholder="e.g., Exchange, Log4j")
                
                # Filter
                filtered_kevs = list(kev_checker.kev_catalog.values())
                
                if vendor_search:
                    filtered_kevs = [k for k in filtered_kevs if vendor_search.lower() in k.vendor_project.lower()]
                if product_search:
                    filtered_kevs = [k for k in filtered_kevs if product_search.lower() in k.product.lower()]
                
                # Sort options
                sort_by = st.selectbox("Sort by", ["Date Added (newest)", "Due Date", "Vendor", "CVE ID"])
                
                if sort_by == "Date Added (newest)":
                    filtered_kevs.sort(key=lambda x: x.date_added, reverse=True)
                elif sort_by == "Due Date":
                    filtered_kevs.sort(key=lambda x: x.due_date)
                elif sort_by == "Vendor":
                    filtered_kevs.sort(key=lambda x: x.vendor_project)
                else:
                    filtered_kevs.sort(key=lambda x: x.cve_id)
                
                st.markdown(f"**Showing {len(filtered_kevs)} KEVs**")
                
                # Display as table
                if filtered_kevs:
                    import pandas as pd
                    df_data = [{
                        'CVE ID': k.cve_id,
                        'Vendor': k.vendor_project,
                        'Product': k.product,
                        'Added': k.date_added.strftime('%Y-%m-%d'),
                        'Due': k.due_date.strftime('%Y-%m-%d'),
                        'Status': 'OVERDUE' if k.is_overdue else f'{k.days_until_due}d',
                        'Ransomware': k.known_ransomware_campaign_use,
                    } for k in filtered_kevs[:100]]  # Limit to 100 for performance
                    
                    df = pd.DataFrame(df_data)
                    st.dataframe(df, use_container_width=True, hide_index=True)
                    
                    if len(filtered_kevs) > 100:
                        st.caption(f"Showing first 100 of {len(filtered_kevs)} results")
                    
                    # Download option
                    csv_data = df.to_csv(index=False)
                    st.download_button(
                        "📥 Download as CSV",
                        data=csv_data,
                        file_name=f"kev_catalog_{datetime.now().strftime('%Y%m%d')}.csv",
                        mime="text/csv"
                    )
        
        with kev_tab4:
            st.markdown("### 📄 BOD 22-01 Compliance Report")
            
            if 'kev_findings' in st.session_state and st.session_state.kev_findings:
                kevs = st.session_state.kev_findings
                
                # Generate assessment
                assessment = BOD2201Assessment(
                    assessment_date=datetime.now(),
                    organization=st.session_state.get('system_info', {}).get('system_name', 'Organization'),
                    total_kevs_in_catalog=len(kev_checker.kev_catalog),
                )
                
                from cisa_kev_checker import KEVFinding, RemediationStatus
                for kev in kevs:
                    assessment.kevs_found.append(KEVFinding(kev=kev))
                
                # Generate report
                report = kev_checker.generate_report(assessment)
                
                st.code(report, language="text")
                
                col1, col2 = st.columns(2)
                with col1:
                    st.download_button(
                        "📥 Download Report (TXT)",
                        data=report,
                        file_name=f"bod_2201_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt",
                        mime="text/plain",
                        use_container_width=True
                    )
                with col2:
                    json_data = json.dumps(assessment.to_dict(), indent=2, default=str)
                    st.download_button(
                        "📥 Download Report (JSON)",
                        data=json_data,
                        file_name=f"bod_2201_assessment_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json",
                        use_container_width=True
                    )
            else:
                st.info("👆 Use the 'Check CVEs' tab first to identify KEVs in your environment, then return here to generate a compliance report.")
            
            st.markdown("---")
            st.markdown("""
            ### About BOD 22-01
            
            **Binding Operational Directive 22-01** requires federal agencies to:
            
            1. **Track** all Known Exploited Vulnerabilities (KEVs) in their environment
            2. **Remediate** within specified timeframes:
               - CVEs from 2021 or later: **14 days** from addition to KEV catalog
               - CVEs from before 2021: **6 months** from addition to KEV catalog
            3. **Report** remediation status to CISA
            
            **References:**
            - [BOD 22-01 Directive](https://www.cisa.gov/news-events/directives/bod-22-01-reducing-significant-risk-known-exploited-vulnerabilities)
            - [KEV Catalog](https://www.cisa.gov/known-exploited-vulnerabilities-catalog)
            """)

    def render_threat_modeling_page(self):
        """Render the Threat Modeling page as a dedicated tab."""
        st.markdown("""
        <div style="background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 50%, #3d7ab5 100%); padding: 1.5rem;
                    border-radius: 12px; margin-bottom: 1.5rem; border-left: 5px solid #f59e0b;">
            <h2 style="color: #ffffff; margin: 0;">🎯 Threat Modeling</h2>
            <p style="color: #b8d4e8; margin: 0.5rem 0 0 0;">
                Map security controls to threats, assets, and analytical frameworks like MITRE ATT&CK
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # Get assessment data from session state
        assessment_data = st.session_state.get('results', None)
        findings = None
        if hasattr(st.session_state, 'risk_assessment') and st.session_state.risk_assessment:
            findings = st.session_state.risk_assessment.get('findings', None)
        
        # Render the threat modeling section from nist_dashboard
        render_threat_modeling_section(assessment_data=assessment_data, findings=findings)

    def render_documentation_page(self):
        """Render the Documentation page for storing security documents like SSPs."""
        from datetime import datetime
        
        st.markdown("""
        <div style="background: linear-gradient(135deg, #1e3a5f 0%, #2d5a87 100%); 
                    color: white; padding: 1.5rem; border-radius: 10px; margin-bottom: 1.5rem;">
            <h2 style="margin: 0; color: white;">📁 Security Artifacts Repository</h2>
            <p style="margin: 0.5rem 0 0 0; opacity: 0.9;">
                Store and manage security-relevant documents such as SSPs, POA&Ms, and other artifacts
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        # S3 Configuration - use configured bucket, env var, or imported constant
        s3_bucket = st.session_state.get('s3_bucket') or os.environ.get('S3_BUCKET_NAME') or S3_BUCKET_NAME
        docs_folder = "Documentation"
        
        if not s3_bucket:
            st.warning("⚠️ S3 bucket not configured. Please set the S3_BUCKET_NAME environment variable or configure it in the assessment settings.")
            return
        
        # Store in session state for consistency
        st.session_state.s3_bucket = s3_bucket
        
        # Show connected bucket info
        st.success(f"✅ Connected to S3 bucket: **{s3_bucket}**")
        
        # Initialize S3 client - use session state credentials (from login) or environment variables
        try:
            aws_access_key = st.session_state.get('aws_access_key') or os.environ.get('AWS_ACCESS_KEY_ID')
            aws_secret_key = st.session_state.get('aws_secret_key') or os.environ.get('AWS_SECRET_ACCESS_KEY')
            aws_region = st.session_state.get('aws_region') or os.environ.get('AWS_DEFAULT_REGION', 'us-east-1')
            
            if aws_access_key and aws_secret_key:
                s3_client = boto3.client('s3',
                    aws_access_key_id=aws_access_key,
                    aws_secret_access_key=aws_secret_key,
                    region_name=aws_region
                )
            else:
                # Try default credentials (IAM role, etc.)
                s3_client = boto3.client('s3', region_name=aws_region)
        except Exception as e:
            st.error(f"❌ Failed to initialize S3 client: {str(e)}")
            return
        
        # Create two columns for upload and listing
        col_upload, col_list = st.columns([1, 1])
        
        with col_upload:
            st.markdown("### 📤 Upload Document")
            st.markdown("*Supported formats: PDF, DOCX, XLSX, TXT, JSON, CSV, MD*")
            
            # Document type selector
            doc_type = st.selectbox(
                "Document Type",
                ["System Security Plan (SSP)", "Plan of Action & Milestones (POA&M)", 
                 "Security Assessment Report (SAR)", "Authorization to Operate (ATO)",
                 "Continuous Monitoring Report", "Incident Response Plan",
                 "Configuration Management Plan", "Contingency Plan", "Other"]
            )
            
            # File uploader
            uploaded_file = st.file_uploader(
                "Choose a file",
                type=['pdf', 'docx', 'xlsx', 'txt', 'json', 'csv', 'md'],
                key="doc_uploader"
            )
            
            # Optional description
            doc_description = st.text_area("Description (optional)", height=80)
            
            if uploaded_file is not None:
                if st.button("📤 Upload to S3", type="primary", use_container_width=True):
                    try:
                        # Create a clean filename with document type prefix
                        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                        doc_type_prefix = doc_type.split('(')[0].strip().replace(' ', '_')
                        file_extension = uploaded_file.name.split('.')[-1]
                        s3_key = f"{docs_folder}/{doc_type_prefix}_{timestamp}.{file_extension}"
                        
                        # Add metadata
                        metadata = {
                            'original_filename': uploaded_file.name,
                            'document_type': doc_type,
                            'uploaded_by': st.session_state.get('aws_account_id', 'Unknown'),
                            'description': doc_description[:256] if doc_description else ''
                        }
                        
                        # Upload to S3
                        s3_client.upload_fileobj(
                            uploaded_file,
                            s3_bucket,
                            s3_key,
                            ExtraArgs={'Metadata': metadata}
                        )
                        
                        st.success(f"✅ Document uploaded successfully!")
                        st.markdown(f"**Location:** `s3://{s3_bucket}/{s3_key}`")
                        st.rerun()
                        
                    except Exception as e:
                        st.error(f"❌ Upload failed: {str(e)}")
        
        with col_list:
            st.markdown("### 📂 Stored Documents")
            
            # Refresh button
            if st.button("🔄 Refresh List", use_container_width=True):
                st.rerun()
            
            try:
                # List objects in the Documentation folder
                response = s3_client.list_objects_v2(
                    Bucket=s3_bucket,
                    Prefix=f"{docs_folder}/"
                )
                
                if 'Contents' not in response or len(response['Contents']) == 0:
                    st.info("📭 No documents stored yet. Upload your first document!")
                else:
                    # Filter out the folder itself
                    documents = [obj for obj in response['Contents'] if not obj['Key'].endswith('/')]
                    
                    if not documents:
                        st.info("📭 No documents stored yet. Upload your first document!")
                    else:
                        st.markdown(f"**{len(documents)} document(s) found**")
                        
                        for doc in sorted(documents, key=lambda x: x['LastModified'], reverse=True):
                            file_key = doc['Key']
                            file_name = file_key.split('/')[-1]
                            file_size = doc['Size'] / 1024  # KB
                            last_modified = doc['LastModified'].strftime("%Y-%m-%d %H:%M")
                            
                            # Try to get metadata
                            try:
                                head_response = s3_client.head_object(Bucket=s3_bucket, Key=file_key)
                                metadata = head_response.get('Metadata', {})
                                original_name = metadata.get('original_filename', file_name)
                                doc_type_meta = metadata.get('document_type', 'Unknown')
                            except:
                                original_name = file_name
                                doc_type_meta = 'Unknown'
                            
                            # Display document card
                            with st.expander(f"📄 {original_name}", expanded=False):
                                st.markdown(f"""
                                - **Type:** {doc_type_meta}
                                - **Size:** {file_size:.1f} KB
                                - **Uploaded:** {last_modified}
                                - **S3 Key:** `{file_key}`
                                """)
                                
                                # Download button
                                col_dl, col_del = st.columns(2)
                                with col_dl:
                                    try:
                                        # Generate presigned URL for download
                                        presigned_url = s3_client.generate_presigned_url(
                                            'get_object',
                                            Params={'Bucket': s3_bucket, 'Key': file_key},
                                            ExpiresIn=3600
                                        )
                                        st.markdown(f'<a href="{presigned_url}" target="_blank"><button style="background: #1e3a5f; color: white; border: none; padding: 0.5rem 1rem; border-radius: 5px; cursor: pointer; width: 100%;">⬇️ Download</button></a>', unsafe_allow_html=True)
                                    except Exception as e:
                                        st.error(f"Download error: {str(e)}")
                                
                                with col_del:
                                    if st.button(f"🗑️ Delete", key=f"del_{file_key}", use_container_width=True):
                                        try:
                                            s3_client.delete_object(Bucket=s3_bucket, Key=file_key)
                                            st.success("Document deleted!")
                                            st.rerun()
                                        except Exception as e:
                                            st.error(f"Delete failed: {str(e)}")
                            
            except s3_client.exceptions.NoSuchBucket:
                st.error(f"❌ S3 bucket '{s3_bucket}' does not exist.")
            except Exception as e:
                st.error(f"❌ Failed to list documents: {str(e)}")
        
        # Document types reference
        st.markdown("---")
        st.markdown("### 📚 Document Types Reference")
        
        doc_types_data = {
            "Document Type": ["SSP", "POA&M", "SAR", "ATO", "CM Report", "IRP", "CMP", "CP"],
            "Full Name": [
                "System Security Plan",
                "Plan of Action & Milestones", 
                "Security Assessment Report",
                "Authorization to Operate",
                "Continuous Monitoring Report",
                "Incident Response Plan",
                "Configuration Management Plan",
                "Contingency Plan"
            ],
            "Purpose": [
                "Documents system security controls implementation",
                "Tracks security weaknesses and remediation timeline",
                "Results of security control assessments",
                "Formal authorization decision document",
                "Ongoing security status monitoring",
                "Procedures for security incident handling",
                "Baseline configuration and change management",
                "Business continuity and disaster recovery"
            ]
        }
        
        import pandas as pd
        df = pd.DataFrame(doc_types_data)
        st.dataframe(df, use_container_width=True, hide_index=True)
    
    def render_aws_login_page(self):
        """Render AWS credentials login page."""
        # Hide sidebar
        st.markdown("""
        <style>
            [data-testid="stSidebar"] { display: none; }
        </style>
        """, unsafe_allow_html=True)
        
        # Apply custom CSS
        apply_custom_css()
        
        # Header with SAELAR logo centered
        logo_base64 = get_base64_image(LOGO_PATH)
        if logo_base64:
            st.markdown(f"""
            <div style="display: flex; justify-content: center; align-items: center; width: 100%;">
                <img src="data:image/png;base64,{logo_base64}" style="width: 600px; max-width: 95%; height: auto;" />
            </div>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # AWS logo and title
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            # AWS Logo (using official AWS orange color scheme)
            st.markdown("""
            <div style="text-align: center; margin-bottom: 1rem;">
                <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 304 182" width="120" style="margin: 0 auto;">
                    <path fill="#252F3E" d="M86.4 66.4c0 3.7.4 6.7 1.1 8.9.8 2.2 1.8 4.6 3.2 7.2.5.8.7 1.6.7 2.3 0 1-.6 2-1.9 3l-6.3 4.2c-.9.6-1.8.9-2.6.9-1 0-2-.5-3-1.4-1.4-1.5-2.6-3.1-3.6-4.7-1-1.7-2-3.6-3.1-5.9-7.8 9.2-17.6 13.8-29.4 13.8-8.4 0-15.1-2.4-20-7.2-4.9-4.8-7.4-11.2-7.4-19.2 0-8.5 3-15.4 9.1-20.6 6.1-5.2 14.2-7.8 24.5-7.8 3.4 0 6.9.3 10.6.8 3.7.5 7.5 1.3 11.5 2.2v-7.3c0-7.6-1.6-12.9-4.7-16-3.2-3.1-8.6-4.6-16.3-4.6-3.5 0-7.1.4-10.8 1.3-3.7.9-7.3 2-10.8 3.4-1.6.7-2.8 1.1-3.5 1.3-.7.2-1.2.3-1.6.3-1.4 0-2.1-1-2.1-3.1v-4.9c0-1.6.2-2.8.7-3.5.5-.7 1.4-1.4 2.8-2.1 3.5-1.8 7.7-3.3 12.6-4.5 4.9-1.3 10.1-1.9 15.6-1.9 11.9 0 20.6 2.7 26.2 8.1 5.5 5.4 8.3 13.6 8.3 24.6v32.4zM45.8 81.6c3.3 0 6.7-.6 10.3-1.8 3.6-1.2 6.8-3.4 9.5-6.4 1.6-1.9 2.8-4 3.4-6.4.6-2.4 1-5.3 1-8.7v-4.2c-2.9-.7-6-1.3-9.2-1.7-3.2-.4-6.3-.6-9.4-.6-6.7 0-11.6 1.3-14.9 4-3.3 2.7-4.9 6.5-4.9 11.5 0 4.7 1.2 8.2 3.7 10.6 2.4 2.5 5.9 3.7 10.5 3.7zm80.3 10.8c-1.8 0-3-.3-3.8-1-.8-.6-1.5-2-2.1-3.9l-23.5-77.3c-.6-2-.9-3.3-.9-4 0-1.6.8-2.5 2.4-2.5h9.8c1.9 0 3.2.3 3.9 1 .8.6 1.4 2 2 3.9l16.8 66.2 15.6-66.2c.5-2 1.1-3.3 1.9-3.9.8-.6 2.2-1 4-1h8c1.9 0 3.2.3 4 1 .8.6 1.5 2 1.9 3.9l15.8 67 17.3-67c.6-2 1.3-3.3 2-3.9.8-.6 2.1-1 3.9-1h9.3c1.6 0 2.5.8 2.5 2.5 0 .5-.1 1-.2 1.6-.1.6-.3 1.4-.7 2.5l-24.1 77.3c-.6 2-1.3 3.3-2.1 3.9-.8.6-2.1 1-3.8 1h-8.6c-1.9 0-3.2-.3-4-1-.8-.7-1.5-2-1.9-4l-15.5-64.5-15.4 64.4c-.5 2-1.1 3.3-1.9 4-.8.7-2.2 1-4 1h-8.6zm128.5 2.7c-5.2 0-10.4-.6-15.4-1.8-5-1.2-8.9-2.5-11.5-4-1.6-.9-2.7-1.9-3.1-2.8-.4-.9-.6-1.9-.6-2.8v-5.1c0-2.1.8-3.1 2.3-3.1.6 0 1.2.1 1.8.3.6.2 1.5.6 2.5 1 3.4 1.5 7.1 2.7 11 3.5 4 .8 7.9 1.2 11.9 1.2 6.3 0 11.2-1.1 14.6-3.3 3.4-2.2 5.2-5.4 5.2-9.5 0-2.8-.9-5.1-2.7-7-1.8-1.9-5.2-3.6-10.1-5.2l-14.5-4.5c-7.3-2.3-12.7-5.7-16-10.2-3.3-4.4-5-9.3-5-14.5 0-4.2.9-7.9 2.7-11.1 1.8-3.2 4.2-6 7.2-8.2 3-2.3 6.4-4 10.4-5.2 4-1.2 8.2-1.7 12.6-1.7 2.2 0 4.5.1 6.7.4 2.3.3 4.4.7 6.5 1.1 2 .5 3.9 1 5.7 1.6 1.8.6 3.2 1.2 4.2 1.8 1.4.8 2.4 1.6 3 2.5.6.8.9 1.9.9 3.3v4.7c0 2.1-.8 3.2-2.3 3.2-.8 0-2.1-.4-3.8-1.2-5.7-2.6-12.1-3.9-19.2-3.9-5.7 0-10.2.9-13.3 2.8-3.1 1.9-4.7 4.8-4.7 8.9 0 2.8 1 5.2 3 7.1 2 1.9 5.7 3.8 11 5.5l14.2 4.5c7.2 2.3 12.4 5.5 15.5 9.6 3.1 4.1 4.6 8.8 4.6 14 0 4.3-.9 8.2-2.6 11.6-1.8 3.4-4.2 6.4-7.3 8.8-3.1 2.5-6.8 4.3-11.1 5.6-4.5 1.4-9.2 2.1-14.3 2.1z"/>
                    <path fill="#FF9900" d="M273.5 143.7c-32.9 24.3-80.7 37.2-121.8 37.2-57.6 0-109.5-21.3-148.7-56.7-3.1-2.8-.3-6.6 3.4-4.4 42.4 24.6 94.7 39.5 148.8 39.5 36.5 0 76.6-7.6 113.5-23.2 5.5-2.5 10.2 3.6 4.8 7.6z"/>
                    <path fill="#FF9900" d="M287.2 128.1c-4.2-5.4-27.8-2.6-38.5-1.3-3.2.4-3.7-2.4-.8-4.5 18.8-13.2 49.7-9.4 53.3-5 3.6 4.5-1 35.4-18.6 50.2-2.7 2.3-5.3 1.1-4.1-1.9 4-9.9 12.9-32.2 8.7-37.5z"/>
                </svg>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("""
            <h2 style="color: #1e3a5f; text-align: center; margin-bottom: 0;">AWS Configuration</h2>
            <p style="color: #64748b; text-align: center; margin-top: 0.5rem;">Enter the AWS credentials to the environment to be assessed</p>
            """, unsafe_allow_html=True)
        
        st.markdown("---")
        
        # Center the AWS credentials form
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            # Check for saved credentials
            saved_creds = load_saved_credentials()
            
            if saved_creds and not st.session_state.get('show_new_aws_form', False):
                st.success(f"✅ Saved credentials found for Account: `{saved_creds.get('account_id', 'N/A')}`")
                
                col_use, col_new = st.columns(2)
                with col_use:
                    if st.button("🔗 Saved Credentials", use_container_width=True):
                        access_key = saved_creds.get('access_key_id', '')
                        secret_key = saved_creds.get('secret_key', '')
                        region = saved_creds.get('region', 'us-east-1')
                        account_id = saved_creds.get('account_id', '')
                        
                        # Validate credentials are complete
                        if not access_key or not secret_key:
                            st.error("❌ Saved credentials are incomplete. Please enter new credentials.")
                            st.session_state.show_new_aws_form = True
                        else:
                            # Validate with AWS
                            with st.spinner("Validating saved credentials..."):
                                success, message = validate_aws_credentials(access_key, secret_key, account_id, region)
                            
                            if success:
                                set_aws_credentials(access_key, secret_key, region)
                                st.session_state.aws_validated = True
                                st.session_state.aws_account_id = account_id
                                st.session_state.aws_region = region
                                st.session_state.aws_iam_username = saved_creds.get('iam_username', '')
                                st.rerun()
                            else:
                                st.error(f"❌ {message}")
                                st.info("💡 Your saved credentials may have expired or been rotated. Please enter new credentials.")
                with col_new:
                    if st.button("🔄 New Credentials", use_container_width=True):
                        st.session_state.show_new_aws_form = True
                        st.rerun()
                
                # Show PROCEED button after using saved credentials
                if st.session_state.get('aws_validated', False) and not st.session_state.get('aws_configured', False):
                    st.markdown("---")
                    st.success(f"✅ **AWS Connection Successful!**\n\nAccount: `{st.session_state.aws_account_id}`  |  Region: `{st.session_state.aws_region}`")
                    
                    if st.button("🚀 PROCEED TO THE ASSESSMENT SECTION", use_container_width=True, key="proceed_saved_btn"):
                        st.session_state.aws_configured = True
                        st.rerun()
            else:
                st.markdown("### ☁️ Enter AWS Credentials")
                
                with st.form("aws_credentials_form"):
                    account_id = st.text_input("AWS Account ID", placeholder="123456789012", max_chars=12, help="Your 12-digit AWS Account ID")
                    iam_username = st.text_input("IAM Username", placeholder="security-auditor", help="Your IAM username (optional)")
                    
                    col_key, col_secret = st.columns(2)
                    with col_key:
                        access_key = st.text_input("Access Key ID", placeholder="AKIA...", help="Your IAM Access Key ID")
                    with col_secret:
                        secret_key = st.text_input("Secret Access Key", type="password", placeholder="Secret...", help="Your IAM Secret Access Key")
                    
                    region = st.selectbox("AWS Region", options=[
                        "us-east-1", "us-east-2", "us-west-1", "us-west-2",
                        "eu-west-1", "eu-west-2", "eu-central-1",
                        "ap-southeast-1", "ap-southeast-2", "ap-northeast-1"
                    ])
                    
                    save_creds = st.checkbox("💾 Save credentials for future use", value=True)
                    
                    aws_submitted = st.form_submit_button("🔗 Connect to AWS", use_container_width=True)
                    
                    if aws_submitted:
                        if not account_id or not access_key or not secret_key:
                            st.error("Please fill in Account ID, Access Key, and Secret Key")
                        elif len(account_id) != 12 or not account_id.isdigit():
                            st.error("Account ID must be exactly 12 digits")
                        else:
                            with st.spinner("Validating AWS credentials..."):
                                success, message = validate_aws_credentials(access_key, secret_key, account_id, region)
                            
                            if success:
                                if save_creds:
                                    save_aws_credentials_to_file(
                                        access_key, secret_key, region, account_id, iam_username, "", "saelar"
                                    )
                                set_aws_credentials(access_key, secret_key, region)
                                st.session_state.aws_validated = True
                                st.session_state.aws_account_id = account_id
                                st.session_state.aws_region = region
                                st.session_state.aws_iam_username = iam_username
                                st.session_state.show_new_aws_form = False
                                st.rerun()
                            else:
                                st.error(f"❌ {message}")
                
                # Show PROCEED button after new credentials validation
                if st.session_state.get('aws_validated', False) and not st.session_state.get('aws_configured', False):
                    st.markdown("---")
                    st.success(f"✅ **AWS Connection Successful!**\n\nAccount: `{st.session_state.aws_account_id}`  |  Region: `{st.session_state.aws_region}`")
                    
                    if st.button("🚀 PROCEED TO THE ASSESSMENT SECTION", use_container_width=True, key="proceed_new_btn"):
                        st.session_state.aws_configured = True
                        st.rerun()
    
    def run(self):
        """Main application entry point."""
        # Check if splash page has been accepted
        if not st.session_state.splash_accepted:
            if render_splash_page():
                st.session_state.splash_accepted = True
                st.rerun()
            return
        
        # Check if AWS is configured
        if not st.session_state.aws_configured:
            self.render_aws_login_page()
            return
        
        # Ensure AWS credentials are in environment (they may have been lost on rerun)
        if not os.environ.get('AWS_ACCESS_KEY_ID') or not os.environ.get('AWS_SECRET_ACCESS_KEY'):
            saved_creds = load_saved_credentials()
            if saved_creds and saved_creds.get('access_key_id') and saved_creds.get('secret_key'):
                set_aws_credentials(
                    saved_creds.get('access_key_id', ''),
                    saved_creds.get('secret_key', ''),
                    saved_creds.get('region', st.session_state.get('aws_region', 'us-east-1'))
                )
            else:
                # Credentials not available, need to re-authenticate
                st.session_state.aws_configured = False
                st.warning("⚠️ Session expired. Please re-authenticate.")
                st.rerun()
                return
        
        # Apply CSS styling
        apply_custom_css()
        
        # Expand sidebar - ensure it's visible with !important to override any hidden state
        st.markdown("""
        <style>
            [data-testid="stSidebar"] {
                display: block !important;
                min-width: 300px !important;
            }
            [data-testid="stSidebarNav"] {
                display: block !important;
            }
            section[data-testid="stSidebar"] > div {
                display: block !important;
            }
        </style>
        """, unsafe_allow_html=True)
        
        # Render header
        render_header()
        
        # Render main content (sidebar is rendered inside render_assessment_page)
        self.render_main_content()


# =============================================================================
# MAIN ENTRY POINT
# =============================================================================

def main():
    """
    Main entry point for the NIST 800-53 Rev 5 Assessment Application.
    
    This function initializes the page configuration and starts the
    application controller.
    """
    # Configure the page (must be first Streamlit command)
    configure_page()
    
    # Create and run the application controller
    app = NISTApplicationController()
    app.run()


# =============================================================================
# RUN APPLICATION
# =============================================================================

if __name__ == "__main__":
    main()
