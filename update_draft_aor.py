#!/usr/bin/env python3
"""
Update DRAFT_AOR_SAE_CSTA-SANDBOX.docx with SAELAR/SOPRA data leakage content.
"""
from docx import Document
from docx.shared import Pt
from pathlib import Path

ROOT = Path(__file__).resolve().parent
DRAFT_PATH = ROOT / "DRAFT_AOR_SAE_CSTA-SANDBOX.docx"
OUT_PATH = ROOT / "DRAFT_AOR_SAE_CSTA-SANDBOX_UPDATED.docx"

RISKS = [
    "AI Model Interaction Exposure (Bedrock/Ollama): User prompts and AI responses sent to AWS Bedrock or local Ollama may include sensitive assessment context, system names, account identifiers, or remediation details. Cloud providers may retain or log this data. Even with local Ollama, prompts and responses may be written to disk or logs during testing.",
    "AWS Credentials in Assessment Flows: SAELAR prompts for IAM access keys, secret keys, and account IDs when configuring AWS-based assessments. These values may be stored in session state, logs, or temporary files. In a sandbox environment, test credentials should be used rather than production credentials.",
    "Third-Party API and External Service Transmission: Integration with CISA KEV, Nessus, Splunk, and similar services transmits product names, CVE identifiers, and other context over the network. Third-party providers may log requests and responses. Data may pass through ngrok or other tunneling services with their own logging and retention policies.",
    "Exported Documents and Reports: SSPs, POA&Ms, risk assessments, and other exports may contain system names, owners, findings, remediation details, and organizational information. These files may be stored on local or shared drives in the sandbox with less restrictive access controls than production environments.",
    "Authentication and Configuration Data at Rest: User credentials (hashed) and API tokens are stored in configuration files such as .nist_users.json and application settings. These files may be backed up, copied, or accessible to other processes or users on the same sandbox host.",
]

REMEDIATION = [
    "Replace SHA-256 password hashing with bcrypt or argon2 for credential storage.",
    "Implement session-scoped credential handling; avoid persisting AWS keys to disk. Use AWS credential chain or temporary session tokens where possible.",
    "Ensure all external API calls use TLS 1.2 or higher. Restrict third-party integrations to sandbox or trial API keys during testing.",
    "Implement access controls and audit logging for exported document storage. Apply data classification and retention policies for sandbox exports.",
    "Migrate credential and configuration storage to a secrets manager (e.g., AWS Secrets Manager) or encrypted vault. Apply least-privilege file permissions.",
]

REMEDIATION_COSTS = (
    "Estimated effort: 40–80 hours development and testing. Timeline: 2–4 months following AOR approval, "
    "contingent on resource availability. Costs may include engineering labor, potential migration to managed "
    "secrets services, and validation testing. A reassessment will be conducted upon completion of remediation."
)

MITIGATIONS = [
    ("AI Model Interaction Exposure", "Use air-gapped Ollama (SAELAR_AIRGAPPED=true) for testing when feasible to keep prompts and responses on-premises. When using Bedrock, restrict prompts to synthetic or de-identified data only. Avoid pasting production system names or real account identifiers into AI conversations."),
    ("AWS Credentials in Assessment Flows", "Dedicated test IAM users and roles with minimal permissions (read-only where possible) will be used for sandbox assessments. No production credentials will be entered. Test credentials will be rotated or deleted after the testing period."),
    ("Third-Party API and External Service Transmission", "Sandbox or trial API keys will be used for Nessus, Splunk, and similar integrations where available. Network egress from the sandbox can be restricted to approved endpoints. CISA KEV and similar public APIs transmit only product/CVE identifiers, not organization-specific data."),
    ("Exported Documents and Reports", "Sandbox storage will be isolated from production systems. Exports will use synthetic system names, placeholder owners, and de-identified findings. Access to sandbox file shares will be limited to authorized testers. Exported files will be purged at the end of the testing period."),
    ("Authentication and Configuration Data at Rest", "Restrictive file permissions (e.g., 600) on credential and configuration files. Sandbox hosts will be dedicated to testing and not used for production workloads. No automated backups of sandbox credential files; if backups exist, they will be excluded from off-site replication."),
]


def main():
    doc = Document(DRAFT_PATH)
    paras = list(doc.paragraphs)

    risk_start = risk_end = mitig_start = mitig_end = None
    for i, para in enumerate(paras):
        t = para.text.strip()
        if "Risk Description:" in t:
            risk_start = i
        elif risk_start is not None and "Remediation:" in t:
            risk_end = i
            break

    remed_start = remed_end = costs_start = costs_end = None
    for i, para in enumerate(paras):
        t = para.text.strip()
        if "Remediation:" in t and "Costs" not in t:
            remed_start = i
        elif remed_start is not None and "Remediation Costs and Schedule:" in t:
            remed_end = i
            costs_start = i
        elif costs_start is not None and "Mitigating Factors:" in t:
            costs_end = i
            break

    for i, para in enumerate(paras):
        t = para.text.strip()
        if "Mitigating Factors:" in t:
            mitig_start = i
        elif mitig_start is not None and ("AO Validation:" in t or "AODR Validation:" in t or "Co-AO Validation:" in t):
            mitig_end = i
            break

    # Update Subject (para 13) to include SAELAR/SOPRA
    for para in paras:
        if "Subject:" in para.text and "Tenable" in para.text:
            para.clear()
            para.add_run(
                "Subject: Request for Acceptance of Risk (AOR) for Potential Data Leakage/Spillage "
                "During Sandbox Testing of SAELAR and SOPRA"
            )
            break

    # Replace Risk Description content
    if risk_start is not None and risk_end is not None:
        for idx in range(risk_start + 1, risk_end):
            paras[idx].clear()
        for idx, risk_text in enumerate(RISKS):
            j = risk_start + 1 + idx
            if j < risk_end:
                paras[j].clear()
                paras[j].add_run(f"{idx + 1}. {risk_text}")

    # Replace Remediation content
    if remed_start is not None and remed_end is not None:
        for idx in range(remed_start + 1, remed_end):
            paras[idx].clear()
        for idx, remed_text in enumerate(REMEDIATION):
            j = remed_start + 1 + idx
            if j < remed_end:
                paras[j].clear()
                paras[j].add_run(f"{idx + 1}. {remed_text}")

    # Replace Remediation Costs and Schedule content
    if costs_start is not None and costs_end is not None:
        for idx in range(costs_start + 1, costs_end):
            paras[idx].clear()
        if costs_start + 1 < costs_end:
            paras[costs_start + 1].clear()
            paras[costs_start + 1].add_run(REMEDIATION_COSTS)

    # Replace Mitigating Factors content
    if mitig_start is not None and mitig_end is not None:
        for idx in range(mitig_start + 1, mitig_end):
            paras[idx].clear()
        for idx, (title, desc) in enumerate(MITIGATIONS):
            j = mitig_start + 1 + idx
            if j < mitig_end:
                paras[j].clear()
                p = paras[j]
                r1 = p.add_run(f"{idx + 1}. {title}\n")
                r1.bold = True
                r2 = p.add_run(desc)

    doc.save(OUT_PATH)
    print(f"Updated: {OUT_PATH}")


if __name__ == "__main__":
    main()
