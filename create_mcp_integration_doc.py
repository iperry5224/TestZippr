"""
Generate SAELAR MCP Integration Options Document as Word Document
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

def create_document():
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Title
    title = doc.add_heading('SAELAR MCP Integration Options', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph('Leveraging Model Context Protocol for Enhanced Security Operations')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.italic = True
    
    doc.add_paragraph()
    
    # Document info table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('Document Version:', '1.0'),
        ('Date:', 'January 21, 2026'),
        ('Prepared For:', 'SAELAR Development Team'),
        ('Classification:', 'Internal Use'),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
        info_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document outlines strategic options for integrating Model Context Protocol (MCP) into the SAELAR platform. '
        'MCP is Anthropic\'s open protocol that enables AI models to securely connect to external data sources and tools, '
        'providing a standardized way for AI assistants to access real-time information and execute actions.'
    )
    doc.add_paragraph(
        'By implementing MCP servers, SAELAR can expose its security assessment capabilities to any MCP-compatible AI client, '
        'enabling conversational security analysis, automated remediation workflows, and intelligent compliance documentation.'
    )
    
    # What is MCP Section
    doc.add_heading('What is Model Context Protocol (MCP)?', level=1)
    doc.add_paragraph(
        'MCP is an open protocol developed by Anthropic that standardizes how AI applications connect to external data sources '
        'and tools. It follows a client-server architecture where:'
    )
    
    mcp_components = [
        'MCP Servers expose resources (data), tools (functions), and prompts (templates) to AI clients',
        'MCP Clients (Claude Desktop, Cursor, custom apps) connect to servers and use exposed capabilities',
        'Resources provide read access to data (files, databases, APIs, live system state)',
        'Tools allow AI to execute actions (run commands, call APIs, modify systems)',
        'Prompts offer reusable templates for common workflows',
    ]
    for comp in mcp_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_paragraph()
    benefits_para = doc.add_paragraph()
    benefits_para.add_run('Key Benefits: ').bold = True
    benefits_para.add_run(
        'Standardized integration, security controls, real-time data access, '
        'tool execution with approval workflows, and multi-server orchestration.'
    )
    
    # MCP Integration Options
    doc.add_heading('MCP Integration Options for SAELAR', level=1)
    
    # Option 1
    option1 = doc.add_heading('Option 1: SAELAR Data MCP Server', level=2)
    option1.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph(
        'Create an MCP server that exposes SAELAR\'s assessment data, findings, and compliance status as resources '
        'that any MCP-compatible AI can query.'
    )
    
    doc.add_paragraph('Exposed Resources:', style='Normal').runs[0].font.bold = True
    resources_1 = [
        '/assessments/{id} - Full assessment results with all control statuses',
        '/findings/critical - List of critical findings requiring immediate attention',
        '/findings/by-family/{family} - Findings grouped by NIST control family',
        '/controls/{control_id} - Status and details for specific controls',
        '/poam/active - Active Plan of Action & Milestones items',
        '/risk/summary - Overall risk posture and scores',
        '/compliance/score - Current compliance percentage by family',
    ]
    for res in resources_1:
        doc.add_paragraph(res, style='List Bullet')
    
    doc.add_paragraph('Exposed Tools:', style='Normal').runs[0].font.bold = True
    tools_1 = [
        'run_assessment(families, region) - Trigger a new NIST assessment',
        'get_finding_details(finding_id) - Get detailed information about a finding',
        'calculate_risk(finding_id) - Calculate risk score for a finding',
        'export_results(format) - Export assessment results to JSON/CSV/PDF',
    ]
    for tool in tools_1:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph('Use Cases:', style='Normal').runs[0].font.bold = True
    usecases_1 = [
        'Natural language queries: "What are our top 5 security risks right now?"',
        'Cross-referencing: "Which findings affect our production environment?"',
        'Trend analysis: "How has our compliance score changed over the last month?"',
        'Report generation: "Create an executive summary of our security posture"',
    ]
    for uc in usecases_1:
        doc.add_paragraph(uc, style='List Bullet')
    
    # Effort/Impact table for Option 1
    doc.add_paragraph()
    effort_table1 = doc.add_table(rows=2, cols=3)
    effort_table1.style = 'Table Grid'
    effort_data1 = [
        ('Effort', 'Impact', 'Priority'),
        ('Medium (2-3 weeks)', 'High', 'Recommended'),
    ]
    for i, row_data in enumerate(effort_data1):
        for j, cell_text in enumerate(row_data):
            effort_table1.rows[i].cells[j].text = cell_text
            if i == 0:
                effort_table1.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Option 2
    option2 = doc.add_heading('Option 2: AWS Security Tools MCP Server', level=2)
    option2.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph(
        'Build an MCP server that wraps AWS security services, allowing AI to directly investigate '
        'and remediate security issues in your AWS environment.'
    )
    
    doc.add_paragraph('Exposed Tools:', style='Normal').runs[0].font.bold = True
    tools_2 = [
        'get_iam_users_without_mfa() - Find all IAM users without MFA enabled',
        'get_public_s3_buckets() - List S3 buckets with public access',
        'get_overly_permissive_security_groups() - Find security groups allowing 0.0.0.0/0',
        'get_unencrypted_volumes() - List EBS volumes without encryption',
        'get_security_hub_findings(severity) - Fetch Security Hub findings by severity',
        'apply_remediation(finding_id, action) - Execute approved remediation',
        'verify_remediation(finding_id) - Confirm fix was applied successfully',
        'rollback_remediation(finding_id) - Revert a remediation if needed',
    ]
    for tool in tools_2:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph('Use Cases:', style='Normal').runs[0].font.bold = True
    usecases_2 = [
        'Interactive investigation: "Show me all resources exposed to the internet"',
        'Guided remediation: "Fix the public S3 bucket issue and verify it\'s resolved"',
        'What-if analysis: "What would happen if we restrict this security group?"',
        'Compliance verification: "Verify all production databases are encrypted"',
    ]
    for uc in usecases_2:
        doc.add_paragraph(uc, style='List Bullet')
    
    doc.add_paragraph('Security Considerations:', style='Normal').runs[0].font.bold = True
    security_2 = [
        'Implement approval workflows for destructive operations',
        'Use IAM roles with least-privilege permissions',
        'Log all tool executions for audit trail',
        'Support dry-run mode for remediation testing',
    ]
    for sec in security_2:
        doc.add_paragraph(sec, style='List Bullet')
    
    # Effort/Impact table for Option 2
    doc.add_paragraph()
    effort_table2 = doc.add_table(rows=2, cols=3)
    effort_table2.style = 'Table Grid'
    effort_data2 = [
        ('Effort', 'Impact', 'Priority'),
        ('Medium-High (3-4 weeks)', 'Very High', 'Highly Recommended'),
    ]
    for i, row_data in enumerate(effort_data2):
        for j, cell_text in enumerate(row_data):
            effort_table2.rows[i].cells[j].text = cell_text
            if i == 0:
                effort_table2.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Option 3
    option3 = doc.add_heading('Option 3: Compliance Documentation MCP Server', level=2)
    option3.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph(
        'Expose NIST 800-53 controls, FedRAMP requirements, and compliance documentation templates '
        'as MCP resources for AI-powered document generation.'
    )
    
    doc.add_paragraph('Exposed Resources:', style='Normal').runs[0].font.bold = True
    resources_3 = [
        '/nist/800-53/rev5/controls - Complete NIST 800-53 Rev 5 control catalog',
        '/nist/800-53/rev5/{control_id} - Detailed control requirements and guidance',
        '/nist/800-53/rev5/baselines/{level} - Low/Moderate/High baseline controls',
        '/fedramp/requirements - FedRAMP-specific requirements',
        '/cis/aws/benchmarks - CIS AWS Foundations Benchmark mappings',
        '/saelar/templates/ssp - System Security Plan templates',
        '/saelar/templates/poam - POA&M document templates',
        '/saelar/templates/sar - Security Assessment Report templates',
    ]
    for res in resources_3:
        doc.add_paragraph(res, style='List Bullet')
    
    doc.add_paragraph('Exposed Tools:', style='Normal').runs[0].font.bold = True
    tools_3 = [
        'map_finding_to_controls(finding) - Map a security finding to relevant NIST controls',
        'generate_implementation_statement(control_id, system_info) - Create control implementation narrative',
        'check_control_inheritance(control_id) - Determine if control is inherited from provider',
        'generate_ssp_section(section, data) - Generate SSP section content',
        'generate_poam_entry(finding) - Create POA&M entry from finding',
        'validate_documentation(doc_type) - Check documentation completeness',
    ]
    for tool in tools_3:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph('Use Cases:', style='Normal').runs[0].font.bold = True
    usecases_3 = [
        'SSP generation: "Generate the Access Control section of our SSP based on current AWS config"',
        'Control mapping: "Which NIST controls does this MFA finding relate to?"',
        'Gap analysis: "What controls are we missing for FedRAMP Moderate?"',
        'Documentation updates: "Update our POA&M with the latest assessment findings"',
    ]
    for uc in usecases_3:
        doc.add_paragraph(uc, style='List Bullet')
    
    # Effort/Impact table for Option 3
    doc.add_paragraph()
    effort_table3 = doc.add_table(rows=2, cols=3)
    effort_table3.style = 'Table Grid'
    effort_data3 = [
        ('Effort', 'Impact', 'Priority'),
        ('Low-Medium (1-2 weeks)', 'High', 'Recommended'),
    ]
    for i, row_data in enumerate(effort_data3):
        for j, cell_text in enumerate(row_data):
            effort_table3.rows[i].cells[j].text = cell_text
            if i == 0:
                effort_table3.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Option 4
    option4 = doc.add_heading('Option 4: Threat Intelligence MCP Server', level=2)
    option4.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph(
        'Connect to threat intelligence feeds and vulnerability databases, enabling AI to correlate '
        'findings with real-world threats and prioritize based on active exploitation.'
    )
    
    doc.add_paragraph('Exposed Resources:', style='Normal').runs[0].font.bold = True
    resources_4 = [
        '/cisa/kev - CISA Known Exploited Vulnerabilities catalog',
        '/cisa/kev/{cve_id} - Specific KEV entry details',
        '/mitre/attack/techniques - MITRE ATT&CK technique catalog',
        '/mitre/attack/techniques/{id} - Specific technique details',
        '/nvd/cves/recent - Recently published CVEs',
        '/nvd/cves/{cve_id} - Specific CVE details and CVSS scores',
    ]
    for res in resources_4:
        doc.add_paragraph(res, style='List Bullet')
    
    doc.add_paragraph('Exposed Tools:', style='Normal').runs[0].font.bold = True
    tools_4 = [
        'check_kev_status(cve_id) - Check if CVE is in CISA KEV (actively exploited)',
        'get_attack_technique(technique_id) - Get ATT&CK technique details and mitigations',
        'correlate_finding_to_threats(finding) - Map finding to threat actors and campaigns',
        'get_exploitation_likelihood(cve_id) - Estimate likelihood of exploitation',
        'get_threat_context(finding) - Get threat intelligence context for a finding',
    ]
    for tool in tools_4:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph('Use Cases:', style='Normal').runs[0].font.bold = True
    usecases_4 = [
        'Threat-based prioritization: "Which of our findings are being actively exploited in the wild?"',
        'Attack path analysis: "What attack techniques could exploit this misconfiguration?"',
        'Threat actor correlation: "Are any known threat groups targeting vulnerabilities we have?"',
        'Risk adjustment: "Recalculate risk scores based on current threat intelligence"',
    ]
    for uc in usecases_4:
        doc.add_paragraph(uc, style='List Bullet')
    
    # Effort/Impact table for Option 4
    doc.add_paragraph()
    effort_table4 = doc.add_table(rows=2, cols=3)
    effort_table4.style = 'Table Grid'
    effort_data4 = [
        ('Effort', 'Impact', 'Priority'),
        ('Medium (2-3 weeks)', 'High', 'Recommended'),
    ]
    for i, row_data in enumerate(effort_data4):
        for j, cell_text in enumerate(row_data):
            effort_table4.rows[i].cells[j].text = cell_text
            if i == 0:
                effort_table4.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Option 5
    option5 = doc.add_heading('Option 5: Ticketing & Workflow MCP Server', level=2)
    option5.runs[0].font.color.rgb = RGBColor(0, 112, 192)
    
    doc.add_paragraph(
        'Integrate with ticketing systems (Jira, ServiceNow, Azure DevOps) to enable AI-driven '
        'workflow automation for security findings.'
    )
    
    doc.add_paragraph('Exposed Tools:', style='Normal').runs[0].font.bold = True
    tools_5 = [
        'create_ticket(finding, project, assignee) - Create ticket from security finding',
        'update_ticket(ticket_id, status, comment) - Update ticket status',
        'link_finding_to_ticket(finding_id, ticket_id) - Associate finding with ticket',
        'get_ticket_status(ticket_id) - Check remediation ticket status',
        'get_overdue_tickets() - List tickets past their due date',
        'escalate_ticket(ticket_id, reason) - Escalate overdue or critical tickets',
    ]
    for tool in tools_5:
        doc.add_paragraph(tool, style='List Bullet')
    
    doc.add_paragraph('Use Cases:', style='Normal').runs[0].font.bold = True
    usecases_5 = [
        'Auto-ticketing: "Create Jira tickets for all critical findings and assign to the security team"',
        'Status tracking: "What\'s the status of remediation for our MFA finding?"',
        'Escalation: "Escalate all overdue critical findings to the security manager"',
        'Reporting: "Generate a report of all open security tickets by team"',
    ]
    for uc in usecases_5:
        doc.add_paragraph(uc, style='List Bullet')
    
    # Effort/Impact table for Option 5
    doc.add_paragraph()
    effort_table5 = doc.add_table(rows=2, cols=3)
    effort_table5.style = 'Table Grid'
    effort_data5 = [
        ('Effort', 'Impact', 'Priority'),
        ('Low (1-2 weeks)', 'Medium-High', 'Optional'),
    ]
    for i, row_data in enumerate(effort_data5):
        for j, cell_text in enumerate(row_data):
            effort_table5.rows[i].cells[j].text = cell_text
            if i == 0:
                effort_table5.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Option 6
    option6 = doc.add_heading('Option 6: SAELAR Autonomous Security Agent', level=2)
    option6.runs[0].font.color.rgb = RGBColor(192, 0, 0)
    
    doc.add_paragraph(
        'Build an autonomous security agent that orchestrates multiple MCP servers to perform '
        'continuous security monitoring, automated remediation, and intelligent alerting.'
    )
    
    doc.add_paragraph('Agent Capabilities:', style='Normal').runs[0].font.bold = True
    capabilities_6 = [
        'Scheduled Assessments: Run daily/weekly security assessments automatically',
        'Continuous Monitoring: Watch for configuration changes that affect security posture',
        'Auto-Remediation: Automatically fix low-risk issues with pre-approved playbooks',
        'Intelligent Alerting: Correlate findings with threat intel before alerting',
        'Workflow Orchestration: Create tickets, assign owners, track remediation',
        'Executive Reporting: Generate daily/weekly security briefings automatically',
    ]
    for cap in capabilities_6:
        doc.add_paragraph(cap, style='List Bullet')
    
    doc.add_paragraph('Orchestration Flow:', style='Normal').runs[0].font.bold = True
    doc.add_paragraph(
        '1. Agent connects to multiple MCP servers (SAELAR, AWS, Threat Intel, Ticketing)\n'
        '2. Runs scheduled assessment via SAELAR MCP\n'
        '3. For each finding, queries Threat Intel MCP for context\n'
        '4. Auto-remediates low-risk items via AWS MCP (with logging)\n'
        '5. Creates tickets for high-risk items via Ticketing MCP\n'
        '6. Generates and distributes daily security report'
    )
    
    doc.add_paragraph('Security Guardrails:', style='Normal').runs[0].font.bold = True
    guardrails_6 = [
        'Human approval required for high-risk remediations',
        'All actions logged to immutable audit trail',
        'Configurable auto-remediation policies',
        'Emergency stop capability',
        'Anomaly detection on agent behavior',
    ]
    for guard in guardrails_6:
        doc.add_paragraph(guard, style='List Bullet')
    
    # Effort/Impact table for Option 6
    doc.add_paragraph()
    effort_table6 = doc.add_table(rows=2, cols=3)
    effort_table6.style = 'Table Grid'
    effort_data6 = [
        ('Effort', 'Impact', 'Priority'),
        ('High (6-8 weeks)', 'Very High', 'Future Phase'),
    ]
    for i, row_data in enumerate(effort_data6):
        for j, cell_text in enumerate(row_data):
            effort_table6.rows[i].cells[j].text = cell_text
            if i == 0:
                effort_table6.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Summary Comparison Table
    doc.add_heading('Summary Comparison', level=1)
    
    summary_table = doc.add_table(rows=7, cols=5)
    summary_table.style = 'Table Grid'
    summary_data = [
        ('Option', 'Description', 'Effort', 'Impact', 'Priority'),
        ('1. SAELAR Data Server', 'Expose assessment data as resources', 'Medium', 'High', '★★★'),
        ('2. AWS Security Tools', 'Direct AWS investigation & remediation', 'Medium-High', 'Very High', '★★★★'),
        ('3. Compliance Docs', 'NIST controls & doc generation', 'Low-Medium', 'High', '★★★'),
        ('4. Threat Intelligence', 'CISA KEV, MITRE ATT&CK integration', 'Medium', 'High', '★★★'),
        ('5. Ticketing Integration', 'Jira/ServiceNow workflow automation', 'Low', 'Medium-High', '★★'),
        ('6. Autonomous Agent', 'Full orchestration & auto-remediation', 'High', 'Very High', '★★★★★'),
    ]
    for i, row_data in enumerate(summary_data):
        for j, cell_text in enumerate(row_data):
            summary_table.rows[i].cells[j].text = cell_text
            if i == 0:
                summary_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True
    
    doc.add_paragraph()
    
    # Recommended Implementation Roadmap
    doc.add_heading('Recommended Implementation Roadmap', level=1)
    
    doc.add_heading('Phase 1: Foundation (Weeks 1-4)', level=2)
    phase1 = [
        'Build SAELAR Data MCP Server (Option 1) - Core data exposure',
        'Build Compliance Documentation MCP Server (Option 3) - Leverage existing NIST data',
        'Test integration with Claude Desktop and Cursor',
    ]
    for item in phase1:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Phase 2: AWS Integration (Weeks 5-8)', level=2)
    phase2 = [
        'Build AWS Security Tools MCP Server (Option 2) - Investigation & remediation',
        'Implement approval workflows for destructive operations',
        'Add audit logging for all tool executions',
    ]
    for item in phase2:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Phase 3: Intelligence (Weeks 9-12)', level=2)
    phase3 = [
        'Build Threat Intelligence MCP Server (Option 4) - CISA KEV, MITRE ATT&CK',
        'Integrate threat context into risk scoring',
        'Optional: Add Ticketing Integration (Option 5)',
    ]
    for item in phase3:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Phase 4: Automation (Weeks 13-20)', level=2)
    phase4 = [
        'Build SAELAR Autonomous Agent (Option 6)',
        'Implement scheduled assessments and auto-remediation',
        'Deploy continuous monitoring capabilities',
    ]
    for item in phase4:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_paragraph()
    
    # Technical Requirements
    doc.add_heading('Technical Requirements', level=1)
    
    doc.add_paragraph('Development Stack:', style='Normal').runs[0].font.bold = True
    tech_stack = [
        'Python 3.10+ with asyncio support',
        'MCP SDK (pip install mcp)',
        'boto3 for AWS integration',
        'FastAPI or similar for HTTP endpoints (optional)',
        'SQLite or PostgreSQL for state management',
    ]
    for tech in tech_stack:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_paragraph('Infrastructure:', style='Normal').runs[0].font.bold = True
    infra = [
        'MCP servers can run locally or containerized',
        'Recommend Docker deployment for production',
        'Support for stdio (local) and SSE (remote) transports',
        'AWS IAM roles with least-privilege permissions',
    ]
    for inf in infra:
        doc.add_paragraph(inf, style='List Bullet')
    
    doc.add_paragraph('Security Requirements:', style='Normal').runs[0].font.bold = True
    security = [
        'Authentication for MCP server connections',
        'Encrypted transport (TLS) for remote connections',
        'Audit logging for all tool executions',
        'Rate limiting to prevent abuse',
        'Approval workflows for sensitive operations',
    ]
    for sec in security:
        doc.add_paragraph(sec, style='List Bullet')
    
    doc.add_paragraph()
    
    # Next Steps
    doc.add_heading('Next Steps', level=1)
    next_steps = [
        'Review this document and select priority options',
        'Identify team members for MCP development',
        'Set up MCP SDK development environment',
        'Define security requirements and approval workflows',
        'Begin Phase 1 implementation',
    ]
    for i, step in enumerate(next_steps, 1):
        doc.add_paragraph(f'{i}. {step}')
    
    doc.add_paragraph()
    
    # Conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'Implementing MCP integration for SAELAR represents a significant opportunity to transform security operations '
        'through AI-powered automation. By exposing SAELAR\'s capabilities through standardized MCP servers, we enable '
        'conversational security analysis, automated remediation, and intelligent compliance workflows.'
    )
    doc.add_paragraph(
        'The recommended phased approach allows for incremental value delivery while building toward '
        'a fully autonomous security operations capability. Starting with data exposure (Option 1) and '
        'compliance documentation (Option 3) provides immediate value with lower risk, while the AWS '
        'Security Tools (Option 2) and Autonomous Agent (Option 6) deliver transformative capabilities.'
    )
    
    doc.add_paragraph()
    
    # Footer
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('Document prepared for SAELAR MCP Integration Initiative').italic = True
    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer2.add_run('Generated: January 21, 2026').italic = True
    
    # Save
    output_path = r'C:\Users\iperr\OneDrive\Desktop\AI-Guides\SAELAR_MCP_Integration_Options.docx'
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path

if __name__ == '__main__':
    create_document()
