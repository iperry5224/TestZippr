"""
WORDY - Word Document Generator for SAELAR
==========================================
A utility script to generate Word (.docx) documents from Python.

Usage:
    python wordy.py                           # Generate comparison doc
    python wordy.py --convert-folder <path>   # Convert all .txt/.md files in folder

This script can be modified to generate various Word documents.
Currently configured to generate the SAELAR vs AWS Audit Manager comparison.

Requirements:
    pip install python-docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import re
import sys
import glob
import tempfile
from datetime import datetime


def _default_output_path(prefix: str, safe_name: str = "") -> str:
    """Return a cross-platform writable path for generated documents."""
    output_dir = os.environ.get("SAELAR_OUTPUT_DIR")
    if not output_dir:
        output_dir = os.path.join(tempfile.gettempdir(), "saelar_docs")
    os.makedirs(output_dir, exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d")
    filename = f"{prefix}_{safe_name}_{timestamp}.docx" if safe_name else f"{prefix}_{timestamp}.docx"
    return os.path.join(output_dir, filename)


def create_comparison_document(output_path: str = None):
    """Generate SAELAR vs AWS Audit Manager comparison as Word document."""
    
    if output_path is None:
        output_path = _default_output_path("SAELAR_vs_AWS_Audit_Manager_Comparison")
    
    doc = Document()

    # Title
    title = doc.add_heading('SAELAR vs. AWS Audit Manager', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading('Comparative Analysis', level=1)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Author: AI Analysis for Ira Perry')
    doc.add_paragraph()

    # Executive Summary
    doc.add_heading('Executive Summary', level=2)
    
    doc.add_paragraph(
        'SAELAR (Security Architecture and Evaluation Linear Assessment Reporting Tool) is a '
        'purpose-built, self-hosted security assessment platform that delivers comprehensive '
        'NIST 800-53 Rev 5 compliance evaluation at zero licensing cost. Unlike AWS Audit '
        'Manager—which charges approximately $1.25 per resource per month and spreads its '
        'capabilities across multiple compliance frameworks—SAELAR focuses exclusively on deep '
        'NIST 800-53 Rev 5 coverage, assessing all 13 cloud-relevant control families with '
        'real-time AWS infrastructure scanning, integrated quantitative risk scoring, and '
        'ISSO-ready remediation guidance. For organizations where NIST 800-53 compliance is the '
        'primary concern, SAELAR provides superior depth, actionable intelligence, and significant '
        'cost savings while maintaining full deployment flexibility through Docker containerization—'
        'including support for air-gapped environments where cloud-dependent services cannot operate.'
    )
    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Aspect'
    hdr[1].text = 'SAELAR'
    hdr[2].text = 'AWS Audit Manager'

    data = [
        ('Type', 'Self-hosted, open-source tool', 'AWS Managed Service'),
        ('Cost', 'Free (self-hosted)', 'Pay-per-use (~$1.25/resource/month)'),
        ('Focus', 'NIST 800-53 Rev 5 specialist', 'Multi-framework compliance'),
        ('Deployment', 'Container/local', 'AWS native'),
    ]
    for i, (aspect, saelar, aws) in enumerate(data, 1):
        table.rows[i].cells[0].text = aspect
        table.rows[i].cells[1].text = saelar
        table.rows[i].cells[2].text = aws

    doc.add_paragraph()

    # Compliance Frameworks
    doc.add_heading('1. Compliance Frameworks', level=2)

    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Feature'
    hdr[1].text = 'SAELAR'
    hdr[2].text = 'AWS Audit Manager'

    data = [
        ('NIST 800-53 Rev 5', '✓ Deep, specialized (13 families)', '✓ Supported (prebuilt)'),
        ('PCI DSS', '✗', '✓'),
        ('SOC 2', '✗', '✓'),
        ('HIPAA', '✗', '✓'),
        ('GDPR', '✗', '✓'),
        ('Custom Frameworks', '✗', '✓'),
    ]
    for i, (feature, saelar, aws) in enumerate(data, 1):
        table.rows[i].cells[0].text = feature
        table.rows[i].cells[1].text = saelar
        table.rows[i].cells[2].text = aws

    doc.add_paragraph()
    doc.add_paragraph('Winner: AWS Audit Manager for breadth; SAELAR for NIST 800-53 depth')

    # Control Families
    doc.add_heading('2. NIST 800-53 Control Families (SAELAR)', level=2)

    table = doc.add_table(rows=14, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Code'
    hdr[1].text = 'Control Family'

    families = [
        ('AC', 'Access Control'),
        ('AU', 'Audit & Accountability'),
        ('CA', 'Assessment, Authorization & Monitoring'),
        ('CM', 'Configuration Management'),
        ('CP', 'Contingency Planning'),
        ('IA', 'Identification & Authentication'),
        ('IR', 'Incident Response'),
        ('MP', 'Media Protection'),
        ('RA', 'Risk Assessment'),
        ('SA', 'System & Services Acquisition'),
        ('SC', 'System & Communications Protection'),
        ('SI', 'System & Information Integrity'),
        ('SR', 'Supply Chain Risk Management'),
    ]
    for i, (code, family) in enumerate(families, 1):
        table.rows[i].cells[0].text = code
        table.rows[i].cells[1].text = family

    doc.add_paragraph()

    # Evidence Collection
    doc.add_heading('3. Evidence Collection', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Capability'
    hdr[1].text = 'SAELAR'
    hdr[2].text = 'AWS Audit Manager'

    data = [
        ('Automation', '✓ Real-time API scanning', '✓ Continuous collection'),
        ('AWS Services Checked', '15+ services', 'CloudTrail, Config, Security Hub'),
        ('Manual Evidence Upload', '✗', '✓'),
        ('Evidence Retention', 'S3 export', 'Managed by AWS'),
    ]
    for i, (cap, saelar, aws) in enumerate(data, 1):
        table.rows[i].cells[0].text = cap
        table.rows[i].cells[1].text = saelar
        table.rows[i].cells[2].text = aws

    doc.add_paragraph()

    # Reporting
    doc.add_heading('4. Reporting & Output', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Feature'
    hdr[1].text = 'SAELAR'
    hdr[2].text = 'AWS Audit Manager'

    data = [
        ('Real-time Dashboard', '✓ Streamlit UI', '✓ AWS Console'),
        ('Risk Scoring', '✓ Integrated Risk Calculator', '✗'),
        ('Export Formats', 'JSON, Markdown, S3', 'Assessment reports'),
        ('ISSO-Friendly Reports', '✓ Designed for ISSOs', 'General compliance'),
        ('Remediation Guidance', '✓ Per-finding recommendations', 'Limited'),
    ]
    for i, (feature, saelar, aws) in enumerate(data, 1):
        table.rows[i].cells[0].text = feature
        table.rows[i].cells[1].text = saelar
        table.rows[i].cells[2].text = aws

    doc.add_paragraph()
    doc.add_paragraph('Winner: SAELAR for actionable, ISSO-focused output')

    # Architecture
    doc.add_heading('5. Architecture & Deployment', level=2)

    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Aspect'
    hdr[1].text = 'SAELAR'
    hdr[2].text = 'AWS Audit Manager'

    data = [
        ('Hosting', 'Self-hosted (Docker, local)', 'AWS Managed'),
        ('Multi-Account', 'Single account focus', '✓ AWS Organizations'),
        ('Multi-Cloud', '✗ AWS only', '✗ AWS only'),
        ('Internet Exposure', 'Optional (ngrok)', 'AWS Console only'),
        ('Air-gapped Support', '✓ Possible', '✗ Requires AWS connectivity'),
    ]
    for i, (aspect, saelar, aws) in enumerate(data, 1):
        table.rows[i].cells[0].text = aspect
        table.rows[i].cells[1].text = saelar
        table.rows[i].cells[2].text = aws

    doc.add_paragraph()

    # Cost Analysis
    doc.add_heading('6. Cost Analysis', level=2)

    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Scenario'
    hdr[1].text = 'SAELAR'
    hdr[2].text = 'AWS Audit Manager'

    data = [
        ('100 AWS resources', '$0', '~$125/month'),
        ('500 AWS resources', '$0', '~$625/month'),
        ('1,000 AWS resources', '$0', '~$1,250/month'),
        ('Infrastructure cost', 'EC2/container (~$20-50/mo)', 'Included'),
    ]
    for i, (scenario, saelar, aws) in enumerate(data, 1):
        table.rows[i].cells[0].text = scenario
        table.rows[i].cells[1].text = saelar
        table.rows[i].cells[2].text = aws

    doc.add_paragraph()
    doc.add_paragraph('Winner: SAELAR for cost-conscious organizations')

    # Strengths
    doc.add_heading('SAELAR Strengths', level=2)
    strengths = [
        'Zero licensing cost - completely free',
        'NIST 800-53 specialist - deeper, more specific checks',
        'Integrated Risk Calculator - quantitative risk scoring',
        'ISSO-friendly - designed for security officers',
        'Self-hosted - full control, air-gap capable',
        'Real-time transparency - immediate results',
        'Remediation guidance - actionable recommendations per finding',
        'Portable - Docker containerized',
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('AWS Audit Manager Strengths', level=2)
    strengths = [
        'Multi-framework - one tool for many compliance needs',
        'Multi-account - enterprise-scale via Organizations',
        'Continuous monitoring - always-on evidence collection',
        'Manual evidence - upload supporting documents',
        'Delegation workflows - team collaboration',
        'AWS managed - no infrastructure to maintain',
    ]
    for s in strengths:
        doc.add_paragraph(s, style='List Bullet')

    # Recommendations
    doc.add_heading('Recommendation Matrix', level=2)

    table = doc.add_table(rows=9, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Use Case'
    hdr[1].text = 'Recommended Tool'

    data = [
        ('NIST 800-53 Rev 5 focused assessment', 'SAELAR'),
        ('Multiple compliance frameworks', 'AWS Audit Manager'),
        ('Cost-sensitive organization', 'SAELAR'),
        ('Large enterprise (1000s of accounts)', 'AWS Audit Manager'),
        ('ISSO preparing for audits', 'SAELAR'),
        ('Risk quantification needed', 'SAELAR'),
        ('Air-gapped environments', 'SAELAR'),
        ('Quick deployment, no infra management', 'AWS Audit Manager'),
    ]
    for i, (use_case, tool) in enumerate(data, 1):
        table.rows[i].cells[0].text = use_case
        table.rows[i].cells[1].text = tool

    doc.add_paragraph()

    # Conclusion
    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'SAELAR excels as a specialized, cost-effective tool for organizations focused on '
        'NIST 800-53 Rev 5 compliance with AWS infrastructure. Its integrated risk calculator, '
        'ISSO-friendly reporting, and zero licensing cost make it ideal for security teams '
        'needing actionable, real-time assessments.'
    )
    doc.add_paragraph(
        'AWS Audit Manager is better suited for enterprises requiring multi-framework compliance '
        '(PCI DSS, SOC 2, HIPAA, GDPR) across large multi-account environments with team '
        'collaboration features.'
    )
    doc.add_paragraph(
        'For organizations primarily concerned with NIST 800-53 Rev 5 compliance and seeking '
        'cost efficiency, SAELAR is the recommended choice.'
    )

    doc.add_paragraph()
    doc.add_paragraph('Analysis generated for Security Architecture and Evaluation (SAE) purposes.', style='Intense Quote')

    # Save
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path


def create_simple_document(title: str, content: str, output_path: str):
    """
    Create a simple Word document with a title and content.
    
    Args:
        title: Document title
        content: Body text content
        output_path: Full path to save the .docx file
    """
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()
    
    for paragraph in content.split('\n\n'):
        doc.add_paragraph(paragraph)
    
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path


def create_table_document(title: str, headers: list, rows: list, output_path: str):
    """
    Create a Word document with a table.
    
    Args:
        title: Document title
        headers: List of column headers
        rows: List of row data (each row is a list)
        output_path: Full path to save the .docx file
    """
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()
    
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    
    # Add headers
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    
    # Add data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = str(cell_data)
    
    doc.save(output_path)
    print(f'Document saved to: {output_path}')
    return output_path


def create_report_from_content(title: str, content: str, output_path: str):
    """
    Create a Word document from markdown-style content (e.g. from Chad/SOPRA AI).
    Handles # headings, **bold**, lists, etc. Compatible with Google Docs when uploaded.

    Args:
        title: Document title
        content: Markdown-style body content
        output_path: Full path to save the .docx file

    Returns:
        Path to the created document
    """
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph()

    in_code_block = False
    current_paragraph = []

    for line in content.split('\n'):
        if line.strip().startswith('```'):
            if in_code_block:
                if current_paragraph:
                    code_text = '\n'.join(current_paragraph)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    current_paragraph = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            current_paragraph.append(line)
            continue

        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            doc.add_paragraph(line.strip()[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\. ', line.strip()):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line.strip()).strip(), style='List Number')
        elif line.strip().startswith('---'):
            doc.add_paragraph('─' * 50)
        elif line.strip() == '':
            continue
        else:
            # Preserve **bold** as bold runs
            parts = re.split(r'(\*\*[^*]+\*\*)', line)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    clean = re.sub(r'\*(.+?)\*', r'\1', part)
                    clean = re.sub(r'`(.+?)`', r'\1', clean)
                    p.add_run(clean)

    doc.save(output_path)
    return output_path


def convert_text_to_docx(input_path: str, output_path: str = None):
    """
    Convert a text or markdown file to Word document.
    
    Args:
        input_path: Path to .txt or .md file
        output_path: Path for output .docx (auto-generated if None)
    
    Returns:
        Path to the created document
    """
    if output_path is None:
        base = os.path.splitext(input_path)[0]
        output_path = base + '.docx'
    
    # Read the source file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Extract title from first line or filename
    lines = content.split('\n')
    first_line = lines[0].strip() if lines else ''
    
    # Check if first line is a markdown heading
    if first_line.startswith('# '):
        title = first_line[2:].strip()
        content = '\n'.join(lines[1:])
    elif first_line.startswith('==') or (len(lines) > 1 and lines[1].startswith('==')):
        title = first_line
        content = '\n'.join(lines[2:]) if len(lines) > 1 and lines[1].startswith('==') else '\n'.join(lines[1:])
    else:
        # Use filename as title
        title = os.path.splitext(os.path.basename(input_path))[0].replace('_', ' ')
    
    # Add title
    doc.add_heading(title, 0)
    
    # Process content
    in_code_block = False
    current_paragraph = []
    
    for line in content.split('\n'):
        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                if current_paragraph:
                    code_text = '\n'.join(current_paragraph)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    current_paragraph = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            continue
        
        if in_code_block:
            current_paragraph.append(line)
            continue
        
        # Handle markdown headings
        if line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line).strip(), style='List Number')
        elif line.strip().startswith('---'):
            doc.add_paragraph('─' * 50)
        elif line.strip() == '':
            continue
        else:
            # Regular paragraph - clean up markdown formatting
            clean_line = line
            clean_line = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_line)  # Bold
            clean_line = re.sub(r'\*(.+?)\*', r'\1', clean_line)      # Italic
            clean_line = re.sub(r'`(.+?)`', r'\1', clean_line)        # Code
            doc.add_paragraph(clean_line)
    
    doc.save(output_path)
    print(f'[OK] Converted: {os.path.basename(input_path)} -> {os.path.basename(output_path)}')
    return output_path


def convert_folder(folder_path: str, extensions: list = None):
    """
    Convert all text/markdown files in a folder to Word documents.
    
    Args:
        folder_path: Path to folder containing files to convert
        extensions: List of extensions to convert (default: ['.txt', '.md'])
    
    Returns:
        List of paths to created documents
    """
    if extensions is None:
        extensions = ['.txt', '.md']
    
    converted = []
    skipped = []
    
    print(f'\nConverting files in: {folder_path}')
    print('=' * 60)
    
    for ext in extensions:
        pattern = os.path.join(folder_path, f'*{ext}')
        for file_path in glob.glob(pattern):
            # Skip temp files
            if os.path.basename(file_path).startswith('~$'):
                continue
            
            # Check if .docx already exists
            docx_path = os.path.splitext(file_path)[0] + '.docx'
            if os.path.exists(docx_path):
                skipped.append(os.path.basename(file_path))
                print(f'[SKIP] Already exists: {os.path.basename(file_path)}')
                continue
            
            try:
                result = convert_text_to_docx(file_path, docx_path)
                converted.append(result)
            except Exception as e:
                print(f'[ERROR] Converting {os.path.basename(file_path)}: {e}')
    
    print('=' * 60)
    print(f'Converted: {len(converted)} files')
    if skipped:
        print(f'Skipped (already had .docx): {len(skipped)} files')
    
    return converted


def create_ssp_document(ssp_data: dict, output_path: str = None):
    """
    Generate a System Security Plan (SSP) Word document.
    
    Args:
        ssp_data: Dictionary from SSPGenerator.to_dict()
        output_path: Path for output .docx file
    
    Returns:
        Path to created document
    """
    if output_path is None:
        system_name = ssp_data.get('system_info', {}).get('system_name', 'System')
        safe_name = system_name.replace(' ', '_').replace('/', '-')[:30]
        output_path = _default_output_path("SSP", safe_name)
    
    doc = Document()
    
    system_info = ssp_data.get('system_info', {})
    summary = ssp_data.get('summary', {})
    family_summary = ssp_data.get('family_summary', {})
    controls = ssp_data.get('controls', [])
    poam = ssp_data.get('poam', [])
    stats = ssp_data.get('statistics', {})
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('System Security Plan', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    subtitle = doc.add_paragraph(system_info.get('system_name', 'Information System'))
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph(f"Security Categorization: {system_info.get('categorization', 'Moderate')}")
    doc.add_paragraph(f"Document Version: {ssp_data.get('metadata', {}).get('version', '1.0')}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%B %d, %Y')}")
    doc.add_paragraph()
    
    # Document control table
    doc.add_heading('Document Control', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    doc_control = [
        ('System Owner', system_info.get('system_owner', 'TBD')),
        ('Authorizing Official', system_info.get('authorizing_official', 'TBD')),
        ('ISSO', system_info.get('isso', 'TBD')),
        ('System Acronym', system_info.get('system_acronym', 'N/A')),
        ('Classification', 'Unclassified'),
    ]
    for i, (label, value) in enumerate(doc_control):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS (placeholder)
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. System Identification',
        '3. Security Control Summary',
        '4. Control Implementation Details',
        '5. Plan of Action and Milestones (POA&M)',
        'Appendix A: Risk Assessment Summary'
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 1: EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('1. Executive Summary', level=1)
    
    doc.add_heading('1.1 Security Posture Overview', level=2)
    
    posture = summary.get('security_posture', 'Under Review')
    compliance_pct = summary.get('compliance_percentage', 0)
    
    doc.add_paragraph(
        f"The {system_info.get('system_name', 'system')} has been assessed against NIST 800-53 Rev 5 "
        f"security controls applicable to a {system_info.get('categorization', 'Moderate')} impact system. "
        f"The current security posture is rated as {posture} with {compliance_pct}% of applicable "
        f"controls fully implemented."
    )
    
    doc.add_paragraph(summary.get('posture_description', ''))
    
    # Summary statistics table
    doc.add_heading('1.2 Assessment Statistics', level=2)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    stat_rows = [
        ('Total Controls Assessed', str(stats.get('total_controls', 0))),
        ('Fully Implemented', str(stats.get('implemented', 0))),
        ('Partially Implemented', str(stats.get('partial', 0))),
        ('Planned', str(stats.get('planned', 0))),
        ('Not Implemented', str(stats.get('not_implemented', 0))),
        ('POA&M Items', str(len(poam))),
    ]
    for i, (label, value) in enumerate(stat_rows):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    # =========================================================================
    # SECTION 2: SYSTEM IDENTIFICATION
    # =========================================================================
    doc.add_heading('2. System Identification', level=1)
    
    doc.add_heading('2.1 System Name and Identifier', level=2)
    doc.add_paragraph(f"System Name: {system_info.get('system_name', 'TBD')}")
    doc.add_paragraph(f"System Acronym: {system_info.get('system_acronym', 'N/A')}")
    
    doc.add_heading('2.2 System Description', level=2)
    doc.add_paragraph(system_info.get('system_description', 'System description to be provided.'))
    
    doc.add_heading('2.3 Authorization Boundary', level=2)
    doc.add_paragraph(system_info.get('authorization_boundary', 'Authorization boundary to be defined.'))
    
    doc.add_heading('2.4 Security Categorization', level=2)
    doc.add_paragraph(
        f"In accordance with FIPS 199 and NIST SP 800-60, this system has been categorized as "
        f"{system_info.get('categorization', 'Moderate')} impact for confidentiality, integrity, and availability."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 3: SECURITY CONTROL SUMMARY
    # =========================================================================
    doc.add_heading('3. Security Control Summary', level=1)
    
    doc.add_paragraph(
        "The following table summarizes the implementation status of security controls by control family."
    )
    
    # Family summary table
    if family_summary:
        table = doc.add_table(rows=len(family_summary) + 1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        headers = ['Family', 'Total', 'Implemented', 'Partial', 'Compliance %']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # Data rows
        for row_idx, (family_code, data) in enumerate(family_summary.items(), 1):
            table.rows[row_idx].cells[0].text = f"{family_code} - {data.get('name', '')}"
            table.rows[row_idx].cells[1].text = str(data.get('total', 0))
            table.rows[row_idx].cells[2].text = str(data.get('implemented', 0))
            table.rows[row_idx].cells[3].text = str(data.get('partial', 0))
            table.rows[row_idx].cells[4].text = f"{data.get('compliance_pct', 0)}%"
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 4: CONTROL IMPLEMENTATION DETAILS
    # =========================================================================
    doc.add_heading('4. Control Implementation Details', level=1)
    
    doc.add_paragraph(
        "This section provides detailed implementation statements for each assessed security control."
    )
    
    # Group controls by family
    controls_by_family = {}
    for control in controls:
        family = control.get('family', 'XX')
        if family not in controls_by_family:
            controls_by_family[family] = []
        controls_by_family[family].append(control)
    
    for family_code in sorted(controls_by_family.keys()):
        family_name = family_summary.get(family_code, {}).get('name', family_code)
        doc.add_heading(f'4.{list(controls_by_family.keys()).index(family_code) + 1} {family_code} - {family_name}', level=2)
        
        for control in controls_by_family[family_code]:
            doc.add_heading(f"{control.get('control_id', '')} - {control.get('control_name', '')}", level=3)
            
            # Status
            doc.add_paragraph(f"Implementation Status: {control.get('status', 'Unknown')}")
            
            # Implementation statement
            doc.add_paragraph(control.get('implementation_statement', 'Implementation details to be provided.'))
            
            # Recommendations if any
            recommendations = control.get('recommendations', [])
            if recommendations and recommendations[0]:
                doc.add_paragraph('Recommendations:', style='Heading 4')
                for rec in recommendations[:3]:
                    if rec:
                        doc.add_paragraph(rec, style='List Bullet')
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 5: POA&M
    # =========================================================================
    doc.add_heading('5. Plan of Action and Milestones (POA&M)', level=1)
    
    if poam:
        doc.add_paragraph(
            f"The following {len(poam)} items require remediation action to achieve full compliance."
        )
        
        table = doc.add_table(rows=len(poam) + 1, cols=5)
        table.style = 'Table Grid'
        
        # Header
        headers = ['POA&M ID', 'Control', 'Risk Level', 'Weakness', 'Target Date']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        
        # Data rows
        for row_idx, item in enumerate(poam, 1):
            table.rows[row_idx].cells[0].text = item.get('poam_id', '')[:15]
            table.rows[row_idx].cells[1].text = item.get('control_id', '')
            table.rows[row_idx].cells[2].text = item.get('risk_level', '')
            weakness = item.get('weakness', '')[:50]
            table.rows[row_idx].cells[3].text = weakness + '...' if len(item.get('weakness', '')) > 50 else weakness
            target = item.get('scheduled_completion', '')
            if target:
                target = target[:10] if isinstance(target, str) else target
            table.rows[row_idx].cells[4].text = str(target) if target else 'TBD'
        
        doc.add_paragraph()
        
        # Detailed POA&M items
        doc.add_heading('5.1 POA&M Details', level=2)
        
        for item in poam:
            doc.add_heading(f"{item.get('poam_id', 'POAM-XXX')}", level=3)
            doc.add_paragraph(f"Control: {item.get('control_id', '')}")
            doc.add_paragraph(f"Risk Level: {item.get('risk_level', '')}")
            doc.add_paragraph(f"Weakness: {item.get('weakness', '')}")
            doc.add_paragraph(f"Remediation Plan: {item.get('remediation_plan', '')}")
            doc.add_paragraph(f"Status: {item.get('status', 'Open')}")
    else:
        doc.add_paragraph("No POA&M items identified. All assessed controls are fully implemented.")
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX A: RISK ASSESSMENT SUMMARY
    # =========================================================================
    doc.add_heading('Appendix A: Risk Assessment Summary', level=1)
    
    total_risk = stats.get('total_risk_score', 0)
    
    doc.add_paragraph(
        f"Total Risk Score: {total_risk:.1f}"
    )
    
    doc.add_paragraph(
        "Risk scores are calculated using the SAELAR Risk Calculator based on likelihood and impact "
        "assessments for each finding. Higher scores indicate greater risk requiring prioritized remediation."
    )
    
    # Risk distribution
    high_risk = [c for c in controls if c.get('risk_level') == 'High']
    if high_risk:
        doc.add_heading('High Risk Findings', level=2)
        for control in high_risk[:10]:
            doc.add_paragraph(
                f"- {control.get('control_id', '')}: {control.get('control_name', '')} "
                f"(Score: {control.get('risk_score', 0):.1f})",
                style='List Bullet'
            )
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph(
        f"Generated by SAELAR SSP Generator | {datetime.now().strftime('%B %d, %Y')}",
        style='Intense Quote'
    )
    
    # Save
    doc.save(output_path)
    print(f'[OK] SSP Document saved to: {output_path}')
    return output_path


def create_ssp_markdown(ssp_data: dict) -> str:
    """Generate a System Security Plan as a Markdown string (no file I/O)."""
    system_info = ssp_data.get('system_info', {})
    summary = ssp_data.get('summary', {})
    family_summary = ssp_data.get('family_summary', {})
    controls = ssp_data.get('controls', [])
    poam = ssp_data.get('poam', [])
    stats = ssp_data.get('statistics', {})

    lines = []
    lines.append("# System Security Plan")
    lines.append(f"\n**{system_info.get('system_name', 'Information System')}**")
    lines.append(f"\nSecurity Categorization: {system_info.get('categorization', 'Moderate')}")
    lines.append(f"Document Version: {ssp_data.get('metadata', {}).get('version', '1.0')}")
    lines.append(f"Generated: {datetime.now().strftime('%B %d, %Y')}")

    lines.append("\n## Document Control")
    lines.append("\n| Field | Value |")
    lines.append("|---|---|")
    lines.append(f"| System Owner | {system_info.get('system_owner', 'TBD')} |")
    lines.append(f"| ISSO | {system_info.get('isso_name', 'TBD')} |")
    lines.append(f"| Authorizing Official | {system_info.get('authorizing_official', 'TBD')} |")
    lines.append(f"| Categorization | {system_info.get('categorization', 'Moderate')} |")
    lines.append(f"| Status | {system_info.get('operational_status', 'Operational')} |")

    lines.append("\n## Executive Summary")
    lines.append(f"\n- **Total Controls Assessed:** {summary.get('total_controls', 0)}")
    lines.append(f"- **Implemented:** {summary.get('implemented', 0)}")
    lines.append(f"- **Partially Implemented:** {summary.get('partial', 0)}")
    lines.append(f"- **Not Implemented:** {summary.get('not_implemented', 0)}")
    lines.append(f"- **Compliance Percentage:** {summary.get('compliance_percentage', 0)}%")
    lines.append(f"- **Security Posture:** {summary.get('security_posture', 'Unknown')}")

    if system_info.get('system_description'):
        lines.append("\n## System Description")
        lines.append(f"\n{system_info['system_description']}")

    if system_info.get('authorization_boundary'):
        lines.append("\n## Authorization Boundary")
        lines.append(f"\n{system_info['authorization_boundary']}")

    if family_summary:
        lines.append("\n## Control Family Summary")
        lines.append("\n| Family | Implemented | Partial | Not Implemented | Total | Compliance |")
        lines.append("|---|---|---|---|---|---|")
        for family_code, data in family_summary.items():
            total = data.get('total', 0)
            impl = data.get('implemented', 0)
            pct = f"{(impl/total*100):.0f}%" if total > 0 else "N/A"
            lines.append(f"| {family_code} | {impl} | {data.get('partial', 0)} | {data.get('not_implemented', 0)} | {total} | {pct} |")

    if controls:
        lines.append("\n## Control Implementation Details")
        for ctrl in controls:
            status = ctrl.get('status', 'Unknown')
            lines.append(f"\n### {ctrl.get('control_id', '')} — {ctrl.get('control_name', '')}")
            lines.append(f"\n**Status:** {status}")
            if ctrl.get('findings'):
                lines.append("\n**Findings:**")
                for f in ctrl['findings']:
                    lines.append(f"- {f}")
            if ctrl.get('recommendations'):
                lines.append("\n**Recommendations:**")
                for r in ctrl['recommendations']:
                    lines.append(f"- {r}")

    if poam:
        lines.append("\n## Plan of Action & Milestones (POA&M)")
        lines.append("\n| # | Control | Weakness | Risk | Milestone |")
        lines.append("|---|---|---|---|---|")
        for i, item in enumerate(poam, 1):
            lines.append(f"| {i} | {item.get('control_id', '')} | {item.get('weakness', '')} | {item.get('risk_level', '')} | {item.get('milestone', '')} |")

    lines.append(f"\n---\n\n*Generated by SAELAR SSP Generator | {datetime.now().strftime('%B %d, %Y')}*")
    return "\n".join(lines)


def create_poam_markdown(ssp_data: dict) -> str:
    """Generate a POA&M document as a Markdown string (no file I/O)."""
    system_info = ssp_data.get('system_info', {})
    poam_items = ssp_data.get('poam', [])
    summary = ssp_data.get('summary', {})

    lines = []
    lines.append("# Plan of Action & Milestones (POA&M)")
    lines.append(f"\n**{system_info.get('system_name', 'Information System')}**")
    lines.append(f"\nGenerated: {datetime.now().strftime('%B %d, %Y')}")
    lines.append(f"Security Categorization: {system_info.get('categorization', 'Moderate')}")
    lines.append(f"Total POA&M Items: {len(poam_items)}")

    if poam_items:
        high = sum(1 for p in poam_items if p.get('risk_level') == 'High')
        med = sum(1 for p in poam_items if p.get('risk_level') == 'Medium')
        low = sum(1 for p in poam_items if p.get('risk_level') == 'Low')
        lines.append(f"\n## Summary")
        lines.append(f"\n- **High Risk:** {high}")
        lines.append(f"- **Medium Risk:** {med}")
        lines.append(f"- **Low Risk:** {low}")

        lines.append(f"\n## POA&M Items")
        lines.append("\n| # | Control ID | Weakness | Risk Level | Remediation | Milestone | Status |")
        lines.append("|---|---|---|---|---|---|---|")
        for i, item in enumerate(poam_items, 1):
            lines.append(f"| {i} | {item.get('control_id', '')} | {item.get('weakness', '')[:50]} | {item.get('risk_level', '')} | {item.get('remediation', '')[:50]} | {item.get('milestone', '')} | {item.get('status', 'Open')} |")
    else:
        lines.append("\n*No POA&M items — all controls implemented.*")

    lines.append(f"\n---\n\n*Generated by SAELAR POA&M Generator | {datetime.now().strftime('%B %d, %Y')}*")
    return "\n".join(lines)


def create_rar_markdown(rar_data: dict) -> str:
    """Generate a Risk Assessment Report as a Markdown string (no file I/O)."""
    system_info = rar_data.get('system_info', {})
    assessment_results = rar_data.get('assessment_results', {})
    vulnerabilities = rar_data.get('vulnerabilities', [])
    risk_summary = rar_data.get('risk_summary', {})
    recommendations = rar_data.get('recommendations', [])
    poam = rar_data.get('poam', [])
    stats = rar_data.get('statistics', {})

    system_name = system_info.get('system_name', 'System Name')
    system_acronym = system_info.get('system_acronym', 'NOAA50xx')
    categorization = system_info.get('categorization', 'Moderate')

    lines = []
    lines.append(f"# {system_name} ({system_acronym})")
    lines.append(f"\n# Risk Assessment Report")
    lines.append(f"\nGenerated: {datetime.now().strftime('%B %d, %Y')}")
    lines.append(f"Security Categorization: {categorization}")

    lines.append("\n## Document Control")
    lines.append("\n| Field | Value |")
    lines.append("|---|---|")
    lines.append(f"| System Owner | {system_info.get('system_owner', 'TBD')} |")
    lines.append(f"| ISSO | {system_info.get('isso_name', 'TBD')} |")
    lines.append(f"| Authorizing Official | {system_info.get('authorizing_official', 'TBD')} |")

    lines.append("\n## Executive Summary")
    lines.append(f"\n- **Overall Risk Level:** {risk_summary.get('overall_risk_level', 'Unknown')}")
    lines.append(f"- **Total Risk Score:** {risk_summary.get('total_risk_score', 0)}")
    lines.append(f"- **Total Controls Assessed:** {stats.get('total_controls', 0)}")
    lines.append(f"- **Implemented:** {stats.get('implemented', 0)}")
    lines.append(f"- **Not Implemented:** {stats.get('not_implemented', 0)}")

    if vulnerabilities:
        lines.append("\n## Identified Vulnerabilities")
        lines.append("\n| # | Vulnerability | Severity | Likelihood | Impact |")
        lines.append("|---|---|---|---|---|")
        for i, vuln in enumerate(vulnerabilities, 1):
            lines.append(f"| {i} | {vuln.get('title', '')} | {vuln.get('severity', '')} | {vuln.get('likelihood', '')} | {vuln.get('impact', '')} |")

    if recommendations:
        lines.append("\n## Recommendations")
        for i, rec in enumerate(recommendations, 1):
            if isinstance(rec, dict):
                lines.append(f"\n### {i}. {rec.get('title', 'Recommendation')}")
                lines.append(f"\n{rec.get('description', '')}")
                if rec.get('priority'):
                    lines.append(f"\n**Priority:** {rec['priority']}")
            else:
                lines.append(f"\n{i}. {rec}")

    if poam:
        lines.append("\n## Associated POA&M Items")
        lines.append("\n| # | Control | Weakness | Risk Level |")
        lines.append("|---|---|---|---|")
        for i, item in enumerate(poam, 1):
            lines.append(f"| {i} | {item.get('control_id', '')} | {item.get('weakness', '')[:50]} | {item.get('risk_level', '')} |")

    lines.append(f"\n---\n\n*Generated by SAELAR Risk Assessment Report Generator | {datetime.now().strftime('%B %d, %Y')}*")
    return "\n".join(lines)


def create_poam_document(ssp_data: dict, output_path: str = None):
    """
    Generate a standalone Plan of Action & Milestones (POA&M) Word document.
    
    Args:
        ssp_data: Dictionary from SSPGenerator.to_dict()
        output_path: Path for output .docx file
    
    Returns:
        Path to created document
    """
    if output_path is None:
        system_name = ssp_data.get('system_info', {}).get('system_name', 'System')
        safe_name = system_name.replace(' ', '_').replace('/', '-')[:30]
        output_path = _default_output_path("POAM", safe_name)
    
    doc = Document()
    
    system_info = ssp_data.get('system_info', {})
    poam_items = ssp_data.get('poam', [])
    summary = ssp_data.get('summary', {})
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('Plan of Action & Milestones (POA&M)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    system_name = system_info.get('system_name', 'Information System')
    subtitle = doc.add_paragraph(system_name)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.bold = True
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    # Generation info
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y')}\n").italic = True
    info_para.add_run(f"Security Categorization: {system_info.get('categorization', 'Moderate')}\n")
    info_para.add_run(f"Total POA&M Items: {len(poam_items)}")
    
    doc.add_page_break()
    
    # =========================================================================
    # POA&M SUMMARY
    # =========================================================================
    doc.add_heading('1. POA&M Summary', level=1)
    
    # Count by risk level
    high_count = sum(1 for p in poam_items if p.get('risk_level') == 'High')
    medium_count = sum(1 for p in poam_items if p.get('risk_level') == 'Medium')
    low_count = sum(1 for p in poam_items if p.get('risk_level') == 'Low')
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    headers = ['Category', 'Count']
    for i, header in enumerate(headers):
        cell = summary_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
    
    data = [
        ('High Risk Items', str(high_count)),
        ('Medium Risk Items', str(medium_count)),
        ('Low Risk Items', str(low_count)),
    ]
    
    for row_idx, (label, value) in enumerate(data, start=1):
        summary_table.rows[row_idx].cells[0].text = label
        summary_table.rows[row_idx].cells[1].text = value
    
    doc.add_paragraph()
    
    # =========================================================================
    # POA&M ITEMS
    # =========================================================================
    doc.add_heading('2. POA&M Items', level=1)
    
    if not poam_items:
        doc.add_paragraph('No POA&M items identified. All controls are fully implemented.')
    else:
        # Group by risk level
        for risk_level in ['High', 'Medium', 'Low']:
            level_items = [p for p in poam_items if p.get('risk_level') == risk_level]
            if not level_items:
                continue
            
            doc.add_heading(f'{risk_level} Risk Items ({len(level_items)})', level=2)
            
            # Create table for this risk level
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Headers
            headers = ['POA&M ID', 'Control', 'Weakness', 'Remediation', 'Due Date']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            
            # Data rows
            for item in level_items[:20]:  # Limit to 20 per section
                row = table.add_row()
                row.cells[0].text = item.get('poam_id', 'N/A')[:15]
                row.cells[1].text = item.get('control_id', 'N/A')
                row.cells[2].text = item.get('weakness', 'N/A')[:100] + ('...' if len(item.get('weakness', '')) > 100 else '')
                row.cells[3].text = item.get('remediation_plan', 'N/A')[:100] + ('...' if len(item.get('remediation_plan', '')) > 100 else '')
                due = item.get('scheduled_completion', 'TBD')
                if due and due != 'TBD':
                    try:
                        due = datetime.fromisoformat(due.replace('Z', '+00:00')).strftime('%Y-%m-%d')
                    except:
                        pass
                row.cells[4].text = str(due) if due else 'TBD'
                
                # Set font size
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.size = Pt(8)
            
            doc.add_paragraph()
    
    # =========================================================================
    # SIGNATURE BLOCK
    # =========================================================================
    doc.add_page_break()
    doc.add_heading('3. Approval', level=1)
    
    doc.add_paragraph('This POA&M has been reviewed and approved by:')
    doc.add_paragraph()
    
    approval_table = doc.add_table(rows=3, cols=3)
    approval_table.style = 'Table Grid'
    
    headers = ['Role', 'Name', 'Signature/Date']
    for i, header in enumerate(headers):
        approval_table.rows[0].cells[i].text = header
        approval_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    roles = [
        ('System Owner', system_info.get('system_owner', '')),
        ('ISSO', system_info.get('isso_name', '')),
    ]
    
    for row_idx, (role, name) in enumerate(roles, start=1):
        approval_table.rows[row_idx].cells[0].text = role
        approval_table.rows[row_idx].cells[1].text = name
        approval_table.rows[row_idx].cells[2].text = ''
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph(
        f"Generated by SAELAR POA&M Generator | {datetime.now().strftime('%B %d, %Y')}",
        style='Intense Quote'
    )
    
    # Save
    doc.save(output_path)
    print(f'[OK] POA&M Document saved to: {output_path}')
    return output_path


def create_rar_document(rar_data: dict, output_path: str = None):
    """
    Generate a Risk Assessment Report (RAR) Word document based on NOAA/NESDIS template.
    
    Args:
        rar_data: Dictionary containing:
            - system_info: System identification info
            - assessment_results: NIST 800-53 assessment findings
            - vulnerabilities: List of identified vulnerabilities
            - risk_summary: Risk scores and levels
            - recommendations: Remediation recommendations
            - poam: POA&M items
        output_path: Path for output .docx file
    
    Returns:
        Path to created document
    """
    if output_path is None:
        system_name = rar_data.get('system_info', {}).get('system_name', 'System')
        safe_name = system_name.replace(' ', '_').replace('/', '-')[:30]
        output_path = _default_output_path("RAR", safe_name)
    
    doc = Document()
    
    system_info = rar_data.get('system_info', {})
    assessment_results = rar_data.get('assessment_results', {})
    vulnerabilities = rar_data.get('vulnerabilities', [])
    risk_summary = rar_data.get('risk_summary', {})
    recommendations = rar_data.get('recommendations', [])
    poam = rar_data.get('poam', [])
    stats = rar_data.get('statistics', {})
    
    system_name = system_info.get('system_name', 'System Name')
    system_acronym = system_info.get('system_acronym', 'NOAA50xx')
    categorization = system_info.get('categorization', 'Moderate')
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(f"{system_name} ({system_acronym})")
    run.bold = True
    run.font.size = Pt(18)
    
    doc.add_paragraph()
    
    main_title = doc.add_heading('Risk Assessment Report', 0)
    main_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    org_para.add_run(system_info.get('organization', 'Organization Name')).bold = True
    
    doc.add_paragraph()
    
    version_para = doc.add_paragraph()
    version_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_para.add_run(f"Version {rar_data.get('metadata', {}).get('version', '1.0')}")
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.add_run(datetime.now().strftime('%m/%d/%Y'))
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    doc.add_paragraph(f"Prepared by: {system_info.get('preparer', 'SAELAR Risk Assessment Generator')}")
    
    doc.add_page_break()
    
    # =========================================================================
    # RECORD OF CHANGES
    # =========================================================================
    doc.add_heading('Record of Changes', level=1)
    
    changes_table = doc.add_table(rows=3, cols=4)
    changes_table.style = 'Table Grid'
    
    headers = ['Date Of Change', 'Version', 'Individual', 'Reason for Change']
    for i, header in enumerate(headers):
        changes_table.rows[0].cells[i].text = header
        changes_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    changes_table.rows[1].cells[0].text = datetime.now().strftime('%m/%d/%Y')
    changes_table.rows[1].cells[1].text = 'V1.0'
    changes_table.rows[1].cells[2].text = system_info.get('preparer', 'SAELAR')
    changes_table.rows[1].cells[3].text = 'Initial Risk Assessment based on SCA'
    
    doc.add_page_break()
    
    # =========================================================================
    # EXECUTIVE SUMMARY
    # =========================================================================
    doc.add_heading('Executive Summary', level=1)
    
    total_controls = stats.get('total_controls', 0)
    satisfied = stats.get('implemented', 0)
    not_satisfied = stats.get('not_implemented', 0) + stats.get('partial', 0)
    known_risk = len([p for p in poam if p.get('status') == 'Accepted Risk'])
    
    doc.add_paragraph(
        f"The risk assessment for the {system_name} ({system_acronym}) is updated annually in accordance "
        f"with DOC IT Security Program Policy. This update incorporates results of the FY{datetime.now().year} "
        f"independent Security Controls Assessment conducted by SAELAR."
    )
    
    doc.add_paragraph()
    
    # Summary of Technical Risk Analysis
    doc.add_heading('Summary of Technical Risk Analysis', level=2)
    
    overall_risk = risk_summary.get('overall_risk_level', 'Moderate')
    risk_score = risk_summary.get('total_risk_score', 0)
    
    doc.add_paragraph(
        f"The {system_acronym} has been assessed against NIST 800-53 Rev 5 security controls. "
        f"Based on the security control assessment and vulnerability analysis, the overall "
        f"residual risk level has been determined to be {overall_risk}."
    )
    
    doc.add_paragraph()
    
    # Risk Findings Table
    doc.add_heading(f'{system_acronym} Risk Findings', level=2)
    
    doc.add_paragraph('Table 1 summarizes the Security Control Assessment results:')
    
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = 'Table Grid'
    
    headers = ['Security Controls', 'Total (includes Enhancements)', 'Satisfied', 'Not Satisfied', 'Not Satisfied (Known Risk)']
    for i, header in enumerate(headers):
        summary_table.rows[0].cells[i].text = header
        summary_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
    
    summary_table.rows[1].cells[0].text = 'Total'
    summary_table.rows[1].cells[1].text = str(total_controls)
    summary_table.rows[1].cells[2].text = str(satisfied)
    summary_table.rows[1].cells[3].text = str(not_satisfied)
    summary_table.rows[1].cells[4].text = str(known_risk)
    
    doc.add_paragraph()
    
    # Security Categorization
    doc.add_heading(f'{system_acronym} Security Categorization', level=2)
    
    doc.add_paragraph(
        f"In accordance with Federal Information Processing Standards (FIPS) Publication 199 and "
        f"NIST SP 800-60, the {system_acronym} has been categorized as follows:"
    )
    
    cat_table = doc.add_table(rows=4, cols=2)
    cat_table.style = 'Table Grid'
    
    cat_data = [
        ('Confidentiality', system_info.get('confidentiality_level', categorization)),
        ('Integrity', system_info.get('integrity_level', categorization)),
        ('Availability', system_info.get('availability_level', categorization)),
    ]
    
    cat_table.rows[0].cells[0].text = 'Impact Level'
    cat_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    cat_table.rows[0].cells[1].text = ''
    
    for i, (category, level) in enumerate(cat_data, 1):
        cat_table.rows[i].cells[0].text = category
        cat_table.rows[i].cells[1].text = level
    
    doc.add_paragraph()
    doc.add_paragraph(
        f"Using the highest impact level among the three security goals (High Water Mark), "
        f"the overall impact level of {system_acronym} is {categorization}."
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # RECOMMENDATIONS FOR MITIGATING RISKS
    # =========================================================================
    doc.add_heading('Recommendations for Mitigating Risks', level=2)
    
    doc.add_paragraph(
        f"To reduce the {overall_risk} level of residual risk from exploit of vulnerabilities "
        f"that existed in {system_acronym}, the following risk mitigation strategies are recommended:"
    )
    
    if recommendations:
        rec_table = doc.add_table(rows=len(recommendations[:10]) + 1, cols=2)
        rec_table.style = 'Table Grid'
        
        rec_table.rows[0].cells[0].text = 'Control Number'
        rec_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        rec_table.rows[0].cells[1].text = 'Recommended Mitigation'
        rec_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
        
        for i, rec in enumerate(recommendations[:10], 1):
            rec_table.rows[i].cells[0].text = rec.get('control_id', '')
            rec_table.rows[i].cells[1].text = rec.get('recommendation', '')[:200]
    else:
        doc.add_paragraph('See POA&M section for detailed remediation recommendations.')
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        'Executive Summary',
        '1. Introduction',
        '   1.1 Purpose',
        '   1.2 Scope',
        '   1.3 References',
        '2. Assessment Overview',
        '   2.1 Assessment Team',
        '   2.2 NIST Methodology',
        '3. System General Description/Purpose',
        '4. Threat Identification',
        '5. Vulnerability Identification',
        '6. Security Control Analysis',
        '7. Risk Determination',
        '8. Recommendations',
        'Appendix A: Acronyms and Abbreviations',
        'Appendix B: Definitions',
        'Appendix C: References',
        'Appendix D: Risk Assessment Results'
    ]
    
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 1: INTRODUCTION
    # =========================================================================
    doc.add_heading('1. INTRODUCTION', level=1)
    
    doc.add_paragraph(
        f"This document explains the results of the Risk Assessment completed for the {system_name} "
        f"({system_acronym}). The Risk Assessment Report (RAR) was updated to incorporate findings "
        f"from the FY{datetime.now().year} Security Controls Assessment conducted by SAELAR."
    )
    
    doc.add_paragraph(
        f"This risk assessment was conducted as part of an overall Security Assessment and Authorization "
        f"(A&A) process in accordance with Federal laws and mandates. A risk assessment identifies "
        f"security vulnerabilities, threats, and an estimate of the consequences if a threat exploits "
        f"a security vulnerability."
    )
    
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        f"The purpose of this risk assessment is to identify the risks to system security and determine "
        f"the probability of occurrence, the resulting impact, and additional safeguards that would mitigate "
        f"this impact. Federal law and regulations, including OMB Circular A-130, require all federally "
        f"funded systems to adhere to a security program that incorporates risk management."
    )
    
    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        f"This assessment addresses the organizational requirements; network, application, data architectures; "
        f"and the threats and vulnerabilities that pertain to the components of {system_acronym} within its "
        f"system boundaries, as described in the System Security Plan (SSP)."
    )
    
    doc.add_heading('1.3 References', level=2)
    doc.add_paragraph('References can be found listed in Appendix C.')
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 2: ASSESSMENT OVERVIEW
    # =========================================================================
    doc.add_heading('2. ASSESSMENT OVERVIEW', level=1)
    
    doc.add_heading('2.1 Assessment Team', level=2)
    
    team_table = doc.add_table(rows=3, cols=5)
    team_table.style = 'Table Grid'
    
    headers = ['Name', 'Role', 'Organization', 'Phone Number', 'E-mail Address']
    for i, header in enumerate(headers):
        team_table.rows[0].cells[i].text = header
        team_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    team_table.rows[1].cells[0].text = system_info.get('system_owner', 'TBD')
    team_table.rows[1].cells[1].text = 'System Owner'
    team_table.rows[1].cells[2].text = system_info.get('organization', 'TBD')
    team_table.rows[1].cells[3].text = system_info.get('owner_phone', 'TBD')
    team_table.rows[1].cells[4].text = system_info.get('owner_email', 'TBD')
    
    team_table.rows[2].cells[0].text = system_info.get('isso_name', 'TBD')
    team_table.rows[2].cells[1].text = 'ISSO'
    team_table.rows[2].cells[2].text = system_info.get('organization', 'TBD')
    team_table.rows[2].cells[3].text = system_info.get('isso_phone', 'TBD')
    team_table.rows[2].cells[4].text = system_info.get('isso_email', 'TBD')
    
    doc.add_paragraph()
    
    doc.add_heading('2.2 NIST Methodology', level=2)
    doc.add_paragraph(
        f"The risk assessment methodology employed for {system_acronym} followed the guidance in "
        f"NIST SP 800-30, Revision 1, Guide for Conducting Risk Assessments."
    )
    
    doc.add_paragraph('The formula used to calculate the risk is:')
    doc.add_paragraph('Risk = (Likelihood of a Threat Occurring) x (Magnitude of Impact)', style='Intense Quote')
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 3: SYSTEM DESCRIPTION
    # =========================================================================
    doc.add_heading('3. SYSTEM GENERAL DESCRIPTION/PURPOSE', level=1)
    
    doc.add_paragraph(
        f"Details of the system description, significant security roles, system interconnections, "
        f"data flow diagrams, user base, system categorization, and security authorization boundary "
        f"including control requirements can be found in the {system_acronym} System Security Plan."
    )
    
    doc.add_paragraph(system_info.get('system_description', 'System description to be provided.'))
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 4: THREAT IDENTIFICATION
    # =========================================================================
    doc.add_heading('4. THREAT IDENTIFICATION', level=1)
    
    doc.add_paragraph(
        f"A threat is the potential for a particular threat-source to successfully exercise a "
        f"particular vulnerability. The following threat-sources were considered for their potential "
        f"to harm the {system_acronym} system."
    )
    
    doc.add_heading('4.1 System-Specific Threat Sources', level=2)
    
    # Standard threat table
    threat_table = doc.add_table(rows=6, cols=3)
    threat_table.style = 'Table Grid'
    
    threat_headers = ['Threat-Source', 'Motivation', 'Threat Actions']
    for i, header in enumerate(threat_headers):
        threat_table.rows[0].cells[i].text = header
        threat_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    
    threats = [
        ('Hacker, cracker', 'Challenge, Ego, Rebellion', 'Hacking, Social engineering, System intrusion'),
        ('Computer criminal', 'Monetary gain, Data theft', 'Fraud, Spoofing, System intrusion'),
        ('Terrorist', 'Destruction, Exploitation', 'System attack, DDoS, Tampering'),
        ('Industrial espionage', 'Competitive advantage', 'Information theft, System penetration'),
        ('Insider threat', 'Curiosity, Revenge, Monetary gain', 'Data theft, Sabotage, Abuse'),
    ]
    
    for i, (source, motivation, actions) in enumerate(threats, 1):
        threat_table.rows[i].cells[0].text = source
        threat_table.rows[i].cells[1].text = motivation
        threat_table.rows[i].cells[2].text = actions
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 5: VULNERABILITY IDENTIFICATION
    # =========================================================================
    doc.add_heading('5. VULNERABILITY IDENTIFICATION', level=1)
    
    doc.add_paragraph(
        f"The analysis of the risk to {system_acronym} includes an analysis of the vulnerabilities "
        f"associated with the system. Vulnerabilities were identified through Security Controls Assessment "
        f"and vulnerability scanning."
    )
    
    doc.add_heading('5.1 Vulnerabilities Assessed', level=2)
    
    if vulnerabilities:
        vuln_count = len(vulnerabilities)
        high_vulns = len([v for v in vulnerabilities if v.get('severity') == 'High'])
        med_vulns = len([v for v in vulnerabilities if v.get('severity') == 'Medium'])
        low_vulns = len([v for v in vulnerabilities if v.get('severity') == 'Low'])
        
        doc.add_paragraph(f'Total vulnerabilities identified: {vuln_count}')
        doc.add_paragraph(f'- High: {high_vulns}', style='List Bullet')
        doc.add_paragraph(f'- Medium: {med_vulns}', style='List Bullet')
        doc.add_paragraph(f'- Low: {low_vulns}', style='List Bullet')
    else:
        doc.add_paragraph('Vulnerability details are documented in the Security Assessment Report.')
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 6: SECURITY CONTROL ANALYSIS
    # =========================================================================
    doc.add_heading('6. SECURITY CONTROL ANALYSIS', level=1)
    
    doc.add_paragraph(
        f"The threat and vulnerability results are analyzed in this section. Controls were assessed "
        f"against NIST 800-53 Rev 5 requirements."
    )
    
    # Family breakdown if available
    family_summary = assessment_results.get('family_summary', {})
    if family_summary:
        control_table = doc.add_table(rows=len(family_summary) + 1, cols=4)
        control_table.style = 'Table Grid'
        
        headers = ['Control Family', 'Total', 'Satisfied', 'Compliance %']
        for i, header in enumerate(headers):
            control_table.rows[0].cells[i].text = header
            control_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for row_idx, (family_code, data) in enumerate(family_summary.items(), 1):
            control_table.rows[row_idx].cells[0].text = f"{family_code} - {data.get('name', '')}"
            control_table.rows[row_idx].cells[1].text = str(data.get('total', 0))
            control_table.rows[row_idx].cells[2].text = str(data.get('implemented', 0))
            control_table.rows[row_idx].cells[3].text = f"{data.get('compliance_pct', 0)}%"
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 7: RISK DETERMINATION
    # =========================================================================
    doc.add_heading('7. RISK DETERMINATION', level=1)
    
    doc.add_heading('Risk Summary', level=2)
    
    doc.add_paragraph(
        f"Based on the security control assessment and analysis, the overall risk level for "
        f"{system_acronym} has been determined to be {overall_risk}."
    )
    
    doc.add_paragraph(f"Total Risk Score: {risk_score:.1f}")
    
    # Risk findings table
    high_risk_items = [p for p in poam if p.get('risk_level') == 'High']
    if high_risk_items:
        doc.add_heading('High Risk Findings', level=2)
        
        risk_table = doc.add_table(rows=min(len(high_risk_items), 10) + 1, cols=3)
        risk_table.style = 'Table Grid'
        
        headers = ['Control #', 'Threats and Vulnerabilities', 'Risk Level']
        for i, header in enumerate(headers):
            risk_table.rows[0].cells[i].text = header
            risk_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for i, item in enumerate(high_risk_items[:10], 1):
            risk_table.rows[i].cells[0].text = item.get('control_id', '')
            risk_table.rows[i].cells[1].text = item.get('weakness', '')[:100]
            risk_table.rows[i].cells[2].text = 'High'
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 8: RECOMMENDATIONS
    # =========================================================================
    doc.add_heading('8. RECOMMENDATIONS', level=1)
    
    doc.add_paragraph(
        f"To reduce the {overall_risk} level of residual risk from exploit of vulnerabilities "
        f"that existed in {system_acronym}, the following remediation strategies are recommended:"
    )
    
    if poam:
        rec_table = doc.add_table(rows=min(len(poam), 15) + 1, cols=2)
        rec_table.style = 'Table Grid'
        
        rec_table.rows[0].cells[0].text = 'Control No'
        rec_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
        rec_table.rows[0].cells[1].text = 'Remediation'
        rec_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
        
        for i, item in enumerate(poam[:15], 1):
            rec_table.rows[i].cells[0].text = item.get('control_id', '')
            rec_table.rows[i].cells[1].text = item.get('remediation_plan', item.get('recommendation', ''))[:200]
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX A: ACRONYMS
    # =========================================================================
    doc.add_heading('APPENDIX A: ACRONYMS AND ABBREVIATIONS', level=1)
    
    acronyms = [
        ('A&A', 'Assessment and Authorization'),
        ('AO', 'Authorizing Official'),
        ('ATO', 'Authorization To Operate'),
        ('FIPS', 'Federal Information Processing Standard'),
        ('FISMA', 'Federal Information Security Management Act'),
        ('GSS', 'General Support System'),
        ('ISSO', 'Information System Security Officer'),
        ('IT', 'Information Technology'),
        ('NIST', 'National Institute of Standards and Technology'),
        ('OMB', 'Office of Management and Budget'),
        ('POA&M', 'Plan of Action & Milestones'),
        ('RAR', 'Risk Assessment Report'),
        ('SCA', 'Security Controls Assessment'),
        ('SP', 'Special Publication'),
        ('SSP', 'System Security Plan'),
    ]
    
    acr_table = doc.add_table(rows=len(acronyms) + 1, cols=2)
    acr_table.style = 'Table Grid'
    
    acr_table.rows[0].cells[0].text = 'Abbreviation / Acronym'
    acr_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
    acr_table.rows[0].cells[1].text = 'Definition'
    acr_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
    
    for i, (abbr, definition) in enumerate(acronyms, 1):
        acr_table.rows[i].cells[0].text = abbr
        acr_table.rows[i].cells[1].text = definition
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX C: REFERENCES
    # =========================================================================
    doc.add_heading('APPENDIX C: REFERENCES', level=1)
    
    references = [
        'Federal Information Security Modernization Act (FISMA) of 2014',
        'OMB Circular No. A-130, Appendix III, Security of Federal Automated Information Resources',
        'FIPS 199, Standards for Security Categorization of Federal Information and Information Systems',
        'FIPS 200, Minimum Security Requirements for Federal Information and Information Systems',
        'NIST SP 800-18, Rev 1, Guide for Developing Security Plans for Federal Information Systems',
        'NIST SP 800-30, Rev 1, Guide for Conducting Risk Assessments',
        'NIST SP 800-37, Rev 2, Risk Management Framework for Information Systems and Organizations',
        'NIST SP 800-53, Rev 5, Security and Privacy Controls for Information Systems and Organizations',
        'NIST SP 800-53A, Rev 5, Assessing Security and Privacy Controls in Information Systems',
        'NIST SP 800-60, Rev 1, Guide for Mapping Types of Information and Information Systems',
    ]
    
    for ref in references:
        doc.add_paragraph(ref, style='List Bullet')
    
    doc.add_page_break()
    
    # =========================================================================
    # APPENDIX D: RISK ASSESSMENT RESULTS
    # =========================================================================
    doc.add_heading('APPENDIX D: RISK ASSESSMENT RESULTS', level=1)
    
    doc.add_paragraph(f'Summary of {system_acronym} Vulnerabilities Identified by SCA:')
    
    # Summary by severity
    controls_data = assessment_results.get('controls', [])
    
    if controls_data:
        high_count = len([c for c in controls_data if c.get('risk_level') == 'High'])
        med_count = len([c for c in controls_data if c.get('risk_level') == 'Medium'])
        low_count = len([c for c in controls_data if c.get('risk_level') == 'Low'])
        
        result_table = doc.add_table(rows=2, cols=5)
        result_table.style = 'Table Grid'
        
        headers = ['System', 'Critical', 'High', 'Moderate', 'Low']
        for i, header in enumerate(headers):
            result_table.rows[0].cells[i].text = header
            result_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        result_table.rows[1].cells[0].text = system_acronym
        result_table.rows[1].cells[1].text = '0'
        result_table.rows[1].cells[2].text = str(high_count)
        result_table.rows[1].cells[3].text = str(med_count)
        result_table.rows[1].cells[4].text = str(low_count)
    
    doc.add_paragraph()
    
    # SAR Results table
    if controls_data:
        doc.add_paragraph('SAR Results:')
        
        sar_table = doc.add_table(rows=min(len(controls_data), 20) + 1, cols=4)
        sar_table.style = 'Table Grid'
        
        headers = ['Control No.', 'Control Name', 'Vulnerability/Deficiency', 'Risk Level']
        for i, header in enumerate(headers):
            sar_table.rows[0].cells[i].text = header
            sar_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            sar_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)
        
        for i, control in enumerate(controls_data[:20], 1):
            sar_table.rows[i].cells[0].text = control.get('control_id', '')
            sar_table.rows[i].cells[1].text = control.get('control_name', '')[:30]
            weakness = control.get('finding', control.get('weakness', ''))
            sar_table.rows[i].cells[2].text = weakness[:60] if weakness else 'N/A'
            sar_table.rows[i].cells[3].text = control.get('risk_level', 'N/A')
            
            for cell in sar_table.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)
    
    # Footer
    doc.add_paragraph()
    doc.add_paragraph('─' * 50)
    doc.add_paragraph(
        f"Generated by SAELAR Risk Assessment Report Generator | {datetime.now().strftime('%B %d, %Y')}",
        style='Intense Quote'
    )
    
    # Save
    doc.save(output_path)
    print(f'[OK] RAR Document saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    if len(sys.argv) > 1:
        if sys.argv[1] == '--convert-folder' and len(sys.argv) > 2:
            convert_folder(sys.argv[2])
        elif sys.argv[1] == '--convert' and len(sys.argv) > 2:
            convert_text_to_docx(sys.argv[2])
        else:
            print('Usage:')
            print('  python wordy.py                           # Generate comparison doc')
            print('  python wordy.py --convert-folder <path>   # Convert all .txt/.md in folder')
            print('  python wordy.py --convert <file>          # Convert single file')
    else:
        # Default: Generate the SAELAR vs AWS Audit Manager comparison
        create_comparison_document()

