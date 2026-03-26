"""SOPRA ISSO — POA&M Tracker (AI-Enhanced)"""
import streamlit as st
from datetime import datetime
import datetime as dt_module
import io

from sopra_controls import get_control_by_id
from sopra.persistence import _load_json, _save_json, _next_id, load_results_from_file
from sopra.utils import calculate_risk_score
from sopra.isso.ai_engine import ai_generate_poam_entries

def render_poam_tracker():
    """POA&M Tracker with milestones, due dates, responsible parties, overdue alerts."""
    st.markdown("## 📋 Plan of Action & Milestones (POA&M) Tracker")
    st.caption("View and manage POA&M items that will be included in the SSP document.")

    poam_items = _load_json("poam.json", [])
    results = st.session_state.get("opra_assessment_results")
    if not results or not results.get("findings"):
        results = load_results_from_file()
        if results:
            st.session_state.opra_assessment_results = results

    # ── AI POA&M Generation ──────────────────────────────────────────
    if results and results.get("findings"):
        failed = [f for f in results["findings"] if f["status"] == "Failed"]
        existing_cids = {p.get("control_id") for p in poam_items}
        new_failed = [f for f in failed if f.get("control_id") and f["control_id"] not in existing_cids]

        if new_failed:
            ai_col1, ai_col2 = st.columns([3, 1])
            with ai_col1:
                st.info(f"**{len(new_failed)}** failed findings have no POA&M entry yet.")
            with ai_col2:
                ai_generate = st.button("🤖 AI Generate POA&M", type="primary", key="ai_gen_poam")

            if ai_generate:
                with st.spinner("AI is analyzing findings and generating POA&M entries..."):
                    ai_entries = ai_generate_poam_entries(new_failed)

                if ai_entries:
                    new_count = 0
                    ai_lookup = {e["control_id"]: e for e in ai_entries if "control_id" in e}
                    for f in new_failed:
                        cid = f.get("control_id", "")
                        ai_data = ai_lookup.get(cid, {})
                        sev = f.get("severity", "Medium")
                        days = ai_data.get("estimated_days", {"Critical": 30, "High": 60, "Medium": 90, "Low": 180}.get(sev, 90))
                        due = (datetime.now() + __import__('datetime').timedelta(days=days)).strftime("%Y-%m-%d")

                        milestones = []
                        for mk in ["milestone_1", "milestone_2", "milestone_3"]:
                            if ai_data.get(mk):
                                milestones.append(ai_data[mk])

                        poam_items.append({
                            "id": _next_id(poam_items, "POAM"),
                            "control_id": cid,
                            "finding": ai_data.get("finding", f.get("control_name", "")),
                            "severity": sev,
                            "status": "Open",
                            "responsible_party": ai_data.get("responsible_party", "TBD"),
                            "due_date": due,
                            "milestone_pct": 0,
                            "milestones": milestones,
                            "risk_if_delayed": ai_data.get("risk_if_delayed", ""),
                            "notes": "\n".join(milestones) if milestones else "",
                            "created": datetime.now().strftime("%Y-%m-%d"),
                            "updated": datetime.now().strftime("%Y-%m-%d"),
                            "ai_generated": True,
                        })
                        existing_cids.add(cid)
                        new_count += 1
                    _save_json("poam.json", poam_items)
                    st.success(f"AI generated **{new_count}** POA&M entries with smart milestones, owners, and due dates.")
                    st.rerun()
        else:
            # Standard auto-gen for any missed (non-AI fallback)
            existing_cids = {p.get("control_id") for p in poam_items}
            new_count = 0
            for f in failed:
                cid = f.get("control_id", "")
                if cid and cid not in existing_cids:
                    sev = f.get("severity", "Medium")
                    days = {"Critical": 30, "High": 60, "Medium": 90, "Low": 180}.get(sev, 90)
                    due = (datetime.now() + __import__('datetime').timedelta(days=days)).strftime("%Y-%m-%d")
                    poam_items.append({
                        "id": _next_id(poam_items, "POAM"),
                        "control_id": cid,
                        "finding": f.get("control_name", ""),
                        "severity": sev,
                        "status": "Open",
                        "responsible_party": "TBD",
                        "due_date": due,
                        "milestone_pct": 0,
                        "milestones": [],
                        "notes": "",
                        "created": datetime.now().strftime("%Y-%m-%d"),
                        "updated": datetime.now().strftime("%Y-%m-%d")
                    })
                    existing_cids.add(cid)
                    new_count += 1
            if new_count > 0:
                _save_json("poam.json", poam_items)

    if not poam_items:
        st.info("No POA&M items yet. Run an assessment to auto-generate items from failed controls.")
        return

    # Summary KPIs
    open_ct = len([p for p in poam_items if p["status"] in ("Open", "In Progress")])
    overdue_ct = len([p for p in poam_items if p["status"] != "Closed" and p.get("due_date", "9999") < datetime.now().strftime("%Y-%m-%d")])
    closed_ct = len([p for p in poam_items if p["status"] == "Closed"])
    ai_ct = len([p for p in poam_items if p.get("ai_generated")])
    k1, k2, k3, k4 = st.columns(4)
    with k1:
        st.metric("Total Items", len(poam_items))
    with k2:
        st.metric("Open", open_ct)
    with k3:
        st.metric("Overdue", overdue_ct, delta=f"-{overdue_ct}" if overdue_ct > 0 else "0", delta_color="inverse")
    with k4:
        st.metric("Closed", closed_ct)

    if ai_ct:
        st.caption(f"🤖 {ai_ct} items were AI-generated with smart milestones and risk assessments")

    # Risk Acceptance tile — navigates to Acceptance of Risk process and document creation
    st.markdown("---")
    st.markdown(
        '<div style="background: linear-gradient(145deg, rgba(255,193,7,0.12), rgba(255,193,7,0.04)); '
        'border: 1px solid rgba(255,193,7,0.35); border-radius: 12px; padding: 1rem; margin-bottom: 0.5rem; display: inline-block;">'
        '<div style="font-size: 1.8rem; margin-bottom: 0.25rem;">⚖️</div>'
        '<div style="color: #ffc107; font-weight: 700; font-size: 1rem;">Risk Acceptance</div>'
        '<div style="color: #6b839e; font-size: 0.75rem;">Acceptance of Risk process & document</div>'
        '</div>',
        unsafe_allow_html=True
    )
    st.link_button("⚖️ Open Risk Acceptance", "?view=risk_acceptance", use_container_width=False, help="Perform Acceptance of Risk process and create document")
    st.markdown("---")

    # Mark overdue items
    today = datetime.now().strftime("%Y-%m-%d")
    for p in poam_items:
        if p["status"] not in ("Closed",) and p.get("due_date", "9999") < today:
            p["status"] = "Overdue"
    _save_json("poam.json", poam_items)

    # Table display
    for p in poam_items:
        sev_color = {"Critical": "#e94560", "High": "#ff6b6b", "Medium": "#ffc107", "Low": "#00d9ff"}.get(p["severity"], "#888")
        status_color = {"Open": "#ffc107", "In Progress": "#00d9ff", "Overdue": "#e94560", "Closed": "#00ff88"}.get(p["status"], "#888")
        overdue_badge = ' <span style="background:#e94560;color:#fff;padding:2px 8px;border-radius:10px;font-size:0.65rem;font-weight:700;">OVERDUE</span>' if p["status"] == "Overdue" else ""
        ai_badge = " 🤖" if p.get("ai_generated") else ""

        with st.expander(f"{p['id']} — {p['control_id']} | {p['finding'][:50]}... | {p['status']}{' ⚠️' if p['status'] == 'Overdue' else ''}{ai_badge}"):
            st.markdown(
                '<div style="display: flex; gap: 1rem; flex-wrap: wrap; margin-bottom: 0.5rem;">'
                '<span style="color: ' + sev_color + '; font-weight: 700;">Severity: ' + p['severity'] + '</span>'
                '<span style="color: ' + status_color + '; font-weight: 700;">Status: ' + p['status'] + overdue_badge + '</span>'
                '<span style="color: #6b839e;">Due: ' + p.get('due_date', 'N/A') + '</span>'
                '<span style="color: #6b839e;">Owner: ' + p.get('responsible_party', 'TBD') + '</span>'
                '</div>',
                unsafe_allow_html=True
            )

            # Risk warning from AI
            if p.get("risk_if_delayed"):
                st.warning(f"**Risk if Delayed:** {p['risk_if_delayed']}")

            # Progress bar
            pct = p.get("milestone_pct", 0)
            bar_color = "#00ff88" if pct >= 80 else "#ffc107" if pct >= 40 else "#e94560"
            st.markdown(
                '<div style="background: rgba(255,255,255,0.06); border-radius: 6px; height: 10px; overflow: hidden; margin: 0.3rem 0;">'
                '<div style="width: ' + str(pct) + '%; height: 100%; background: ' + bar_color + '; border-radius: 6px; transition: width 0.4s;"></div>'
                '</div>'
                '<div style="text-align: right; font-size: 0.7rem; color: #6b839e;">' + str(pct) + '% complete</div>',
                unsafe_allow_html=True
            )

            c1, c2, c3 = st.columns(3)
            with c1:
                new_status = st.selectbox("Status", ["Open", "In Progress", "Closed"], index=["Open", "In Progress", "Closed"].index(p["status"]) if p["status"] in ("Open", "In Progress", "Closed") else 0, key=f"poam_st_{p['id']}")
            with c2:
                new_owner = st.text_input("Responsible Party", value=p.get("responsible_party", "TBD"), key=f"poam_own_{p['id']}")
            with c3:
                new_pct = st.slider("Milestone %", 0, 100, pct, key=f"poam_pct_{p['id']}")

            new_notes = st.text_area("Notes", value=p.get("notes", ""), key=f"poam_note_{p['id']}", height=68)

            if st.button("💾 Save", key=f"poam_save_{p['id']}"):
                p["status"] = new_status
                p["responsible_party"] = new_owner
                p["milestone_pct"] = new_pct
                p["notes"] = new_notes
                p["updated"] = datetime.now().strftime("%Y-%m-%d")
                _save_json("poam.json", poam_items)
                st.success(f"Updated {p['id']}")
                st.rerun()

    # Export POA&M as DOCX
    st.markdown("---")
    if st.button("📥 Export POA&M (.docx)", use_container_width=True, type="primary"):
        _export_poam_docx(poam_items)


def _export_poam_docx(poam_items):
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = DocxDocument()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Plan of Action & Milestones (POA&M)')
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | Total Items: {len(poam_items)}")
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = 'Light Grid Accent 1'
    for ci, h in enumerate(['ID', 'Control', 'Finding', 'Severity', 'Status', 'Due Date', 'Owner']):
        tbl.rows[0].cells[ci].text = h
        for run in tbl.rows[0].cells[ci].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for p in poam_items:
        row = tbl.add_row()
        for ci, v in enumerate([p['id'], p['control_id'], p['finding'][:40], p['severity'], p['status'], p.get('due_date',''), p.get('responsible_party','TBD')]):
            row.cells[ci].text = str(v)
            for run in row.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(7)
                if p['status'] == 'Overdue':
                    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button("📥 Download POA&M", data=buf.getvalue(),
                       file_name=f"SOPRA_POAM_{datetime.now().strftime('%Y%m%d')}.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
