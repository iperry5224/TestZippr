"""SOPRA Reports Page"""
import streamlit as st
import plotly.graph_objects as go
import pandas as pd
from datetime import datetime
import io

from sopra_controls import get_control_by_id, get_remediation_script, ControlFamily
from sopra.theme import (
    COLOR_PASSED, COLOR_FAILED, COLOR_CRITICAL, COLOR_HIGH, COLOR_MEDIUM, COLOR_LOW,
    SEV_COLORS, SEV_ICONS, FAMILY_ABBREV, SEV_ORDER, BAR_COLORS_SEV,
    CHART_FONT_COLOR, CHART_GRID_COLOR, chart_layout
)
from sopra.persistence import load_results_from_file
from sopra.utils import (
    aggregate_findings, calculate_risk_score,
    create_status_donut, create_category_bar, create_risk_gauge
)

def render_reports_page():
    """Render the reports page with enhanced visualizations"""
    st.markdown("## 📊 Assessment Reports")
    
    # Load from shared file if session state is empty (e.g. opened in new tab)
    if not st.session_state.get('opra_assessment_results'):
        st.session_state.opra_assessment_results = load_results_from_file()

    if not st.session_state.opra_assessment_results:
        st.warning("⚠️ No assessment data available. Please run an assessment first.")
        if st.button("▶️ Go to Run Security Assessment"):
            st.session_state.opra_active_tab = "Run Security Assessment"
            st.rerun()
        return
    
    results = st.session_state.opra_assessment_results
    findings = results.get("findings", [])
    agg = aggregate_findings(findings)
    total = agg["total"]
    passed = agg["passed"]
    failed = agg["failed"]
    not_assessed = agg["not_assessed"]
    
    score = calculate_risk_score(findings)
    
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        if st.button(f"📋 {total}\nTOTAL CONTROLS", key="rpt_metric_total", use_container_width=True,
                     help="View all controls grouped by category"):
            st.session_state.rpt_drill = "total"
    with col2:
        if st.button(f"✅ {passed}\nPASSED", key="rpt_metric_passed", use_container_width=True,
                     help="View all passed controls with evidence"):
            st.session_state.rpt_drill = "passed"
    with col3:
        if st.button(f"❌ {failed}\nFAILED", key="rpt_metric_failed", use_container_width=True,
                     help="View all failed controls by severity"):
            st.session_state.rpt_drill = "failed"
    with col4:
        if st.button(f"🎯 {score}%\nRISK SCORE", key="rpt_metric_risk", use_container_width=True,
                     help="View risk breakdown and scoring details"):
            st.session_state.rpt_drill = "risk"
    
    # Initialize drill state
    if "rpt_drill" not in st.session_state:
        st.session_state.rpt_drill = None
    
    rpt_panel = st.session_state.get("rpt_drill")
    
    if rpt_panel == "total":
        st.markdown("---")
        st.markdown("### 📋 All Controls — Grouped by Category")
        for cat_name in sorted(set(f.get("category", "Unknown") for f in findings)):
            cat_findings = [f for f in findings if f.get("category") == cat_name]
            cat_passed = len([f for f in cat_findings if f["status"] == "Passed"])
            cat_failed = len([f for f in cat_findings if f["status"] == "Failed"])
            cat_pct = int(cat_passed / len(cat_findings) * 100) if cat_findings else 0
            with st.expander(f"{'✅' if cat_pct >= 70 else '⚠️' if cat_pct >= 40 else '🔴'} {cat_name} — {len(cat_findings)} controls ({cat_pct}% compliant)", expanded=False):
                df_cat = pd.DataFrame(cat_findings)[["control_id", "control_name", "status", "severity", "family"]]
                df_cat.columns = ["ID", "Control", "Status", "Severity", "Family"]
                
                # Color-code status
                st.dataframe(df_cat, use_container_width=True, hide_index=True, height=min(len(df_cat) * 38 + 40, 400))
                
                st.markdown(f"""
                <div style="display: flex; gap: 1rem; margin-top: 0.5rem;">
                    <span style="color: #00d9ff;">✅ {cat_passed} Passed</span>
                    <span style="color: #e94560;">❌ {cat_failed} Failed</span>
                    <span style="color: #4a6fa5;">⬜ {len(cat_findings) - cat_passed - cat_failed} Not Assessed</span>
                </div>
                """, unsafe_allow_html=True)
    
    elif rpt_panel == "passed":
        st.markdown("---")
        st.markdown("### ✅ All Passed Controls")
        passed_findings = [f for f in findings if f["status"] == "Passed"]
        if passed_findings:
            df_passed = pd.DataFrame(passed_findings)
            display_cols = ["control_id", "control_name", "category"]
            if "evidence" in df_passed.columns:
                display_cols.append("evidence")
            if "family" in df_passed.columns:
                display_cols.append("family")
            df_display = df_passed[display_cols]
            df_display.columns = ["ID", "Control", "Category"] + (["Evidence"] if "evidence" in display_cols else []) + (["Family"] if "family" in display_cols else [])
            st.dataframe(df_display, use_container_width=True, hide_index=True, height=min(len(df_display) * 38 + 40, 600))
            
            # Per-category passed summary
            st.markdown("#### Passed by Category")
            for cat in sorted(set(f.get("category", "Unknown") for f in passed_findings)):
                cat_count = len([f for f in passed_findings if f.get("category") == cat])
                cat_total = len([f for f in findings if f.get("category") == cat])
                pct = int(cat_count / cat_total * 100) if cat_total else 0
                bar_color = "#00d9ff" if pct >= 70 else "#ffc107" if pct >= 40 else "#e94560"
                st.markdown(f"""
                <div style="margin-bottom: 0.5rem;">
                    <div style="display: flex; justify-content: space-between;">
                        <span style="color: #f5f5f5; font-size: 0.85rem;">{cat}</span>
                        <span style="color: {bar_color}; font-weight: 700;">{cat_count}/{cat_total} ({pct}%)</span>
                    </div>
                    <div style="background: rgba(0,217,255,0.1); border-radius: 6px; height: 14px; overflow: hidden;">
                        <div style="background: {bar_color}; width: {pct}%; height: 100%; border-radius: 6px; box-shadow: 0 0 6px {bar_color}66;"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
        else:
            st.info("No passed controls found.")
    
    elif rpt_panel == "failed":
        st.markdown("---")
        st.markdown("### 🔴 All Failed Controls — Prioritized by Severity")
        failed_findings = agg["failed_findings"]
        if failed_findings:
            severity_order = SEV_ORDER
            sev_colors = SEV_COLORS
            sev_icons = SEV_ICONS
            
            for sev in severity_order:
                sev_findings = [f for f in failed_findings if f.get("severity") == sev]
                if sev_findings:
                    color = sev_colors.get(sev, "#c0c0c0")
                    icon = sev_icons.get(sev, "⚪")
                    with st.expander(f"{icon} {sev} — {len(sev_findings)} findings", expanded=(sev == "Critical")):
                        for sf in sev_findings:
                            evidence = sf.get('evidence', 'No evidence recorded')
                            notes = sf.get('notes', '')
                            st.markdown(f"""
                            <div style="background: rgba({','.join(str(int(color.lstrip('#')[i:i+2], 16)) for i in (0,2,4))}, 0.08); 
                                        border-left: 3px solid {color}; padding: 0.6rem 1rem; margin: 0.4rem 0; border-radius: 0 8px 8px 0;">
                                <span style="color: {color}; font-weight: 700;">{sf.get('control_id', 'N/A')}</span>
                                <span style="color: #f5f5f5;"> — {sf.get('control_name', 'Unknown')}</span><br>
                                <span style="color: #8899bb; font-size: 0.85rem;">📁 {sf.get('category', '')} &nbsp;|&nbsp; 🏷️ {sf.get('family', '')}</span><br>
                                <span style="color: #c8e6ff; font-size: 0.85rem;">📝 {evidence}</span>
                                {'<br><span style="color: #a0a0c0; font-size: 0.8rem;">💡 ' + notes + '</span>' if notes else ''}
                            </div>
                            """, unsafe_allow_html=True)
        else:
            st.success("🎉 No failed controls!")
    
    elif rpt_panel == "risk":
        st.markdown("---")
        st.markdown("### 🎯 Risk Score Breakdown")
        
        failed_findings = agg["failed_findings"]
        sev_weights = {"Critical": 10, "High": 7, "Medium": 4, "Low": 1}
        max_possible = total * 10
        actual_risk = sum(sev_weights.get(f.get("severity", "Medium"), 4) for f in failed_findings)
        risk_pct = int((1 - actual_risk / max_possible) * 100) if max_possible > 0 else 100
        
        st.markdown(f"""
        <div style="background: linear-gradient(145deg, rgba(22, 36, 64, 0.8), rgba(26, 45, 80, 0.6)); 
                    border: 1px solid rgba(0, 217, 255, 0.25); border-radius: 12px; padding: 1.5rem; margin-bottom: 1rem;">
            <h4 style="color: #00d9ff; margin: 0 0 1rem 0;">How the Risk Score is Calculated</h4>
            <p style="color: #f5f5f5;">The risk score uses <strong>weighted severity scoring</strong>:</p>
            <table style="color: #f5f5f5; width: 100%; border-collapse: collapse; margin: 0.5rem 0;">
                <tr style="border-bottom: 1px solid rgba(0,217,255,0.15);">
                    <td style="padding: 0.4rem; color: #e94560; font-weight: 600;">Critical</td>
                    <td style="padding: 0.4rem;">10 points per finding</td>
                    <td style="padding: 0.4rem; text-align: right;">{len([f for f in failed_findings if f.get('severity')=='Critical'])} findings = {len([f for f in failed_findings if f.get('severity')=='Critical']) * 10} pts</td>
                </tr>
                <tr style="border-bottom: 1px solid rgba(0,217,255,0.15);">
                    <td style="padding: 0.4rem; color: #ff6b6b; font-weight: 600;">High</td>
                    <td style="padding: 0.4rem;">7 points per finding</td>
                    <td style="padding: 0.4rem; text-align: right;">{len([f for f in failed_findings if f.get('severity')=='High'])} findings = {len([f for f in failed_findings if f.get('severity')=='High']) * 7} pts</td>
                </tr>
                <tr style="border-bottom: 1px solid rgba(0,217,255,0.15);">
                    <td style="padding: 0.4rem; color: #ffc107; font-weight: 600;">Medium</td>
                    <td style="padding: 0.4rem;">4 points per finding</td>
                    <td style="padding: 0.4rem; text-align: right;">{len([f for f in failed_findings if f.get('severity')=='Medium'])} findings = {len([f for f in failed_findings if f.get('severity')=='Medium']) * 4} pts</td>
                </tr>
                <tr>
                    <td style="padding: 0.4rem; color: #00d9ff; font-weight: 600;">Low</td>
                    <td style="padding: 0.4rem;">1 point per finding</td>
                    <td style="padding: 0.4rem; text-align: right;">{len([f for f in failed_findings if f.get('severity')=='Low'])} findings = {len([f for f in failed_findings if f.get('severity')=='Low'])} pts</td>
                </tr>
            </table>
            <div style="margin-top: 1rem; padding-top: 0.8rem; border-top: 1px solid rgba(0,217,255,0.2);">
                <span style="color: #8899bb;">Total Risk Points: <strong style="color: #e94560;">{actual_risk}</strong> / {max_possible} possible</span><br>
                <span style="color: #f5f5f5; font-size: 1.1rem;">Compliance Score: <strong style="color: #00d9ff; font-size: 1.3rem;">{score}%</strong></span>
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # Per-category risk contribution
        st.markdown("#### Risk Contribution by Category")
        cat_risks = {}
        for f in failed_findings:
            cat = f.get("category", "Unknown")
            weight = sev_weights.get(f.get("severity", "Medium"), 4)
            cat_risks[cat] = cat_risks.get(cat, 0) + weight
        
        if cat_risks:
            sorted_risks = sorted(cat_risks.items(), key=lambda x: x[1], reverse=True)
            max_risk = sorted_risks[0][1] if sorted_risks else 1
            for cat_n, risk_pts in sorted_risks:
                bar_w = int(risk_pts / max_risk * 100)
                st.markdown(f"""
                <div style="margin-bottom: 0.5rem;">
                    <div style="display: flex; justify-content: space-between;">
                        <span style="color: #f5f5f5; font-size: 0.85rem;">{cat_n}</span>
                        <span style="color: #e94560; font-weight: 700;">{risk_pts} risk points</span>
                    </div>
                    <div style="background: rgba(233,69,96,0.1); border-radius: 6px; height: 14px; overflow: hidden;">
                        <div style="background: linear-gradient(90deg, #e94560, #ff6b6b); width: {bar_w}%; height: 100%; border-radius: 6px; box-shadow: 0 0 6px rgba(233,69,96,0.5);"></div>
                    </div>
                </div>
                """, unsafe_allow_html=True)
    
    st.markdown("<br>", unsafe_allow_html=True)
    
    # ── Enhanced Charts Row 1 ──
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### Control Status Distribution")
        st.plotly_chart(create_status_donut(passed, failed, not_assessed, total, height=350), use_container_width=True)
    
    with col2:
        st.markdown("### Findings by Category")
        fig_bar = create_category_bar(agg["cat_passed"], agg["cat_failed"], height=350, barmode='group')
        if fig_bar:
            st.plotly_chart(fig_bar, use_container_width=True)
        else:
            st.info("No category data to display")
    
    # ── Enhanced Charts Row 2 ──
    chart2_col1, chart2_col2 = st.columns(2)
    
    with chart2_col1:
        st.markdown("### 🎯 Severity Distribution")
        sev_counts = agg["sev_counts"]
        if sum(sev_counts.values()) > 0:
            fig_sev = go.Figure(data=[go.Bar(
                x=list(sev_counts.keys()),
                y=list(sev_counts.values()),
                marker_color=BAR_COLORS_SEV,
                marker_line=dict(color=["#ff5a75", "#ff8888", "#ffd54f", "#00b8d4"], width=1),
                text=list(sev_counts.values()),
                textposition='auto',
                textfont=dict(color="#ffffff", size=15)
            )])
            fig_sev.update_layout(**chart_layout(
                height=300,
                xaxis=dict(tickfont=dict(color=CHART_FONT_COLOR, size=13), gridcolor='rgba(0, 217, 255, 0.06)'),
                yaxis=dict(tickfont=dict(color=CHART_FONT_COLOR), gridcolor=CHART_GRID_COLOR, title="Count", zeroline=False),
                margin=dict(t=20, b=40, l=40, r=20)
            ))
            st.plotly_chart(fig_sev, use_container_width=True)
        else:
            st.success("🎉 No failed controls to display!")
    
    with chart2_col2:
        st.markdown("### 🔒 Risk Score Gauge")
        st.plotly_chart(create_risk_gauge(score, height=300), use_container_width=True)
    
    st.markdown("---")
    
    # Detailed findings table
    st.markdown("### 📋 Detailed Findings")
    
    failed_findings = agg["failed_findings"]
    if failed_findings:
        df = pd.DataFrame(failed_findings)
        df = df[["control_id", "control_name", "category", "severity", "family"]]
        df.columns = ["Control ID", "Control Name", "Category", "Severity", "Family"]
        st.dataframe(df, use_container_width=True, hide_index=True)
        
        # Remediation guidance section
        if results.get("include_remediation", True):
            st.markdown("### 🔧 Remediation Guidance")
            
            # Group by severity for prioritization
            severity_order = ["Critical", "High", "Medium", "Low"]
            for sev in severity_order:
                sev_findings = [f for f in failed_findings if f.get("severity") == sev]
                if sev_findings:
                    sev_color = {"Critical": "#e94560", "High": "#ff6b6b", "Medium": "#ffc107", "Low": "#00d9ff"}.get(sev, "#c0c0c0")
                    st.markdown(f"#### <span style='color: {sev_color};'>🔴 {sev} Severity ({len(sev_findings)} findings)</span>", unsafe_allow_html=True)
                    
                    for finding in sev_findings:
                        control = get_control_by_id(finding["control_id"])
                        if control:
                            with st.expander(f"**{finding['control_id']}**: {finding['control_name']}"):
                                st.markdown(f"**Category:** {finding['category']}")
                                st.markdown(f"**Expected Result:** {control.expected_result}")
                                
                                if control.remediation_steps:
                                    st.markdown("**Remediation Steps:**")
                                    for step in control.remediation_steps:
                                        downtime_badge = " ⚠️ *Requires Downtime*" if step.requires_downtime else ""
                                        time_est = f" ({step.estimated_time})" if step.estimated_time else ""
                                        st.markdown(f"{step.step_number}. {step.description}{time_est}{downtime_badge}")
                                        if step.command:
                                            st.code(step.command, language=step.script_type or "powershell")
                                
                                # Generate PowerShell script button
                                ps_script = get_remediation_script(finding["control_id"], "powershell")
                                if ps_script and len(ps_script) > 100:
                                    st.download_button(
                                        label="📥 Download PowerShell Script",
                                        data=ps_script,
                                        file_name=f"remediate_{finding['control_id']}.ps1",
                                        mime="text/plain",
                                        key=f"dl_{finding['control_id']}"
                                    )
                                
                                if control.references:
                                    st.markdown(f"**References:** {', '.join(control.references)}")
    else:
        st.success("🎉 No failed controls found!")
    
    # Export options
    st.markdown("### 📥 Export Reports")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📄 Export Executive Summary", use_container_width=True):
            export_executive_summary(results, findings)
    
    with col2:
        if st.button("📊 Export Full Report", use_container_width=True):
            export_full_report(results, findings)
    
    with col3:
        if st.button("📋 Export POA&M", use_container_width=True):
            export_poam(results, findings)


def export_executive_summary(results, findings):
    """Generate and download executive summary as Word document"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    total = len(findings)
    passed = len([f for f in findings if f["status"] == "Passed"])
    failed = len([f for f in findings if f["status"] == "Failed"])
    score = calculate_risk_score(findings)

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('SOPRA Executive Summary')
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('On-Premise Risk Assessment Report')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

    doc.add_paragraph()
    doc.add_paragraph(f"Assessment Date: {results.get('timestamp', datetime.now().isoformat())[:10]}")
    doc.add_paragraph("Generated By: SOPRA v2.0.0")
    doc.add_paragraph()

    # Overview table
    doc.add_heading('Assessment Overview', level=2)
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate([
        ('Metric', 'Value'),
        ('Total Controls Assessed', str(total)),
        ('Controls Passed', str(passed)),
        ('Controls Failed', str(failed)),
        ('Overall Risk Score', f'{score}%'),
    ]):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = val
        if i == 0:
            for c in range(2):
                for run in tbl.rows[0].cells[c].paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()

    # Risk Summary
    doc.add_heading('Risk Summary', level=2)
    severity_counts = {}
    for f in findings:
        if f["status"] == "Failed":
            sev = f.get("severity", "Unknown")
            severity_counts[sev] = severity_counts.get(sev, 0) + 1
    for sev in ["Critical", "High", "Medium", "Low"]:
        count = severity_counts.get(sev, 0)
        if count > 0:
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(f"{sev}: ")
            r.bold = True
            p.add_run(f"{count} finding(s)")

    doc.add_paragraph()
    doc.add_heading('Key Recommendations', level=2)
    critical_high = [f for f in findings if f["status"] == "Failed" and f.get("severity") in ["Critical", "High"]][:5]
    for i, f in enumerate(critical_high, 1):
        p = doc.add_paragraph(style='List Number')
        r = p.add_run(f"{f['control_id']}: ")
        r.bold = True
        p.add_run(f"{f['control_name']} ({f['category']})")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('This report was generated by SOPRA — SAE On-Premise Risk Assessment Tool')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.download_button(
        label="📥 Download Executive Summary",
        data=buf.getvalue(),
        file_name=f"SOPRA_Executive_Summary_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.success("✅ Executive summary generated (.docx — open in Google Docs)!")


def export_full_report(results, findings):
    """Generate and download full assessment report as Word document"""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Title page
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('SOPRA Full Assessment Report')
    r.font.size = Pt(26)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('On-Premise Risk Assessment')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)

    doc.add_paragraph()
    doc.add_paragraph(f"Assessment Date: {results.get('timestamp', datetime.now().isoformat())[:10]}")
    doc.add_paragraph("Generated By: SOPRA v2.0.0")
    doc.add_page_break()

    # ── All Findings by Category ──
    doc.add_heading('All Findings', level=1)

    categories = {}
    for f in findings:
        cat = f.get("category", "Unknown")
        if cat not in categories:
            categories[cat] = []
        categories[cat].append(f)

    for cat, cat_findings in categories.items():
        doc.add_heading(cat, level=2)

        # Findings table for this category
        tbl = doc.add_table(rows=1 + len(cat_findings), cols=7)
        tbl.style = 'Light Grid Accent 1'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, hdr in enumerate(['Control ID', 'Control Name', 'Status', 'Severity',
                                   'NIST 800-53', 'CIS Controls', 'Control Family']):
            cell = tbl.rows[0].cells[ci]
            cell.text = hdr
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(8)

        for ri, f in enumerate(cat_findings):
            ctrl = get_control_by_id(f.get('control_id', ''))
            tbl.rows[ri + 1].cells[0].text = f.get('control_id', '')
            tbl.rows[ri + 1].cells[1].text = f.get('control_name', '')
            tbl.rows[ri + 1].cells[2].text = f.get('status', '')
            sev_text = f.get('severity', '') if f.get('status') == 'Failed' else ''
            tbl.rows[ri + 1].cells[3].text = sev_text or ''
            tbl.rows[ri + 1].cells[4].text = ', '.join(ctrl.nist_mapping) if ctrl and ctrl.nist_mapping else ''
            tbl.rows[ri + 1].cells[5].text = ctrl.cis_mapping if ctrl and ctrl.cis_mapping else ''
            tbl.rows[ri + 1].cells[6].text = ctrl.family.value if ctrl else f.get('family', '')
            for ci in range(7):
                for run in tbl.rows[ri + 1].cells[ci].paragraphs[0].runs:
                    run.font.size = Pt(8)
                    if f.get('status') == 'Failed':
                        run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        doc.add_paragraph()

    # ── Remediation Guidance ──
    failed = [f for f in findings if f["status"] == "Failed"]
    if failed:
        doc.add_page_break()
        doc.add_heading('Remediation Guidance', level=1)

        for f in failed:
            control = get_control_by_id(f["control_id"])
            if not control:
                continue

            doc.add_heading(f"{f['control_id']}: {f['control_name']}", level=2)

            p = doc.add_paragraph()
            r = p.add_run("Severity: ")
            r.bold = True
            sev = f.get('severity', 'Unknown')
            sr = p.add_run(sev)
            sr.bold = True
            sev_color_map = {"Critical": 0xc0392b, "High": 0xe74c3c, "Medium": 0xf39c12, "Low": 0x3498db}
            sr.font.color.rgb = RGBColor((sev_color_map.get(sev, 0x333333) >> 16) & 0xFF,
                                          (sev_color_map.get(sev, 0x333333) >> 8) & 0xFF,
                                          sev_color_map.get(sev, 0x333333) & 0xFF)

            p2 = doc.add_paragraph()
            r2 = p2.add_run("Expected Result: ")
            r2.bold = True
            p2.add_run(control.expected_result)

            # Full compliance mapping block
            if control.nist_mapping or control.cis_mapping or control.references:
                cm_tbl = doc.add_table(rows=1, cols=2)
                cm_tbl.style = 'Light Grid Accent 1'
                cm_tbl.rows[0].cells[0].text = 'Framework'
                cm_tbl.rows[0].cells[1].text = 'Mapping'
                for c in range(2):
                    for run in cm_tbl.rows[0].cells[c].paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(8)
                if control.nist_mapping:
                    row = cm_tbl.add_row()
                    row.cells[0].text = 'NIST 800-53 Rev 5'
                    row.cells[1].text = ', '.join(control.nist_mapping)
                if control.cis_mapping:
                    row = cm_tbl.add_row()
                    row.cells[0].text = 'CIS Controls v8'
                    row.cells[1].text = control.cis_mapping
                # Derive NIST CSF function from control family
                csf_map = {
                    "Access Control": "PR.AC — Protect: Access Control",
                    "Audit & Accountability": "DE.AE — Detect: Anomalies & Events",
                    "Audit and Accountability": "DE.AE — Detect: Anomalies & Events",
                    "Configuration Management": "PR.IP — Protect: Info Protection",
                    "Contingency Planning": "PR.IP — Protect: Info Protection",
                    "Identification & Authentication": "PR.AC — Protect: Access Control",
                    "Identification and Authentication": "PR.AC — Protect: Access Control",
                    "Incident Response": "RS.RP — Respond: Response Planning",
                    "Maintenance": "PR.MA — Protect: Maintenance",
                    "Media Protection": "PR.DS — Protect: Data Security",
                    "Physical & Environmental Protection": "PR.AC — Protect: Access Control",
                    "Physical & Environmental": "PR.AC — Protect: Access Control",
                    "Planning": "ID.GV — Identify: Governance",
                    "Personnel Security": "PR.AC — Protect: Access Control",
                    "Risk Assessment": "ID.RA — Identify: Risk Assessment",
                    "System & Services Acquisition": "ID.SC — Identify: Supply Chain",
                    "System & Communications Protection": "PR.DS — Protect: Data Security",
                    "System & Comm Protection": "PR.DS — Protect: Data Security",
                    "System & Information Integrity": "DE.CM — Detect: Continuous Monitoring",
                    "System & Info Integrity": "DE.CM — Detect: Continuous Monitoring",
                }
                family_name = control.family.value if control.family else ''
                csf_func = csf_map.get(family_name, '')
                if csf_func:
                    row = cm_tbl.add_row()
                    row.cells[0].text = 'NIST CSF 2.0'
                    row.cells[1].text = csf_func
                # Map NIST family abbreviation
                fam_abbr = FAMILY_ABBREV.get(family_name, '')
                if fam_abbr:
                    row = cm_tbl.add_row()
                    row.cells[0].text = 'Control Family'
                    row.cells[1].text = f"{fam_abbr} — {family_name}"
                if control.references:
                    row = cm_tbl.add_row()
                    row.cells[0].text = 'References'
                    row.cells[1].text = '; '.join(control.references)
                # Style all data rows
                for row in cm_tbl.rows[1:]:
                    for c in range(2):
                        for run in row.cells[c].paragraphs[0].runs:
                            run.font.size = Pt(8)
                doc.add_paragraph()

            if control.remediation_steps:
                doc.add_paragraph()
                rh = doc.add_paragraph()
                rr = rh.add_run("Remediation Steps:")
                rr.bold = True
                rr.font.size = Pt(11)

                for step in control.remediation_steps:
                    step_text = f"Step {step.step_number}: {step.description}"
                    if step.estimated_time:
                        step_text += f"  ({step.estimated_time})"
                    if step.requires_downtime:
                        step_text += "  [REQUIRES DOWNTIME]"
                    doc.add_paragraph(step_text, style='List Number')

                    if step.command:
                        cp = doc.add_paragraph()
                        cr = cp.add_run(step.command)
                        cr.font.name = 'Consolas'
                        cr.font.size = Pt(8.5)
                        cr.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
                        cp.paragraph_format.left_indent = Cm(1)

            doc.add_paragraph()

    # ── Appendix: Full Compliance Cross-Walk ──
    doc.add_page_break()
    doc.add_heading('Appendix A: Compliance Cross-Walk', level=1)
    doc.add_paragraph(
        'The following table provides a complete cross-walk of every assessed control '
        'to applicable compliance frameworks. Analysts can use this mapping to trace '
        'findings across NIST 800-53 Rev 5, CIS Controls v8, and NIST CSF 2.0.'
    )

    # Build cross-walk for ALL findings (passed and failed)
    csf_map = {
        "Access Control": "PR.AC",
        "Audit & Accountability": "DE.AE",
        "Audit and Accountability": "DE.AE",
        "Configuration Management": "PR.IP",
        "Contingency Planning": "PR.IP",
        "Identification & Authentication": "PR.AC",
        "Identification and Authentication": "PR.AC",
        "Incident Response": "RS.RP",
        "Maintenance": "PR.MA",
        "Media Protection": "PR.DS",
        "Physical & Environmental Protection": "PR.AC",
        "Physical & Environmental": "PR.AC",
        "Planning": "ID.GV",
        "Personnel Security": "PR.AC",
        "Risk Assessment": "ID.RA",
        "System & Services Acquisition": "ID.SC",
        "System & Communications Protection": "PR.DS",
        "System & Comm Protection": "PR.DS",
        "System & Information Integrity": "DE.CM",
        "System & Info Integrity": "DE.CM",
    }

    xw_headers = ['Control ID', 'Control Name', 'Status', 'Severity',
                   'Family', 'NIST 800-53', 'CIS v8', 'NIST CSF 2.0']
    xw_tbl = doc.add_table(rows=1, cols=len(xw_headers))
    xw_tbl.style = 'Light Grid Accent 1'
    xw_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, hdr in enumerate(xw_headers):
        xw_tbl.rows[0].cells[ci].text = hdr
        for run in xw_tbl.rows[0].cells[ci].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(7.5)

    for f in findings:
        ctrl = get_control_by_id(f.get('control_id', ''))
        row = xw_tbl.add_row()
        row.cells[0].text = f.get('control_id', '')
        row.cells[1].text = f.get('control_name', '')
        row.cells[2].text = f.get('status', '')
        row.cells[3].text = f.get('severity', '') if f.get('status') == 'Failed' else ''
        family_name = ctrl.family.value if ctrl else f.get('family', '')
        fam_abbr = FAMILY_ABBREV.get(family_name, '')
        row.cells[4].text = f"{fam_abbr}-{family_name}" if fam_abbr else family_name
        row.cells[5].text = ', '.join(ctrl.nist_mapping) if ctrl and ctrl.nist_mapping else ''
        row.cells[6].text = ctrl.cis_mapping if ctrl and ctrl.cis_mapping else ''
        row.cells[7].text = csf_map.get(family_name, '')
        is_failed = f.get('status') == 'Failed'
        for ci in range(len(xw_headers)):
            for run in row.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(7.5)
                if is_failed:
                    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)

    doc.add_paragraph()

    # Cross-walk legend
    doc.add_heading('Cross-Walk Legend', level=2)
    legend_items = [
        ('NIST 800-53 Rev 5', 'Security and Privacy Controls for Information Systems (e.g., AC-2, IA-5, SC-8)'),
        ('CIS Controls v8', 'Center for Internet Security Critical Security Controls (e.g., 5.1, 5.4, 8.2)'),
        ('NIST CSF 2.0', 'Cybersecurity Framework Functions — ID (Identify), PR (Protect), DE (Detect), RS (Respond), RC (Recover)'),
        ('Control Family', 'NIST 800-53 family abbreviation — AC, AU, CM, CP, IA, IR, MA, MP, PE, PL, PS, RA, SA, SC, SI'),
    ]
    for framework, desc in legend_items:
        p = doc.add_paragraph()
        r = p.add_run(f"{framework}: ")
        r.bold = True
        r.font.size = Pt(9)
        dr = p.add_run(desc)
        dr.font.size = Pt(9)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('This report was generated by SOPRA — SAE On-Premise Risk Assessment Tool')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    st.download_button(
        label="📥 Download Full Report",
        data=buf.getvalue(),
        file_name=f"SOPRA_Full_Report_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.success("✅ Full report generated (.docx — open in Google Docs)!")


def export_poam(results, findings):
    """Generate and download POA&M (Plan of Action & Milestones)"""
    failed = [f for f in findings if f["status"] == "Failed"]
    
    if not failed:
        st.info("No failed controls to include in POA&M")
        return
    
    # Create POA&M in CSV format for easy import
    poam_data = []
    for i, f in enumerate(failed, 1):
        control = get_control_by_id(f["control_id"])
        
        # Calculate estimated completion based on remediation steps
        total_time = ""
        if control and control.remediation_steps:
            total_time = " + ".join([s.estimated_time for s in control.remediation_steps if s.estimated_time])
        
        poam_data.append({
            "POA&M ID": f"SOPRA-{i:03d}",
            "Control ID": f["control_id"],
            "Weakness Name": f["control_name"],
            "Category": f["category"],
            "Severity": f.get("severity", "Unknown"),
            "Status": "Open",
            "Expected Completion": total_time or "TBD",
            "Responsible Party": "TBD",
            "NIST Mapping": ", ".join(control.nist_mapping) if control and control.nist_mapping else "",
            "Notes": control.expected_result if control else ""
        })
    
    df = pd.DataFrame(poam_data)
    csv = df.to_csv(index=False)
    
    st.download_button(
        label="📥 Download POA&M (CSV)",
        data=csv,
        file_name=f"SOPRA_POAM_{datetime.now().strftime('%Y%m%d')}.csv",
        mime="text/csv"
    )
    st.success("✅ POA&M generated!")


