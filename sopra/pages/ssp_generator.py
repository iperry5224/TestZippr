"""SOPRA SSP Generator Page (AI-Enhanced)"""
import streamlit as st
from datetime import datetime
import io

from sopra_controls import get_control_by_id
from sopra.theme import FAMILY_ABBREV
from sopra.utils import calculate_risk_score
from sopra.persistence import load_results_from_file
from sopra.isso.ai_engine import ai_generate_ssp_narrative

def render_ssp_generator():
    """Render the System Security Plan Generator page"""
    st.markdown("## 📜 System Security Plan (SSP) Generator")
    st.markdown("Auto-generate an SSP from your latest assessment results, mapped to NIST 800-53 control families.")

    results = st.session_state.get("opra_assessment_results")
    if not results or not results.get("findings"):
        results = load_results_from_file()
        if results:
            st.session_state.opra_assessment_results = results
    if not results or not results.get("findings"):
        st.warning("No assessment data available. Run an assessment first to generate an SSP.")
        return

    findings = results["findings"]
    total = len(findings)
    passed = len([f for f in findings if f["status"] == "Passed"])
    failed = len([f for f in findings if f["status"] == "Failed"])
    score = calculate_risk_score(findings)

    # SSP Overview cards
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        st.metric("Total Controls", total)
    with c2:
        st.metric("Implemented", passed, delta=f"{round(passed/max(total,1)*100)}%")
    with c3:
        st.metric("Not Implemented", failed, delta=f"-{round(failed/max(total,1)*100)}%", delta_color="inverse")
    with c4:
        st.metric("Risk Score", f"{score}%")

    st.markdown("---")

    # Control family implementation status
    st.markdown("### Control Family Implementation Status")
    family_status = {}
    for f in findings:
        ctrl = get_control_by_id(f.get("control_id", ""))
        fam = ctrl.family.value if ctrl else f.get("family", "Unknown")
        if fam not in family_status:
            family_status[fam] = {"Implemented": 0, "Partially Implemented": 0, "Not Implemented": 0, "controls": []}
        if f["status"] == "Passed":
            family_status[fam]["Implemented"] += 1
        elif f["status"] == "Failed":
            family_status[fam]["Not Implemented"] += 1
        else:
            family_status[fam]["Partially Implemented"] += 1
        family_status[fam]["controls"].append(f)

    for fam_name, data in sorted(family_status.items()):
        fam_total = data["Implemented"] + data["Not Implemented"] + data["Partially Implemented"]
        fam_pct = round((data["Implemented"] / max(fam_total, 1)) * 100)
        abbr = FAMILY_ABBREV.get(fam_name, "??")
        bar_color = "#00ff88" if fam_pct >= 80 else "#ffc107" if fam_pct >= 50 else "#e94560"

        with st.expander(f"{abbr} — {fam_name}  ({data['Implemented']}/{fam_total} implemented — {fam_pct}%)", expanded=False):
            st.markdown(f"""
            <div style="display: flex; gap: 1rem; margin-bottom: 0.5rem;">
                <div style="flex: 1; height: 8px; background: rgba(255,255,255,0.08); border-radius: 4px; overflow: hidden;">
                    <div style="width: {fam_pct}%; height: 100%; background: {bar_color}; border-radius: 4px;"></div>
                </div>
                <span style="color: {bar_color}; font-weight: 700; font-size: 0.85rem;">{fam_pct}%</span>
            </div>
            """, unsafe_allow_html=True)

            col_a, col_b, col_c = st.columns(3)
            with col_a:
                st.markdown(f"**Implemented:** {data['Implemented']}")
            with col_b:
                st.markdown(f"**Partial:** {data['Partially Implemented']}")
            with col_c:
                st.markdown(f"**Not Implemented:** {data['Not Implemented']}")

            for ctrl_f in data["controls"]:
                status_icon = {"Passed": "✅", "Failed": "❌"}.get(ctrl_f["status"], "⚠️")
                impl_status = "Implemented" if ctrl_f["status"] == "Passed" else "Not Implemented"
                ctrl_obj = get_control_by_id(ctrl_f.get("control_id", ""))
                nist = ", ".join(ctrl_obj.nist_mapping) if ctrl_obj and ctrl_obj.nist_mapping else ""
                st.markdown(f"{status_icon} **{ctrl_f.get('control_id', '')}** — {ctrl_f.get('control_name', '')}  |  _{impl_status}_  |  NIST: {nist}")

    st.markdown("---")

    # ── AI Narrative Generation ──────────────────────────────────────
    st.markdown("### 🤖 AI Control Implementation Narratives")
    st.caption("Generate professional SSP control implementation descriptions using AI. These are included in the exported document.")

    if st.button("🤖 Generate AI Narratives for All Controls", key="ssp_ai_narratives"):
        ai_narratives = {}
        progress = st.progress(0)
        status_text = st.empty()
        total_controls = len(findings)

        for idx, f in enumerate(findings):
            cid = f.get("control_id", "")
            ctrl_obj = get_control_by_id(cid)
            family = ctrl_obj.family.value if ctrl_obj else f.get("family", "")
            category = f.get("category", "")
            status = f.get("status", "Not Assessed")

            with st.spinner(f"Generating narrative for {cid}..."):
                narrative = ai_generate_ssp_narrative(cid, f.get("control_name", ""), family, category, status)
            ai_narratives[cid] = narrative

            progress.progress((idx + 1) / max(total_controls, 1))
            status_text.text(f"Generated {idx + 1} / {total_controls} narratives...")

        st.session_state["_ssp_ai_narratives"] = ai_narratives
        progress.empty()
        status_text.empty()
        st.success(f"AI generated **{len(ai_narratives)}** control implementation narratives. They will be included in the SSP export.")

    if st.session_state.get("_ssp_ai_narratives"):
        narr_count = len(st.session_state["_ssp_ai_narratives"])
        st.info(f"🤖 {narr_count} AI narratives ready for export. Click Generate SSP below to include them.")

        # Preview a sample
        with st.expander("Preview AI Narratives (first 3)"):
            for cid, narr in list(st.session_state["_ssp_ai_narratives"].items())[:3]:
                st.markdown(f"**{cid}:**")
                st.markdown(narr)
                st.markdown("---")

    st.markdown("---")

    # Generate SSP Document
    st.markdown("### 📥 Export SSP Document")
    if st.button("📜 Generate System Security Plan (.docx)", use_container_width=True, type="primary"):
        ai_narratives = st.session_state.get("_ssp_ai_narratives", {})
        _generate_ssp_docx(results, findings, family_status, ai_narratives)


def _generate_ssp_docx(results, findings, family_status, ai_narratives=None):
    """Generate and download SSP as Word document with optional AI narratives"""
    if ai_narratives is None:
        ai_narratives = {}
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    total = len(findings)
    passed = len([f for f in findings if f["status"] == "Passed"])
    failed = len([f for f in findings if f["status"] == "Failed"])
    score = calculate_risk_score(findings)

    doc = DocxDocument()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    # Title page
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('System Security Plan (SSP)')
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('SOPRA — SAE On-Premise Risk Assessment')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x2c, 0x3e, 0x50)
    doc.add_paragraph()
    for label, val in [("Assessment Date:", results.get('timestamp', '')[:10]),
                       ("Generated By:", "SOPRA v2.0.0"),
                       ("Classification:", "Internal Use Only"),
                       ("System Type:", "On-Premise Infrastructure")]:
        p = doc.add_paragraph()
        r = p.add_run(f"{label} ")
        r.bold = True
        p.add_run(val)
    doc.add_page_break()

    # Section 1: System Overview
    doc.add_heading('1. System Overview', level=1)
    doc.add_paragraph(
        'This System Security Plan documents the security controls implemented for the '
        'on-premise information system as assessed by SOPRA. It provides a comprehensive '
        'view of security control implementation status across all NIST 800-53 control families.'
    )
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = 'Light Grid Accent 1'
    for i, (k, v) in enumerate([('Total Controls Assessed', str(total)),
                                 ('Controls Implemented', f'{passed} ({round(passed/max(total,1)*100)}%)'),
                                 ('Controls Not Implemented', f'{failed} ({round(failed/max(total,1)*100)}%)'),
                                 ('Overall Risk Score', f'{score}%'),
                                 ('Assessment Source', results.get('source', 'N/A'))]):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v
        for run in tbl.rows[i].cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
        for run in tbl.rows[i].cells[1].paragraphs[0].runs:
            run.font.size = Pt(9)
    doc.add_paragraph()

    # Section 2: System Description
    doc.add_page_break()
    doc.add_heading('2. System Description', level=1)

    doc.add_heading('2.1 System Name and Identifier', level=2)
    doc.add_paragraph(
        'System Name: Satellite Telemetry and Environmental Reconnaissance System (STARS)\n'
        'System Identifier: STARS-OPS-2026\n'
        'System Owner: National Weather Intelligence Directorate (NWID)'
    )

    doc.add_heading('2.2 System Purpose and Mission', level=2)
    doc.add_paragraph(
        'The Satellite Telemetry and Environmental Reconnaissance System (STARS) is a mission-critical, '
        'on-premise information system that ingests, processes, and disseminates real-time satellite '
        'telemetry data for weather surveillance and severe storm detection. STARS serves as the primary '
        'ground-station processing node for the NOAA GOES-East and GOES-West geostationary satellite '
        'constellation, providing continuous atmospheric monitoring across the continental United States, '
        'the Gulf of Mexico, and the Atlantic hurricane basin.'
    )
    doc.add_paragraph(
        'The system receives raw spectral, infrared, and microwave sounding data at approximately '
        '14.2 terabytes per day via dedicated S-band and X-band downlink antennas. This data is '
        'processed through a multi-stage ingestion pipeline that performs radiometric calibration, '
        'geolocation correction, and quality assurance before being fed into operational storm-tracking '
        'algorithms and numerical weather prediction (NWP) models.'
    )

    doc.add_heading('2.3 System Architecture', level=2)
    doc.add_paragraph(
        'STARS operates within a physically isolated, SCIF-adjacent facility located at the National '
        'Weather Intelligence Center (NWIC) in Norman, Oklahoma. The system architecture comprises '
        'the following major components:'
    )
    arch_items = [
        ('Satellite Ground Station', 'Dual 9.1-meter parabolic antennas with automated tracking '
         'mounts, low-noise amplifiers (LNAs), and redundant RF downconverters for S-band (2.2 GHz) '
         'and X-band (8.2 GHz) reception.'),
        ('Data Ingestion Cluster', 'A 12-node high-performance computing (HPC) cluster running Red Hat '
         'Enterprise Linux 9, responsible for real-time frame synchronization, decommutation, and '
         'Level-0 data product generation. Sustained throughput: 850 Mbps.'),
        ('Processing & Analytics Tier', 'A 48-node GPU-accelerated compute farm (NVIDIA A100) executing '
         'mesoscale convective complex (MCC) detection, tropical cyclone intensity estimation (Dvorak '
         'technique), and ensemble storm-track forecasting at 2-minute refresh intervals.'),
        ('Storage Area Network (SAN)', 'A 4-petabyte Lustre parallel file system providing low-latency '
         'storage for active processing, backed by a 20-petabyte tape library (IBM TS4500) for long-term '
         'archival of historical telemetry records dating back to 2004.'),
        ('Dissemination Gateway', 'Secure, one-way data diode connections to the National Weather Service '
         '(NWS) Advanced Weather Interactive Processing System (AWIPS), the Federal Emergency Management '
         'Agency (FEMA) Integrated Public Alert and Warning System (IPAWS), and Department of Defense '
         'weather support channels via SIPRNet.'),
        ('Operations Center', 'A 24/7 staffed watch floor with 14 analyst workstations, a 12-panel '
         'video wall displaying real-time satellite imagery composites, and dedicated voice circuits to '
         'the Storm Prediction Center (SPC) and National Hurricane Center (NHC).'),
    ]
    arch_tbl = doc.add_table(rows=1, cols=2)
    arch_tbl.style = 'Light Grid Accent 1'
    arch_tbl.rows[0].cells[0].text = 'Component'
    arch_tbl.rows[0].cells[1].text = 'Description'
    for run in arch_tbl.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    for run in arch_tbl.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    for comp, desc in arch_items:
        row = arch_tbl.add_row()
        row.cells[0].text = comp
        row.cells[1].text = desc
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(8)
    doc.add_paragraph()

    doc.add_heading('2.4 Information Types and Data Sensitivity', level=2)
    doc.add_paragraph(
        'STARS processes the following categories of information:'
    )
    data_items = [
        ('Raw Satellite Telemetry', 'Unclassified // For Official Use Only (U//FOUO). Raw spectral and '
         'infrared radiance measurements received directly from GOES satellite instruments including the '
         'Advanced Baseline Imager (ABI) and Geostationary Lightning Mapper (GLM).'),
        ('Derived Storm Products', 'Unclassified. Processed meteorological products including mesocyclone '
         'detection signatures, tornado vortex signatures (TVS), quantitative precipitation estimates (QPE), '
         'and tropical cyclone track/intensity forecasts.'),
        ('Telemetry Housekeeping Data', 'Controlled Unclassified Information (CUI). Satellite bus health '
         'metrics, orbital ephemeris data, and ground station calibration coefficients that could reveal '
         'sensor capabilities and orbital parameters.'),
        ('Dissemination Routing Tables', 'CUI // NOFORN. Network routing configurations and priority '
         'schemas that govern data distribution to military and civilian consumers, including real-time '
         'feeds to deployed military weather teams.'),
    ]
    data_tbl = doc.add_table(rows=1, cols=2)
    data_tbl.style = 'Light Grid Accent 1'
    data_tbl.rows[0].cells[0].text = 'Information Type'
    data_tbl.rows[0].cells[1].text = 'Classification & Description'
    for run in data_tbl.rows[0].cells[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    for run in data_tbl.rows[0].cells[1].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    for dtype, ddesc in data_items:
        row = data_tbl.add_row()
        row.cells[0].text = dtype
        row.cells[1].text = ddesc
        for run in row.cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        for run in row.cells[1].paragraphs[0].runs:
            run.font.size = Pt(8)
    doc.add_paragraph()

    doc.add_heading('2.5 System Boundary and Interconnections', level=2)
    doc.add_paragraph(
        'The STARS authorization boundary encompasses all hardware, software, and network infrastructure '
        'within the NWIC Building 7 secure processing facility, including the antenna farm located within '
        'the fenced compound perimeter. The system maintains the following external interconnections:'
    )
    interconn = [
        ('NOAA NESDIS Command & Data Acquisition (CDA)', 'Satellite command uplink and telemetry '
         'coordination. Dedicated T3 circuit with Type-1 encryption (NSA-approved).'),
        ('NWS AWIPS II', 'One-way data diode for dissemination of derived storm products. '
         'Unidirectional fiber-optic link via Waterfall Security hardware enforcer.'),
        ('FEMA IPAWS', 'Emergency alert triggering interface for tornado and hurricane warnings. '
         'API-based connection over a dedicated MPLS circuit with mutual TLS authentication.'),
        ('DoD Joint Meteorological & Oceanographic (METOC)', 'SIPRNet connection for military weather '
         'support data feeds. Cross-domain solution (Raise the Bar approved) with content filtering.'),
    ]
    for name, desc in interconn:
        p = doc.add_paragraph()
        r = p.add_run(f'{name}: ')
        r.bold = True
        r.font.size = Pt(9)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9)

    doc.add_heading('2.6 Operational Environment', level=2)
    doc.add_paragraph(
        'STARS operates in a 24/7/365 continuous operations tempo to ensure uninterrupted weather '
        'surveillance capability. The system supports an average of 42 active users across three '
        'operational shifts, including satellite meteorologists, data analysts, system administrators, '
        'and network engineers. During severe weather events (hurricane landfall, tornado outbreaks, '
        'derecho events), surge staffing increases the active user count to approximately 85 personnel. '
        'The system has a Recovery Time Objective (RTO) of 15 minutes and a Recovery Point Objective '
        '(RPO) of zero data loss for real-time telemetry streams, reflecting its FIPS 199 HIGH '
        'availability impact categorization.'
    )

    # Section 3: Control Family Implementation (renumbered from 2)
    doc.add_page_break()
    doc.add_heading('3. Control Family Implementation Status', level=1)
    fam_tbl = doc.add_table(rows=1, cols=5)
    fam_tbl.style = 'Light Grid Accent 1'
    fam_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, hdr in enumerate(['Family', 'Implemented', 'Not Implemented', 'Total', 'Compliance %']):
        fam_tbl.rows[0].cells[ci].text = hdr
        for run in fam_tbl.rows[0].cells[ci].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for fam_name in sorted(family_status.keys()):
        data = family_status[fam_name]
        fam_total = data["Implemented"] + data["Not Implemented"] + data["Partially Implemented"]
        pct = round((data["Implemented"] / max(fam_total, 1)) * 100)
        abbr = FAMILY_ABBREV.get(fam_name, "??")
        row = fam_tbl.add_row()
        row.cells[0].text = f"{abbr} — {fam_name}"
        row.cells[1].text = str(data["Implemented"])
        row.cells[2].text = str(data["Not Implemented"])
        row.cells[3].text = str(fam_total)
        row.cells[4].text = f"{pct}%"
        for ci in range(5):
            for run in row.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(9)
                if pct < 50:
                    run.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
    doc.add_paragraph()

    # Section 4: Detailed Control Descriptions per Family
    doc.add_page_break()
    doc.add_heading('4. Control Implementation Details', level=1)
    for fam_name in sorted(family_status.keys()):
        data = family_status[fam_name]
        abbr = FAMILY_ABBREV.get(fam_name, "??")
        doc.add_heading(f"{abbr} — {fam_name}", level=2)
        for ctrl_f in data["controls"]:
            ctrl_obj = get_control_by_id(ctrl_f.get("control_id", ""))
            impl = "Implemented" if ctrl_f["status"] == "Passed" else "Not Implemented"
            p = doc.add_paragraph()
            r = p.add_run(f"{ctrl_f.get('control_id', '')}: {ctrl_f.get('control_name', '')}")
            r.bold = True
            r.font.size = Pt(10)

            det_tbl = doc.add_table(rows=0, cols=2)
            det_tbl.style = 'Light Grid Accent 1'
            for label, val in [("Implementation Status", impl),
                                ("Severity", ctrl_f.get("severity", "N/A") if ctrl_f["status"] == "Failed" else "N/A"),
                                ("NIST 800-53", ", ".join(ctrl_obj.nist_mapping) if ctrl_obj and ctrl_obj.nist_mapping else ""),
                                ("CIS Controls v8", ctrl_obj.cis_mapping if ctrl_obj and ctrl_obj.cis_mapping else ""),
                                ("Description", ctrl_obj.description if ctrl_obj else ""),
                                ("Expected Result", ctrl_obj.expected_result if ctrl_obj else ""),
                                ("Responsible Party", "TBD")]:
                if val:
                    row = det_tbl.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = str(val)
                    for run in row.cells[0].paragraphs[0].runs:
                        run.bold = True
                        run.font.size = Pt(8)
                    for run in row.cells[1].paragraphs[0].runs:
                        run.font.size = Pt(8)

            # Add AI-generated implementation narrative if available
            cid = ctrl_f.get("control_id", "")
            if cid in ai_narratives:
                p = doc.add_paragraph()
                r = p.add_run("Implementation Narrative (AI-Generated):")
                r.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0x00, 0x64, 0xc8)
                narr_p = doc.add_paragraph(ai_narratives[cid])
                narr_p.style.font.size = Pt(9)

            doc.add_paragraph()

    # Footer
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('System Security Plan generated by SOPRA — SAE On-Premise Risk Assessment')
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    st.download_button(
        label="📥 Download SSP (.docx)",
        data=buf.getvalue(),
        file_name=f"SOPRA_SSP_{datetime.now().strftime('%Y%m%d')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.success("✅ System Security Plan generated!")


