"""SSP Generator tab – build and export a System Security Plan document.

This module pays special attention to I/O error handling because the original
bug manifested as [Errno 5] Input/output error when reading uploaded files or
writing generated documents to disk/memory buffers.
"""

import io
import json
import datetime
from typing import Optional

import streamlit as st

from utils.file_io import safe_read_upload, safe_write_bytes, get_temp_dir

SSP_SECTIONS = [
    "System Name and Identifier",
    "System Description",
    "System Environment",
    "System Boundary",
    "System Interconnections",
    "Laws, Regulations, and Policies",
    "Roles and Responsibilities",
    "Information Types",
    "Security Categorization",
    "Security Control Implementation",
    "Continuous Monitoring",
    "Incident Response",
    "Contingency Plan Summary",
]


def _init_state():
    if "ssp_data" not in st.session_state:
        st.session_state.ssp_data = {section: "" for section in SSP_SECTIONS}
    if "ssp_metadata" not in st.session_state:
        st.session_state.ssp_metadata = {
            "organization": "",
            "author": "",
            "version": "1.0",
            "date": datetime.date.today().isoformat(),
        }


def render_ssp_generator():
    _init_state()
    st.header("System Security Plan Generator")
    st.markdown(
        "Complete each section below to build your SSP document. "
        "You can import a previously saved draft or export to JSON/DOCX."
    )

    _render_metadata()
    st.divider()
    _render_import()
    st.divider()
    _render_sections()
    st.divider()
    _render_export()


def _render_metadata():
    st.subheader("Document Metadata")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.ssp_metadata["organization"] = st.text_input(
            "Organization Name",
            value=st.session_state.ssp_metadata.get("organization", ""),
            key="ssp_org",
        )
        st.session_state.ssp_metadata["author"] = st.text_input(
            "Author",
            value=st.session_state.ssp_metadata.get("author", ""),
            key="ssp_author",
        )
    with col2:
        st.session_state.ssp_metadata["version"] = st.text_input(
            "Version",
            value=st.session_state.ssp_metadata.get("version", "1.0"),
            key="ssp_version",
        )
        st.session_state.ssp_metadata["date"] = st.text_input(
            "Date",
            value=st.session_state.ssp_metadata.get("date", ""),
            key="ssp_date",
        )


def _render_import():
    st.subheader("Import Draft")
    uploaded = st.file_uploader(
        "Upload a previously saved SSP draft (JSON)",
        type=["json"],
        key="ssp_upload",
    )
    if uploaded is not None:
        raw = safe_read_upload(uploaded)
        if raw is None:
            return
        try:
            data = json.loads(raw)
        except (json.JSONDecodeError, UnicodeDecodeError) as exc:
            st.error(f"Could not parse SSP draft: {exc}")
            return

        if "sections" in data:
            st.session_state.ssp_data.update(data["sections"])
        if "metadata" in data:
            st.session_state.ssp_metadata.update(data["metadata"])
        st.success("Draft imported successfully.")


def _render_sections():
    st.subheader("SSP Sections")
    for section in SSP_SECTIONS:
        st.session_state.ssp_data[section] = st.text_area(
            section,
            value=st.session_state.ssp_data.get(section, ""),
            height=120,
            key=f"ssp_section_{section}",
        )


def _render_export():
    st.subheader("Export SSP")
    col_json, col_docx = st.columns(2)

    with col_json:
        if st.button("Export as JSON", key="ssp_export_json"):
            _export_json()

    with col_docx:
        if st.button("Export as DOCX", key="ssp_export_docx"):
            _export_docx()


def _export_json():
    """Export SSP data as a JSON download."""
    try:
        payload = json.dumps(
            {
                "metadata": st.session_state.ssp_metadata,
                "sections": st.session_state.ssp_data,
            },
            indent=2,
        )
        st.download_button(
            "Download SSP JSON",
            data=payload,
            file_name="ssp_document.json",
            mime="application/json",
            key="ssp_dl_json",
        )
    except (IOError, OSError) as exc:
        st.error(f"Export failed (I/O error): {exc}")
    except Exception as exc:
        st.error(f"Export failed: {exc}")


def _export_docx():
    """Export SSP data as a DOCX file, with guarded I/O.

    Uses an in-memory BytesIO buffer to avoid filesystem [Errno 5] errors
    that occur on unreliable temp storage or container volume mounts.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        st.error(
            "python-docx is required for DOCX export. "
            "Install it with: `pip install python-docx`"
        )
        return

    try:
        doc = Document()

        style = doc.styles["Title"]
        style.font.size = Pt(24)

        meta = st.session_state.ssp_metadata
        doc.add_heading("System Security Plan", level=0)
        doc.add_paragraph(
            f"Organization: {meta.get('organization', 'N/A')}\n"
            f"Author: {meta.get('author', 'N/A')}\n"
            f"Version: {meta.get('version', '1.0')}\n"
            f"Date: {meta.get('date', 'N/A')}"
        )

        for section, content in st.session_state.ssp_data.items():
            doc.add_heading(section, level=1)
            doc.add_paragraph(content if content else "(Not yet completed)")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        docx_bytes = buf.getvalue()

        st.download_button(
            "Download SSP DOCX",
            data=docx_bytes,
            file_name="ssp_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="ssp_dl_docx",
        )
    except (IOError, OSError) as exc:
        st.error(
            f"DOCX generation encountered an I/O error: {exc}. "
            "Try exporting as JSON instead, or check that the temp directory is writable."
        )
    except Exception as exc:
        st.error(f"DOCX generation failed: {exc}")
