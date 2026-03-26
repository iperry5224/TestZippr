"""
Convert SAELAR_SOPRA_Consolidated_Briefing.md to Word (.docx) for Google Docs.
Run: python md_to_docx_briefing.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_para(doc, text, bold_parts=None):
    """Add paragraph, optionally bolding specific phrases."""
    p = doc.add_paragraph()
    if bold_parts:
        remaining = text
        for phrase in bold_parts:
            if phrase in remaining:
                before, _, after = remaining.partition(phrase)
                if before:
                    p.add_run(before)
                r = p.add_run(phrase)
                r.bold = True
                remaining = after
        if remaining:
            p.add_run(remaining)
    else:
        p.add_run(text)
    return p

def main():
    base = Path(__file__).parent
    doc = Document()

    # Title
    title = doc.add_heading('SAELAR & SOPRA — Consolidated Team Briefing', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Security Assessment & ISSO Automation Platform').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Prepared by: SOPSAEL Development Team | Date: February 22, 2026 | Classification: Internal Use Only').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Executive Overview
    doc.add_heading('Executive Overview', level=1)
    doc.add_paragraph(
        'SAELAR (Security Assessment Engine for Live AWS Resources) and SOPRA (SAE On-Premise Risk Assessment) '
        'form an integrated platform for automated security assessment and Governance, Risk & Compliance (GRC) workflow automation. '
        'Together they address the full cycle from live AWS assessment to ATO-ready documentation and remediation.'
    )
    t = doc.add_table(rows=3, cols=3)
    t.style = 'Table Grid'
    h = t.rows[0].cells
    h[0].text, h[1].text, h[2].text = 'Component', 'Role', 'Key Deliverable'
    for c in h:
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
    t.rows[1].cells[0].text, t.rows[1].cells[1].text, t.rows[1].cells[2].text = 'SAELAR', 'Automated NIST 800-53 assessment of live AWS resources', 'Real-time control pass/fail status, risk scores, evidence'
    t.rows[2].cells[0].text, t.rows[2].cells[1].text, t.rows[2].cells[2].text = 'SOPRA', 'ISSO workflow automation with 13 AI capabilities', 'POA&Ms, risk acceptances, SSP narratives, remediation scripts'
    doc.add_paragraph()
    doc.add_paragraph('This document consolidates technical and executive briefing materials for team distribution.')
    doc.add_paragraph()

    # Part 1: SAELAR
    doc.add_heading('Part 1: SAELAR — Automated AWS Assessment', level=1)
    doc.add_heading('What SAELAR Does', level=2)
    doc.add_paragraph(
        'SAELAR-53 performs automated NIST 800-53 Rev 5 security controls assessment against live AWS resources. '
        'It queries AWS APIs in real time to evaluate configuration state (e.g., CloudTrail enabled, IAM password policy, '
        'S3 encryption, GuardDuty status) and produces pass/fail results, risk scores, and audit-ready documentation.'
    )
    doc.add_paragraph(
        'Unlike static snapshots or manual checklists, SAELAR provides accurate, up-to-date security posture for continuous monitoring and ATO preparation.'
    )
    doc.add_paragraph()

    doc.add_heading('SAELAR Team Q&A', level=2)
    doc.add_paragraph('Q1: Mapping of APIs and mechanisms used to retrieve each piece of data — See the Data Source Crosswalk table below.')
    doc.add_paragraph('Q2: How does SAELAR align with NIST SP 800-53A Rev 5? — SAELAR aligns primarily through the Examine method. It does not perform Interview or active Test. See 800-53A Alignment below.')
    doc.add_paragraph()

    doc.add_heading('Data Source Crosswalk', level=2)
    crosswalk = [
        ('Account identity & credentials', 'sts.get_caller_identity()', 'AWS STS'),
        ('IAM users, policies, MFA', 'iam.list_users(), iam.list_access_keys(), etc.', 'AWS IAM'),
        ('S3 buckets, encryption, policies', 's3.list_buckets(), s3.get_bucket_encryption(), etc.', 'AWS S3'),
        ('EC2, VPCs, security groups', 'ec2.describe_instances(), ec2.describe_vpcs(), etc.', 'AWS EC2'),
        ('CloudTrail trails & status', 'cloudtrail.describe_trails(), etc.', 'AWS CloudTrail'),
        ('CloudWatch alarms', 'cloudwatch.describe_alarms()', 'AWS CloudWatch'),
        ('AWS Config', 'config.describe_configuration_recorders(), etc.', 'AWS Config'),
        ('Security Hub', 'securityhub.get_findings()', 'AWS Security Hub'),
        ('GuardDuty', 'guardduty.list_findings()', 'AWS GuardDuty'),
        ('Inspector', 'inspector2.list_findings()', 'AWS Inspector v2'),
        ('KMS, Secrets Manager', 'kms.list_keys(), secretsmanager.list_secrets()', 'AWS KMS / Secrets Manager'),
        ('RDS, Backup, ELB, ACM', 'rds., backup., elbv2., acm. APIs', 'AWS RDS/Backup/ELB/ACM'),
        ('CISA KEV', 'HTTP GET to CISA KEV Catalog', 'CISA (external)'),
        ('AI remediation', 'bedrock-runtime.InvokeModel()', 'AWS Bedrock'),
    ]
    tbl = doc.add_table(rows=len(crosswalk)+1, cols=3)
    tbl.style = 'Table Grid'
    for j, hdr in enumerate(['Data / Capability', 'API or Mechanism', 'AWS Service / Source']):
        tbl.rows[0].cells[j].text = hdr
        for r in tbl.rows[0].cells[j].paragraphs[0].runs:
            r.bold = True
    for i, (d, a, s) in enumerate(crosswalk):
        tbl.rows[i+1].cells[0].text, tbl.rows[i+1].cells[1].text, tbl.rows[i+1].cells[2].text = d, a, s
    doc.add_paragraph()

    doc.add_heading('NIST SP 800-53A Rev 5 Alignment', level=2)
    align_tbl = doc.add_table(rows=4, cols=3)
    align_tbl.style = 'Table Grid'
    for j, h in enumerate(['Assessment Method', 'Description', 'SAELAR Support']):
        align_tbl.rows[0].cells[j].text = h
        for r in align_tbl.rows[0].cells[j].paragraphs[0].runs:
            r.bold = True
    align_data = [
        ('Examine', 'Review documentation, design specs, mechanism behavior', 'Primary. SAELAR uses AWS APIs to examine configuration state.'),
        ('Interview', 'Discuss with personnel', 'Not supported.'),
        ('Test', 'Execute mechanisms or simulate attacks', 'Limited. No active tests (e.g., pen tests).'),
    ]
    for i, row in enumerate(align_data):
        for j, val in enumerate(row):
            align_tbl.rows[i+1].cells[j].text = val
    doc.add_paragraph('Summary: SAELAR performs automated examine-based assessments. For full 800-53A compliance, supplement with interview and test activities.')
    doc.add_paragraph()

    # Part 2: SOPRA
    doc.add_heading('Part 2: SOPRA — ISSO Workflow Automation with AI', level=1)
    doc.add_heading('The Problem', level=2)
    doc.add_paragraph(
        'Information System Security Officers (ISSOs) spend 60–70% of their time on repetitive, document-heavy tasks. '
        'Conservative estimate for a single ATO cycle: 400–600 hours of ISSO labor.'
    )
    doc.add_paragraph()

    doc.add_heading('The Solution: SOPRA AI', level=2)
    doc.add_paragraph('SOPRA integrates 13 distinct AI capabilities powered by AWS Bedrock. Every AI feature:')
    for b in [
        'Operates with a single button click',
        'Produces auditor-ready output using formal RMF/ATO language',
        'Tags all AI-generated content with a visible badge for transparency',
        'Falls back to deterministic templates when AI is unavailable',
        'Sends only aggregate assessment data to AI — no PII, no raw evidence',
    ]:
        doc.add_paragraph(b, style='List Bullet')
    doc.add_paragraph()

    doc.add_heading('AI Feature Catalog', level=2)
    features = [
        ('1. AI-Powered POA&M Generation', '4–6 hours saved per cycle'),
        ('2. AI-Drafted Risk Acceptance Justifications', '15–25 min per acceptance'),
        ('3. AI Evidence Sufficiency Analysis', '6–10 hours per audit prep'),
        ('4. AI Control Inheritance Auto-Classification', '4–6 hours per boundary review'),
        ('5. AI STIG-to-SOPRA Auto-Mapping', '2–3 hours per STIG import'),
        ('6. AI Incident-to-Finding NLP Correlation', '20–40 min per incident'),
        ('7. AI Natural Language Crosswalk Queries', '15–25 min per query'),
        ('8. AI-Written SSP Control Implementation Narratives', '150–300+ hours per ATO cycle'),
        ('9. AI Risk-Based Assessment Schedule Optimization', '2–3 hours per quarterly review'),
        ('10. AI Approval Package Summary Drafting', '15–25 min per approval'),
        ('11. AI Chat Assistant ("Chad")', 'Replaces hours of manual research'),
        ('12. AI Remediation Plans & Attack Chain Detection', '8–15 hours per remediation cycle'),
        ('13. AI Automated Ticket Generation', '2–4 hours per assessment'),
    ]
    for name, saved in features:
        p = doc.add_paragraph()
        p.add_run(name + ': ').bold = True
        p.add_run(saved)
    doc.add_paragraph()

    doc.add_heading('Total Impact (SOPRA AI)', level=2)
    impact_tbl = doc.add_table(rows=8, cols=2)
    impact_tbl.style = 'Table Grid'
    impact_data = [
        ('Total ISSO hours saved per ATO cycle', '~360 hours'),
        ('Percentage reduction in manual effort', '~90%'),
        ('Full-time equivalent freed up', '~2.25 FTEs per ATO cycle'),
        ('Time to generate full SSP (200 controls)', 'Minutes vs. weeks'),
        ('Assessment categories covered', '20 (200 controls)'),
        ('Compliance frameworks mapped', 'NIST 800-53 Rev 5 + CIS v8'),
        ('FIPS 199 baselines supported', 'Low (75) / Moderate (155) / High (200)'),
    ]
    for i, (k, v) in enumerate(impact_data):
        impact_tbl.rows[i].cells[0].text, impact_tbl.rows[i].cells[1].text = k, v
        for r in impact_tbl.rows[i].cells[0].paragraphs[0].runs:
            r.bold = True
    doc.add_paragraph()

    doc.add_heading('Security & Compliance Assurance', level=2)
    sec_tbl = doc.add_table(rows=8, cols=2)
    sec_tbl.style = 'Table Grid'
    sec_data = [
        ('Does AI have access to classified data?', 'No. Only aggregate metrics. No PII, evidence, or classified content.'),
        ('What if AI is wrong?', 'Human review required. All AI content tagged with badges.'),
        ('What if AI is unavailable?', 'Deterministic fallback. Zero functionality lost.'),
        ('Does this meet RMF requirements?', 'Yes. Formatted to NIST RMF standards.'),
        ('Is the AI auditable?', 'Yes. Full audit trail of approvals.'),
        ('Can this run air-gapped?', 'Yes. AI is the only optional network dependency.'),
        ('What about FedRAMP?', 'AWS Bedrock is FedRAMP-authorized.'),
    ]
    for i, (q, a) in enumerate(sec_data):
        sec_tbl.rows[i].cells[0].text, sec_tbl.rows[i].cells[1].text = q, a
    doc.add_paragraph()

    doc.add_heading('Implementation Roadmap', level=2)
    road_tbl = doc.add_table(rows=6, cols=3)
    road_tbl.style = 'Table Grid'
    for j, h in enumerate(['Phase', 'Timeline', 'Deliverable']):
        road_tbl.rows[0].cells[j].text = h
        for r in road_tbl.rows[0].cells[j].paragraphs[0].runs:
            r.bold = True
    road_data = [
        ('Phase 1: Deploy', 'Week 1', 'Install, configure AWS, run first assessment'),
        ('Phase 2: Baseline', 'Week 2', 'FIPS 199 categorization, full 200-control assessment'),
        ('Phase 3: ISSO Onboard', 'Weeks 3–4', 'Train on AI features, begin automation'),
        ('Phase 4: ATO Prep', 'Weeks 5–8', 'AI-assisted SSP, evidence collection, approval packages'),
        ('Phase 5: Continuous', 'Ongoing', 'AI-optimized scheduling, continuous monitoring'),
    ]
    for i, row in enumerate(road_data):
        for j, val in enumerate(row):
            road_tbl.rows[i+1].cells[j].text = val
    doc.add_paragraph()

    doc.add_heading('Bottom Line', level=2)
    doc.add_paragraph('SAELAR provides real-time, automated NIST 800-53 assessment of live AWS resources.')
    doc.add_paragraph('SOPRA transforms the ISSO role from document author to security decision-maker through 13 AI capabilities.')
    p = doc.add_paragraph()
    p.add_run('Together: One platform. 200 controls. ~360 hours saved per ATO cycle.').bold = True
    doc.add_paragraph()

    doc.add_heading('SAELAR vs. AWS Native Tools — Comparative Analysis', level=3)
    comp_tbl = doc.add_table(rows=7, cols=3)
    comp_tbl.style = 'Table Grid'
    for j, h in enumerate(['Dimension', 'AWS Native Tools (Config, Security Hub, Inspector, Audit Manager)', 'SAELAR / Time & Cost Advantage']):
        comp_tbl.rows[0].cells[j].text = h
        for r in comp_tbl.rows[0].cells[j].paragraphs[0].runs:
            r.bold = True
    comp_data = [
        ('Assessment framework', 'CIS, AWS Foundational — requires manual NIST 800-53 mapping', 'Native NIST 800-53 Rev 5. Saves 20–40 hrs mapping findings.'),
        ('Tool integration', '4–5 separate consoles', 'Single platform, one dashboard. Saves 10–15 hrs consolidating findings.'),
        ('ATO deliverables', 'Evidence only — ISSO manually writes SSP, POA&M', 'Auto-generates SSP, POA&M, risk drafts. Saves 200–360 hrs.'),
        ('Setup & configuration', 'Multiple services, cross-account setup', 'Single credential setup. Saves 8–16 hrs setup; 2–4 hrs per account.'),
        ('Ongoing monitoring', 'Custom dashboards, manual correlation', 'Unified risk score, control-family view. Saves 5–10 hrs/month.'),
        ('Licensing & cost', 'Config + Security Hub + Inspector + Audit Manager fees', 'Leverages same AWS APIs; runs on-premise. Lower TCO.'),
    ]
    for i, (d, aws, saelar) in enumerate(comp_data):
        comp_tbl.rows[i+1].cells[0].text = d
        comp_tbl.rows[i+1].cells[1].text = aws
        comp_tbl.rows[i+1].cells[2].text = saelar
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Summary: ').bold = True
    p.add_run('SAELAR consolidates what AWS native tools do separately into one NIST 800-53–focused workflow. Combined with SOPRA AI, expect 250–400+ hours saved per ATO cycle versus AWS native tools alone, plus reduced licensing and operational overhead.')
    doc.add_paragraph()
    doc.add_paragraph('SOPSAEL — SAELAR + SOPRA Integrated Platform')
    doc.add_paragraph('For questions or demonstration requests, contact the development team.')

    out_path = base / 'SAELAR_SOPRA_Consolidated_Briefing.docx'
    doc.save(out_path)
    print(f"Created: {out_path}")
    print("Upload to Google Drive (drive.google.com) and open with Google Docs.")

if __name__ == '__main__':
    main()
